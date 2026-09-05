from __future__ import annotations

import copy
from contextlib import redirect_stdout
import hashlib
import io
import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

import install_adapters as adapters  # noqa: E402
from check_dependencies import check_dependencies  # noqa: E402
from content_contract import load_practice_task_contract, practice_tasks_to_content  # noqa: E402
from content_quality import validate_collection, validate_content  # noqa: E402
from cross_artifact_quality import validate_cross_artifact  # noqa: E402
from generate_work_orders import DEFAULT_TEMPLATE, generate, main as generate_work_orders_main  # noqa: E402
from install import install as install_skill  # noqa: E402
from validate_output import validate_document  # noqa: E402


def _digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def _handoff(domain: str = "software") -> dict:
    nursing = domain == "nursing"
    return {
        "contract_version": "1.0",
        "course_name": "基础护理技术" if nursing else "数据库技术",
        "practice_hours": 2,
        "granularity": "per_task",
        "tasks": [
            {
                "task_id": "PT-REAL-01",
                "project_id": "P-REAL-01",
                "title": "完成生命体征测量记录" if nursing else "设计客户订单数据模型",
                "lesson_ids": ["L03", "L04"],
                "practice_hours": 2,
                "scenario": "在护理模拟环境中完成测量并形成记录。" if nursing else "根据客户订单业务说明设计数据模型。",
                "objectives": ["按规范完成记录" if nursing else "识别业务对象和关系"],
                "required_inputs": ["操作任务单" if nursing else "订单业务说明"],
                "tools_or_materials": ["护理模拟设备", "护理记录单"] if nursing else ["建模工具", "订单业务说明"],
                "steps": ["核对任务和准备条件", "实施操作并记录过程", "整理材料并复核提交"],
                "deliverables": ["生命体征测量记录表", "交接说明"] if nursing else ["客户订单概念模型图", "模型说明页"],
                "acceptance_criteria": ["记录项目完整", "记录可追溯到任务单"] if nursing else ["模型对象覆盖业务说明", "关系和说明保持一致"],
                "safety_or_compliance": ["遵守操作安全和隐私要求"] if nursing else ["遵守数据保密要求"],
            }
        ],
    }


def _valid_content() -> dict:
    value = json.loads((ROOT / "examples" / "software.example.json").read_text(encoding="utf-8"))[0]
    value["practice_task_id"] = "PT-REAL-01"
    value["task_title"] = value["project_name"]
    value["project_id"] = "P-REAL-01"
    value["safety_or_compliance"] = ["遵守数据保密要求"]
    return value


class WorkOrderPhase2Tests(unittest.TestCase):
    def test_one_canonical_practice_schema_and_handoff_trace_fields(self) -> None:
        shared = json.loads((ROOT.parents[1] / "schemas" / "shared" / "practice-task-contract.schema.json").read_text(encoding="utf-8"))
        lesson_entry = json.loads((ROOT.parent.parent / "教案生成器" / "lesson-plan-docx-generator" / "schemas" / "practice-task-contract.schema.json").read_text(encoding="utf-8"))
        self.assertIn("$defs", shared)
        self.assertEqual(lesson_entry["$ref"], shared["$id"])
        self.assertNotIn("$defs", lesson_entry)
        handoff = _handoff()
        content = practice_tasks_to_content(handoff, major="软件技术", class_or_audience="高职一年级", allow_non_production=True)[0]
        self.assertEqual(content["practice_task_id"], "PT-REAL-01")
        self.assertEqual(content["lesson_ids"], ["L03", "L04"])
        self.assertEqual(content["practice_hours"], 2)
        self.assertEqual(content["project_id"], "P-REAL-01")
        self.assertEqual(content["safety_or_compliance"], ["遵守数据保密要求"])

    def test_cross_artifact_valid_and_each_hard_gate_rejects_mismatch(self) -> None:
        handoff = _handoff()
        content = practice_tasks_to_content(handoff, major="软件技术", class_or_audience="高职一年级", allow_non_production=True)[0]
        report = validate_cross_artifact(handoff, content)
        self.assertEqual(report["status"], "pass", report)
        single_list_report = validate_cross_artifact(handoff, [content])
        self.assertEqual(single_list_report["status"], "pass", single_list_report)
        for field, replacement in (
            ("practice_task_id", "PT-WRONG"),
            ("lesson_ids", ["L99"]),
            ("practice_hours", 1),
            ("task_title", "编制会计凭证"),
            ("safety_or_compliance", []),
        ):
            broken = copy.deepcopy(content)
            broken[field] = replacement
            self.assertEqual(validate_cross_artifact(handoff, broken)["status"], "fail", field)

        missing_deliverable = copy.deepcopy(content)
        for item in missing_deliverable["task_items"]:
            item["deliverables"] = ["无关材料"]
        self.assertEqual(validate_cross_artifact(handoff, missing_deliverable)["status"], "fail")
        nursing = _handoff("nursing")
        nursing_content = practice_tasks_to_content(nursing, major="护理", class_or_audience="高职一年级", allow_non_production=True)[0]
        nursing_content["task_items"][0]["tools_or_materials"] = ["MySQL Workbench"]
        self.assertEqual(validate_cross_artifact(nursing, nursing_content)["status"], "fail")

    def test_collection_cross_artifact_enforces_one_to_one_workorders(self) -> None:
        handoff = _handoff()
        second = copy.deepcopy(handoff["tasks"][0])
        second["task_id"] = "PT-REAL-02"
        second["project_id"] = "P-REAL-02"
        second["title"] = "设计库存商品数据模型"
        second["scenario"] = "根据库存商品业务资料设计数据模型。"
        handoff["practice_hours"] = 4
        handoff["tasks"] = [handoff["tasks"][0], second]
        contents = practice_tasks_to_content(
            handoff,
            major="软件技术",
            class_or_audience="高职一年级",
            allow_non_production=True,
        )
        report = validate_cross_artifact(handoff, contents)
        self.assertEqual(report["status"], "pass", report)
        self.assertEqual(report["metrics"]["practice_task_count"], 2)
        self.assertEqual(report["metrics"]["work_order_count"], 2)
        self.assertTrue(report["checks"]["one_to_one_mapping"]["status"] == "pass")
        missing = validate_cross_artifact(handoff, contents[:1])
        self.assertEqual(missing["status"], "fail")
        self.assertEqual(missing["checks"]["one_to_one_mapping"]["status"], "fail")

    def test_cli_collection_handoff_emits_one_workorder_per_practice_task(self) -> None:
        handoff = _handoff()
        second = copy.deepcopy(handoff["tasks"][0])
        second["task_id"] = "PT-REAL-02"
        second["project_id"] = "P-REAL-02"
        second["title"] = "设计库存商品数据模型"
        second["scenario"] = "根据库存商品业务资料设计数据模型。"
        handoff["practice_hours"] = 4
        handoff["tasks"] = [handoff["tasks"][0], second]
        contents = practice_tasks_to_content(
            handoff,
            major="软件技术",
            class_or_audience="高职一年级",
            allow_non_production=True,
        )
        with tempfile.TemporaryDirectory(prefix="workorder-phase2-cli-collection-") as temp_name:
            root = Path(temp_name)
            handoff_path = root / "handoff.json"
            content_path = root / "content.json"
            output_dir = root / "output"
            handoff_path.write_text(json.dumps(handoff, ensure_ascii=False), encoding="utf-8")
            content_path.write_text(json.dumps(contents, ensure_ascii=False), encoding="utf-8")
            stdout = io.StringIO()
            with redirect_stdout(stdout):
                status = generate_work_orders_main(
                    [
                        "--content-json",
                        str(content_path),
                        "--practice-task-json",
                        str(handoff_path),
                        "--output-dir",
                        str(output_dir),
                        "--replace",
                        "--json",
                    ]
                )
            self.assertEqual(status, 0, stdout.getvalue())
            report = json.loads(stdout.getvalue())
            self.assertEqual(report["status"], "pass", report)
            self.assertEqual(len(report["outputs"]), 2)
            self.assertEqual(
                {item["practice_task_id"] for item in report["outputs"]},
                {"PT-REAL-01", "PT-REAL-02"},
            )
            self.assertEqual(len(list(output_dir.glob("*.docx"))), 2)

    def test_executability_deliverable_acceptance_domain_and_repetition_categories(self) -> None:
        valid = _valid_content()
        report = validate_content(valid)
        self.assertEqual(report["status"], "pass", report)
        self.assertEqual(report["categories"]["executability"], "pass")
        vague = copy.deepcopy(valid)
        vague["task_items"][0]["description"] = "认真完成任务"
        self.assertEqual(validate_content(vague)["categories"]["executability"], "fail")
        vague_deliverable = copy.deepcopy(valid)
        vague_deliverable["task_items"][0]["deliverables"] = ["完成任务"]
        self.assertEqual(validate_content(vague_deliverable)["categories"]["deliverable"], "fail")
        vague_criterion = copy.deepcopy(valid)
        vague_criterion["task_items"][0]["acceptance_criteria"] = ["认真完成任务"]
        self.assertEqual(validate_content(vague_criterion)["categories"]["acceptance"], "fail")
        nursing = _valid_content()
        nursing["major"] = "护理"
        nursing["course_name"] = "基础护理技术"
        nursing["task_items"][0]["tools_or_materials"] = ["MySQL Workbench"]
        self.assertEqual(validate_content(nursing)["categories"]["cross_domain"], "fail")

    def test_workorder_content_is_one_two_hour_practice_task(self) -> None:
        invalid = _valid_content()
        invalid["practice_hours"] = 4
        report = validate_content(invalid)
        self.assertEqual(report["categories"]["practice_hours_unit"], "fail")
        self.assertTrue(any("exactly one 2-hour" in error for error in report["errors"]))

    def test_dynamic_one_to_five_rows_keep_scores_results_and_template_hash(self) -> None:
        before = _digest(DEFAULT_TEMPLATE)
        base = _valid_content()
        with tempfile.TemporaryDirectory(prefix="workorder-phase2-rows-") as temp_name:
            for count in range(1, 6):
                content = copy.deepcopy(base)
                items = []
                for index in range(count):
                    item = copy.deepcopy(base["task_items"][index % len(base["task_items"])])
                    item["title"] = f"{item['title']}{index + 1}"
                    item["description"] += f"（第{index + 1}项）"
                    item["deliverables"] = [f"{value}{index + 1}" for value in item["deliverables"]]
                    item["acceptance_criteria"] = [f"{value}{index + 1}" for value in item["acceptance_criteria"]]
                    item["score"] = 90 // count + (1 if index < 90 % count else 0)
                    items.append(item)
                content["task_items"] = items
                report = generate([content], output_dir=Path(temp_name) / str(count), template=DEFAULT_TEMPLATE)
                self.assertEqual(report["status"], "pass", report)
                output = Path(report["outputs"][0]["path"])
                document_report = validate_document(output, content)
                self.assertEqual(document_report["status"], "pass", document_report)
                self.assertEqual(document_report["metrics"]["task_rows"], count)
                self.assertEqual(document_report["metrics"]["student_result_cells_blank"], count)
        self.assertEqual(_digest(DEFAULT_TEMPLATE), before)

    def test_output_qa_accepts_reserved_result_lines_but_keeps_fixed_attendance_title(self) -> None:
        content = _valid_content()
        with tempfile.TemporaryDirectory(prefix="workorder-phase2-output-boundary-") as temp_name:
            output = Path(generate([content], output_dir=Path(temp_name), template=DEFAULT_TEMPLATE)["outputs"][0]["path"])
            document = Document(str(output))
            document.tables[0].rows[2].cells[3].paragraphs[0].text = "________________"
            document.save(str(output))
            reserved = validate_document(output, content)
            self.assertEqual(reserved["status"], "pass", reserved)
            self.assertEqual(reserved["metrics"]["student_result_cells_unanswered"], 3)

            document = Document(str(output))
            document.tables[0].rows[1].cells[1].paragraphs[0].text = "出勤"
            document.save(str(output))
            renamed = validate_document(output, content)
            self.assertEqual(renamed["status"], "fail")
            self.assertTrue(any("课堂考勤" in error for error in renamed["errors"]))

    def test_install_adapters_and_runtime_modes_are_transactional(self) -> None:
        with tempfile.TemporaryDirectory(prefix="workorder-phase2-install-") as temp_name:
            root = Path(temp_name)
            installed = install_skill(ROOT, root / "skills")
            self.assertTrue((installed / "schemas" / "shared" / "practice-task-contract.schema.json").is_file())
            with self.assertRaises(FileExistsError):
                install_skill(ROOT, root / "skills")
            install_skill(ROOT, root / "skills", replace=True)

            minimal = root / "minimal"
            adapters.install(ROOT, minimal, adapters=["all"])
            minimal_engine = minimal / adapters.ENGINE_NAME
            self.assertEqual(adapters._detect_existing_engine_mode(minimal_engine, ROOT), "minimal")
            for name, paths in adapters.ADAPTER_PATHS.items():
                for relative in paths:
                    self.assertTrue((minimal / relative).is_file(), name)
            with self.assertRaises(ValueError):
                adapters.install(ROOT, minimal, adapters=["agents"], copy_engine=True)
            adapters.install(ROOT, minimal, adapters=["agents"], copy_engine=True, replace=True)
            self.assertEqual(adapters._detect_existing_engine_mode(minimal_engine, ROOT), "full-current")
            self.assertTrue((minimal_engine / adapters.SHARED_SCHEMA).is_file())
            (minimal_engine / adapters.ENGINE_STATE_FILE).unlink()
            self.assertEqual(adapters._detect_existing_engine_mode(minimal_engine, ROOT), "full-stale")
            with self.assertRaises(ValueError):
                adapters.install(ROOT, minimal, adapters=["agents"])

            full = root / "full"
            adapters.install(ROOT, full, adapters=["all"], copy_engine=True)
            full_engine = full / adapters.ENGINE_NAME
            self.assertEqual(adapters._detect_existing_engine_mode(full_engine, ROOT), "full-current")
            self.assertIn(adapters.MARKER_START, (full / "CLAUDE.md").read_text(encoding="utf-8"))
            self.assertIn(adapters.ENGINE_NAME, (full / ".aider.conf.yml").read_text(encoding="utf-8"))
            before = (full_engine / adapters.ENGINE_STATE_FILE).read_bytes()
            adapters.install(ROOT, full, adapters=["agents"])
            self.assertEqual((full_engine / adapters.ENGINE_STATE_FILE).read_bytes(), before)

    def test_dependency_doctor_is_read_only_and_reports_install_hint_contract(self) -> None:
        report = check_dependencies()
        self.assertTrue(report["read_only"])
        self.assertEqual(report["status"], "pass", report)
        self.assertIsNone(report["install_hint"])

    def test_collection_repetition_scope_does_not_include_fixed_rubric(self) -> None:
        values = [_valid_content()]
        values[0]["practice_task_id"] = "PT-01"
        duplicate = copy.deepcopy(values[0])
        duplicate["practice_task_id"] = "PT-02"
        duplicate["project_name"] = "建立图书借阅数据表"
        for index, item in enumerate(duplicate["task_items"], start=1):
            item["title"] += f"（借阅{index}）"
            item["description"] += f"（借阅任务{index}）"
            item["deliverables"] = [f"{value}（借阅{index}）" for value in item["deliverables"]]
            item["acceptance_criteria"] = [f"{value}（借阅{index}）" for value in item["acceptance_criteria"]]
        report = validate_collection([values[0], duplicate])
        self.assertEqual(report["status"], "pass", report)
        self.assertTrue(all("rubric" not in item for item in report["metrics"]["repetition_scope"]))


if __name__ == "__main__":
    unittest.main()
