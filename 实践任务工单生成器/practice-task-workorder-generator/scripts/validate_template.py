"""Validate the real canonical Work Order template and its manifest."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path
from typing import Any

import yaml
from docx import Document


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _text(cell: Any) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def validate(template: Path, manifest_path: Path) -> dict[str, Any]:
    errors: list[str] = []
    try:
        manifest = yaml.safe_load(manifest_path.read_text(encoding="utf-8"))
    except Exception as exc:
        return {"status": "fail", "errors": [f"manifest read failed: {exc}"]}
    actual_hash = _sha256(template) if template.is_file() else ""
    expected_hash = str(manifest.get("fingerprint", {}).get("sha256", "")).upper()
    if actual_hash != expected_hash:
        errors.append(f"template fingerprint mismatch: expected {expected_hash}, got {actual_hash}")
    try:
        document = Document(str(template))
    except Exception as exc:
        return {"status": "fail", "errors": errors + [f"DOCX cannot be opened: {exc}"]}
    if len(document.tables) != 3:
        errors.append(f"expected 3 top-level tables, got {len(document.tables)}")
        return {"status": "fail", "errors": errors}
    work, student, teacher = document.tables
    headers = [_text(cell) for cell in work.rows[0].cells]
    required = ["任务对象", "任务项", "任务描述", "任务结果", "分值"]
    if headers != required:
        errors.append(f"work-order headers mismatch: {headers}")
    if len(work.rows) < 5:
        errors.append("canonical template needs header, attendance and sample task rows")
    score_values: list[int] = []
    for row in work.rows[1:]:
        try:
            score_values.append(int(_text(row.cells[4])))
        except ValueError:
            errors.append("work-order score column contains a non-integer")
    if len(work.rows) > 1 and _text(work.rows[1].cells[1]) != "课堂考勤":
        errors.append("canonical attendance title must be 课堂考勤")
    if score_values and score_values[0] != 10:
        errors.append("canonical attendance score must be 10")
    if score_values and sum(score_values) != 100:
        errors.append(f"canonical score total must be 100, got {sum(score_values)}")
    student_text = "\n".join(_text(cell) for row in student.rows for cell in row.cells)
    document_text = "\n".join(paragraph.text for paragraph in document.paragraphs) + "\n" + student_text
    for marker in ("自我评价", "小组评价", "教师评价"):
        if marker not in student_text:
            errors.append(f"student evaluation marker missing: {marker}")
    for marker in ("20%", "30%", "50%"):
        if marker not in document_text:
            errors.append(f"student evaluation weight missing: {marker}")
    teacher_text = "\n".join(_text(cell) for row in teacher.rows for cell in row.cells)
    if len(teacher.rows) != 4:
        errors.append(f"teacher evaluation rows must be 4, got {len(teacher.rows)}")
    for marker in ("A", "B", "C"):
        if teacher_text.count(marker) < 3:
            errors.append(f"teacher evaluation criterion marker incomplete: {marker}")
    return {
        "status": "pass" if not errors else "fail",
        "errors": errors,
        "metrics": {
            "sha256": actual_hash,
            "tables": len(document.tables),
            "rows": len(work.rows),
            "scores": score_values,
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--manifest", required=True, type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    report = validate(args.template, args.manifest)
    print(json.dumps(report, ensure_ascii=False, indent=2) if args.json else report["status"])
    if report["status"] != "pass" and not args.json:
        for error in report["errors"]:
            print(f"- {error}")
    return 0 if report["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
