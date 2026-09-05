"""Generate Practice Work Order Content V1 DOCX files transactionally."""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable

from docx import Document

PACKAGE_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(Path(__file__).resolve().parent))

from content_contract import (  # noqa: E402
    WorkOrderContractError,
    load_practice_task_contract,
    load_work_order_content,
    practice_tasks_to_authoring_skeleton,
    safe_filename,
)
from content_quality import validate_collection, validate_content  # noqa: E402
from cross_artifact_quality import validate_cross_artifact  # noqa: E402
from validate_output import validate_document  # noqa: E402


DEFAULT_TEMPLATE = PACKAGE_ROOT / "assets" / "templates" / "practice-work-order" / "v1.0.0" / "template.docx"


def _replace_paragraph_text(paragraph: Any, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _replace_cell_text(cell: Any, text: str) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(text)
        return
    first = paragraphs[0]
    for child in list(first._p):
        if child.tag.endswith("}pPr"):
            continue
        first._p.remove(child)
    first.add_run(text)
    for paragraph in paragraphs[1:]:
        cell._tc.remove(paragraph._p)


def _replace_in_paragraph(paragraph: Any, old: str, new: str) -> None:
    if old not in paragraph.text:
        return
    for run in paragraph.runs:
        run.text = run.text.replace(old, new)
    if old in paragraph.text:
        _replace_paragraph_text(paragraph, paragraph.text.replace(old, new))


def _replace_in_table(table: Any, old: str, new: str) -> None:
    seen: set[int] = set()
    for row in table.rows:
        for cell in row.cells:
            identity = id(cell._tc)
            if identity in seen:
                continue
            seen.add(identity)
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, old, new)


def _format_content_description(item: dict[str, Any]) -> str:
    def joined(values: Iterable[str]) -> str:
        return "；".join(str(value).strip() for value in values if str(value).strip())

    return (
        f"{item['description'].strip()}\n"
        f"工具/材料：{joined(item['tools_or_materials'])}\n"
        f"步骤：{joined(item['steps'])}\n"
        f"交付物：{joined(item['deliverables'])}\n"
        f"验收：{joined(item['acceptance_criteria'])}"
    )


def _source_project_name(document: Any) -> str:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    match = re.search(r"《([^》]+)》", text)
    if match:
        return match.group(1)
    for table in document.tables:
        match = re.search(r"项目实训[^\n：:]+", "\n".join(cell.text for row in table.rows for cell in row.cells))
        if match:
            return match.group(0).strip()
    return "项目实训2 设计数据库"


def _set_title_paragraphs(document: Any, old_project: str, project_name: str) -> None:
    for paragraph in document.paragraphs:
        if "学习工单" in paragraph.text:
            _replace_paragraph_text(paragraph, f"《{project_name}》学习工单")
        elif "课堂评价表-学生" in paragraph.text:
            _replace_paragraph_text(paragraph, f"《{project_name}》课堂评价表-学生")
        elif "课堂评价表-教师" in paragraph.text:
            _replace_paragraph_text(paragraph, f"《{project_name}》课堂评价表-教师")
        elif old_project in paragraph.text:
            _replace_in_paragraph(paragraph, old_project, project_name)


def _clone_task_rows(table: Any, items: list[dict[str, Any]]) -> None:
    prototype = deepcopy(table.rows[2]._tr)
    for row in list(table.rows[2:]):
        table._tbl.remove(row._tr)
    for item in items:
        table._tbl.append(deepcopy(prototype))
        row = table.rows[-1]
        _replace_cell_text(row.cells[1], item["title"])
        _replace_cell_text(row.cells[2], _format_content_description(item))
        _replace_cell_text(row.cells[3], "")
        _replace_cell_text(row.cells[4], str(item["score"]))


def _write_document(template: Path, output: Path, content: dict[str, Any]) -> None:
    document = Document(str(template))
    if len(document.tables) != 3:
        raise WorkOrderContractError("canonical template must contain exactly three top-level tables")
    old_project = _source_project_name(document)
    project_name = content["project_name"].strip()
    _set_title_paragraphs(document, old_project, project_name)
    work_table, student_table, teacher_table = document.tables
    _clone_task_rows(work_table, content["task_items"])
    metadata_parts = [
        content["course_name"],
        f"专业：{content['major']}",
        f"对象：{content['class_or_audience']}",
        f"实践任务ID：{content['practice_task_id']}",
        f"任务标题：{content.get('task_title', content['project_name'])}",
    ]
    if content.get("project_id"):
        metadata_parts.append(f"项目ID：{content['project_id']}")
    metadata_parts.extend(
        [
            f"实践学时：{content['practice_hours']}",
            content["group"]["name"],
            content["group"]["leader_placeholder"],
            content["group"]["member_placeholder"],
        ]
    )
    metadata = "\n".join(metadata_parts)
    if content.get("safety_or_compliance"):
        metadata += "\n安全/合规：" + "；".join(content["safety_or_compliance"])
    _replace_cell_text(work_table.rows[1].cells[0], metadata)
    _replace_in_table(student_table, old_project, project_name)
    _replace_in_table(teacher_table, old_project, project_name)
    document.core_properties.title = f"《{project_name}》学习工单"
    document.core_properties.subject = (
        f"Practice Task {content['practice_task_id']} / "
        f"{content.get('task_title', project_name)}"
    )
    document.core_properties.keywords = " ".join(
        value
        for value in (
            content["course_name"],
            content["practice_task_id"],
            content.get("project_id", ""),
        )
        if value
    )
    document.save(str(output))


def _render_optional(path: Path, output_dir: Path) -> dict[str, Any]:
    from render_qa import render_file

    return render_file(path, output_dir)


def _path_exists(path: Path) -> bool:
    return path.exists() or path.is_symlink()


def _remove_path(path: Path | None) -> None:
    if path is None or not _path_exists(path):
        return
    if path.is_dir() and not path.is_symlink():
        shutil.rmtree(path)
    else:
        path.unlink()


def _relocate_render_report(report: dict[str, Any], final_render_dir: Path) -> dict[str, Any]:
    relocated = dict(report)
    pdf = relocated.get("pdf")
    if pdf:
        relocated["pdf"] = str(final_render_dir / Path(str(pdf)).name)
    return relocated


def _publish_batch(
    publications: list[tuple[Path, Path]],
    *,
    output_dir: Path,
    stage: Path,
) -> None:
    """Publish all candidates with rollback for every touched target."""

    output_was_existing = _path_exists(output_dir)
    render_root = output_dir / "render"
    render_root_was_existing = _path_exists(render_root)
    backup_dir = stage / "backups"
    backup_dir.mkdir()
    committed: list[tuple[Path, Path | None]] = []
    rollback_errors: list[str] = []
    try:
        output_dir.mkdir(parents=True, exist_ok=True)
        for index, (target, staged) in enumerate(publications):
            target.parent.mkdir(parents=True, exist_ok=True)
            backup: Path | None = None
            if _path_exists(target):
                backup = backup_dir / f"{index:04d}-{safe_filename(target.name)}"
                os.replace(str(target), str(backup))
            committed.append((target, backup))
            os.replace(str(staged), str(target))
    except BaseException as exc:
        for target, backup in reversed(committed):
            try:
                _remove_path(target)
                if backup is not None and _path_exists(backup):
                    os.replace(str(backup), str(target))
            except Exception as rollback_error:  # pragma: no cover - filesystem-specific
                rollback_errors.append(f"{target}: {rollback_error}")
        if not render_root_was_existing and _path_exists(render_root):
            try:
                render_root.rmdir()
            except OSError:
                pass
        if not output_was_existing and _path_exists(output_dir):
            try:
                output_dir.rmdir()
            except OSError:
                pass
        if rollback_errors:
            detail = "; ".join(rollback_errors)
            raise RuntimeError(f"batch publication failed and rollback was incomplete: {detail}") from exc
        raise


def generate(
    contents: list[dict[str, Any]],
    *,
    output_dir: Path,
    template: Path,
    replace: bool = False,
    render: bool = False,
) -> dict[str, Any]:
    if not template.is_file():
        raise FileNotFoundError(f"template not found: {template}")
    if not contents:
        raise WorkOrderContractError("at least one WorkOrder Content V1 item is required")
    collection_report = validate_collection(contents)
    if collection_report["status"] != "pass":
        raise WorkOrderContractError("Content QA failed: " + "; ".join(collection_report["errors"]))
    output_dir = output_dir.expanduser().absolute()
    if _path_exists(output_dir) and not output_dir.is_dir():
        raise NotADirectoryError(f"output path is not a directory: {output_dir}")

    plans: list[dict[str, Any]] = []
    seen_targets: set[Path] = set()
    for content in contents:
        final = output_dir / f"{safe_filename(content['practice_task_id'])}_{safe_filename(content['project_name'])}.docx"
        targets = [final]
        if render:
            targets.append(output_dir / "render" / final.stem)
        for target in targets:
            if target in seen_targets:
                raise WorkOrderContractError(f"batch contains duplicate output target: {target}")
            seen_targets.add(target)
            if _path_exists(target) and not replace:
                raise FileExistsError(f"output exists; use --replace to overwrite: {target}")
        plans.append({"content": content, "final": final})

    parent = output_dir.parent
    parent.mkdir(parents=True, exist_ok=True)
    stage = Path(tempfile.mkdtemp(prefix=f".{output_dir.name}.batch-", dir=str(parent)))
    results: list[dict[str, Any]] = []
    publications: list[tuple[Path, Path]] = []
    try:
        candidate_dir = stage / "candidates"
        candidate_dir.mkdir()
        render_dir = stage / "render"
        for index, plan in enumerate(plans):
            content = plan["content"]
            final = plan["final"]
            candidate = candidate_dir / f"{index:04d}-{final.name}"
            _write_document(template, candidate, content)
            plan["candidate"] = candidate
            plan["output_qa"] = None
            plan["render"] = None

        for plan in plans:
            content = plan["content"]
            output_report = validate_document(plan["candidate"], content)
            if output_report["status"] != "pass":
                raise WorkOrderContractError(
                    f"Output QA failed for {content['practice_task_id']}: "
                    + "; ".join(output_report["errors"])
                )
            plan["output_qa"] = output_report

        for plan in plans:
            content = plan["content"]
            final = plan["final"]
            result: dict[str, Any] = {
                "practice_task_id": content["practice_task_id"],
                "path": str(final),
                "content_qa": validate_content(content),
                "output_qa": plan["output_qa"],
            }
            if render:
                staged_render_dir = render_dir / final.stem
                render_report = _render_optional(plan["candidate"], staged_render_dir)
                if render_report.get("status") != "pass":
                    status = render_report.get("status", "unknown")
                    reason = render_report.get("reason", "no reason supplied")
                    raise WorkOrderContractError(
                        f"Render QA failed for {content['practice_task_id']}: status={status}; {reason}"
                    )
                plan["render"] = _relocate_render_report(
                    render_report, output_dir / "render" / final.stem
                )
                publications.append((output_dir / "render" / final.stem, staged_render_dir))
                result["render"] = plan["render"]
            publications.append((final, plan["candidate"]))
            results.append(result)

        _publish_batch(publications, output_dir=output_dir, stage=stage)
        return {"status": "pass", "content_qa": collection_report, "outputs": results}
    finally:
        _remove_path(stage)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--content-json", type=Path)
    parser.add_argument("--practice-task-json", type=Path)
    parser.add_argument("--major", help="legacy handoff option; not used to author WorkOrder Content")
    parser.add_argument("--class-or-audience", help="legacy handoff option; not used to author WorkOrder Content")
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument(
        "--authoring-skeleton-output",
        type=Path,
        help="write an Agent authoring skeleton for handoff-only input; never writes DOCX",
    )
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--replace", action="store_true")
    parser.add_argument("--render", action="store_true")
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args(argv)
    try:
        if not args.content_json and not args.practice_task_json:
            raise WorkOrderContractError("one of --content-json or --practice-task-json is required")
        if args.authoring_skeleton_output and args.content_json:
            raise WorkOrderContractError("--authoring-skeleton-output is only valid for handoff-only input")
        if args.practice_task_json and not args.content_json:
            handoff = load_practice_task_contract(args.practice_task_json)
            if args.render:
                raise WorkOrderContractError(
                    "handoff-only input cannot render or generate DOCX; "
                    "Agent-authored --content-json is required"
                )
            skeleton = practice_tasks_to_authoring_skeleton(handoff)
            if args.authoring_skeleton_output:
                args.authoring_skeleton_output.parent.mkdir(parents=True, exist_ok=True)
                args.authoring_skeleton_output.write_text(
                    json.dumps(skeleton, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
                )
            report = {
                "status": "handoff_ready",
                "docx_generated": False,
                "practice_task_contract_validated": True,
                "authoring_skeleton": skeleton,
                "authoring_skeleton_path": str(args.authoring_skeleton_output)
                if args.authoring_skeleton_output
                else None,
            }
        else:
            contents = load_work_order_content(args.content_json)
            if not args.output_dir:
                raise WorkOrderContractError("--output-dir is required with Agent-authored --content-json")
            cross_reports: list[dict[str, Any]] = []
            if args.practice_task_json:
                handoff = load_practice_task_contract(args.practice_task_json)
                collection_report = validate_cross_artifact(handoff, contents)
                if collection_report["status"] != "pass":
                    raise WorkOrderContractError(
                        "Cross-Artifact QA failed: " + "; ".join(collection_report["errors"])
                    )
                nested_reports = collection_report.get("reports")
                cross_reports = (
                    nested_reports
                    if isinstance(nested_reports, list)
                    else [collection_report]
                )
            report = generate(
                contents,
                output_dir=args.output_dir,
                template=args.template,
                replace=args.replace,
                render=args.render,
            )
            if cross_reports:
                report["cross_artifact_qa"] = cross_reports
    except Exception as exc:
        report = {"status": "fail", "errors": [str(exc)]}
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        if report["status"] in {"pass", "handoff_ready"}:
            if report["status"] == "pass":
                for item in report["outputs"]:
                    print(item["path"])
            elif report.get("authoring_skeleton_path"):
                print(report["authoring_skeleton_path"])
            else:
                print("Practice Task handoff validated; Agent authoring is required before DOCX generation.")
        else:
            for error in report["errors"]:
                print(f"ERROR: {error}", file=sys.stderr)
    return 0 if report["status"] in {"pass", "handoff_ready"} else 1


if __name__ == "__main__":
    raise SystemExit(main())
