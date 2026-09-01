"""Output structure and Content-to-DOCX fidelity checks for Work Order V1."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document


class WorkOrderOutputError(ValueError):
    pass


def _cell_text(cell: Any) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _table_text(table: Any) -> str:
    return "\n".join(_cell_text(cell) for row in table.rows for cell in row.cells)


def _normalise(value: str) -> str:
    return " ".join(value.split())


def validate_document(path: Path, content: dict[str, Any] | None = None) -> dict[str, Any]:
    errors: list[str] = []
    try:
        document = Document(str(path))
    except Exception as exc:
        return {"status": "fail", "errors": [f"DOCX cannot be opened: {exc}"], "metrics": {}}
    tables = document.tables
    if len(tables) != 3:
        errors.append(f"expected exactly 3 top-level tables, got {len(tables)}")
        return {"status": "fail", "errors": errors, "metrics": {"tables": len(tables)}}
    headers = [_normalise(_cell_text(cell)) for cell in tables[0].rows[0].cells]
    expected_headers = ["任务对象", "任务项", "任务描述", "任务结果", "分值"]
    if headers != expected_headers:
        errors.append(f"work-order headers differ: expected {expected_headers}, got {headers}")
    if len(tables[0].rows) < 2:
        errors.append("work-order table has no attendance row")
    else:
        attendance = tables[0].rows[1]
        if not any(marker in _cell_text(attendance.cells[1]) for marker in ("出勤", "考勤")):
            errors.append("attendance row is missing")
        if _normalise(_cell_text(attendance.cells[4])) != "10":
            errors.append("attendance score must be 10")
    expected_items = content.get("task_items", []) if content else []
    if content and len(tables[0].rows) != 2 + len(expected_items):
        errors.append(
            f"task row count differs: expected {len(expected_items)}, got {max(0, len(tables[0].rows) - 2)}"
        )
    task_rows = tables[0].rows[2:]
    for index, row in enumerate(task_rows):
        if index >= len(expected_items):
            break
        item = expected_items[index]
        if _normalise(_cell_text(row.cells[1])) != _normalise(item["title"]):
            errors.append(f"task row {index + 1} title does not match Content V1")
        description_text = _normalise(_cell_text(row.cells[2]))
        if _normalise(item["description"]) not in description_text:
            errors.append(f"task row {index + 1} description does not match Content V1")
        for key in ("tools_or_materials", "steps", "deliverables", "acceptance_criteria"):
            for value in item[key]:
                if _normalise(value) not in description_text:
                    errors.append(f"task row {index + 1} missing {key} value")
        if _normalise(_cell_text(row.cells[4])) != str(item["score"]):
            errors.append(f"task row {index + 1} score does not match Content V1")
        if _cell_text(row.cells[3]):
            errors.append(f"task row {index + 1} student result cell must be blank")
    work_order_text = _table_text(tables[0])
    for field in ("course_name", "major", "class_or_audience"):
        if content and _normalise(str(content[field])) not in _normalise(work_order_text):
            errors.append(f"work-order metadata missing {field}")
    if content and f"实践学时：{content['practice_hours']}" not in work_order_text:
        errors.append("work-order metadata missing practice_hours")
    student_text = _table_text(tables[1])
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs) + "\n" + student_text
    for marker in ("自我评价", "小组评价", "教师评价"):
        if marker not in student_text:
            errors.append(f"student evaluation fixed marker missing: {marker}")
    for marker in ("20%", "30%", "50%"):
        if marker not in full_text:
            errors.append(f"student evaluation fixed weight missing: {marker}")
    teacher_text = _table_text(tables[2])
    if len(tables[2].rows) != 4:
        errors.append(f"teacher evaluation table must have 4 rows, got {len(tables[2].rows)}")
    if teacher_text.count("A") < 3 or teacher_text.count("B") < 3 or teacher_text.count("C") < 3:
        errors.append("teacher evaluation fixed A/B/C criteria are incomplete")
    return {
        "status": "pass" if not errors else "fail",
        "errors": errors,
        "metrics": {
            "tables": len(tables),
            "task_rows": len(task_rows),
            "student_result_cells_blank": sum(not _cell_text(row.cells[3]) for row in task_rows),
        },
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--content-json", type=Path)
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    content = None
    if args.content_json:
        value = json.loads(args.content_json.read_text(encoding="utf-8"))
        content = value[0] if isinstance(value, list) else value
    report = validate_document(args.input, content)
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        print(report["status"])
        for error in report["errors"]:
            print(f"- {error}")
    return 0 if report["status"] == "pass" else 1


if __name__ == "__main__":
    raise SystemExit(main())
