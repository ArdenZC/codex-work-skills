from __future__ import annotations

import argparse
import copy
import json
import math
import re
import sys
from datetime import datetime, timezone
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from lxml import etree

from bookmark_utils import bookmark_boundary_locations, bookmark_location, bookmark_parent_cell, bookmark_parent_paragraph, find_bookmark, validate_bookmark_inventory
from package_common import (
    DEFAULT_MANIFEST,
    DEFAULT_SCHEMA,
    EVALUATION_MAX_POINTS,
    bookmark_containers,
    evaluation_cell_values,
    field_bookmark,
    field_spec,
    generated_lesson_fields,
    implementation_bookmarks,
    implementation_cell_values,
    is_semantic_manifest,
    layout_manifest,
    load_manifest,
    manifest_template_path,
    resolve_template_package,
    reflection_bookmarks,
    reflection_cell_values,
    required_bookmarks,
    validate_composed_fields,
    validate_input,
)


LESSON_FILE_PATTERN = re.compile(r"^教案(?P<sequence>\d+)_")


def actual_cells(row) -> list[_Cell]:
    return [_Cell(tc, row._parent) for tc in row._tr.tc_lst]


def cell_text(table, row_index: int, cell_index: int) -> str:
    return actual_cells(table.rows[row_index])[cell_index].text.strip()


def manifest_field_text(document, main_table, manifest: dict[str, Any], name: str) -> str:
    spec = field_spec(manifest, name)
    if is_semantic_manifest(manifest):
        bookmark_name = field_bookmark(manifest, name)
        record = find_bookmark(document, bookmark_name)
        if record is None:
            raise ValueError(f"Semantic bookmark {bookmark_name} for field {name} is missing")
        cell_element = bookmark_parent_cell(document, record)
        if cell_element is not None:
            return _Cell(cell_element, document).text.strip()
        paragraph_element = bookmark_parent_paragraph(document, record)
        if paragraph_element is not None:
            return Paragraph(paragraph_element, document).text.strip()
        raise ValueError(f"Semantic bookmark {bookmark_name} for field {name} has no writable container")
    if not all(key in spec for key in ("table", "row", "cell")):
        raise ValueError(f"Manifest field {name} is not a table-cell field")
    table_index = int(spec["table"])
    target_table = main_table if table_index == 0 else document.tables[table_index]
    return cell_text(target_table, int(spec["row"]), int(spec["cell"]))


def parse_number(value: Any, label: str) -> float:
    try:
        return float(str(value).strip())
    except (TypeError, ValueError) as exc:
        raise ValueError(f"{label} is not numeric: {value!r}") from exc


def parse_decimal(value: Any, label: str) -> Decimal:
    try:
        number = Decimal(str(value).strip())
    except (InvalidOperation, TypeError, ValueError) as exc:
        raise ValueError(f"{label} is not numeric: {value!r}") from exc
    if not number.is_finite():
        raise ValueError(f"{label} must be finite: {value!r}")
    return number


def _lesson_file_sort_key(path: Path) -> tuple[int, int, str]:
    match = LESSON_FILE_PATTERN.match(path.name)
    if match:
        return 0, int(match.group("sequence")), path.name
    return 1, 0, path.name


def _xml(element) -> str:
    return etree.tostring(copy.deepcopy(element), encoding="unicode")


def _canonical_xml(element) -> str:
    if element is None:
        return ""
    return etree.tostring(copy.deepcopy(element), method="c14n", with_comments=False).decode("utf-8")


def _evaluation_writable_coordinates(manifest: dict[str, Any]) -> tuple[set[int], set[int]]:
    structure = manifest["structure"]["evaluation_table"]
    spec = manifest.get("fields", {}).get("evaluation", {})
    rows = {int(row) for row in spec.get("writable_rows", [])}
    cells = {int(cell) for cell in spec.get("writable_cells", [])}
    if not rows:
        rows = set(range(1, int(structure["rows"])))
    if not cells:
        cells = {2, 3}
    return rows, cells


def _evaluation_table(document, manifest: dict[str, Any]):
    main = document.tables[int(manifest["structure"]["main_table"]["index"])]
    spec = manifest["structure"]["evaluation_table"]
    parent = actual_cells(main.rows[int(spec["row"])])[int(spec["cell"])]
    return parent.tables[0] if parent.tables else None


def _evaluation_table_layout_signature(document, manifest: dict[str, Any]) -> dict[str, Any]:
    """Sign only nested-table layout XML; declared score/remark text is excluded."""
    nested = _evaluation_table(document, manifest)
    if nested is None:
        return {"tblPr": "", "tblGrid": "", "rows": []}
    rows = []
    for row in nested.rows:
        rows.append(
            {
                "trPr": _canonical_xml(row._tr.trPr),
                "cells": [
                    {"tcPr": _canonical_xml(cell._tc.tcPr)}
                    for cell in actual_cells(row)
                ],
            }
        )
    return {
        "tblPr": _canonical_xml(nested._tbl.tblPr),
        "tblGrid": _canonical_xml(nested._tbl.tblGrid),
        "rows": rows,
    }


def _settings_xml(document) -> str:
    value = etree.tostring(copy.deepcopy(document.settings._element), encoding="unicode")
    return value.encode("unicode_escape").decode("ascii")


def _cell_text_signature(cell) -> list[Any]:
    values: list[Any] = [paragraph.text for paragraph in cell.paragraphs]
    for nested in cell.tables:
        values.append(
            [
                [_cell_text_signature(nested_cell) for nested_cell in actual_cells(row)]
                for row in nested.rows
            ]
        )
    return values


def _theme_signature(document) -> list[dict[str, Any]]:
    return [
        {"partname": str(part.partname), "blob": part.blob}
        for part in document.part.package.iter_parts()
        if str(part.partname).startswith("/word/theme/")
    ]


def _body_paragraph_signature(document, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    title_spec = manifest.get("fields", {}).get("title", {})
    title_index = int(title_spec.get("paragraph", manifest["structure"]["title"]["paragraph"]))
    return [
        {
            "index": index,
            "xml": _xml(paragraph._p),
        }
        for index, paragraph in enumerate(document.paragraphs)
        if index != title_index
    ]


def _title_format_signature(document, manifest: dict[str, Any]) -> dict[str, Any]:
    title_spec = manifest.get("fields", {}).get("title", {})
    title_index = int(title_spec.get("paragraph", manifest["structure"]["title"]["paragraph"]))
    if title_index < 0 or title_index >= len(document.paragraphs):
        return {"paragraph": "", "run_formats": [], "run_primary": ""}
    paragraph = document.paragraphs[title_index]
    ppr = copy.deepcopy(paragraph._p.pPr)
    if ppr is not None:
        for child in list(ppr):
            if child.tag == qn("w:jc"):
                ppr.remove(child)
    run_values = [
        etree.tostring(copy.deepcopy(run._r.rPr), encoding="unicode")
        for run in paragraph.runs
        if run._r.rPr is not None
    ]
    return {
        "paragraph": etree.tostring(ppr, encoding="unicode") if ppr is not None else "",
        "run_formats": sorted({value for value in run_values if value}),
        "run_primary": next((value for value in run_values if value), ""),
    }


def _protected_text_signature(document, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    manifest = layout_manifest(manifest)
    table = document.tables[int(manifest["structure"]["main_table"]["index"])]
    writable, evaluation = _writable_main_table_cells(manifest)

    values: list[dict[str, Any]] = []
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(actual_cells(row)):
            coordinate = (row_index, cell_index)
            if coordinate in writable or coordinate == evaluation:
                continue
            values.append(
                {
                    "scope": "main",
                    "row": row_index,
                    "cell": cell_index,
                    "text": _cell_text_signature(cell),
                    "format": _direct_format_signature(cell),
                }
            )

    if evaluation is not None:
        parent = actual_cells(table.rows[evaluation[0]])[evaluation[1]]
        nested = parent.tables[0] if parent.tables else None
        if nested is not None:
            writable_rows, writable_cells = _evaluation_writable_coordinates(manifest)
            for row_index, row in enumerate(nested.rows):
                for cell_index, cell in enumerate(actual_cells(row)):
                    if row_index not in writable_rows or cell_index not in writable_cells:
                        values.append(
                            {
                                "scope": "evaluation",
                                "row": row_index,
                                "cell": cell_index,
                                "text": _cell_text_signature(cell),
                                "format": _direct_format_signature(cell),
                            }
                        )
    return values


def _writable_main_table_cells(manifest: dict[str, Any]) -> tuple[set[tuple[int, int]], tuple[int, int] | None]:
    writable: set[tuple[int, int]] = set()
    evaluation: tuple[int, int] | None = None
    for spec in manifest.get("fields", {}).values():
        if spec.get("mode") == "row_cells":
            writable.update(
                (int(row), int(cell))
                for row in spec.get("rows", [])
                for cell in spec.get("cells", [])
            )
            continue
        if spec.get("mode") == "nested_table":
            if int(spec.get("table", -1)) == 0:
                evaluation = (int(spec["row"]), int(spec["cell"]))
            continue
        if int(spec.get("table", -1)) == 0 and "row" in spec and "cell" in spec:
            writable.add((int(spec["row"]), int(spec["cell"])))
    return writable, evaluation


def _normalized_direct_xml(element) -> str:
    if element is None:
        return ""
    if len(element) == 0 and not element.attrib and not (element.text or "").strip():
        return ""
    return _xml(element)


def _direct_format_signature(cell) -> dict[str, Any]:
    paragraphs = []
    for paragraph in cell.paragraphs:
        paragraph_value = _normalized_direct_xml(paragraph._p.pPr)
        run_values = [
            _normalized_direct_xml(run._r.rPr)
            for run in paragraph.runs
        ]
        paragraphs.append({"paragraph": paragraph_value, "runs": run_values})
    paragraph_values = [item["paragraph"] for item in paragraphs]
    run_values = [run for item in paragraphs for run in item["runs"]]
    paragraph_formats = sorted({value for value in paragraph_values if value})
    run_formats = sorted({value for value in run_values if value})
    return {
        "paragraphs": paragraphs,
        "paragraph_formats": paragraph_formats,
        "paragraph_primary": paragraphs[0]["paragraph"] if paragraphs else "",
        "run_formats": run_formats,
        "run_primary": paragraphs[0]["runs"][0] if paragraphs and paragraphs[0]["runs"] else "",
    }


def _writable_direct_format_signature(document, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    manifest = layout_manifest(manifest)
    table = document.tables[int(manifest["structure"]["main_table"]["index"])]
    writable, evaluation = _writable_main_table_cells(manifest)
    values: list[dict[str, Any]] = []
    for row_index, cell_index in sorted(writable):
        cell = actual_cells(table.rows[row_index])[cell_index]
        values.append(
            {
                "scope": "main",
                "row": row_index,
                "cell": cell_index,
                "format": _direct_format_signature(cell),
            }
        )

    if evaluation is not None:
        parent = actual_cells(table.rows[evaluation[0]])[evaluation[1]]
        values.append(
            {
                "scope": "evaluation_parent",
                "row": evaluation[0],
                "cell": evaluation[1],
                "format": _direct_format_signature(parent),
            }
        )
    return values


def _evaluation_writable_format_signature(document, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    manifest = layout_manifest(manifest)
    table = document.tables[int(manifest["structure"]["main_table"]["index"])]
    _writable, evaluation = _writable_main_table_cells(manifest)
    if evaluation is None:
        return []
    nested = _evaluation_table(document, manifest)
    if nested is None:
        return []
    writable_rows, writable_cells = _evaluation_writable_coordinates(manifest)

    def format_signature(cell) -> dict[str, Any]:
        signature = _direct_format_signature(cell)
        paragraphs = [
            {
                "paragraph": item["paragraph"],
                "runs": [value for value in item["runs"] if value],
            }
            for item in signature["paragraphs"]
        ]
        run_values = [value for item in paragraphs for value in item["runs"]]
        return {
            "paragraphs": paragraphs,
            "paragraph_formats": sorted({item["paragraph"] for item in paragraphs if item["paragraph"]}),
            "paragraph_primary": paragraphs[0]["paragraph"] if paragraphs else "",
            "run_formats": sorted(set(run_values)),
            "run_primary": run_values[0] if run_values else "",
        }

    return [
        {
            "row": row_index,
            "cell": cell_index,
            "format": format_signature(cell),
        }
        for row_index, row in enumerate(nested.rows)
        for cell_index, cell in enumerate(actual_cells(row))
        if row_index in writable_rows and cell_index in writable_cells
    ]


def protected_layout_signature(document, manifest: dict[str, Any]) -> dict[str, Any]:
    layout = layout_manifest(manifest)
    table = document.tables[int(layout["structure"]["main_table"]["index"])]
    rows = []
    for row in table.rows:
        rows.append(
            {
                "trPr": _xml(row._tr.trPr) if row._tr.trPr is not None else "",
                "cells": [_xml(cell._tc.tcPr) if cell._tc.tcPr is not None else "" for cell in actual_cells(row)],
            }
        )
    return {
        "tablePr": _xml(table._tbl.tblPr) if table._tbl.tblPr is not None else "",
        "tblGrid": _xml(table._tbl.tblGrid) if table._tbl.tblGrid is not None else "",
        "rows": rows,
        "evaluation_table": _evaluation_table_layout_signature(document, layout),
        "sections": [
            {
                "page_width": section.page_width.twips,
                "page_height": section.page_height.twips,
                "top_margin": section.top_margin.twips,
                "bottom_margin": section.bottom_margin.twips,
                "left_margin": section.left_margin.twips,
                "right_margin": section.right_margin.twips,
                "sectPr": _xml(section._sectPr),
                "header_footer_refs": [
                    _xml(child)
                    for child in section._sectPr
                    if child.tag.endswith("headerReference") or child.tag.endswith("footerReference")
                ],
            }
            for section in document.sections
        ],
        "headers_footers": [
            {
                "header": _xml(section.header._element),
                "footer": _xml(section.footer._element),
                "first_page_header": _xml(section.first_page_header._element),
                "first_page_footer": _xml(section.first_page_footer._element),
                "even_page_header": _xml(section.even_page_header._element),
                "even_page_footer": _xml(section.even_page_footer._element),
                "odd_and_even_pages_header_footer": bool(document.settings.odd_and_even_pages_header_footer),
            }
            for section in document.sections
        ],
        "styles": _xml(document.styles._element),
        "themes": _theme_signature(document),
        "settings_xml": _settings_xml(document),
        "body_paragraphs": _body_paragraph_signature(document, layout),
        "title_format": _title_format_signature(document, layout),
        "protected_text": _protected_text_signature(document, layout),
        "writable_direct_formats": _writable_direct_format_signature(document, layout),
        "evaluation_writable_formats": _evaluation_writable_format_signature(document, layout),
    }


def _protected_layout_matches(actual: dict[str, Any], expected: dict[str, Any]) -> bool:
    actual_direct = actual.get("writable_direct_formats", [])
    expected_direct = expected.get("writable_direct_formats", [])
    actual_title = actual.get("title_format", {})
    expected_title = expected.get("title_format", {})
    actual_base = {
        key: value
        for key, value in actual.items()
        if key not in {"writable_direct_formats", "title_format"}
    }
    expected_base = {
        key: value
        for key, value in expected.items()
        if key not in {"writable_direct_formats", "title_format"}
    }
    if actual_base != expected_base or len(actual_direct) != len(expected_direct):
        return False
    if actual_title.get("paragraph", "") != expected_title.get("paragraph", ""):
        return False
    if not set(actual_title.get("run_formats", [])).issubset(set(expected_title.get("run_formats", []))):
        return False
    if expected_title.get("run_primary") and expected_title["run_primary"] not in actual_title.get("run_formats", []):
        return False

    # Text replacement can collapse source runs/paragraphs, so compare each output paragraph with the source paragraph format.
    expected_by_coordinate = {
        (item["scope"], item.get("row"), item.get("cell")): item["format"]
        for item in expected_direct
    }
    for item in actual_direct:
        coordinate = (item["scope"], item.get("row"), item.get("cell"))
        expected_format = expected_by_coordinate.get(coordinate)
        if expected_format is None:
            return False
        actual_format = item["format"]
        actual_paragraphs = actual_format.get("paragraphs", [])
        expected_paragraphs = expected_format.get("paragraphs", [])
        if not actual_paragraphs or not expected_paragraphs:
            if actual_paragraphs != expected_paragraphs:
                return False
            continue
        expected_paragraph = expected_paragraphs[0].get("paragraph", "")
        expected_runs = expected_paragraphs[0].get("runs", [])
        expected_run = expected_runs[0] if expected_runs else ""
        for paragraph in actual_paragraphs:
            if paragraph.get("paragraph", "") != expected_paragraph:
                return False
            actual_runs = paragraph.get("runs", [])
            if expected_run:
                if not actual_runs or any(run != expected_run for run in actual_runs):
                    return False
            elif any(actual_runs):
                return False
    return True


def _validate_text_length(name: str, value: str, spec: dict[str, Any], errors: list[str]) -> None:
    max_chars = spec.get("max_chars")
    if max_chars is not None and len(value) > int(max_chars):
        errors.append(f"{name} exceeds manifest max_chars={max_chars}: {len(value)}")


def _field_targets(document, table, spec: dict[str, Any], name: str = "", manifest: dict[str, Any] | None = None):
    if manifest is not None and is_semantic_manifest(manifest):
        if spec.get("bookmark"):
            bookmark_name = str(spec["bookmark"])
            record = find_bookmark(document, bookmark_name)
            if record is None:
                return []
            cell_element = bookmark_parent_cell(document, record)
            if cell_element is not None:
                cell = _Cell(cell_element, document)
                return [(cell.text, len(cell.paragraphs))]
            paragraph_element = bookmark_parent_paragraph(document, record)
            if paragraph_element is not None:
                return [(Paragraph(paragraph_element, document).text, 1)]
            return []
        if name == "implementation":
            values = []
            for group in implementation_bookmarks(manifest):
                for bookmark_name in group:
                    record = find_bookmark(document, bookmark_name)
                    if record is None:
                        continue
                    cell_element = bookmark_parent_cell(document, record)
                    if cell_element is not None:
                        cell = _Cell(cell_element, document)
                        values.append((cell.text, len(cell.paragraphs)))
            return values
        if name == "reflection":
            values = []
            for bookmark_name in reflection_bookmarks(manifest):
                record = find_bookmark(document, bookmark_name)
                if record is None:
                    continue
                cell_element = bookmark_parent_cell(document, record)
                if cell_element is not None:
                    cell = _Cell(cell_element, document)
                    values.append((cell.text, len(cell.paragraphs)))
            return values
        return []
    if spec.get("target") == "document_paragraph" or "paragraph" in spec:
        index = int(spec.get("paragraph", -1))
        if 0 <= index < len(document.paragraphs):
            return [(document.paragraphs[index].text, 1)]
        return []
    if spec.get("mode") == "row_cells":
        values = []
        for row_index in spec.get("rows", []):
            if row_index >= len(table.rows):
                continue
            cells = actual_cells(table.rows[int(row_index)])
            for cell_index in spec.get("cells", []):
                if cell_index < len(cells):
                    values.append((cells[int(cell_index)].text, len(cells[int(cell_index)].paragraphs)))
        return values
    if "table" in spec and "row" in spec and "cell" in spec:
        table_index = int(spec["table"])
        if table_index == 0 and int(spec["row"]) < len(table.rows):
            cells = actual_cells(table.rows[int(spec["row"])])
            cell_index = int(spec["cell"])
            if cell_index < len(cells):
                cell = cells[cell_index]
                return [(cell.text, len(cell.paragraphs))]
    return []


def _document_text(document, table) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for row in table.rows:
        for cell in actual_cells(row):
            values.append(cell.text)
            for nested in cell.tables:
                for nested_row in nested.rows:
                    values.extend(nested_cell.text for nested_cell in nested_row.cells)
    for section in document.sections:
        values.extend(paragraph.text for paragraph in section.header.paragraphs)
        values.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(values)


def _evaluation_score_errors(score_values: list[Decimal]) -> list[str]:
    errors: list[str] = []
    for row, value in enumerate(score_values, 1):
        maximum = EVALUATION_MAX_POINTS[row - 1] if row <= len(EVALUATION_MAX_POINTS) else None
        if value < 0:
            errors.append(f"evaluation score row {row} is negative: {value}")
        elif maximum is not None and value > maximum:
            errors.append(f"evaluation score row {row} exceeds rubric maximum {maximum}: {value}")
    return errors


def _base_qa_report(
    out_dir: Path,
    manifest: dict[str, Any],
    qa_report_path: Path | str | None,
    template_path: Path | str | None,
    custom_template: bool | None,
    engine: str | None,
    template_validation: bool,
    output_validation: bool,
    extra_warnings: list[str] | None,
) -> dict[str, Any]:
    canonical_template = manifest_template_path(manifest)
    selected_template = (
        Path(template_path).expanduser().resolve() if template_path else canonical_template
    )
    is_custom_template = (
        bool(custom_template) if custom_template is not None else selected_template != canonical_template
    )
    skipped = []
    if not template_validation:
        skipped.append("template")
    if not output_validation:
        skipped.append("output")
    warnings = list(extra_warnings or [])
    if is_custom_template:
        warnings.append("Custom template selected; output was validated against the supplied manifest.")
    if not template_validation:
        warnings.append("Template validation skipped by explicit flag.")
    if not output_validation:
        warnings.append("Output validation skipped by explicit flag.")
    warnings = list(dict.fromkeys(warnings))
    validation = {"template": template_validation, "output": output_validation, "skipped": skipped}
    report_path = Path(qa_report_path).expanduser().resolve() if qa_report_path else out_dir / "qa-report.json"
    report: dict[str, Any] = {
        "status": "failed",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "template_id": manifest.get("template", {}).get("id"),
        "template_version": manifest.get("template", {}).get("version"),
        "generator_version": manifest.get("generator", {}).get("version"),
        "template_path": str(selected_template),
        "custom_template": is_custom_template,
        "engine": engine or "unknown",
        "validation": validation,
        "validation_skipped": skipped,
        "output_dir": str(out_dir),
        "errors": [],
        "warnings": warnings,
        "checks": {"validation": validation},
        "files_checked": 0,
        "qa_report": str(report_path),
    }
    if is_semantic_manifest(manifest):
        report.update(
            {
                "anchor_mode": "word_bookmark",
                "required_anchor_count": len(required_bookmarks(manifest)),
                "preserved_anchor_count": 0,
                "missing_anchors": [],
                "duplicate_anchors": [],
                "invalid_anchor_names": [],
                "unexpected_anchor_names": [],
                "invalid_anchor_ids": [],
                "anchor_boundary_errors": [],
            }
        )
    else:
        report["anchor_mode"] = "legacy_coordinates"
    return report


def _write_qa_report(report: dict[str, Any]) -> dict[str, Any]:
    report_path = Path(report["qa_report"])
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return report


def write_skipped_report(
    output_dir: Path | str,
    data: dict[str, Any],
    manifest: dict[str, Any] | None = None,
    qa_report_path: Path | str | None = None,
    schema_path: Path | str = DEFAULT_SCHEMA,
    *,
    template_path: Path | str | None = None,
    custom_template: bool | None = None,
    engine: str | None = None,
    template_validation: bool = True,
    warnings: list[str] | None = None,
) -> dict[str, Any]:
    out_dir = Path(output_dir).expanduser().resolve()
    manifest = manifest or load_manifest()
    validate_input(data, schema_path)
    out_dir.mkdir(parents=True, exist_ok=True)
    report = _base_qa_report(
        out_dir,
        manifest,
        qa_report_path,
        template_path,
        custom_template,
        engine,
        template_validation,
        False,
        warnings,
    )
    report["status"] = "skipped"
    report["checks"]["file_count"] = {"expected": len(data["lessons"]), "actual": len(list(out_dir.glob("*.docx")))}
    return _write_qa_report(report)


def validate_output_dir(
    output_dir: Path | str,
    data: dict[str, Any],
    manifest: dict[str, Any] | None = None,
    qa_report_path: Path | str | None = None,
    schema_path: Path | str = DEFAULT_SCHEMA,
    *,
    template_path: Path | str | None = None,
    custom_template: bool | None = None,
    engine: str | None = None,
    template_validation: bool = True,
    output_validation: bool = True,
    extra_warnings: list[str] | None = None,
) -> dict[str, Any]:
    out_dir = Path(output_dir).expanduser().resolve()
    manifest = manifest or load_manifest()
    validate_input(data, schema_path)
    validate_composed_fields(data, manifest)
    lessons = data["lessons"]
    files = sorted(out_dir.glob("*.docx"), key=_lesson_file_sort_key)
    report = _base_qa_report(
        out_dir,
        manifest,
        qa_report_path,
        template_path,
        custom_template,
        engine,
        template_validation,
        output_validation,
        extra_warnings,
    )
    errors: list[str] = report["errors"]
    checks: dict[str, Any] = report["checks"]

    if not files:
        errors.append(f"No DOCX files generated in {out_dir}")
    if len(files) != len(lessons):
        errors.append(f"Output count mismatch: expected {len(lessons)}, got {len(files)}")

    main_spec = manifest["structure"]["main_table"]
    course_expected = str(data["course_name"])
    total_hours = 0.0
    lesson_checks = []
    anchor_results: list[dict[str, Any]] = []
    for index, (path, item) in enumerate(zip(files, lessons), start=1):
        item_errors: list[str] = []
        if not path.is_file() or path.stat().st_size == 0:
            item_errors.append("file is missing or empty")
            errors.extend(f"file {index}: {message}" for message in item_errors)
            lesson_checks.append({"file_index": index, "errors": item_errors})
            continue
        try:
            document = Document(str(path))
        except Exception as exc:  # pragma: no cover - library-specific parse errors
            item_errors.append(f"DOCX could not be opened: {exc}")
            lesson_checks.append({"file_index": index, "errors": item_errors})
            errors.extend(f"file {index}: {message}" for message in item_errors)
            continue
        if is_semantic_manifest(manifest):
            anchor_inventory = validate_bookmark_inventory(
                document,
                required_bookmarks(manifest),
                bookmark_containers(manifest),
            )
            anchor_results.append(anchor_inventory)
            if not anchor_inventory["valid"]:
                item_errors.extend(
                    f"semantic bookmark protection failed: {error}"
                    for error in anchor_inventory["errors"]
                )
        if len(document.tables) <= int(main_spec["index"]):
            item_errors.append("main table is missing")
            lesson_checks.append({"file_index": index, "errors": item_errors})
            errors.extend(f"file {index}: {message}" for message in item_errors)
            continue
        if is_semantic_manifest(manifest) and not anchor_inventory["valid"]:
            errors.extend(f"file {index}: {message}" for message in item_errors)
            lesson_checks.append({"file_index": index, "errors": item_errors, "fields_checked": []})
            continue
        table = document.tables[int(main_spec["index"])]
        canonical_template = manifest_template_path(manifest)
        if canonical_template.exists():
            canonical_document = Document(str(canonical_template))
            if not _protected_layout_matches(
                protected_layout_signature(document, manifest),
                protected_layout_signature(canonical_document, manifest),
            ):
                item_errors.append("protected DOCX layout changed")
            if is_semantic_manifest(manifest) and anchor_inventory["valid"]:
                location_changes = [
                    name
                    for name in required_bookmarks(manifest)
                    if bookmark_boundary_locations(document, name) != bookmark_boundary_locations(canonical_document, name)
                ]
                if location_changes:
                    item_errors.append("semantic bookmark location changed: " + ", ".join(location_changes))
        if len(table.rows) != int(main_spec["rows"]):
            item_errors.append(f"main table rows expected {main_spec['rows']}, got {len(table.rows)}")
        if len(table.columns) != int(main_spec["columns"]):
            item_errors.append(f"main table columns expected {main_spec['columns']}, got {len(table.columns)}")

        expected_course = str(item.get("course_name") or data["course_name"])
        expected_major = str(item.get("major") or data.get("major", "软件技术"))
        expected_audience = str(item.get("audience") or data.get("audience", "高职二年级"))
        field_values: dict[str, str] = {}
        for field_name in ("course_name", "major", "audience", "unit", "task", "hours"):
            try:
                field_values[field_name] = manifest_field_text(document, table, manifest, field_name)
            except (KeyError, ValueError, IndexError) as exc:
                field_values[field_name] = ""
                item_errors.append(f"semantic bookmark protection failed for {field_name}: {exc}")
        if field_values["course_name"] != expected_course:
            item_errors.append("course field mismatch")
        if field_values["major"] != expected_major:
            item_errors.append("major field mismatch")
        if field_values["audience"] != expected_audience:
            item_errors.append("audience field mismatch")
        if field_values["unit"] != str(item["unit"]):
            item_errors.append("unit field mismatch")
        if field_values["task"] != str(item["task"]):
            item_errors.append("task field mismatch")
        generated = generated_lesson_fields(
            str(item["unit"]),
            str(item["task"]),
            item.get("flows", []),
            item.get("knowledge", []),
            item.get("tools", "课程PPT、微课视频、任务单、评分表和成果模板"),
        )
        for name, expected in generated.items():
            try:
                actual = manifest_field_text(document, table, manifest, name)
            except (KeyError, ValueError, IndexError) as exc:
                actual = ""
                item_errors.append(f"semantic bookmark protection failed for {name}: {exc}")
            if actual != expected:
                item_errors.append(f"{name} content mismatch")
        try:
            hours = parse_number(field_values["hours"], "hours")
            expected_hours = parse_number(item["hours"], "input hours")
            total_hours += hours
            if not math.isclose(hours, expected_hours, abs_tol=0.01):
                item_errors.append(f"hours mismatch: expected {expected_hours}, got {hours}")
        except ValueError as exc:
            item_errors.append(str(exc))
        if not field_values["unit"].startswith("项目"):
            item_errors.append("unit is not projectized")

        expected_title = f"{index} 《{expected_course}》教学单元设计：{item['task']}"
        if is_semantic_manifest(manifest):
            title_record = find_bookmark(document, field_bookmark(manifest, "title"))
            title_element = bookmark_parent_paragraph(document, title_record) if title_record else None
            actual_title = Paragraph(title_element, document).text if title_element is not None else ""
        else:
            title_spec = field_spec(manifest, "title")
            title_index = int(title_spec.get("paragraph", manifest["structure"]["title"]["paragraph"]))
            actual_title = document.paragraphs[title_index].text if 0 <= title_index < len(document.paragraphs) else ""
        if actual_title != expected_title:
            item_errors.append("title field mismatch")

        for name, spec in manifest.get("fields", {}).items():
            if name in {"evaluation"}:
                continue
            for value, paragraph_count in _field_targets(document, table, spec, name, manifest):
                _validate_text_length(name, value, spec, item_errors)
                max_paragraphs = spec.get("max_paragraphs")
                if max_paragraphs is not None and paragraph_count > int(max_paragraphs):
                    item_errors.append(f"{name} exceeds manifest max_paragraphs={max_paragraphs}: {paragraph_count}")

        if is_semantic_manifest(manifest):
            for bookmark_group, expected_values in zip(
                implementation_bookmarks(manifest),
                implementation_cell_values(str(item["task"]), item.get("flows", [])),
            ):
                for bookmark_name, expected_value in zip(bookmark_group, expected_values.values()):
                    record = find_bookmark(document, bookmark_name)
                    cell_element = bookmark_parent_cell(document, record) if record else None
                    actual_value = _Cell(cell_element, document).text if cell_element is not None else ""
                    expected_text = str(expected_value).rstrip("\r\n")
                    if actual_value != expected_text:
                        item_errors.append(f"implementation cell mismatch at bookmark {bookmark_name}")
            for bookmark_name, expected_value in zip(
                reflection_bookmarks(manifest),
                reflection_cell_values(str(item["task"])),
            ):
                record = find_bookmark(document, bookmark_name)
                cell_element = bookmark_parent_cell(document, record) if record else None
                actual_value = _Cell(cell_element, document).text if cell_element is not None else ""
                if actual_value != expected_value:
                    item_errors.append(f"reflection cell mismatch at bookmark {bookmark_name}")
        else:
            implementation_rows = [int(row) for row in manifest["fields"]["implementation"]["rows"]]
            for row_index, expected_values in zip(
                implementation_rows,
                implementation_cell_values(str(item["task"]), item.get("flows", [])),
            ):
                cells = actual_cells(table.rows[row_index]) if row_index < len(table.rows) else []
                for cell_index, expected_value in expected_values.items():
                    actual_value = cells[cell_index].text if cell_index < len(cells) else ""
                    expected_text = str(expected_value).rstrip("\r\n")
                    if actual_value != expected_text:
                        item_errors.append(f"implementation cell mismatch at row {row_index} cell {cell_index}")

            reflection_rows = [int(row) for row in manifest["fields"]["reflection"]["rows"]]
            for row_index, expected_value in zip(reflection_rows, reflection_cell_values(str(item["task"]))):
                cells = actual_cells(table.rows[row_index]) if row_index < len(table.rows) else []
                actual_value = cells[2].text if len(cells) > 2 else ""
                if actual_value != expected_value:
                    item_errors.append(f"reflection cell mismatch at row {row_index} cell 2")

        nested_spec = manifest["structure"]["evaluation_table"]
        try:
            if is_semantic_manifest(manifest):
                evaluation_record = find_bookmark(document, field_bookmark(manifest, "evaluation"))
                evaluation_element = bookmark_parent_cell(document, evaluation_record) if evaluation_record else None
                if evaluation_element is None:
                    raise ValueError("evaluation parent bookmark is missing or not in a table cell")
                eval_cell = _Cell(evaluation_element, document)
            else:
                eval_cell = actual_cells(table.rows[int(nested_spec["row"])])[int(nested_spec["cell"])]
            nested = eval_cell.tables[0]
            if len(nested.rows) != int(nested_spec["rows"]) or len(nested.columns) != int(nested_spec["columns"]):
                item_errors.append("evaluation table structure changed")
            score_values = [
                parse_decimal(nested.cell(row, 2).text, f"evaluation score row {row}")
                for row in range(1, int(nested_spec["rows"]))
            ]
            item_errors.extend(_evaluation_score_errors(score_values))
            target = parse_decimal(item.get("score", 89 + ((index - 1) % 6) * 0.5), f"evaluation target {index}")
            score_sum = sum(score_values, Decimal("0"))
            if score_sum != target:
                item_errors.append(f"evaluation total mismatch: expected {target}, got {score_sum}")
            for row_index, expected_values in enumerate(evaluation_cell_values(float(target), index), start=1):
                for cell_index, expected_value in expected_values.items():
                    actual_value = nested.cell(row_index, cell_index).text.strip()
                    if actual_value != expected_value:
                        item_errors.append(f"evaluation cell mismatch at row {row_index} cell {cell_index}")
        except (IndexError, ValueError) as exc:
            item_errors.append(f"evaluation table validation failed: {exc}")

        all_text = _document_text(document, table)
        for forbidden in manifest.get("validation", {}).get("forbidden_template_text", []):
            if forbidden in {course_expected, expected_course, str(item["unit"]), str(item["task"])}:
                continue
            if forbidden in all_text:
                item_errors.append(f"forbidden template text remains: {forbidden}")
        if expected_course != "Linux操作系统应用" and "Linux操作系统应用" in all_text:
            item_errors.append("template course-name placeholder Linux操作系统应用 remains")
        if item_errors:
            errors.extend(f"file {index}: {message}" for message in item_errors)
        lesson_checks.append({"file_index": index, "errors": item_errors, "fields_checked": sorted(field_values)})

    if is_semantic_manifest(manifest):
        missing_anchors = sorted({name for result in anchor_results for name in result["missing"]})
        duplicate_anchors = sorted({name for result in anchor_results for name in result["duplicates"]})
        report["preserved_anchor_count"] = min(
            (result["preserved_count"] for result in anchor_results),
            default=0,
        )
        report["missing_anchors"] = missing_anchors
        report["duplicate_anchors"] = duplicate_anchors
        checks["anchors"] = {
            "mode": "word_bookmark",
            "required": report["required_anchor_count"],
            "preserved": report["preserved_anchor_count"],
            "missing": missing_anchors,
            "duplicates": duplicate_anchors,
            "duplicate_ids": sorted({name for result in anchor_results for name in result["duplicate_ids"]}),
            "invalid_names": sorted({name for result in anchor_results for name in result["invalid_names"]}),
            "unexpected_names": sorted({name for result in anchor_results for name in result["unexpected_names"]}),
            "invalid_ids": sorted({name for result in anchor_results for name in result["invalid_ids"]}),
            "boundary_errors": sorted({name for result in anchor_results for name in result["boundary_errors"]}),
        }

    expected_total = data.get("total_hours")
    if expected_total is not None:
        try:
            if not math.isclose(total_hours, parse_number(expected_total, "total_hours"), abs_tol=0.01):
                errors.append(f"Total hours mismatch: expected {expected_total}, got {total_hours:g}")
        except ValueError as exc:
            errors.append(str(exc))
    checks["file_count"] = {"expected": len(lessons), "actual": len(files)}
    checks["total_hours"] = {"expected": expected_total, "actual": total_hours}
    checks["lessons"] = lesson_checks
    report["files_checked"] = len(files)
    if not errors:
        report["status"] = "skipped" if report["validation_skipped"] else "passed"

    _write_qa_report(report)
    if errors:
        raise RuntimeError("Output validation failed: " + "; ".join(errors[:8]))
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate generated lesson-plan DOCX files and write a QA report.")
    parser.add_argument("--input-json", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--manifest", default="")
    parser.add_argument("--schema", default=str(DEFAULT_SCHEMA))
    parser.add_argument("--qa-report", default="")
    parser.add_argument("--template-path", default="")
    parser.add_argument("--custom-template", action="store_true")
    parser.add_argument("--engine", default="")
    parser.add_argument("--skip-template-validation", action="store_true")
    parser.add_argument("--skip-validation", action="store_true")
    args = parser.parse_args()
    try:
        data = json.loads(Path(args.input_json).read_text(encoding="utf-8-sig"))
        template_path, _manifest_path, manifest = resolve_template_package(
            args.template_path or None,
            args.manifest or None,
        )
        custom_template = args.custom_template if args.custom_template else None
        if args.skip_validation:
            report = write_skipped_report(
                args.output_dir,
                data,
                manifest,
                args.qa_report or None,
                args.schema,
                template_path=template_path,
                custom_template=custom_template,
                engine=args.engine or None,
                template_validation=not args.skip_template_validation,
            )
        else:
            report = validate_output_dir(
                args.output_dir,
                data,
                manifest,
                args.qa_report or None,
                args.schema,
                template_path=template_path,
                custom_template=custom_template,
                engine=args.engine or None,
                template_validation=not args.skip_template_validation,
            )
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    for warning in report.get("warnings", []):
        print(f"WARNING: {warning}")
    action = "skipped validation" if report["status"] == "skipped" else "validated"
    print(f"{action} files={report['checks']['file_count']['actual']} qa={report['qa_report']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
