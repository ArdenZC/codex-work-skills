from __future__ import annotations

import argparse
import copy
import json
import os
import re
import shutil
import sys
import tempfile
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import _Cell

from bookmark_utils import bookmark_parent_cell, bookmark_parent_paragraph, find_bookmark
from content_contract import format_evaluation_values, format_implementation, format_reflection, lesson_content_field_values, lesson_filename, lesson_header_values, format_title
from content_quality import ContentQualityError, validate_content_quality
from package_common import DEFAULT_MANIFEST, DEFAULT_SCHEMA, ensure_supported_major, field_bookmark, field_spec, implementation_bookmarks, is_semantic_manifest, load_manifest, manifest_template_path, reflection_bookmarks, resolve_template_package, score_breakdown, validate_content_v2_input
from path_safety import assert_external_qa_path_safe, assert_output_path_safe, lesson_protected_paths, paths_equal
from validate_output import validate_output_dir, write_skipped_report
from validate_template import validate_template


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "templates" / "lesson-plan" / "v1.1.2" / "template.docx"
UNSAFE_VALIDATION_SKIP_ENV = "LESSON_ALLOW_UNSAFE_VALIDATION_SKIP"



def actual_cells(row):
    return [_Cell(tc, row._parent) for tc in row._tr.tc_lst]


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag not in {qn("w:pPr"), qn("w:bookmarkStart"), qn("w:bookmarkEnd")}:
            paragraph._p.remove(child)


def copy_run_format(src, dst):
    source_rpr = src._element.rPr
    target_rpr = dst._element.rPr
    if source_rpr is None:
        if target_rpr is not None:
            dst._element.remove(target_rpr)
        return
    if target_rpr is None:
        target_rpr = dst._element.get_or_add_rPr()
    for child in list(target_rpr):
        target_rpr.remove(child)
    for child in source_rpr:
        target_rpr.append(copy.deepcopy(child))


def copy_paragraph_format(src, dst):
    source_ppr = src._p.pPr
    target_ppr = dst._p.pPr
    if source_ppr is None:
        if target_ppr is not None:
            dst._p.remove(target_ppr)
        return
    if target_ppr is None:
        target_ppr = dst._p.get_or_add_pPr()
    for child in list(target_ppr):
        target_ppr.remove(child)
    for child in source_ppr:
        target_ppr.append(copy.deepcopy(child))


def set_paragraph_text(paragraph, text, align=None, source_run=None, bookmark_end=None):
    src = source_run or (paragraph.runs[0] if paragraph.runs else None)
    clear_paragraph(paragraph)
    run = paragraph.add_run(str(text))
    if src is not None:
        copy_run_format(src, run)
    if bookmark_end is not None and bookmark_end.getparent() is paragraph._p:
        paragraph._p.remove(run._r)
        paragraph._p.insert(list(paragraph._p).index(bookmark_end), run._r)
    if align is not None:
        paragraph.alignment = align

def set_cell_text(cell, text, align=None, preserve_cell_layout=False, bookmark_end=None):
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    set_paragraph_text(paragraph, text, align, bookmark_end=bookmark_end)
    for extra in list(cell.paragraphs)[1:]:
        extra._element.getparent().remove(extra._element)
    if not preserve_cell_layout:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_multiline(cell, text, align=None, bookmark_end=None):
    lines = str(text).splitlines() or [""]
    paragraphs = list(cell.paragraphs) if cell.paragraphs else [cell.add_paragraph()]
    source_paragraph = paragraphs[0]
    source_run = source_paragraph.runs[0] if source_paragraph.runs else None
    while len(paragraphs) < len(lines):
        paragraph = cell.add_paragraph()
        if source_paragraph.style is not None:
            paragraph.style = source_paragraph.style
        paragraphs.append(paragraph)
    for paragraph, line in zip(paragraphs, lines):
        if paragraph is not source_paragraph:
            copy_paragraph_format(source_paragraph, paragraph)
        set_paragraph_text(
            paragraph,
            line,
            align,
            source_run,
            bookmark_end=bookmark_end if paragraph is source_paragraph else None,
        )
    for extra in paragraphs[len(lines):]:
        extra._element.getparent().remove(extra._element)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row(table, row_idx, values_by_idx):
    cells = actual_cells(table.rows[row_idx])
    for idx, value in values_by_idx.items():
        if idx < len(cells):
            writer = set_cell_multiline if "\n" in str(value) else set_cell_text
            writer(cells[idx], value)


def _semantic_target(document, manifest: dict[str, Any], name: str):
    bookmark_name = field_bookmark(manifest, name)
    record = find_bookmark(document, bookmark_name)
    if record is None:
        raise ValueError(f"Semantic bookmark {bookmark_name} for field {name} is missing")
    spec = field_spec(manifest, name)
    target = spec["target"]
    if target == "document_paragraph":
        paragraph_element = bookmark_parent_paragraph(document, record)
        if paragraph_element is None:
            raise ValueError(f"Semantic bookmark {bookmark_name} for field {name} is not in a document paragraph")
        return Paragraph(paragraph_element, document), record.end
    if target == "table_cell":
        cell_element = bookmark_parent_cell(document, record)
        if cell_element is None:
            raise ValueError(f"Semantic bookmark {bookmark_name} for field {name} is not in a table cell")
        return _Cell(cell_element, document), record.end
    if target == "nested_table":
        raise ValueError(f"Semantic field {name} target nested_table requires the evaluation writer")
    raise ValueError(f"Unsupported semantic field target for fields.{name}: {target}")


def set_manifest_field(document, table, manifest: dict[str, Any], name: str, value: Any, align=None):
    spec = field_spec(manifest, name)
    if is_semantic_manifest(manifest):
        target, bookmark_end = _semantic_target(document, manifest, name)
        mode = spec["mode"]
        if isinstance(target, Paragraph):
            if mode != "replace_text_preserve_style":
                raise ValueError(f"Unsupported semantic field mode for fields.{name}: {mode}")
            set_paragraph_text(target, value, align, bookmark_end=bookmark_end)
            return
        if mode == "replace_single_paragraph":
            set_cell_text(target, value, align, bookmark_end=bookmark_end)
            return
        if mode == "replace_paragraphs":
            set_cell_multiline(target, value, align, bookmark_end=bookmark_end)
            return
        raise ValueError(f"Unsupported semantic field mode for fields.{name}: {mode}")
    if "table" not in spec or "row" not in spec or "cell" not in spec:
        raise ValueError(f"Field {name} is not a single table-cell field")
    target_table = table if int(spec["table"]) == 0 else table._parent.tables[int(spec["table"])]
    cell = actual_cells(target_table.rows[int(spec["row"])])[int(spec["cell"])]
    writer = set_cell_multiline if spec.get("mode") in {"replace_paragraphs", "replace_multiline"} else set_cell_text
    writer(cell, value, align)


def set_anchored_cell(document, manifest: dict[str, Any], bookmark_name: str, value: Any, multiline: bool = False):
    record = find_bookmark(document, bookmark_name)
    if record is None:
        raise ValueError(f"Semantic bookmark {bookmark_name} is missing")
    cell_element = bookmark_parent_cell(document, record)
    if cell_element is None:
        raise ValueError(f"Semantic bookmark {bookmark_name} is not in a table cell")
    cell = _Cell(cell_element, document)
    writer = set_cell_multiline if multiline else set_cell_text
    writer(cell, value, bookmark_end=record.end)


def add_eval_table(cell, target: float, lesson: dict[str, Any], manifest: dict[str, Any]):
    table = cell.tables[0] if cell.tables else cell.add_table(rows=14, cols=4)
    for r_idx, values in enumerate(
        format_evaluation_values(lesson, score_breakdown(target)),
        start=1,
    ):
        set_cell_text(table.cell(r_idx, 2), values[2], preserve_cell_layout=True)
        set_cell_text(table.cell(r_idx, 3), values[3], preserve_cell_layout=True)


def build_lesson(template: Path, out_dir: Path, meta: dict[str, Any], item: dict[str, Any], seq: int, manifest: dict[str, Any]) -> Path:
    header_values = lesson_header_values(meta, item)
    course = header_values["course_name"]
    major = header_values["major"]
    audience = header_values["audience"]
    hours = header_values["hours"]
    lesson_content = lesson_content_field_values(item)
    score = float(item["evaluation"]["score"])

    doc = Document(str(template))
    title_text = format_title(seq, meta, item)
    if is_semantic_manifest(manifest):
        target, bookmark_end = _semantic_target(doc, manifest, "title")
        set_paragraph_text(target, title_text, WD_ALIGN_PARAGRAPH.CENTER, bookmark_end=bookmark_end)
    else:
        title_spec = field_spec(manifest, "title")
        title_index = int(title_spec.get("paragraph", manifest["structure"]["title"]["paragraph"]))
        if title_index >= len(doc.paragraphs):
            raise ValueError(f"Title paragraph coordinate is invalid: {title_index}")
        set_paragraph_text(doc.paragraphs[title_index], title_text, WD_ALIGN_PARAGRAPH.CENTER)

    table = doc.tables[0]
    for name, value in {**header_values, **lesson_content}.items():
        set_manifest_field(doc, table, manifest, name, value)
    if is_semantic_manifest(manifest):
        evaluation_name = field_bookmark(manifest, "evaluation")
        evaluation_record = find_bookmark(doc, evaluation_name)
        if evaluation_record is None:
            raise ValueError(f"Semantic bookmark {evaluation_name} for field evaluation is missing")
        evaluation_element = bookmark_parent_cell(doc, evaluation_record)
        if evaluation_element is None:
            raise ValueError(f"Semantic bookmark {evaluation_name} for field evaluation is not in a table cell")
        evaluation_cell = _Cell(evaluation_element, doc)
    else:
        evaluation_spec = manifest["structure"]["evaluation_table"]
        evaluation_cell = actual_cells(table.rows[int(evaluation_spec["row"])])[int(evaluation_spec["cell"])]
    add_eval_table(evaluation_cell, score, item, manifest)

    if is_semantic_manifest(manifest):
        for bookmark_group, values in zip(implementation_bookmarks(manifest), format_implementation(item["implementation"])):
            for bookmark_name, value in zip(bookmark_group, [values[index] for index in range(5)]):
                set_anchored_cell(doc, manifest, bookmark_name, value, multiline="\n" in str(value))
        for bookmark_name, value in zip(reflection_bookmarks(manifest), format_reflection(item["reflection"])):
            set_anchored_cell(doc, manifest, bookmark_name, value, multiline="\n" in str(value))
    else:
        implementation_rows = [int(row) for row in manifest["fields"]["implementation"]["rows"]]
        for row_index, values in zip(implementation_rows, format_implementation(item["implementation"])):
            set_row(table, row_index, values)
        reflection_rows = [int(row) for row in manifest["fields"]["reflection"]["rows"]]
        for row_index, value in zip(reflection_rows, format_reflection(item["reflection"])):
            set_row(table, row_index, {2: value})

    out = out_dir / lesson_filename(seq, header_values["unit"], header_values["task"])
    doc.save(out)
    return out


def validate_outputs(
    out_dir: Path,
    meta: dict[str, Any],
    manifest: dict[str, Any],
    schema_path: Path,
    qa_report: Path | None = None,
    *,
    template: Path,
    custom_template: bool,
    template_validation: bool,
    template_warnings: list[str],
    render: bool,
) -> dict[str, Any]:
    return validate_output_dir(
        out_dir,
        meta,
        manifest,
        qa_report,
        schema_path,
        template_path=template,
        custom_template=custom_template,
        engine="python-docx",
        template_validation=template_validation,
        extra_warnings=template_warnings,
        render=render,
    )


def _unique_backup_path(out_dir: Path) -> Path:
    for _ in range(100):
        candidate = out_dir.parent / f"_{out_dir.name}_backup_{uuid.uuid4().hex[:12]}"
        if not candidate.exists():
            return candidate
    raise FileExistsError(f"Unable to allocate a backup path beside {out_dir}")


def _remove_path(path: Path) -> None:
    if path.is_dir() and not path.is_symlink():
        shutil.rmtree(path)
    elif path.exists() or path.is_symlink():
        path.unlink()


def _cleanup_path(path: Path, label: str) -> str | None:
    """Attempt cleanup without hiding a primary generation or commit error."""

    try:
        if path.exists() or path.is_symlink():
            _remove_path(path)
    except Exception as exc:  # pragma: no cover - filesystem failures vary by platform
        return f"cleanup failed for {label}: {exc}; residual path: {path}"
    return None


def _cleanup_empty_directory(path: Path, label: str) -> str | None:
    try:
        if path.exists():
            path.rmdir()
    except Exception as exc:  # pragma: no cover - filesystem failures vary by platform
        return f"cleanup failed for {label}: {exc}; residual path: {path}"
    return None


def atomic_commit_candidate(candidate: Path, out_dir: Path, backup_existing: bool) -> Path | None:
    """Swap a fully validated candidate into place and restore on commit failure."""

    displaced: Path | None = None
    try:
        if out_dir.exists():
            if not out_dir.is_dir():
                raise FileExistsError(f"Output path is not a directory: {out_dir}")
            if any(out_dir.iterdir()) and not backup_existing:
                raise FileExistsError(f"Output directory is not empty: {out_dir}")
            backup_path = _unique_backup_path(out_dir)
            os.replace(str(out_dir), str(backup_path))
            displaced = backup_path
        os.replace(str(candidate), str(out_dir))
        if displaced is not None and not backup_existing:
            _remove_path(displaced)
            displaced = None
        return displaced
    except Exception:
        if displaced is not None and out_dir.exists():
            _remove_path(out_dir)
        if displaced is not None and displaced.exists() and not out_dir.exists():
            try:
                os.replace(str(displaced), str(out_dir))
            except Exception as restore_error:  # pragma: no cover - only reachable on a second filesystem failure
                raise RuntimeError(f"Output commit failed and rollback failed: {restore_error}") from restore_error
        raise


def _unique_file_backup_path(path: Path) -> Path:
    for _ in range(100):
        candidate = path.parent / f"_{path.name}_backup_{uuid.uuid4().hex[:12]}"
        if not candidate.exists() and not candidate.is_symlink():
            return candidate
    raise FileExistsError(f"Unable to allocate a backup path beside {path}")


def _best_effort_post_commit_cleanup(path: Path | None, label: str) -> str | None:
    """Clean an old backup after publication without changing its outcome."""

    if path is None:
        return None
    try:
        _remove_path(path)
    except Exception as exc:  # pragma: no cover - filesystem failures vary by platform
        return f"cleanup failed for {label}: {exc}; residual backup path: {path}"
    return None


def atomic_commit_candidate_with_external_qa(
    candidate: Path,
    out_dir: Path,
    external_candidate: Path,
    external_qa: Path,
    backup_existing: bool,
) -> Path | None:
    """Commit output and an external QA report, restoring both on any failure."""

    displaced_output: Path | None = None
    displaced_qa: Path | None = None
    output_swapped = False
    qa_swapped = False
    try:
        if out_dir.exists():
            if not out_dir.is_dir():
                raise FileExistsError(f"Output path is not a directory: {out_dir}")
            if any(out_dir.iterdir()) and not backup_existing:
                raise FileExistsError(f"Output directory is not empty: {out_dir}")
            displaced_output = _unique_backup_path(out_dir)
            os.replace(str(out_dir), str(displaced_output))
        os.replace(str(candidate), str(out_dir))
        output_swapped = True

        if external_qa.exists() or external_qa.is_symlink():
            if external_qa.is_dir():
                raise FileExistsError(f"External QA report path must be a file, not a directory: {external_qa}")
            displaced_qa = _unique_file_backup_path(external_qa)
            os.replace(str(external_qa), str(displaced_qa))
        os.replace(str(external_candidate), str(external_qa))
        qa_swapped = True

        # Publication is complete once both new paths are in place.  Cleanup
        # is intentionally outside rollback: a filesystem failure while
        # deleting an old backup must leave the new output and QA available.
        cleanup_diagnostics = []
        if displaced_output is not None and not backup_existing:
            cleanup_error = _best_effort_post_commit_cleanup(displaced_output, "old output backup")
            if cleanup_error:
                cleanup_diagnostics.append(cleanup_error)
            else:
                displaced_output = None
        if displaced_qa is not None:
            cleanup_error = _best_effort_post_commit_cleanup(displaced_qa, "old external QA backup")
            if cleanup_error:
                cleanup_diagnostics.append(cleanup_error)
            else:
                displaced_qa = None
        for diagnostic in cleanup_diagnostics:
            print(f"WARNING: {diagnostic}", file=sys.stderr)
        return displaced_output
    except Exception as commit_error:
        rollback_errors: list[str] = []

        def rollback(action: str, callback) -> None:
            try:
                callback()
            except Exception as exc:  # pragma: no cover - injected filesystem failures vary by platform
                rollback_errors.append(f"{action}: {exc}")

        if qa_swapped and (external_qa.exists() or external_qa.is_symlink()):
            rollback("remove new external QA", lambda: _remove_path(external_qa))
        if displaced_qa is not None and displaced_qa.exists() and not external_qa.exists():
            rollback("restore external QA", lambda: os.replace(str(displaced_qa), str(external_qa)))
        if output_swapped and (out_dir.exists() or out_dir.is_symlink()):
            rollback("remove new output", lambda: _remove_path(out_dir))
        if displaced_output is not None and displaced_output.exists() and not out_dir.exists():
            rollback("restore output", lambda: os.replace(str(displaced_output), str(out_dir)))

        if rollback_errors:
            raise RuntimeError(
                "External QA commit failed and rollback failed: "
                + "; ".join(rollback_errors)
                + f"; original error: {commit_error}"
            ) from commit_error
        raise


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate template-matched Chinese lesson plan DOCX files.")
    parser.add_argument("--template", default="")
    parser.add_argument("--tasks-json", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--backup-existing", action="store_true")
    parser.add_argument("--manifest", default="")
    parser.add_argument("--schema", default=str(DEFAULT_SCHEMA))
    parser.add_argument("--skip-template-validation", action="store_true")
    parser.add_argument("--skip-output-validation", action="store_true")
    parser.add_argument("--render", action="store_true", help="Render validated DOCX files to disposable PDFs when a renderer is available")
    parser.add_argument("--qa-report", default="")
    args = parser.parse_args()
    if (args.skip_template_validation or args.skip_output_validation) and os.environ.get(UNSAFE_VALIDATION_SKIP_ENV) != "1":
        raise RuntimeError("Unsafe validation bypass is disabled.")

    template, manifest_path, manifest = resolve_template_package(
        args.template or None,
        args.manifest or None,
    )
    ensure_supported_major(manifest)
    if not template.exists():
        raise FileNotFoundError(f"Template not found: {template}")
    out_dir = Path(args.output_dir).expanduser().absolute()
    source_path = Path(args.tasks_json).expanduser().resolve()
    schema_path = Path(args.schema).expanduser().resolve()
    with source_path.open("r", encoding="utf-8") as f:
        meta = json.load(f)
    validate_content_v2_input(meta, schema_path)
    lessons = meta["lessons"]

    package_roots = [manifest_path.parent]
    base_manifest = manifest.get("template", {}).get("base_manifest")
    if base_manifest:
        package_roots.append((manifest_path.parent / str(base_manifest)).resolve().parent)
    protected_paths = lesson_protected_paths(
        skill_dir=SKILL_DIR,
        source=source_path,
        schema=schema_path,
        template=template,
        manifest=manifest_path,
        package_roots=package_roots,
    )
    assert_output_path_safe(out_dir, protected_paths)
    requested_qa = Path(args.qa_report).expanduser().resolve() if args.qa_report else None
    internal_qa = out_dir / "qa-report.json"
    external_qa = None
    qa_parent_created = False
    if requested_qa is not None:
        same_as_internal = paths_equal(requested_qa, internal_qa)
        if not same_as_internal:
            external_qa = assert_external_qa_path_safe(requested_qa, out_dir, protected_paths)
    if out_dir.exists() and not out_dir.is_dir():
        raise FileExistsError(f"Output path is not a directory: {out_dir}")
    if out_dir.exists() and any(out_dir.iterdir()) and not args.backup_existing:
        raise FileExistsError(f"Output directory is not empty: {out_dir}")

    template_warnings: list[str] = []
    if not args.skip_template_validation:
        template_report = validate_template(template, manifest_path)
        template_warnings = template_report.get("warnings", [])
        for warning in template_warnings:
            print(f"WARNING: {warning}")

    out_dir.parent.mkdir(parents=True, exist_ok=True)
    candidate = Path(tempfile.mkdtemp(prefix=f".{out_dir.name}.candidate-", dir=str(out_dir.parent)))
    assert_output_path_safe(candidate, protected_paths)
    external_candidate: Path | None = None
    operation_error: BaseException | None = None
    try:
        content_quality = validate_content_quality(meta, manifest)
        generated_filenames: list[str] = []
        for seq, item in enumerate(lessons, 1):
            generated_filenames.append(build_lesson(template, candidate, meta, item, seq, manifest).name)
        candidate_qa = candidate / "qa-report.json"
        if not args.skip_output_validation:
            report = validate_outputs(
                candidate,
                meta,
                manifest,
                schema_path,
                candidate_qa,
                template=template,
                custom_template=bool(args.template),
                template_validation=not args.skip_template_validation,
                template_warnings=template_warnings,
                render=args.render,
            )
        else:
            report = write_skipped_report(
                candidate,
                meta,
                manifest,
                candidate_qa,
                schema_path,
                template_path=template,
                custom_template=bool(args.template),
                engine="python-docx",
                template_validation=not args.skip_template_validation,
                warnings=template_warnings,
                render=args.render,
            )
        final_qa = requested_qa or internal_qa
        report["output_dir"] = str(out_dir)
        report["qa_report"] = str(final_qa)
        report_text = json.dumps(report, ensure_ascii=False, indent=2) + "\n"
        candidate_qa.write_text(report_text, encoding="utf-8")
        if external_qa is not None:
            if not external_qa.parent.exists():
                external_qa.parent.mkdir(parents=True, exist_ok=True)
                qa_parent_created = True
            fd, external_candidate_name = tempfile.mkstemp(
                prefix=f".{external_qa.name}.candidate-",
                suffix=".tmp",
                dir=str(external_qa.parent),
            )
            os.close(fd)
            external_candidate = Path(external_candidate_name)
            external_candidate.write_text(report_text, encoding="utf-8")
            backup = atomic_commit_candidate_with_external_qa(
                candidate,
                out_dir,
                external_candidate,
                external_qa,
                args.backup_existing,
            )
        else:
            backup = atomic_commit_candidate(candidate, out_dir, args.backup_existing)
        for filename in generated_filenames:
            print(out_dir / filename)
        if backup is not None:
            print(f"backup={backup}")
        action = "skipped validation" if report["status"] == "skipped" else "validated"
        print(f"{action} files={report['checks']['file_count']['actual']} total_hours={report['checks']['total_hours']['actual']:g} qa={report['qa_report']}")
    except BaseException as exc:
        operation_error = exc
        raise
    finally:
        cleanup_diagnostics: list[str] = []
        candidate_error = _cleanup_path(candidate, "candidate directory")
        if candidate_error:
            cleanup_diagnostics.append(candidate_error)
        if external_candidate is not None:
            external_candidate_error = _cleanup_path(external_candidate, "external QA candidate")
            if external_candidate_error:
                cleanup_diagnostics.append(external_candidate_error)
        # A successful commit intentionally leaves the newly created parent
        # containing the external report.  Only remove it after a failed
        # operation, when rollback has left no committed report behind.
        if operation_error is not None and qa_parent_created and external_qa is not None and external_qa.parent.exists():
            parent_error = _cleanup_empty_directory(external_qa.parent, "external QA parent directory")
            if parent_error:
                cleanup_diagnostics.append(parent_error)
        for diagnostic in cleanup_diagnostics:
            print(f"WARNING: {diagnostic}", file=sys.stderr)
            if operation_error is not None:
                operation_error.add_note(diagnostic)


if __name__ == "__main__":
    main()
