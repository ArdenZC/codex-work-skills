from __future__ import annotations

import argparse
import copy
import hashlib
import json
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from lxml import etree

from bookmark_utils import bookmark_boundary_locations, bookmark_location, validate_bookmark_inventory
from package_common import DEFAULT_MANIFEST, bookmark_containers, base_manifest_path, ensure_supported_major, field_spec, is_semantic_manifest, layout_manifest, load_manifest, manifest_template_path, required_bookmarks, resolve_template_package


class TemplateValidationError(RuntimeError):
    def __init__(self, report: dict[str, Any]):
        self.report = report
        message = "; ".join(report.get("errors", [])) or "Template validation failed"
        super().__init__(message)


def actual_cells(row) -> list[_Cell]:
    return [_Cell(tc, row._parent) for tc in row._tr.tc_lst]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def _xml_without_text(element) -> str:
    cloned = copy.deepcopy(element)
    for node in cloned.iter():
        if node.tag in {qn("w:t"), qn("w:instrText"), qn("w:delText")}:
            node.text = ""
    return etree.tostring(cloned, encoding="unicode")


def _clear_text_nodes(element) -> None:
    for node in element.iter():
        if node.tag in {qn("w:t"), qn("w:instrText"), qn("w:delText")}:
            node.text = ""


def _evaluation_writable_coordinates(manifest: dict[str, Any]) -> tuple[set[int], set[int]]:
    structure = manifest["structure"]["evaluation_table"]
    spec = manifest.get("fields", {}).get("evaluation", {})
    rows = {int(row) for row in spec.get("writable_rows", [])}
    cells = {int(cell) for cell in spec.get("writable_cells", [])}
    if not rows:
        rows = set(range(1, int(structure["rows"])))
    if not cells:
        cells = {2, 3}
    return rows, cells


def _settings_xml(document) -> str:
    value = etree.tostring(copy.deepcopy(document.settings._element), encoding="unicode")
    return value.encode("unicode_escape").decode("ascii")


def _cell_signature(
    cell,
    writable: bool = False,
    evaluation: bool = False,
    evaluation_writable_rows: set[int] | None = None,
    evaluation_writable_cells: set[int] | None = None,
) -> str:
    cloned = copy.deepcopy(cell._tc)
    if writable:
        _clear_text_nodes(cloned)
    elif evaluation:
        writable_rows = evaluation_writable_rows or set(range(1, 14))
        writable_cells = evaluation_writable_cells or {2, 3}
        for nested_table in cloned.iter(qn("w:tbl")):
            rows = [child for child in nested_table if child.tag == qn("w:tr")]
            for row_index, row in enumerate(rows):
                if row_index not in writable_rows:
                    continue
                nested_cells = [child for child in row if child.tag == qn("w:tc")]
                for cell_index in writable_cells:
                    if cell_index < len(nested_cells):
                        _clear_text_nodes(nested_cells[cell_index])
    return etree.tostring(cloned, encoding="unicode")


def _writable_main_table_cells(manifest: dict[str, Any]) -> tuple[set[tuple[int, int]], tuple[int, int] | None]:
    writable: set[tuple[int, int]] = set()
    evaluation: tuple[int, int] | None = None
    for spec in manifest.get("fields", {}).values():
        if spec.get("mode") == "row_cells":
            writable.update(
                (int(row), int(cell))
                for row in spec.get("rows", [])
                for cell in spec.get("cells", [])
            )
            continue
        if spec.get("mode") == "nested_table":
            if int(spec.get("table", -1)) == 0:
                evaluation = (int(spec["row"]), int(spec["cell"]))
            continue
        if int(spec.get("table", -1)) == 0 and "row" in spec and "cell" in spec:
            writable.add((int(spec["row"]), int(spec["cell"])))
    return writable, evaluation


def _styles_signature(document) -> str:
    return etree.tostring(copy.deepcopy(document.styles._element), encoding="unicode")


def _theme_signature(document) -> list[dict[str, Any]]:
    return [
        {"partname": str(part.partname), "blob": part.blob}
        for part in document.part.package.iter_parts()
        if str(part.partname).startswith("/word/theme/")
    ]


def _body_paragraph_signature(document, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    title_spec = manifest.get("fields", {}).get("title", {})
    title_index = int(title_spec.get("paragraph", manifest["structure"]["title"]["paragraph"]))
    return [
        {
            "index": index,
            "xml": etree.tostring(copy.deepcopy(paragraph._p), encoding="unicode"),
        }
        for index, paragraph in enumerate(document.paragraphs)
        if index != title_index
    ]


def _title_format_signature(document, manifest: dict[str, Any]) -> dict[str, Any]:
    title_spec = manifest.get("fields", {}).get("title", {})
    title_index = int(title_spec.get("paragraph", manifest["structure"]["title"]["paragraph"]))
    if title_index < 0 or title_index >= len(document.paragraphs):
        return {"paragraph": "", "run_formats": [], "run_primary": ""}
    paragraph = document.paragraphs[title_index]
    ppr = copy.deepcopy(paragraph._p.pPr)
    if ppr is not None:
        for child in list(ppr):
            if child.tag == qn("w:jc"):
                ppr.remove(child)
    run_values = [
        etree.tostring(copy.deepcopy(run._r.rPr), encoding="unicode")
        for run in paragraph.runs
        if run._r.rPr is not None
    ]
    return {
        "paragraph": etree.tostring(ppr, encoding="unicode") if ppr is not None else "",
        "run_formats": sorted({value for value in run_values if value}),
        "run_primary": next((value for value in run_values if value), ""),
    }


def _main_table_signature(document, manifest: dict[str, Any]) -> dict[str, Any]:
    table = document.tables[int(manifest["structure"]["main_table"]["index"])]
    writable_cells, evaluation_cell = _writable_main_table_cells(manifest)
    evaluation_writable_rows, evaluation_writable_cells = _evaluation_writable_coordinates(manifest)
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        for cell_index, cell in enumerate(actual_cells(row)):
            cells.append(
                _cell_signature(
                    cell,
                    writable=(row_index, cell_index) in writable_cells,
                    evaluation=evaluation_cell == (row_index, cell_index),
                    evaluation_writable_rows=evaluation_writable_rows,
                    evaluation_writable_cells=evaluation_writable_cells,
                )
            )
        rows.append({"tr": _xml_without_text(row._tr), "cells": cells})
    return {
        "table": _xml_without_text(table._tbl),
        "rows": rows,
        "row_count": len(table.rows),
        "column_count": len(table.columns),
        "styles": _styles_signature(document),
        "themes": _theme_signature(document),
        "body_paragraphs": _body_paragraph_signature(document, manifest),
        "title_format": _title_format_signature(document, manifest),
    }


def _section_signature(document) -> list[dict[str, Any]]:
    result = []
    for section in document.sections:
        result.append(
            {
                "page_width": section.page_width.twips,
                "page_height": section.page_height.twips,
                "top_margin": section.top_margin.twips,
                "bottom_margin": section.bottom_margin.twips,
                "left_margin": section.left_margin.twips,
                "right_margin": section.right_margin.twips,
                "sectPr": etree.tostring(copy.deepcopy(section._sectPr), encoding="unicode"),
                "header": [p.text for p in section.header.paragraphs],
                "footer": [p.text for p in section.footer.paragraphs],
                "header_xml": etree.tostring(copy.deepcopy(section.header._element), encoding="unicode"),
                "footer_xml": etree.tostring(copy.deepcopy(section.footer._element), encoding="unicode"),
                "first_page_header_xml": etree.tostring(copy.deepcopy(section.first_page_header._element), encoding="unicode"),
                "first_page_footer_xml": etree.tostring(copy.deepcopy(section.first_page_footer._element), encoding="unicode"),
                "even_page_header_xml": etree.tostring(copy.deepcopy(section.even_page_header._element), encoding="unicode"),
                "even_page_footer_xml": etree.tostring(copy.deepcopy(section.even_page_footer._element), encoding="unicode"),
                "odd_and_even_pages_header_footer": bool(document.settings.odd_and_even_pages_header_footer),
                "settings_xml": _settings_xml(document),
            }
        )
    return result


def _all_text(document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in actual_cells(row):
                values.append(cell.text)
                for nested in cell.tables:
                    for nested_row in nested.rows:
                        values.extend(c.text for c in nested_row.cells)
    return "\n".join(values)


def _check_field_coordinates(document, manifest: dict[str, Any], errors: list[str]) -> None:
    if is_semantic_manifest(manifest):
        return
    table_count = len(document.tables)
    for name, spec in manifest.get("fields", {}).items():
        if spec.get("target") == "document_paragraph" or "paragraph" in spec:
            paragraph_index = int(spec.get("paragraph", -1))
            if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
                errors.append(f"Field {name} points to missing document paragraph {paragraph_index}")
            continue
        if spec.get("mode") == "row_cells":
            rows = spec.get("rows", [])
            cells = spec.get("cells", [])
            for row_index in rows:
                for cell_index in cells:
                    if (
                        row_index < 0
                        or cell_index < 0
                        or row_index >= manifest["structure"]["main_table"]["rows"]
                        or cell_index >= manifest["structure"]["main_table"]["columns"]
                    ):
                        errors.append(f"Field {name} points outside the declared main table: row={row_index}, cell={cell_index}")
            continue
        if "table" not in spec or "row" not in spec or "cell" not in spec:
            continue
        table_index = int(spec["table"])
        row_index = int(spec["row"])
        cell_index = int(spec["cell"])
        if table_index < 0 or table_index >= table_count:
            errors.append(f"Field {name} points to missing table {table_index}")
            continue
        table = document.tables[table_index]
        if row_index < 0 or row_index >= len(table.rows):
            errors.append(f"Field {name} points to missing row {row_index}")
            continue
        if cell_index < 0 or cell_index >= len(actual_cells(table.rows[row_index])):
            errors.append(f"Field {name} points to missing cell {cell_index} in row {row_index}")


def _document_xml_without_bookmarks(path: Path) -> bytes:
    with zipfile.ZipFile(path, "r") as package:
        root = etree.fromstring(package.read("word/document.xml"))
    word_namespace = qn("w:body").split("}", 1)[0][1:]
    for node in list(root.xpath(".//w:bookmarkStart | .//w:bookmarkEnd", namespaces={"w": word_namespace})):
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)
    return etree.tostring(root, method="c14n", with_comments=False)


def _docx_equivalent_without_bookmarks(actual: Path, reference: Path) -> bool:
    with zipfile.ZipFile(actual, "r") as actual_zip, zipfile.ZipFile(reference, "r") as reference_zip:
        actual_names = set(actual_zip.namelist())
        reference_names = set(reference_zip.namelist())
        names = actual_names | reference_names
        for name in names:
            if name == "word/document.xml":
                if _document_xml_without_bookmarks(actual) != _document_xml_without_bookmarks(reference):
                    return False
                continue
            if name not in actual_names or name not in reference_names or actual_zip.read(name) != reference_zip.read(name):
                return False
    return True


def validate_template(
    template_path: Path | str,
    manifest_path: Path | str = DEFAULT_MANIFEST,
    compatibility_template: Path | str | None = None,
) -> dict[str, Any]:
    template = Path(template_path).expanduser().resolve()
    manifest = load_manifest(manifest_path)
    errors: list[str] = []
    warnings: list[str] = []
    report: dict[str, Any] = {
        "template": str(template),
        "manifest": str(Path(manifest_path).expanduser().resolve()),
        "template_version": manifest.get("template", {}).get("version"),
        "errors": errors,
        "warnings": warnings,
        "checks": {},
    }

    try:
        ensure_supported_major(manifest)
    except ValueError as exc:
        errors.append(str(exc))

    if not template.exists():
        errors.append(f"Template not found: {template}")
        raise TemplateValidationError(report)

    expected_hash = str(manifest.get("fingerprint", {}).get("sha256", "")).upper()
    actual_hash = sha256(template)
    canonical = manifest_template_path(manifest)
    is_canonical = template == canonical
    if not is_canonical and not canonical.exists():
        errors.append(f"Canonical template not found: {canonical}")
    report["checks"]["sha256"] = {"expected": expected_hash, "actual": actual_hash}
    if is_canonical and actual_hash != expected_hash:
        errors.append(f"Canonical template SHA-256 mismatch: expected {expected_hash}, got {actual_hash}")
    elif not is_canonical and actual_hash != expected_hash:
        warnings.append(
            f"Custom template fingerprint differs from the {manifest.get('template', {}).get('version')} canonical template."
        )

    if compatibility_template is None:
        entries = manifest.get("template", {}).get("compatibility_entries", [])
        if entries:
            compatibility_template = Path(manifest["_path"]).parent / entries[0]
    if compatibility_template:
        compat = Path(compatibility_template).expanduser().resolve()
        if compat.exists():
            compat_hash = sha256(compat)
            report["checks"]["compatibility_sha256"] = {"path": str(compat), "actual": compat_hash}
            if compat_hash != expected_hash:
                errors.append(f"Compatibility template diverges from canonical template: {compat}")
        else:
            warnings.append(f"Compatibility template entry is missing: {compat}")

    try:
        document = Document(str(template))
        report["checks"]["docx_open"] = True
    except Exception as exc:  # pragma: no cover - library-specific parse errors
        errors.append(f"DOCX could not be opened: {exc}")
        raise TemplateValidationError(report)

    structure = manifest.get("structure", {})
    main_spec = structure.get("main_table", {})
    table_index = int(main_spec.get("index", 0))
    expected_top_level_tables = int(structure.get("top_level_tables", 1))
    if len(document.tables) != expected_top_level_tables:
        errors.append(f"Top-level table count mismatch: expected {expected_top_level_tables}, got {len(document.tables)}")
    if len(document.tables) <= table_index:
        errors.append(f"Missing main table at index {table_index}")
    else:
        table = document.tables[table_index]
        if len(table.rows) != int(main_spec.get("rows", 0)):
            errors.append(f"Main table row count mismatch: expected {main_spec.get('rows')}, got {len(table.rows)}")
        if len(table.columns) != int(main_spec.get("columns", 0)):
            errors.append(f"Main table column count mismatch: expected {main_spec.get('columns')}, got {len(table.columns)}")
        report["checks"]["main_table"] = {"rows": len(table.rows), "columns": len(table.columns)}

    evaluation = structure.get("evaluation_table", {})
    try:
        eval_cell = actual_cells(document.tables[int(evaluation["table"])].rows[int(evaluation["row"])])[int(evaluation["cell"])]
        nested = eval_cell.tables[0]
        if len(nested.rows) != int(evaluation["rows"]) or len(nested.columns) != int(evaluation["columns"]):
            errors.append(
                f"Evaluation table size mismatch: expected {evaluation.get('rows')}x{evaluation.get('columns')}, "
                f"got {len(nested.rows)}x{len(nested.columns)}"
            )
        report["checks"]["evaluation_table"] = {"rows": len(nested.rows), "columns": len(nested.columns)}
    except (IndexError, KeyError, TypeError) as exc:
        errors.append(f"Evaluation table coordinate is invalid: {exc}")

    if is_semantic_manifest(manifest):
        inventory = validate_bookmark_inventory(
            document,
            required_bookmarks(manifest),
            bookmark_containers(manifest),
        )
        report["checks"]["bookmarks"] = {
            "required_count": inventory["required_count"],
            "actual_count": inventory["main_count"],
            "missing": inventory["missing"],
            "duplicates": inventory["duplicates"],
            "duplicate_ids": inventory["duplicate_ids"],
            "invalid_names": inventory["invalid_names"],
            "unexpected_names": inventory["unexpected_names"],
            "invalid_ids": inventory["invalid_ids"],
            "orphaned": inventory["orphaned"],
            "outside_main": inventory["outside_main"],
            "container_errors": inventory["container_errors"],
            "boundary_errors": inventory["boundary_errors"],
        }
        errors.extend(f"Semantic bookmark protection failed: {error}" for error in inventory["errors"])

    labels = manifest.get("validation", {}).get("required_labels", [])
    document_text = _all_text(document)
    missing_labels = [label for label in labels if label not in document_text]
    if missing_labels:
        errors.append(f"Missing fixed labels: {', '.join(missing_labels)}")
    _check_field_coordinates(document, manifest, errors)
    report["checks"]["required_labels"] = {"count": len(labels), "missing": missing_labels}
    document_sections = _section_signature(document)
    report["checks"]["sections"] = document_sections

    if is_semantic_manifest(manifest):
        reference = canonical
        if is_canonical:
            base_value = manifest.get("template", {}).get("base_template")
            reference = (Path(manifest["_path"]).parent / str(base_value)).resolve() if base_value else Path()
        if not reference.exists():
            errors.append(f"Semantic template reference not found: {reference}")
        elif not _docx_equivalent_without_bookmarks(template, reference):
            errors.append("Semantic template changed visible content or structure relative to the protected baseline.")
        if not is_canonical and canonical.exists() and inventory["valid"]:
            canonical_document = Document(str(canonical))
            location_changes = [
                name
                for name in required_bookmarks(manifest)
                if bookmark_boundary_locations(document, name) != bookmark_boundary_locations(canonical_document, name)
            ]
            if location_changes:
                errors.append("Semantic bookmark location changed: " + ", ".join(location_changes))
    elif canonical.exists() and not is_canonical:
        canonical_doc = Document(str(canonical))
        if _main_table_signature(canonical_doc, manifest) != _main_table_signature(document, manifest):
            errors.append("Custom template changed protected main-table structure or formatting.")
        if _section_signature(canonical_doc) != document_sections:
            errors.append("Custom template changed protected page, header, footer, or section settings.")

    if errors:
        raise TemplateValidationError(report)
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Validate the versioned DOCX lesson-plan template package.")
    parser.add_argument("--template", default="")
    parser.add_argument("--manifest", default="")
    parser.add_argument("--compatibility-template", default="")
    parser.add_argument("--json", action="store_true", dest="as_json")
    args = parser.parse_args()
    try:
        template, manifest_path, manifest = resolve_template_package(
            args.template or None,
            args.manifest or None,
        )
        report = validate_template(template, manifest_path, args.compatibility_template or None)
    except TemplateValidationError as exc:
        report = exc.report
        if args.as_json:
            print(json.dumps(report, ensure_ascii=False, indent=2))
        else:
            for error in report.get("errors", []):
                print(f"ERROR: {error}", file=sys.stderr)
            for warning in report.get("warnings", []):
                print(f"WARNING: {warning}", file=sys.stderr)
        return 1
    except Exception as exc:
        report = {
            "template": str(Path(args.template).expanduser().resolve()) if args.template else "",
            "manifest": str(Path(args.manifest).expanduser().resolve()) if args.manifest else str(DEFAULT_MANIFEST),
            "template_version": None,
            "errors": [str(exc)],
            "warnings": [],
            "checks": {},
        }
        if args.as_json:
            print(json.dumps(report, ensure_ascii=False, indent=2))
        else:
            print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    if args.as_json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        for warning in report.get("warnings", []):
            print(f"WARNING: {warning}")
        print(f"validated template={template} version={report['template_version']} sha256={report['checks']['sha256']['actual']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
