from __future__ import annotations

import json
import hashlib
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unittest
import zipfile
from copy import copy, deepcopy
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree
from openpyxl import Workbook, load_workbook
from openpyxl.workbook.defined_name import DefinedName
import yaml


ROOT = Path(__file__).resolve().parents[1]
LESSON = ROOT / "教案生成器" / "lesson-plan-docx-generator"
GRADE = ROOT / "平时成绩记分册生成器" / "course-gradebook-generator"
PYTHON = Path(sys.executable)
LESSON_V10_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "manifest.yaml"
LESSON_V11_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "manifest.yaml"
LESSON_V111_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.1" / "manifest.yaml"
GRADE_V10_MANIFEST = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"
GRADE_V10_TEMPLATE = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
GRADE_V11_MANIFEST = GRADE / "assets" / "templates" / "course-gradebook" / "v1.1.0" / "manifest.yaml"
GRADE_V11_TEMPLATE = GRADE / "assets" / "templates" / "course-gradebook" / "v1.1.0" / "template.xls"


_XLSX_NS = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}


def patch_xlsx_cell_font(path: Path, address: str, font_name: str) -> None:
    """Clone one cell style and change only its font name in the XLSX package."""
    with zipfile.ZipFile(path, "r") as source:
        sheet = etree.fromstring(source.read("xl/worksheets/sheet1.xml"))
        styles = etree.fromstring(source.read("xl/styles.xml"))
        cell = sheet.xpath(f".//main:c[@r='{address}']", namespaces=_XLSX_NS)[0]
        cell_xfs = styles.xpath(".//main:cellXfs/main:xf", namespaces=_XLSX_NS)
        fonts = styles.xpath(".//main:fonts/main:font", namespaces=_XLSX_NS)
        style_index = int(cell.get("s", "0"))
        source_xf = cell_xfs[style_index]
        font_index = int(source_xf.get("fontId", "0"))
        changed_font = etree.fromstring(etree.tostring(fonts[font_index]))
        font_node = changed_font.find("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}name")
        if font_node is None:
            font_node = etree.Element("{http://schemas.openxmlformats.org/spreadsheetml/2006/main}name")
            changed_font.insert(0, font_node)
        font_node.set("val", font_name)
        fonts_parent = styles.xpath(".//main:fonts", namespaces=_XLSX_NS)[0]
        fonts_parent.append(changed_font)
        new_font_index = len(fonts)
        changed_xf = etree.fromstring(etree.tostring(source_xf))
        changed_xf.set("fontId", str(new_font_index))
        xfs_parent = styles.xpath(".//main:cellXfs", namespaces=_XLSX_NS)[0]
        xfs_parent.append(changed_xf)
        cell.set("s", str(len(cell_xfs)))
        fonts_parent.set("count", str(len(fonts) + 1))
        xfs_parent.set("count", str(len(cell_xfs) + 1))
        sheet_bytes = etree.tostring(sheet, xml_declaration=True, encoding="UTF-8", standalone=True)
        styles_bytes = etree.tostring(styles, xml_declaration=True, encoding="UTF-8", standalone=True)
        temp_path = path.with_suffix(path.suffix + ".tmp")
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename == "xl/worksheets/sheet1.xml":
                    data = sheet_bytes
                elif info.filename == "xl/styles.xml":
                    data = styles_bytes
                target.writestr(info, data)
    temp_path.replace(path)


def soffice_path() -> str | None:
    candidates = [
        shutil.which("soffice"),
        shutil.which("soffice.com"),
        r"C:\Program Files\LibreOffice\program\soffice.com",
    ]
    return next((item for item in candidates if item and Path(item).exists()), None)


def convert_with_soffice(source: Path, out_dir: Path, target_format: str) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [soffice_path(), "--headless", "--convert-to", target_format, "--outdir", str(out_dir), str(source)],
        check=True,
        capture_output=True,
    )
    target = out_dir / f"{source.stem}.{target_format}"
    if not target.exists():
        raise AssertionError(f"LibreOffice did not create {target}")
    return target


def xlsx_font_signature(path: Path, address: str) -> tuple[object, ...]:
    workbook = load_workbook(path, data_only=False)
    font = workbook["平时成绩"][address].font
    color = font.color
    color_signature = () if color is None else (
        color.type,
        color.rgb,
        color.indexed,
        color.auto,
        color.theme,
        color.tint,
    )
    return (
        font.name,
        font.charset,
        font.family,
        font.scheme,
        font.sz,
        bool(font.b),
        bool(font.i),
        font.u,
        bool(font.strike),
        bool(font.outline),
        bool(font.shadow),
        font.vertAlign,
        color_signature,
    )


def find_roundtrip_font_tamper(
    source_xlsx: Path,
    controlled_baseline_xlsx: Path,
    font_name: str,
    candidates: tuple[str, ...],
    work_dir: Path,
    label: str,
    accept=None,
) -> tuple[Path, str]:
    """Find a font-only XLS round trip that differs from the controlled baseline."""
    work_dir.mkdir(parents=True, exist_ok=True)
    for index, address in enumerate(candidates):
        tampered_xlsx = work_dir / f"{label}-{index}.xlsx"
        shutil.copy2(source_xlsx, tampered_xlsx)
        patch_xlsx_cell_font(tampered_xlsx, address, font_name)
        tampered_xls = convert_with_soffice(
            tampered_xlsx,
            work_dir / f"{label}-{index}-xls",
            "xls",
        )
        final_xlsx = convert_with_soffice(
            tampered_xls,
            work_dir / f"{label}-{index}-final",
            "xlsx",
        )
        signature_differs = xlsx_font_signature(final_xlsx, address) != xlsx_font_signature(
            controlled_baseline_xlsx,
            address,
        )
        if (accept is not None and accept(tampered_xls, address)) or (accept is None and signature_differs):
            return tampered_xls, address
    raise AssertionError(
        f"{font_name} did not produce a distinct font signature in candidates {candidates}"
    )


def run_script(
    script: Path,
    *args: str,
    env: dict[str, str] | None = None,
) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [str(PYTHON), str(script), *args],
        cwd=ROOT,
        env=env,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        check=False,
    )


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def package_file_hashes(root: Path) -> dict[str, str]:
    return {
        path.relative_to(root).as_posix(): file_sha256(path)
        for path in sorted(root.rglob("*"))
        if path.is_file()
    }


def write_gradebook_manifest(
    folder: Path,
    source_manifest: Path,
    source_template: Path,
    mutate=None,
) -> tuple[Path, Path]:
    """Create a self-contained custom gradebook package with a matching fingerprint."""
    package = folder / "gradebook-package"
    package.mkdir(parents=True, exist_ok=True)
    template = package / "template.xls"
    shutil.copy2(source_template, template)
    manifest = yaml.safe_load(source_manifest.read_text(encoding="utf-8"))
    manifest["template"]["file"] = "template.xls"
    manifest["template"].pop("compatibility_entries", None)
    if str(manifest["template"].get("version", "")).startswith("1.1."):
        base = package / "v1.0.0"
        base.mkdir(exist_ok=True)
        shutil.copy2(GRADE_V10_MANIFEST, base / "manifest.yaml")
        shutil.copy2(GRADE_V10_TEMPLATE, base / "template.xls")
        manifest["template"]["base_manifest"] = "v1.0.0/manifest.yaml"
        manifest["template"]["base_template"] = "v1.0.0/template.xls"
    else:
        manifest["template"].pop("base_manifest", None)
        manifest["template"].pop("base_template", None)
    digest = file_sha256(template)
    manifest["fingerprint"]["sha256"] = digest
    manifest["fingerprint"]["value"] = digest
    if mutate is not None:
        mutate(manifest)
    manifest_path = package / "manifest.yaml"
    manifest_path.write_text(yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8")
    return template, manifest_path


def write_declared_template_gradebook_package(folder: Path) -> tuple[Path, Path, Path, dict[str, Path]]:
    """Create a v1.1 package whose selected and declared templates differ."""
    template, manifest_path = write_gradebook_manifest(folder, GRADE_V11_MANIFEST, GRADE_V11_TEMPLATE)
    package = template.parent
    actual_template = package / "actual-template.xls"
    declared_template = package / "declared-template.xls"
    shutil.copy2(template, actual_template)
    shutil.copy2(template, declared_template)
    template.unlink()
    manifest = yaml.safe_load(manifest_path.read_text(encoding="utf-8"))
    manifest["template"]["file"] = declared_template.name
    digest = file_sha256(actual_template)
    manifest["fingerprint"]["sha256"] = digest
    manifest["fingerprint"]["value"] = digest
    manifest_path.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )
    protected = {
        "actual_template": actual_template,
        "declared_template": declared_template,
        "manifest": manifest_path,
        "base_template": package / "v1.0.0" / "template.xls",
        "base_manifest": package / "v1.0.0" / "manifest.yaml",
    }
    return package, actual_template, manifest_path, protected


def patch_xlsx_named_range(
    path: Path,
    name: str,
    *,
    attr_text: str | None = None,
    remove: bool = False,
    new_name: str | None = None,
    hidden: bool | None = None,
    local_sheet_id: int | None | object = None,
) -> None:
    workbook = load_workbook(path)
    if remove:
        if name in workbook.defined_names:
            del workbook.defined_names[name]
    else:
        defined = workbook.defined_names.get(name)
        if defined is None:
            defined = DefinedName(name=name, attr_text=attr_text or "'平时成绩'!$A$1")
            workbook.defined_names.add(defined)
        if new_name is not None:
            del workbook.defined_names[name]
            defined.name = new_name
            workbook.defined_names.add(defined)
        if attr_text is not None:
            defined.attr_text = attr_text
        if hidden is not None:
            defined.hidden = hidden
        if local_sheet_id is not None:
            defined.localSheetId = local_sheet_id
    workbook.save(path)


def tamper_xls_named_range(
    folder: Path,
    source: Path,
    mutate,
    label: str,
) -> Path:
    source_xlsx = convert_with_soffice(source, folder / f"{label}-source-xlsx", "xlsx")
    mutate(source_xlsx)
    return convert_with_soffice(source_xlsx, folder / f"{label}-xls", "xls")


def gradebook_normalized_input(skill: bool = False, count: int = 2) -> dict[str, object]:
    weights = {"regular": 0.5, "theory": 0.3, "skill": 0.2} if skill else {"regular": 0.6, "theory": 0.4, "skill": 0.0}
    rows = []
    for index in range(count):
        regular = [86.5, 91.0, 100.0, 0.0][index % 4]
        theory = [88.0, 90.0, 100.0, 0.0][index % 4]
        score = [92.0, 90.0, 100.0, 0.0][index % 4]
        total = math.floor(regular * weights["regular"] + theory * weights["theory"] + score * weights["skill"] + 0.5)
        rows.append(
            {
                "id": f"240101{index + 1:03d}",
                "name": f"学生{index + 1}",
                "regular": regular,
                "theory": theory,
                "skill": score if skill else 0.0,
                "total": total,
            }
        )
    return {
        "term": "2025-2026-2",
        "course": "软件测试实训",
        "teacher": "张老师",
        "class_name": "软件技术2401班",
        "weights": weights,
        "students": rows,
    }


def write_lesson_manifest_package(
    folder: Path,
    source_manifest: Path,
    source_template: Path,
    version: str,
    mutate=None,
) -> tuple[Path, Path]:
    """Create an isolated versioned package for real validator/generator CLI tests."""
    package = folder / "lesson-package"
    package.mkdir(parents=True, exist_ok=True)
    template = package / "template.docx"
    shutil.copy2(source_template, template)
    manifest = yaml.safe_load(source_manifest.read_text(encoding="utf-8"))
    manifest["template"]["version"] = version
    manifest["template"]["file"] = "template.docx"
    manifest["template"]["compatibility_entries"] = []
    manifest["fingerprint"]["sha256"] = file_sha256(template)
    manifest["fingerprint"]["value"] = file_sha256(template)
    if version.startswith("1.1."):
        base_package = folder / "base-v1.0.0"
        base_package.mkdir(parents=True, exist_ok=True)
        shutil.copy2(LESSON_V10_MANIFEST, base_package / "manifest.yaml")
        shutil.copy2(
            LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            base_package / "template.docx",
        )
        manifest["template"]["base_manifest"] = "../base-v1.0.0/manifest.yaml"
        manifest["template"]["base_template"] = "../base-v1.0.0/template.docx"
    else:
        manifest["template"].pop("base_manifest", None)
        manifest["template"].pop("base_template", None)
    if mutate is not None:
        mutate(manifest)
    manifest_path = package / "manifest.yaml"
    manifest_path.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )
    return template, manifest_path


def write_manifest_for_modified_template(
    folder: Path,
    source_manifest: Path,
    template: Path,
    reference_template: Path,
) -> Path:
    """Create a matching custom manifest while retaining the protected baseline."""
    folder.mkdir(parents=True, exist_ok=True)

    def relative_or_absolute(path: Path) -> str:
        try:
            return os.path.relpath(path.resolve(), folder.resolve())
        except ValueError:
            return str(path.resolve())

    manifest = yaml.safe_load(source_manifest.read_text(encoding="utf-8"))
    manifest["template"]["file"] = relative_or_absolute(reference_template)
    manifest["template"]["compatibility_entries"] = []
    manifest["fingerprint"]["sha256"] = file_sha256(template)
    manifest["fingerprint"]["value"] = file_sha256(template)
    if str(manifest["template"]["version"]).split(".")[:2] == ["1", "1"]:
        manifest["template"]["base_manifest"] = relative_or_absolute(LESSON_V10_MANIFEST)
        manifest["template"]["base_template"] = relative_or_absolute(
            LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"
        )
    else:
        manifest["template"].pop("base_manifest", None)
        manifest["template"].pop("base_template", None)
    manifest_path = folder / "manifest.yaml"
    manifest_path.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )
    return manifest_path


def update_manifest_fingerprint(manifest_path: Path, template: Path) -> None:
    manifest = yaml.safe_load(manifest_path.read_text(encoding="utf-8"))
    digest = file_sha256(template)
    manifest["fingerprint"]["sha256"] = digest
    manifest["fingerprint"]["value"] = digest
    manifest_path.write_text(
        yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )


def patch_docx_document_xml(path: Path, mutate) -> None:
    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                mutate(root)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            target.writestr(info, data)
    temporary.replace(path)


def patch_docx_parts(path: Path, mutate) -> None:
    """Patch multiple DOCX XML stories in one real package mutation."""
    temporary = path.with_suffix(path.suffix + ".tmp")
    with zipfile.ZipFile(path, "r") as source:
        data_by_name = {info.filename: source.read(info.filename) for info in source.infolist()}
        roots = {
            name: etree.fromstring(data)
            for name, data in data_by_name.items()
            if name == "word/document.xml"
            or name.startswith("word/header") and name.endswith(".xml")
            or name.startswith("word/footer") and name.endswith(".xml")
        }
        mutate(roots)
        for name, root in roots.items():
            data_by_name[name] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        infos = source.infolist()
        with zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as target:
            for info in infos:
                target.writestr(info, data_by_name[info.filename])
    temporary.replace(path)


def ensure_header_story(path: Path) -> None:
    document = Document(path)
    document.sections[0].header.paragraphs[0].text = "header test story"
    document.save(path)


def story_first_paragraph(root):
    return next(node for node in root.iter() if node.tag == qn("w:p"))


def add_story_bookmark(root, name: str, bookmark_id: str, *, start: bool = True, end: bool = True) -> None:
    paragraph = story_first_paragraph(root)
    if start:
        bookmark_start = etree.Element(qn("w:bookmarkStart"))
        bookmark_start.set(qn("w:id"), bookmark_id)
        bookmark_start.set(qn("w:name"), name)
        ppr = paragraph.find(qn("w:pPr"))
        paragraph.insert(list(paragraph).index(ppr) + 1 if ppr is not None else 0, bookmark_start)
    if end:
        bookmark_end = etree.Element(qn("w:bookmarkEnd"))
        bookmark_end.set(qn("w:id"), bookmark_id)
        paragraph.append(bookmark_end)


def patch_bookmarked_text(path: Path, name: str, text: str) -> None:
    """Change only text inside a semantic bookmark while preserving its boundaries."""
    def mutate(root) -> None:
        start = bookmark_start(root, name)
        end = bookmark_end(root, start.get(qn("w:id")))
        nodes = list(root.iter())
        start_index = nodes.index(start)
        end_index = nodes.index(end)
        text_nodes = [node for node in nodes[start_index + 1:end_index] if node.tag == qn("w:t")]
        if not text_nodes:
            raise AssertionError(f"bookmark has no text node in test fixture: {name}")
        text_nodes[0].text = text
        for node in text_nodes[1:]:
            node.text = ""

    patch_docx_document_xml(path, mutate)


def bookmark_start(root, name: str):
    nodes = root.xpath(
        ".//w:bookmarkStart[@w:name=$name]",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        name=name,
    )
    if not nodes:
        raise AssertionError(f"bookmark not found in test fixture: {name}")
    return nodes[0]


def bookmark_end(root, bookmark_id: str):
    nodes = root.xpath(
        ".//w:bookmarkEnd[@w:id=$bookmark_id]",
        namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
        bookmark_id=bookmark_id,
    )
    if not nodes:
        raise AssertionError(f"bookmark end not found in test fixture: {bookmark_id}")
    return nodes[0]


def raw_main_table_cell(root, row_index: int, cell_index: int):
    body = root.find(qn("w:body"))
    table = next(node for node in body if node.tag == qn("w:tbl"))
    rows = [node for node in table if node.tag == qn("w:tr")]
    cells = [node for node in rows[row_index] if node.tag == qn("w:tc")]
    return cells[cell_index]


def raw_first_paragraph(cell):
    return next(node for node in cell if node.tag == qn("w:p"))


def move_bookmark_to_cell(root, name: str, row_index: int, cell_index: int) -> None:
    start = bookmark_start(root, name)
    end = bookmark_end(root, start.get(qn("w:id")))
    start.getparent().remove(start)
    end.getparent().remove(end)
    paragraph = raw_first_paragraph(raw_main_table_cell(root, row_index, cell_index))
    ppr = paragraph.find(qn("w:pPr"))
    paragraph.insert(list(paragraph).index(ppr) + 1 if ppr is not None else 0, start)
    paragraph.append(end)


def move_bookmark_end_to_paragraph(root, name: str, paragraph) -> None:
    start = bookmark_start(root, name)
    end = bookmark_end(root, start.get(qn("w:id")))
    end.getparent().remove(end)
    paragraph.append(end)


def move_bookmark_end_before_start(root, name: str) -> None:
    start = bookmark_start(root, name)
    end = bookmark_end(root, start.get(qn("w:id")))
    parent = start.getparent()
    if end.getparent() is not parent:
        raise AssertionError("test fixture bookmark boundaries are not siblings")
    parent.remove(end)
    parent.insert(parent.index(start), end)


class LessonTemplatePackageTests(unittest.TestCase):
    def test_canonical_template_and_compatibility_entry(self) -> None:
        result = run_script(LESSON / "scripts" / "validate_template.py", "--json")
        self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
        report = json.loads(result.stdout)
        self.assertEqual(report["checks"]["main_table"], {"rows": 30, "columns": 10})
        self.assertEqual(report["checks"]["evaluation_table"], {"rows": 14, "columns": 4})

    def test_invalid_input_is_rejected(self) -> None:
        sys.path.insert(0, str(LESSON / "scripts"))
        from package_common import validate_input

        bad = {"course_name": "软件测试", "lessons": [{"unit": "项目一"}]}
        with self.assertRaises(ValueError):
            validate_input(bad)
        bad["weights"] = {"regular": 0.5, "theory": 0.5, "skill": 0.5}
        with self.assertRaises(ValueError):
            validate_input(bad)
        with self.assertRaises(ValueError):
            validate_input({"course_name": "软件测试", "lessons": [{"unit": "", "task": "", "hours": "2"}]})
        with self.assertRaises(ValueError):
            validate_input({"course_name": "课" * 33, "lessons": [{"unit": "项目一", "task": "完成任务", "hours": "2"}]})

    def test_non_projectized_unit_rejects_before_docx_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-unit-guard-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"][0]["unit"] = "第一章 基础测试"
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("projectized teaching", result.stderr)
            self.assertFalse(output.exists())

    def test_score_precision_accepts_half_points_and_defaults(self) -> None:
        sys.modules.pop("package_common", None)
        sys.path.insert(0, str(LESSON / "scripts"))
        from package_common import validate_input

        base = {"course_name": "软件测试实训", "lessons": [{"unit": "项目一", "task": "完成任务", "hours": "2"}]}
        for score in (89, 89.0, 89.5):
            payload = json.loads(json.dumps(base, ensure_ascii=False))
            payload["lessons"][0]["score"] = score
            validate_input(payload)
        validate_input(base)

    def test_nonpositive_hours_rejects_numeric_values_and_strings(self) -> None:
        sys.modules.pop("package_common", None)
        sys.path.insert(0, str(LESSON / "scripts"))
        from package_common import validate_input

        base = {"course_name": "软件测试实训", "lessons": [{"unit": "项目一", "task": "完成任务", "hours": "2"}]}
        for invalid_hours in (-2, "-2", 0, "0"):
            payload = json.loads(json.dumps(base, ensure_ascii=False))
            payload["lessons"][0]["hours"] = invalid_hours
            with self.subTest(invalid_hours=invalid_hours), self.assertRaisesRegex(ValueError, "positive number"):
                validate_input(payload)

    def test_nonpositive_hours_reject_before_docx_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-hours-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"][0]["hours"] = "-2"
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("lessons[0].hours must be a positive number; received -2.", result.stderr)
            self.assertFalse(output.exists())

    def test_long_lesson_hours_reject_before_docx_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-hours-length-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"][0]["hours"] = "1234567890123"
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Input schema validation failed", result.stderr)
            self.assertFalse(output.exists())

    def test_numeric_lesson_hours_length_rejects_before_docx_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-numeric-hours-length-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"][0]["hours"] = 1234567890123
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("hours exceeds manifest max_chars=12", result.stderr)
            self.assertFalse(output.exists())

    def test_invalid_total_hours_reject_before_docx_generation(self) -> None:
        for invalid_total in ("abc", "-2", 0, "0"):
            with self.subTest(invalid_total=invalid_total), tempfile.TemporaryDirectory(prefix="lesson-package-total-hours-") as temp_name:
                folder = Path(temp_name)
                payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
                payload["total_hours"] = invalid_total
                source = folder / "tasks.json"
                output = folder / "output"
                source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn("total_hours must be a positive number", result.stderr)
                self.assertFalse(output.exists())

    def test_teaching_content_capacity_rejects_before_docx_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-content-capacity-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"][0]["flows"] = [f"流程{i}" for i in range(6)]
            payload["lessons"][0]["knowledge"] = [f"知识点{i}" for i in range(3)]
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("flows and knowledge combined must contain at most 8 items", result.stderr)
            self.assertFalse(output.exists())

    def test_composed_lesson_fields_reject_before_docx_generation(self) -> None:
        cases = {
            "teaching_content": {
                "flows": ["流" * 200 for _ in range(8)],
                "knowledge": [],
                "message": "teaching_content exceeds manifest max_chars=1200",
            },
            "knowledge_goal": {
                "flows": [],
                "knowledge": ["知识点" for _ in range(6)],
                "message": "knowledge_goal exceeds manifest max_paragraphs=5",
            },
            "resources": {
                "tools": "\n".join("工具" for _ in range(7)),
                "message": "resources exceeds manifest max_paragraphs=8",
            },
            "implementation": {
                "flows": ["流" * 200 for _ in range(3)],
                "knowledge": [],
                "message": "implementation row 3 cell 1 exceeds manifest max_chars=600",
            },
            "title": {
                "course_name": "课" * 32,
                "task": "任务" * 40,
                "message": "title exceeds manifest max_chars=120",
            },
            "generated_field": {
                "task": "第一行\n第二行\n第三行\n第四行",
                "flows": [],
                "knowledge": ["知识点"],
                "message": "ability_goal exceeds manifest max_paragraphs=5",
            },
        }
        for name, case in cases.items():
            with self.subTest(name=name), tempfile.TemporaryDirectory(prefix="lesson-package-composed-") as temp_name:
                folder = Path(temp_name)
                payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
                payload["lessons"] = [payload["lessons"][0]]
                payload["total_hours"] = 2
                lesson = payload["lessons"][0]
                lesson.update({key: value for key, value in case.items() if key != "message"})
                source = folder / "tasks.json"
                output = folder / "output"
                source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn(case["message"], result.stderr)
                self.assertFalse(output.exists())

    def test_score_precision_rejects_before_docx_generation(self) -> None:
        for invalid_score in (89.2, 89.25):
            with self.subTest(invalid_score=invalid_score), tempfile.TemporaryDirectory(prefix="lesson-package-score-") as temp_name:
                folder = Path(temp_name)
                payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
                payload["lessons"][0]["score"] = invalid_score
                source = folder / "tasks.json"
                output = folder / "output"
                source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn(
                    f"lessons[0].score must use 0.5-point increments; received {invalid_score}.",
                    result.stderr,
                )
                self.assertFalse(output.exists())

    def test_default_score_generates_exact_evaluation_total(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-default-score-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"][0].pop("score")
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = sorted(output.glob("*.docx"))[0]
            nested = Document(generated).tables[0].cell(12, 1).tables[0]
            score_sum = sum(
                (Decimal(nested.cell(row, 2).text.strip()) for row in range(1, 14)),
                Decimal("0"),
            )
            self.assertEqual(score_sum, Decimal("89"))

    def test_two_hour_lesson_preserves_method_labels_and_totals_90_classroom_minutes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-duration-labels-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"] = [payload["lessons"][0]]
            payload["total_hours"] = 2
            payload["lessons"][0]["hours"] = "2"
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")

            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)

            generated = next(output.glob("*.docx"))
            table = Document(generated).tables[0]
            self.assertEqual(table.rows[7].cells[5].text.strip(), "突出方法")
            self.assertEqual(table.rows[8].cells[5].text.strip(), "破解方法")
            self.assertIn("任务驱动、教师示范、分组实训、过程评价", table.rows[7].cells[6].text)
            self.assertIn("提供模板清单、分步演示、同伴互评和教师点评", table.rows[8].cells[6].text)

            classroom_minutes = []
            for row_index in range(18, 25):
                classroom_minutes.extend(
                    int(value)
                    for value in re.findall(r"(?m)^\s*(\d+)min\s*$", table.rows[row_index].cells[0].text)
                )
            self.assertEqual(classroom_minutes, [5, 15, 30, 10, 15, 10, 5])
            self.assertEqual(sum(classroom_minutes), 90)
            self.assertIn("10min", table.rows[16].cells[0].text)
            self.assertIn("15min", table.rows[26].cells[0].text)

            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "passed")

    def test_lesson_course_override_is_validated(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-course-override-") as temp_name:
            folder = Path(temp_name)
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            payload["lessons"] = [payload["lessons"][0]]
            payload["total_hours"] = 2
            payload["lessons"][0]["course_name"] = "接口测试实训"
            payload["lessons"][0]["major"] = "数据科学与大数据技术"
            payload["lessons"][0]["audience"] = "高职三年级"
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(output.glob("*.docx"))
            document = Document(generated)
            self.assertEqual(document.tables[0].cell(0, 1).text.strip(), "接口测试实训")
            self.assertEqual(document.tables[0].cell(0, 5).text.strip(), "数据科学与大数据技术")
            self.assertEqual(document.tables[0].cell(0, 9).text.strip(), "高职三年级")
            self.assertIn("《接口测试实训》", document.paragraphs[0].text)

            patch_bookmarked_text(generated, "lp_major", "错误专业")
            validation = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(validation.returncode, 0)
            self.assertIn("major field mismatch", validation.stderr)

    def test_lesson_filename_is_bounded_by_utf8_bytes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-long-filename-") as temp_name:
            folder = Path(temp_name)
            payload = {
                "course_name": "软件测试",
                "total_hours": 2,
                "lessons": [
                    {
                        "unit": "项目一" + "单元" * 28,
                        "task": "完成" + "测试任务" * 19,
                        "hours": "2",
                        "flows": [],
                        "knowledge": [],
                    }
                ],
            }
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(output.glob("*.docx"))
            self.assertLessEqual(len(generated.name.encode("utf-8")), 255)
            self.assertIn("~", generated.stem)

    def test_long_teaching_content_is_generated(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-long-") as temp_name:
            folder = Path(temp_name)
            payload = {
                "course_name": "软件测试实训" + "课程" * 13,
                "total_hours": 2,
                "lessons": [
                    {
                        "unit": "项目一 测试项目准备",
                        "task": "编制测试计划并完成环境检查",
                        "hours": "2",
                        "flows": [f"流程{i + 1}" + "：" + "检查测试环境、记录问题并提交阶段成果" * 3 for i in range(5)],
                        "knowledge": ["测试计划结构", "环境检查要点"],
                        "score": 90,
                    }
                ],
            }
            source = folder / "tasks.json"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(folder / "output"),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)

    def test_manifest_coordinate_and_major_failures(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-manifest-") as temp_name:
            folder = Path(temp_name)
            template = folder / "template.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", template)
            manifest_text = (LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "manifest.yaml").read_text(encoding="utf-8")
            bad_coordinate = folder / "bad-coordinate.yaml"
            bad_coordinate.write_text(manifest_text.replace("paragraph: 0, mode: replace_text_preserve_style", "paragraph: 999, mode: replace_text_preserve_style", 1), encoding="utf-8")
            result = run_script(LESSON / "scripts" / "validate_template.py", "--template", str(template), "--manifest", str(bad_coordinate), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("missing document paragraph", result.stdout)

            incompatible = folder / "incompatible.yaml"
            incompatible.write_text(manifest_text.replace("version: 1.0.0", "version: 2.0.0", 1), encoding="utf-8")
            result = run_script(LESSON / "scripts" / "validate_template.py", "--template", str(template), "--manifest", str(incompatible), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Unsupported template major version", result.stdout)

    def test_custom_lesson_template_rejects_fixed_label_relocation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-fixed-label-") as temp_name:
            custom = Path(temp_name) / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            table = document.tables[0]
            table.cell(0, 0).text = "授课专业"
            table.cell(0, 3).text = "课程名称"
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                Path(temp_name),
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected main-table structure or formatting", result.stdout)

    def test_custom_lesson_template_rejects_style_definition_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-style-guard-") as temp_name:
            custom = Path(temp_name) / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            document.styles["Normal"].font.size = Pt(13)
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                Path(temp_name),
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected main-table structure or formatting", result.stdout)

    def test_custom_lesson_template_rejects_title_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-title-format-guard-") as temp_name:
            custom = Path(temp_name) / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            document.paragraphs[0].runs[0].font.size = Pt(19)
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                Path(temp_name),
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected main-table structure or formatting", result.stdout)

    def test_custom_lesson_template_rejects_header_footer_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-header-format-guard-") as temp_name:
            custom = Path(temp_name) / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            document.sections[0].header.paragraphs[0].paragraph_format.space_before = Pt(1)
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                Path(temp_name),
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected page, header, footer, or section settings", result.stdout)

    def test_custom_lesson_template_rejects_section_property_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-section-property-guard-") as temp_name:
            folder = Path(temp_name)
            custom = folder / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            columns = OxmlElement("w:cols")
            columns.set(qn("w:num"), "2")
            document.sections[0]._sectPr.append(columns)
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                folder,
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected page, header, footer, or section settings", result.stdout)

    def test_custom_lesson_template_rejects_document_settings_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-settings-guard-") as temp_name:
            folder = Path(temp_name)
            custom = folder / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            default_tab_stop = OxmlElement("w:defaultTabStop")
            default_tab_stop.set(qn("w:val"), "240")
            document.settings._element.append(default_tab_stop)
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                folder,
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected page, header, footer, or section settings", result.stdout)

    def test_custom_lesson_template_rejects_theme_definition_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-theme-guard-") as temp_name:
            folder = Path(temp_name)
            custom = folder / "custom.docx"
            tampered = folder / "tampered.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            with zipfile.ZipFile(custom, "r") as source, zipfile.ZipFile(tampered, "w") as target:
                for info in source.infolist():
                    data = source.read(info.filename)
                    if info.filename == "word/theme/theme1.xml":
                        data = data.replace(b"Office Theme", b"Tampered Theme")
                    target.writestr(info, data)
            manifest = write_manifest_for_modified_template(
                folder,
                LESSON_V10_MANIFEST,
                tampered,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(tampered),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected main-table structure or formatting", result.stdout)

    def test_custom_lesson_template_rejects_protected_body_paragraph_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-body-guard-") as temp_name:
            custom = Path(temp_name) / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", custom)
            document = Document(custom)
            document.paragraphs[1].add_run("未声明正文")
            document.save(custom)
            manifest = write_manifest_for_modified_template(
                Path(temp_name),
                LESSON_V10_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected main-table structure or formatting", result.stdout)

    def test_generation_writes_qa_report_and_preserves_structure(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-test-") as temp_name:
            output = Path(temp_name) / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            self.assertTrue((output / "qa-report.json").exists())
            self.assertEqual(len(list(output.glob("*.docx"))), 2)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["template_id"], "lesson-plan")
            self.assertEqual(report["template_version"], "1.1.1")
            self.assertEqual(report["generator_version"], "1.1.1")
            self.assertEqual(report["anchor_mode"], "word_bookmark")
            self.assertEqual(report["required_anchor_count"], 70)
            self.assertEqual(report["preserved_anchor_count"], 70)
            self.assertEqual(report["engine"], "python-docx")
            self.assertFalse(report["custom_template"])
            self.assertEqual(report["validation_skipped"], [])
            self.assertEqual(report["status"], "passed")
            payload = json.loads((ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8"))
            for path, item in zip(sorted(output.glob("*.docx")), payload["lessons"]):
                self.assertEqual(len(Document(path).tables[0].rows), 30)
                nested = Document(path).tables[0].cell(12, 1).tables[0]
                score_sum = sum(
                    (Decimal(nested.cell(row, 2).text.strip()) for row in range(1, 14)),
                    Decimal("0"),
                )
                self.assertEqual(score_sum, Decimal(str(item["score"])))

    def test_generation_copies_direct_formatting_to_added_multiline_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-multiline-format-") as temp_name:
            output = Path(temp_name) / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            template_cell = Document(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx").tables[0].cell(4, 1)
            output_cell = Document(sorted(output.glob("*.docx"))[0]).tables[0].cell(4, 1)
            self.assertGreater(len(output_cell.paragraphs), 1)
            source_ppr = template_cell.paragraphs[0]._p.pPr.xml
            source_rpr = template_cell.paragraphs[0].runs[0]._r.rPr.xml
            for paragraph in output_cell.paragraphs[1:]:
                self.assertEqual(paragraph._p.pPr.xml, source_ppr)
                self.assertEqual(paragraph.runs[0]._r.rPr.xml, source_rpr)

    def test_output_validation_orders_three_digit_lesson_files_numerically(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-three-digit-") as temp_name:
            folder = Path(temp_name)
            payload = {
                "course_name": "软件测试实训",
                "total_hours": 6,
                "lessons": [
                    {
                        "unit": f"项目{i + 1} 排序验证",
                        "task": f"完成排序验证任务{i + 1}",
                        "hours": "2",
                        "flows": [],
                        "knowledge": [],
                        "score": 89.5,
                    }
                    for i in range(3)
                ],
            }
            source = folder / "tasks.json"
            output = folder / "output"
            source.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
            generate_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(generate_result.returncode, 0, generate_result.stderr or generate_result.stdout)
            generated = sorted(output.glob("*.docx"))
            self.assertEqual(len(generated), 3)

            temporary_names = []
            for index, path in enumerate(generated):
                temporary = output / f".sorting-rename-{index}.docx"
                path.rename(temporary)
                temporary_names.append(temporary)
            for sequence, temporary in zip((2, 10, 100), temporary_names):
                temporary.rename(output / f"教案{sequence}_排序验证.docx")

            validate_result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(validate_result.returncode, 0, validate_result.stderr or validate_result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "passed")
            self.assertEqual(report["files_checked"], 3)

    def test_compatibility_template_and_skipped_validation_leave_qa_metadata(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-compat-qa-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--template",
                str(LESSON / "assets" / "lesson-plan-template.docx"),
                "--manifest",
                str(LESSON_V10_MANIFEST),
                "--tasks-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(output),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "skipped")
            self.assertTrue(report["custom_template"])
            self.assertEqual(report["validation_skipped"], ["output"])
            self.assertIn("output", report["checks"]["validation"]["skipped"])
            self.assertIn("Custom template selected", " ".join(report["warnings"]))

    def test_manifest_loading_failures_are_clear(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-manifest-errors-") as temp_name:
            folder = Path(temp_name)
            template = folder / "template.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", template)
            missing = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(folder / "missing.yaml"),
                "--json",
            )
            self.assertNotEqual(missing.returncode, 0)
            self.assertIn("No such file", missing.stdout)

            malformed = folder / "malformed.yaml"
            malformed.write_text("template: [", encoding="utf-8")
            result = run_script(LESSON / "scripts" / "validate_template.py", "--manifest", str(malformed), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("errors", result.stdout)

            manifest_text = (LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "manifest.yaml").read_text(encoding="utf-8")
            missing_version = folder / "missing-version.yaml"
            missing_version.write_text(manifest_text.replace("  version: 1.0.0\n", "", 1), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(missing_version),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("template.version", result.stdout)

            missing_template = folder / "missing-template.yaml"
            missing_template.write_text(manifest_text.replace("file: template.docx", "file: missing.docx", 1), encoding="utf-8")
            result = run_script(LESSON / "scripts" / "validate_template.py", "--manifest", str(missing_template), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Template not found", result.stdout)

            missing_canonical = folder / "missing-canonical.yaml"
            missing_canonical.write_text(manifest_text.replace("file: template.docx", "file: missing.docx", 1), encoding="utf-8")
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(missing_canonical),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Canonical template not found", result.stdout)

    def test_structure_breaking_docx_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-broken-") as temp_name:
            broken = Path(temp_name) / "broken.docx"
            document = Document(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx")
            document.add_table(rows=1, cols=1)
            document.save(broken)
            manifest = write_manifest_for_modified_template(
                Path(temp_name),
                LESSON_V10_MANIFEST,
                broken,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
            )
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(broken),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Top-level table count mismatch", result.stdout)

    def test_output_residual_template_text_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-residual-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = next(output.glob("*.docx"))
            patch_bookmarked_text(path, "lp_student_base", "Linux操作系统应用残留")
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("forbidden template text", result.stderr)

    def test_output_validation_rejects_evaluation_score_above_rubric_maximum(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-score-cap-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            nested = document.tables[0].cell(12, 1).tables[0]
            original_row_10 = Decimal(nested.cell(10, 2).text.strip())
            nested.cell(1, 2).text = "3.5"
            nested.cell(10, 2).text = str(original_row_10 - Decimal("0.5"))
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("evaluation score row 1 exceeds rubric maximum 3", result.stderr)

    def test_output_validation_rejects_header_footer_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-header-footer-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            document.sections[0].header.paragraphs[0].text = "被篡改页眉"
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_fixed_label_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-fixed-label-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            nested = document.tables[0].cell(12, 1).tables[0]
            nested.cell(1, 1).text = "被篡改的评价要素"
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_fixed_label_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-fixed-label-format-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            document.tables[0].cell(0, 0).paragraphs[0].runs[0].font.size = Pt(19)
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_writable_direct_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-direct-format-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            paragraph = document.tables[0].cell(4, 1).paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1)
            run = paragraph.runs[0]
            run.bold = True
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_title_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-title-output-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            document.paragraphs[0].runs[0].font.size = Pt(19)
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_document_settings_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-settings-output-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            character_spacing = OxmlElement("w:characterSpacingControl")
            character_spacing.set(qn("w:val"), "doNotCompress")
            document.settings._element.append(character_spacing)
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_removed_writable_direct_formatting(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-direct-format-removal-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            cell = document.tables[0].cell(4, 1)
            for paragraph in cell.paragraphs[1:]:
                if paragraph._p.pPr is not None:
                    paragraph._p.remove(paragraph._p.pPr)
                for run in paragraph.runs:
                    if run._r.rPr is not None:
                        run._r.remove(run._r.rPr)
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_composed_content_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-composed-output-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            patch_bookmarked_text(path, "lp_teaching_content", "被替换的教学内容")
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("teaching_content content mismatch", result.stderr)

    def test_output_validation_rejects_implementation_content_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-implementation-output-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            patch_bookmarked_text(path, "lp_impl_prep_content", "被替换的实施内容")
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("implementation cell mismatch", result.stderr)

    def test_output_validation_rejects_deterministic_field_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-deterministic-output-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            patch_bookmarked_text(path, "lp_ability_goal", "被替换的能力目标")
            patch_bookmarked_text(path, "lp_reflection_improvement", "被替换的教学反思")
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("ability_goal content mismatch", result.stderr)
            self.assertIn("reflection cell mismatch", result.stderr)

    def test_output_validation_rejects_evaluation_cell_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-evaluation-output-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            document.tables[0].cell(12, 1).tables[0].cell(1, 3).text = "被篡改评价备注"
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("evaluation cell mismatch", result.stderr)

    def test_output_validation_rejects_evaluation_direct_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-evaluation-format-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            path = sorted(output.glob("*.docx"))[0]
            document = Document(path)
            nested = document.tables[0].cell(12, 1).tables[0]
            nested.cell(1, 2).paragraphs[0].runs[0].font.bold = True
            document.save(path)
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)


    def test_output_validation_rejects_evaluation_layout_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-evaluation-layout-") as temp_name:
            folder = Path(temp_name)
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            pristine_dir = folder / "pristine"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(pristine_dir),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            pristine_files = sorted(pristine_dir.glob("*.docx"))
            self.assertEqual(len(pristine_files), 2)

            def child(parent, tag):
                node = parent.find(qn(tag))
                if node is None:
                    node = OxmlElement(tag)
                    parent.append(node)
                return node

            def change_column_width(nested):
                grid_column = nested._tbl.tblGrid.findall(qn("w:gridCol"))[1]
                grid_column.set(qn("w:w"), str(int(grid_column.get(qn("w:w"), "0")) + 120))

            def change_row_height(nested):
                tr_pr = nested.rows[1]._tr.get_or_add_trPr()
                height = child(tr_pr, "w:trHeight")
                height.set(qn("w:val"), "1440")
                height.set(qn("w:hRule"), "atLeast")

            def change_table_border(nested):
                borders = child(nested._tbl.tblPr, "w:tblBorders")
                top = child(borders, "w:top")
                top.set(qn("w:val"), "double")
                top.set(qn("w:sz"), "16")

            def change_score_shading(nested):
                shading = child(nested.cell(1, 2)._tc.get_or_add_tcPr(), "w:shd")
                shading.set(qn("w:fill"), "FFFF00")

            def change_vertical_alignment(nested):
                vertical = child(nested.cell(1, 2)._tc.get_or_add_tcPr(), "w:vAlign")
                vertical.set(qn("w:val"), "top")

            def change_cell_margin(nested):
                margins = child(nested.cell(1, 2)._tc.get_or_add_tcPr(), "w:tcMar")
                top = child(margins, "w:top")
                top.set(qn("w:w"), "240")
                top.set(qn("w:type"), "dxa")

            def change_table_style(nested):
                style = child(nested._tbl.tblPr, "w:tblStyle")
                style.set(qn("w:val"), "LightShadingAccent1")

            mutations = {
                "column-width": change_column_width,
                "row-height": change_row_height,
                "table-border": change_table_border,
                "score-shading": change_score_shading,
                "vertical-alignment": change_vertical_alignment,
                "cell-margin": change_cell_margin,
                "table-style": change_table_style,
            }
            for name, mutate in mutations.items():
                with self.subTest(layout=name):
                    case_dir = folder / name
                    case_dir.mkdir()
                    case_files = []
                    for pristine in pristine_files:
                        case_file = case_dir / pristine.name
                        shutil.copy2(pristine, case_file)
                        case_files.append(case_file)
                    document = Document(case_files[0])
                    nested = document.tables[0].cell(12, 1).tables[0]
                    original_text = [[cell.text for cell in row.cells] for row in nested.rows]
                    self.assertEqual((len(nested.rows), len(nested.columns)), (14, 4))
                    mutate(nested)
                    self.assertEqual([[cell.text for cell in row.cells] for row in nested.rows], original_text)
                    document.save(case_files[0])
                    result = run_script(
                        LESSON / "scripts" / "validate_output.py",
                        "--input-json",
                        str(source),
                        "--output-dir",
                        str(case_dir),
                    )
                    self.assertNotEqual(result.returncode, 0)
                    self.assertIn("protected DOCX layout changed", result.stderr)

    def test_output_validation_rejects_evaluation_merge_structure_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-evaluation-merge-") as temp_name:
            folder = Path(temp_name)
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            pristine_dir = folder / "pristine"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(pristine_dir),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            pristine_files = sorted(pristine_dir.glob("*.docx"))
            case_dir = folder / "merge"
            case_dir.mkdir()
            case_files = []
            for pristine in pristine_files:
                case_file = case_dir / pristine.name
                shutil.copy2(pristine, case_file)
                case_files.append(case_file)
            document = Document(case_files[0])
            nested = document.tables[0].cell(12, 1).tables[0]
            grid_span = OxmlElement("w:gridSpan")
            grid_span.set(qn("w:val"), "2")
            nested.cell(1, 0)._tc.get_or_add_tcPr().append(grid_span)
            document.save(case_files[0])
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(case_dir),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("protected DOCX layout changed", result.stderr)

    def test_v11_semantic_template_build_is_idempotent_and_preserves_v10_hash(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-semantic-build-") as temp_name:
            folder = Path(temp_name)
            first = folder / "template-first.docx"
            second = folder / "template-second.docx"
            result = run_script(
                LESSON / "scripts" / "build_semantic_template.py",
                "--output",
                str(first),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            self.assertEqual(file_sha256(first), "569B076DE30CD64172EE86F2123C8AE5EA67828F46B51C4280FE32DAF6DE1AD0")
            self.assertEqual(
                file_sha256(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"),
                "11783108468204DD67C9F8EAA1543B67279361ECC842A8B37F8541BDD01D16D5",
            )
            result = run_script(
                LESSON / "scripts" / "build_semantic_template.py",
                "--source",
                str(first),
                "--output",
                str(second),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            self.assertEqual(file_sha256(first), file_sha256(second))
            result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(first),
                "--manifest",
                str(LESSON_V111_MANIFEST),
                "--json",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)

    def test_v11_default_generation_preserves_every_semantic_anchor(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-semantic-inventory-") as temp_name:
            output = Path(temp_name) / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            sys.modules.pop("bookmark_utils", None)
            sys.modules.pop("package_common", None)
            sys.path.insert(0, str(LESSON / "scripts"))
            from bookmark_utils import validate_bookmark_inventory
            from package_common import bookmark_containers, load_manifest, required_bookmarks

            manifest = load_manifest(LESSON_V111_MANIFEST)
            for path in sorted(output.glob("*.docx")):
                inventory = validate_bookmark_inventory(
                    Document(path),
                    required_bookmarks(manifest),
                    bookmark_containers(manifest),
                )
                self.assertTrue(inventory["valid"], inventory["errors"])
                self.assertEqual(inventory["required_count"], 70)
                self.assertEqual(inventory["preserved_count"], 70)
            generated = Document(sorted(output.glob("*.docx"))[0])
            self.assertGreater(len(generated.tables[0].cell(4, 1).paragraphs), 1)
            self.assertGreater(len(generated.tables[0].cell(19, 1).paragraphs), 1)

    def test_v11_template_fault_injections_are_rejected_by_real_validator(self) -> None:
        mutations = {}

        def delete_bookmark(root):
            start = bookmark_start(root, "lp_course_name")
            end = bookmark_end(root, start.get(qn("w:id")))
            start.getparent().remove(start)
            end.getparent().remove(end)

        def duplicate_bookmark(root):
            start = bookmark_start(root, "lp_title")
            end = bookmark_end(root, start.get(qn("w:id")))
            new_id = str(max(int(node.get(qn("w:id"), "0")) for node in root.xpath(".//w:bookmarkStart | .//w:bookmarkEnd", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})) + 1)
            new_start = etree.fromstring(etree.tostring(start))
            new_end = etree.fromstring(etree.tostring(end))
            new_start.set(qn("w:id"), new_id)
            new_end.set(qn("w:id"), new_id)
            start.getparent().append(new_start)
            end.getparent().append(new_end)

        def missing_end(root):
            start = bookmark_start(root, "lp_hours")
            bookmark_end(root, start.get(qn("w:id"))).getparent().remove(bookmark_end(root, start.get(qn("w:id"))))

        def rename_bookmark(root):
            bookmark_start(root, "lp_unit").set(qn("w:name"), "lp_unitt")

        def move_bookmark(root, name: str, row: int, cell: int):
            start = bookmark_start(root, name)
            end = bookmark_end(root, start.get(qn("w:id")))
            start.getparent().remove(start)
            end.getparent().remove(end)
            body = root.find(qn("w:body"))
            table = next(node for node in body if node.tag == qn("w:tbl"))
            rows = [node for node in table if node.tag == qn("w:tr")]
            cells = [node for node in rows[row] if node.tag == qn("w:tc")]
            paragraph = next(node for node in cells[cell] if node.tag == qn("w:p"))
            ppr = paragraph.find(qn("w:pPr"))
            paragraph.insert(list(paragraph).index(ppr) + 1 if ppr is not None else 0, start)
            paragraph.append(end)

        def wrong_cell(root):
            move_bookmark(root, "lp_course_name", 0, 3)

        def moved_evaluation(root):
            move_bookmark(root, "lp_evaluation", 0, 1)

        def fixed_label(root):
            node = root.xpath(".//w:t[text()='课程名称']", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})[0]
            node.text = "错误标签"

        def fixed_format(root):
            node = root.xpath(".//w:t[text()='课程名称']", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})[0]
            run = node.getparent().getparent()
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = etree.Element(qn("w:rPr"))
                run.insert(0, rpr)
            rpr.append(etree.Element(qn("w:b")))

        mutations.update(
            {
                "deleted": (delete_bookmark, "Semantic bookmark protection failed"),
                "duplicate": (duplicate_bookmark, "Semantic bookmark protection failed"),
                "missing-end": (missing_end, "Semantic bookmark protection failed"),
                "typo": (rename_bookmark, "Semantic bookmark protection failed"),
                "wrong-cell": (wrong_cell, "Semantic bookmark location changed"),
                "evaluation-moved": (moved_evaluation, "Semantic bookmark location changed"),
                "fixed-label": (fixed_label, "Semantic template changed visible content or structure"),
                "fixed-format": (fixed_format, "Semantic template changed visible content or structure"),
            }
        )
        for name, (mutate, expected_message) in mutations.items():
            with self.subTest(fault=name), tempfile.TemporaryDirectory(prefix=f"lesson-package-semantic-{name}-") as temp_name:
                custom = Path(temp_name) / "custom.docx"
                shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", custom)
                patch_docx_document_xml(custom, mutate)
                manifest = write_manifest_for_modified_template(
                    Path(temp_name),
                    LESSON_V11_MANIFEST,
                    custom,
                    LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx",
                )
                result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(custom),
                    "--manifest",
                    str(manifest),
                    "--json",
                )
                self.assertNotEqual(result.returncode, 0, result.stdout)
                self.assertIn(expected_message, result.stdout)

    def test_v11_output_validator_rejects_deleted_or_rewritten_bookmarks(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-semantic-output-guard-") as temp_name:
            folder = Path(temp_name)
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            pristine = folder / "pristine"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(pristine),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            for name in ("delete", "rewrite"):
                with self.subTest(fault=name):
                    case_dir = folder / name
                    shutil.copytree(pristine, case_dir)
                    first = sorted(case_dir.glob("*.docx"))[0]
                    if name == "delete":
                        def delete_both(root):
                            start = bookmark_start(root, "lp_teaching_content")
                            end = bookmark_end(root, start.get(qn("w:id")))
                            start.getparent().remove(start)
                            end.getparent().remove(end)
                        patch_docx_document_xml(first, delete_both)
                    else:
                        def write_drops_anchor(path: Path):
                            document = Document(path)
                            document.tables[0].cell(0, 1).text = "被改写课程"
                            document.save(path)
                        write_drops_anchor(first)
                    result = run_script(
                        LESSON / "scripts" / "validate_output.py",
                        "--input-json",
                        str(source),
                        "--output-dir",
                        str(case_dir),
                        "--manifest",
                        str(LESSON_V111_MANIFEST),
                    )
                    self.assertNotEqual(result.returncode, 0)
                    self.assertIn("semantic bookmark protection failed", result.stderr.lower())

    def test_v11_output_qa_report_keeps_anchor_summary_fields_consistent(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        with tempfile.TemporaryDirectory(prefix="lesson-package-semantic-qa-summary-") as temp_name:
            folder = Path(temp_name)
            pristine = folder / "pristine"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(pristine),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)

            field_pairs = (
                ("invalid_anchor_names", "invalid_names"),
                ("unexpected_anchor_names", "unexpected_names"),
                ("invalid_anchor_ids", "invalid_ids"),
                ("anchor_boundary_errors", "boundary_errors"),
            )

            def assert_report_fields(report: dict, *, expected_field: str | None = None) -> None:
                anchors = report["checks"]["anchors"]
                for report_field, anchor_field in field_pairs:
                    self.assertEqual(report[report_field], anchors[anchor_field], report_field)
                if expected_field is not None:
                    self.assertTrue(report[expected_field], expected_field)

            success = folder / "success"
            shutil.copytree(pristine, success)
            success_result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(success),
                "--manifest",
                str(LESSON_V111_MANIFEST),
            )
            self.assertEqual(success_result.returncode, 0, success_result.stderr or success_result.stdout)
            success_report = json.loads((success / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(success_report["status"], "passed")
            assert_report_fields(success_report)
            for report_field, _ in field_pairs:
                self.assertEqual(success_report[report_field], [], report_field)

            def invalid_id(root) -> None:
                start = bookmark_start(root, "lp_hours")
                end = bookmark_end(root, start.get(qn("w:id")))
                start.set(qn("w:id"), "１２")
                end.set(qn("w:id"), "１２")

            def boundary_error(root) -> None:
                move_bookmark_end_before_start(root, "lp_hours")

            def invalid_name(root) -> None:
                bookmark_start(root, "lp_unit").set(qn("w:name"), "lp-unit")

            def unexpected_name(root) -> None:
                add_story_bookmark(root, "lp_unknown", "999")

            faults = (
                ("invalid-id", invalid_id, "invalid_anchor_ids"),
                ("boundary", boundary_error, "anchor_boundary_errors"),
                ("invalid-name", invalid_name, "invalid_anchor_names"),
                ("unexpected-name", unexpected_name, "unexpected_anchor_names"),
            )
            for label, mutate, expected_field in faults:
                with self.subTest(fault=label):
                    case_dir = folder / label
                    shutil.copytree(pristine, case_dir)
                    patch_docx_document_xml(sorted(case_dir.glob("*.docx"))[0], mutate)
                    result = run_script(
                        LESSON / "scripts" / "validate_output.py",
                        "--input-json",
                        str(source),
                        "--output-dir",
                        str(case_dir),
                        "--manifest",
                        str(LESSON_V111_MANIFEST),
                    )
                    self.assertNotEqual(result.returncode, 0)
                    self.assertIn("semantic bookmark", result.stderr.lower())
                    report_path = case_dir / "qa-report.json"
                    self.assertTrue(report_path.exists())
                    report = json.loads(report_path.read_text(encoding="utf-8"))
                    assert_report_fields(report, expected_field=expected_field)

    def test_v11_missing_bookmarks_does_not_downgrade_and_v10_remains_explicitly_supported(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-package-semantic-no-downgrade-") as temp_name:
            folder = Path(temp_name)
            old_template = folder / "old-template.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx", old_template)
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--template",
                str(old_template),
                "--manifest",
                str(LESSON_V11_MANIFEST),
                "--tasks-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(folder / "should-fail"),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("template/manifest mismatch", result.stderr.lower())
            legacy_output = folder / "legacy-output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--template",
                str(LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"),
                "--manifest",
                str(LESSON_V10_MANIFEST),
                "--tasks-json",
                str(ROOT / "tests" / "fixtures" / "lesson-plan-input.json"),
                "--output-dir",
                str(legacy_output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((legacy_output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["anchor_mode"], "legacy_coordinates")

    def test_v11_bookmark_names_are_safe_and_manifest_matches_definition_source(self) -> None:
        sys.modules.pop("semantic_bookmarks", None)
        sys.modules.pop("package_common", None)
        sys.path.insert(0, str(LESSON / "scripts"))
        from package_common import load_manifest, required_bookmarks
        from semantic_bookmarks import BOOKMARK_NAME_PATTERN, managed_bookmark_names

        manifest = load_manifest(LESSON_V11_MANIFEST)
        expected = managed_bookmark_names()
        self.assertEqual(required_bookmarks(manifest), expected)
        self.assertEqual(manifest["anchors"]["required"], expected)
        self.assertTrue(all(BOOKMARK_NAME_PATTERN.fullmatch(name) for name in expected))
        self.assertLessEqual(max(map(len, expected)), 40)
        report = json.loads(
            run_script(
                LESSON / "scripts" / "validate_template.py",
                "--json",
            ).stdout
        )
        self.assertEqual(report["checks"]["bookmarks"]["required_count"], len(expected))
        self.assertEqual(report["checks"]["bookmarks"]["invalid_names"], [])
        self.assertEqual(report["checks"]["bookmarks"]["invalid_ids"], [])

    def test_v11_invalid_bookmark_names_are_rejected_by_real_validator(self) -> None:
        mutations = {
            "long": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "a" * 41),
            "hyphen": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "lp-unit"),
            "space": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "lp unit"),
            "chinese": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "单元"),
            "digit": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "1lp_unit"),
            "reserved": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "_GoBack"),
            "duplicate": lambda root: bookmark_start(root, "lp_unit").set(qn("w:name"), "lp_title"),
        }
        for name, mutate in mutations.items():
            with self.subTest(name=name), tempfile.TemporaryDirectory(prefix=f"lesson-package-name-{name}-") as temp_name:
                custom = Path(temp_name) / "custom.docx"
                shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", custom)
                patch_docx_document_xml(custom, mutate)
                manifest = write_manifest_for_modified_template(
                    Path(temp_name),
                    LESSON_V11_MANIFEST,
                    custom,
                    LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx",
                )
                result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(custom),
                    "--manifest",
                    str(manifest),
                    "--json",
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn("bookmark", result.stdout.lower())

    def test_v11_invalid_bookmark_ids_and_orphans_are_rejected_by_real_validator(self) -> None:
        def invalid_start(root):
            bookmark_start(root, "lp_hours").set(qn("w:id"), "abc")

        def negative_start(root):
            bookmark_start(root, "lp_hours").set(qn("w:id"), "-1")

        def duplicate_start(root):
            bookmark_start(root, "lp_hours").set(qn("w:id"), bookmark_start(root, "lp_unit").get(qn("w:id")))

        def duplicate_end(root):
            hours_start = bookmark_start(root, "lp_hours")
            unit_end = bookmark_end(root, bookmark_start(root, "lp_unit").get(qn("w:id")))
            bookmark_end(root, hours_start.get(qn("w:id"))).set(qn("w:id"), unit_end.get(qn("w:id")))

        def mismatched_end(root):
            start = bookmark_start(root, "lp_hours")
            bookmark_end(root, start.get(qn("w:id"))).set(qn("w:id"), "9999")

        def orphan_start(root):
            start = bookmark_start(root, "lp_hours")
            bookmark_end(root, start.get(qn("w:id"))).getparent().remove(bookmark_end(root, start.get(qn("w:id"))))

        def orphan_end(root):
            start = bookmark_start(root, "lp_hours")
            start.getparent().remove(start)

        mutations = {
            "invalid-start-id": invalid_start,
            "negative-start-id": negative_start,
            "duplicate-start-id": duplicate_start,
            "duplicate-end-id": duplicate_end,
            "mismatched-id": mismatched_end,
            "orphan-start": orphan_start,
            "orphan-end": orphan_end,
        }
        for name, mutate in mutations.items():
            with self.subTest(name=name), tempfile.TemporaryDirectory(prefix=f"lesson-package-id-{name}-") as temp_name:
                custom = Path(temp_name) / "custom.docx"
                shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", custom)
                patch_docx_document_xml(custom, mutate)
                manifest = write_manifest_for_modified_template(
                    Path(temp_name),
                    LESSON_V11_MANIFEST,
                    custom,
                    LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx",
                )
                result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(custom),
                    "--manifest",
                    str(manifest),
                    "--json",
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn("bookmark", result.stdout.lower())

    def test_v11_boundary_faults_are_rejected_by_template_and_output_validators(self) -> None:
        def end_in_other_cell(root):
            start = bookmark_start(root, "lp_unit")
            end = bookmark_end(root, start.get(qn("w:id")))
            end.getparent().remove(end)
            raw_first_paragraph(raw_main_table_cell(root, 0, 3)).append(end)

        def end_in_other_paragraph(root):
            start = bookmark_start(root, "lp_student_base")
            end = bookmark_end(root, start.get(qn("w:id")))
            cell = raw_main_table_cell(root, 3, 1)
            paragraphs = [node for node in cell if node.tag == qn("w:p")]
            target = paragraphs[-1]
            if target is start.getparent():
                target = etree.SubElement(cell, qn("w:p"))
            end.getparent().remove(end)
            target.append(end)

        def end_before_start(root):
            move_bookmark_end_before_start(root, "lp_hours")

        def cross_stage(root):
            start = bookmark_start(root, "lp_impl_prep_content")
            end = bookmark_end(root, start.get(qn("w:id")))
            end.getparent().remove(end)
            raw_first_paragraph(raw_main_table_cell(root, 18, 1)).append(end)

        def cross_evaluation(root):
            start = bookmark_start(root, "lp_course_name")
            end = bookmark_end(root, start.get(qn("w:id")))
            end.getparent().remove(end)
            parent = raw_main_table_cell(root, 12, 1)
            nested = next(node for node in parent if node.tag == qn("w:tbl"))
            nested_row = next(node for node in nested if node.tag == qn("w:tr"))
            nested_cell = next(node for node in nested_row if node.tag == qn("w:tc"))
            raw_first_paragraph(nested_cell).append(end)

        mutations = {
            "end-other-cell": end_in_other_cell,
            "end-other-paragraph": end_in_other_paragraph,
            "end-before-start": end_before_start,
            "cross-stage": cross_stage,
            "cross-evaluation": cross_evaluation,
        }
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        for name, mutate in mutations.items():
            with self.subTest(name=name), tempfile.TemporaryDirectory(prefix=f"lesson-package-boundary-{name}-") as temp_name:
                folder = Path(temp_name)
                custom = folder / "custom.docx"
                shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", custom)
                patch_docx_document_xml(custom, mutate)
                template_manifest = write_manifest_for_modified_template(
                    folder,
                    LESSON_V11_MANIFEST,
                    custom,
                    LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx",
                )
                template_result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(custom),
                    "--manifest",
                    str(template_manifest),
                    "--json",
                )
                self.assertNotEqual(template_result.returncode, 0)
                self.assertIn("bookmark", template_result.stdout.lower())

                pristine = folder / "pristine"
                generated = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(pristine),
                )
                self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)
                output_dir = folder / "output"
                shutil.copytree(pristine, output_dir)
                patch_docx_document_xml(sorted(output_dir.glob("*.docx"))[0], mutate)
                output_result = run_script(
                    LESSON / "scripts" / "validate_output.py",
                    "--input-json",
                    str(source),
                    "--output-dir",
                    str(output_dir),
                    "--manifest",
                    str(LESSON_V11_MANIFEST),
                )
                self.assertNotEqual(output_result.returncode, 0)
                self.assertIn("semantic bookmark", output_result.stderr.lower())

    def test_v11_cross_story_boundaries_are_rejected_by_real_validators(self) -> None:
        def move_end_to_footer(parts):
            document_root = parts["word/document.xml"]
            footer_root = parts["word/footer1.xml"]
            start = bookmark_start(document_root, "lp_hours")
            end = bookmark_end(document_root, start.get(qn("w:id")))
            end.getparent().remove(end)
            footer_paragraph = next(node for node in footer_root.iter() if node.tag == qn("w:p"))
            footer_paragraph.append(end)

        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        with tempfile.TemporaryDirectory(prefix="lesson-package-cross-story-") as temp_name:
            folder = Path(temp_name)
            custom = folder / "custom.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", custom)
            patch_docx_parts(custom, move_end_to_footer)
            template_manifest = write_manifest_for_modified_template(
                folder,
                LESSON_V11_MANIFEST,
                custom,
                LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx",
            )
            template_result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(template_manifest),
                "--json",
            )
            self.assertNotEqual(template_result.returncode, 0)
            self.assertIn("story", template_result.stdout.lower())

            pristine = folder / "pristine"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(pristine),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)
            output_dir = folder / "output"
            shutil.copytree(pristine, output_dir)
            patch_docx_parts(sorted(output_dir.glob("*.docx"))[0], move_end_to_footer)
            output_result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output_dir),
                "--manifest",
                str(LESSON_V11_MANIFEST),
            )
            self.assertNotEqual(output_result.returncode, 0)
            self.assertIn("semantic bookmark", output_result.stderr.lower())

    def test_v11_builder_rejects_complete_templates_with_wrong_semantic_positions(self) -> None:
        def swap_teacher_student(root):
            move_bookmark_to_cell(root, "lp_impl_prep_teacher", 16, 3)
            move_bookmark_to_cell(root, "lp_impl_prep_student", 16, 2)

        def swap_stage_content(root):
            move_bookmark_to_cell(root, "lp_impl_prep_content", 18, 1)
            move_bookmark_to_cell(root, "lp_impl_intro_content", 16, 1)

        def same_physical_cell(root):
            move_bookmark_to_cell(root, "lp_impl_prep_teacher", 16, 1)

        def move_evaluation(root):
            move_bookmark_to_cell(root, "lp_evaluation", 0, 1)

        def move_reflection(root):
            move_bookmark_to_cell(root, "lp_reflection_summary", 28, 2)
            move_bookmark_to_cell(root, "lp_reflection_innovation", 27, 2)

        mutations = {
            "teacher-student": swap_teacher_student,
            "stage-swap": swap_stage_content,
            "same-cell": same_physical_cell,
            "evaluation": move_evaluation,
            "reflection": move_reflection,
        }
        for name, mutate in mutations.items():
            with self.subTest(name=name), tempfile.TemporaryDirectory(prefix=f"lesson-package-builder-{name}-") as temp_name:
                folder = Path(temp_name)
                broken = folder / "broken-v11.docx"
                rebuilt = folder / "rebuilt.docx"
                shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", broken)
                patch_docx_document_xml(broken, mutate)
                result = run_script(
                    LESSON / "scripts" / "build_semantic_template.py",
                    "--source",
                    str(broken),
                    "--output",
                    str(rebuilt),
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn("bookmark", result.stderr.lower())

    def test_v11_builder_rejects_invalid_complete_bookmark_packages_and_preserves_output(self) -> None:
        def footer_duplicate_name(parts):
            add_story_bookmark(parts["word/footer2.xml"], "lp_title", "700")

        def footer_orphan_end(parts):
            add_story_bookmark(parts["word/footer2.xml"], "footer_orphan", "701", start=False)

        def header_invalid_name(parts):
            add_story_bookmark(parts["word/header1.xml"], "header-bad", "702")

        def header_invalid_id(parts):
            add_story_bookmark(parts["word/header1.xml"], "header_bad_id", "１２")

        def footer_duplicate_id(parts):
            add_story_bookmark(parts["word/footer2.xml"], "footer_duplicate_id", "1")

        def footer_managed_name(parts):
            add_story_bookmark(parts["word/footer2.xml"], "lp_impl_prep_content", "703")

        def cross_story_boundary(parts):
            document_root = parts["word/document.xml"]
            footer_root = parts["word/footer2.xml"]
            start = bookmark_start(document_root, "lp_hours")
            end = bookmark_end(document_root, start.get(qn("w:id")))
            end.getparent().remove(end)
            story_first_paragraph(footer_root).append(end)

        mutations = (
            ("footer-duplicate-name", footer_duplicate_name, False, ("bookmark", "duplicate", "footer")),
            ("footer-orphan-end", footer_orphan_end, False, ("bookmark", "orphan", "footer")),
            ("header-invalid-name", header_invalid_name, True, ("bookmark", "header")),
            ("header-invalid-id", header_invalid_id, True, ("bookmark", "header", "id")),
            ("footer-duplicate-id", footer_duplicate_id, False, ("bookmark", "duplicate", "id", "footer")),
            ("footer-managed-name", footer_managed_name, False, ("bookmark", "footer")),
            ("cross-story", cross_story_boundary, False, ("bookmark", "story")),
        )
        canonical = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
        for label, mutate, needs_header, keywords in mutations:
            with self.subTest(name=label), tempfile.TemporaryDirectory(prefix=f"lesson-package-builder-story-{label}-") as temp_name:
                folder = Path(temp_name)
                broken = folder / "broken-v11.docx"
                rebuilt = folder / "rebuilt.docx"
                shutil.copy2(canonical, broken)
                if needs_header:
                    ensure_header_story(broken)
                patch_docx_parts(broken, mutate)
                result = run_script(
                    LESSON / "scripts" / "build_semantic_template.py",
                    "--source",
                    str(broken),
                    "--output",
                    str(rebuilt),
                )
                diagnostic = (result.stdout + result.stderr).lower()
                self.assertNotEqual(result.returncode, 0)
                self.assertTrue(any(keyword in diagnostic for keyword in keywords), diagnostic)
                self.assertFalse(rebuilt.exists())

        with tempfile.TemporaryDirectory(prefix="lesson-package-builder-atomic-output-") as temp_name:
            folder = Path(temp_name)
            broken = folder / "broken-v11.docx"
            target = folder / "existing-output.docx"
            shutil.copy2(canonical, broken)
            patch_docx_parts(broken, footer_orphan_end)
            sentinel = b"existing target must survive validation failure"
            target.write_bytes(sentinel)
            result = run_script(
                LESSON / "scripts" / "build_semantic_template.py",
                "--source",
                str(broken),
                "--output",
                str(target),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertEqual(target.read_bytes(), sentinel)

    def test_v11_manifest_contract_is_strict_before_any_docx_generation(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        canonical = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
        base_manifest = yaml.safe_load(LESSON_V11_MANIFEST.read_text(encoding="utf-8"))

        def remove_path(data, *path):
            current = data
            for key in path[:-1]:
                current = current[key]
            del current[path[-1]]

        mutations = (
            ("anchors-required", "anchors.required", lambda data: remove_path(data, "anchors", "required")),
            ("anchors-containers", "anchors.containers", lambda data: remove_path(data, "anchors", "containers")),
            ("anchors-mode", "anchors.mode", lambda data: remove_path(data, "anchors", "mode")),
            ("course-bookmark", "fields.course_name.bookmark", lambda data: remove_path(data, "fields", "course_name", "bookmark")),
            ("title-target", "fields.title.target", lambda data: remove_path(data, "fields", "title", "target")),
            ("reflection-bookmarks", "fields.reflection.bookmarks", lambda data: remove_path(data, "fields", "reflection", "bookmarks")),
            ("implementation-stages", "fields.implementation.stages", lambda data: remove_path(data, "fields", "implementation", "stages")),
            ("stage-id", "fields.implementation.stages[2].id", lambda data: remove_path(data, "fields", "implementation", "stages", 2, "id")),
            ("stage-code", "fields.implementation.stages[2].code", lambda data: remove_path(data, "fields", "implementation", "stages", 2, "code")),
            ("stage-bookmarks", "fields.implementation.stages[2].bookmarks", lambda data: remove_path(data, "fields", "implementation", "stages", 2, "bookmarks")),
            ("stage-order", "stage id mismatch at index 1", lambda data: data["fields"]["implementation"]["stages"].__setitem__(slice(1, 3), [data["fields"]["implementation"]["stages"][2], data["fields"]["implementation"]["stages"][1]])),
            ("stage-id-mismatch", "stage id mismatch at index 3", lambda data: data["fields"]["implementation"]["stages"][3].__setitem__("id", "wrong_stage")),
            ("stage-code-mismatch", "stage code mismatch at index 3", lambda data: data["fields"]["implementation"]["stages"][3].__setitem__("code", "wrong_code")),
            ("evaluation-bookmark", "fields.evaluation.bookmark", lambda data: remove_path(data, "fields", "evaluation", "bookmark")),
        )

        for label, expected_error, mutate in mutations:
            with self.subTest(name=label), tempfile.TemporaryDirectory(prefix=f"lesson-package-manifest-contract-{label}-") as temp_name:
                folder = Path(temp_name)
                manifest = deepcopy(base_manifest)
                mutate(manifest)
                manifest_path = folder / "manifest.yaml"
                manifest_path.write_text(
                    yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
                    encoding="utf-8",
                )
                validate_result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--json",
                )
                validate_diagnostic = (validate_result.stdout + validate_result.stderr).lower()
                self.assertNotEqual(validate_result.returncode, 0)
                self.assertIn(expected_error.lower(), validate_diagnostic)
                generate_output = folder / "generated"
                generate_result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(generate_output),
                )
                generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
                self.assertNotEqual(generate_result.returncode, 0)
                self.assertIn(expected_error.lower(), generate_diagnostic)
                self.assertFalse(generate_output.exists())

        with tempfile.TemporaryDirectory(prefix="lesson-package-manifest-contract-script-") as temp_name:
            folder = Path(temp_name)
            manifest = deepcopy(base_manifest)
            remove_path(manifest, "anchors", "required")
            manifest_path = folder / "manifest.yaml"
            manifest_path.write_text(yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False), encoding="utf-8")
            shutil.copy2(canonical, folder / "template.docx")
            sys.modules.pop("package_common", None)
            sys.path.insert(0, str(LESSON / "scripts"))
            sys.path.insert(0, str(ROOT / "tests"))
            from validate_template_packages import validate_package

            with self.assertRaisesRegex(ValueError, "anchors.required"):
                validate_package(
                    {
                        "name": "lesson-plan-v1.1.0",
                        "manifest": manifest_path,
                        "schema": LESSON / "schemas" / "lesson-plan-input.schema.json",
                    }
                )

    def test_template_version_anchor_mode_matrix_uses_real_cli_entrypoints(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        v10_template = LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"
        v11_template = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"

        def remove_path(data, *path):
            current = data
            for key in path[:-1]:
                current = current[key]
            del current[path[-1]]

        cases = (
            ("1.0.0-legacy", LESSON_V10_MANIFEST, v10_template, "1.0.0", None, True, "legacy_coordinates"),
            ("1.0.1-legacy", LESSON_V10_MANIFEST, v10_template, "1.0.1", None, True, "legacy_coordinates"),
            ("1.1.0-semantic", LESSON_V11_MANIFEST, v11_template, "1.1.0", None, True, "word_bookmark"),
            ("1.1.1-semantic", LESSON_V11_MANIFEST, v11_template, "1.1.1", None, True, "word_bookmark"),
            (
                "1.1.1-missing-anchors",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
                lambda data: remove_path(data, "anchors"),
                False,
                "anchor mode mismatch",
            ),
            (
                "1.1.1-missing-mode",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
                lambda data: remove_path(data, "anchors", "mode"),
                False,
                "anchor mode mismatch",
            ),
            (
                "1.1.1-legacy-mode",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
                lambda data: data["anchors"].__setitem__("mode", "legacy_coordinates"),
                False,
                "anchor mode mismatch",
            ),
            (
                "1.1.1-old-coordinates-only",
                LESSON_V10_MANIFEST,
                v10_template,
                "1.1.1",
                None,
                False,
                "anchor mode mismatch",
            ),
            (
                "1.0.1-word-bookmark",
                LESSON_V10_MANIFEST,
                v10_template,
                "1.0.1",
                lambda data: data.__setitem__("anchors", {"mode": "word_bookmark"}),
                False,
                "anchor mode mismatch",
            ),
            (
                "1.0.1-semantic-metadata",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.0.1",
                lambda data: remove_path(data, "anchors", "mode"),
                False,
                "legacy manifest anchors.required",
            ),
            ("1.2.0-unsupported", LESSON_V11_MANIFEST, v11_template, "1.2.0", None, False, "unsupported lesson-plan template minor version"),
            ("2.0.0-unsupported", LESSON_V11_MANIFEST, v11_template, "2.0.0", None, False, "unsupported template major version"),
            ("v1.1.0-malformed", LESSON_V11_MANIFEST, v11_template, "v1.1.0", None, False, "template.version"),
            ("1.1-malformed", LESSON_V11_MANIFEST, v11_template, "1.1", None, False, "template.version"),
            ("1.1.0-beta-malformed", LESSON_V11_MANIFEST, v11_template, "1.1.0-beta", None, False, "template.version"),
            ("01.1.0-malformed", LESSON_V11_MANIFEST, v11_template, "01.1.0", None, False, "template.version"),
            ("unicode-digits-malformed", LESSON_V11_MANIFEST, v11_template, "１.１.０", None, False, "template.version"),
        )

        for label, source_manifest, source_template, version, mutate, should_pass, expected in cases:
            with self.subTest(name=label), tempfile.TemporaryDirectory(prefix=f"lesson-version-contract-{label}-") as temp_name:
                folder = Path(temp_name)
                template, manifest = write_lesson_manifest_package(
                    folder,
                    source_manifest,
                    source_template,
                    version,
                    mutate,
                )
                validate_result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(template),
                    "--manifest",
                    str(manifest),
                    "--json",
                )
                validate_diagnostic = (validate_result.stdout + validate_result.stderr).lower()
                output = folder / "generated"
                generate_result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(template),
                    "--manifest",
                    str(manifest),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
                if should_pass:
                    self.assertEqual(validate_result.returncode, 0, validate_diagnostic)
                    self.assertEqual(generate_result.returncode, 0, generate_diagnostic)
                    self.assertTrue(list(output.glob("*.docx")))
                    qa = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
                    self.assertEqual(qa["anchor_mode"], expected)
                else:
                    self.assertNotEqual(validate_result.returncode, 0)
                    self.assertIn(expected.lower(), validate_diagnostic)
                    self.assertNotEqual(generate_result.returncode, 0)
                    self.assertIn(expected.lower(), generate_diagnostic)
                    self.assertFalse(output.exists())

    def test_explicit_manifest_identity_and_patch_fingerprint_use_normal_copies(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        v11_template = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
        v10_template = LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"
        compatibility = LESSON / "assets" / "lesson-plan-template.docx"

        def assert_failure(template: Path, manifest: Path, folder: Path, expected: str) -> None:
            validate_result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(manifest),
                "--json",
            )
            validate_diagnostic = (validate_result.stdout + validate_result.stderr).lower()
            self.assertNotEqual(validate_result.returncode, 0)
            self.assertIn(expected.lower(), validate_diagnostic)
            output = folder / "generated"
            generate_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--template",
                str(template),
                "--manifest",
                str(manifest),
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
            self.assertNotEqual(generate_result.returncode, 0)
            self.assertIn(expected.lower(), generate_diagnostic)
            self.assertFalse(output.exists())

        with tempfile.TemporaryDirectory(prefix="lesson-package-explicit-identity-") as temp_name:
            folder = Path(temp_name)
            custom, manifest = write_lesson_manifest_package(
                folder,
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
            )
            manifest_data = yaml.safe_load(manifest.read_text(encoding="utf-8"))
            self.assertEqual(manifest_data["fingerprint"]["sha256"], file_sha256(custom))
            validate_result = run_script(
                LESSON / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--json",
            )
            self.assertEqual(validate_result.returncode, 0, validate_result.stdout + validate_result.stderr)
            output = folder / "generated"
            generate_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--template",
                str(custom),
                "--manifest",
                str(manifest),
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(generate_result.returncode, 0, generate_result.stdout + generate_result.stderr)
            qa = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(qa["template_version"], "1.1.1")
            self.assertEqual(qa["anchor_mode"], "word_bookmark")

        with tempfile.TemporaryDirectory(prefix="lesson-package-explicit-identity-failures-") as temp_name:
            folder = Path(temp_name)
            canonical_manifest = yaml.safe_load(LESSON_V11_MANIFEST.read_text(encoding="utf-8"))
            canonical_manifest["template"]["version"] = "1.1.1"
            canonical_manifest["template"]["file"] = "template.docx"
            canonical_manifest["template"]["compatibility_entries"] = []
            canonical_manifest_path = folder / "canonical-mismatch.yaml"
            canonical_manifest_path.write_text(
                yaml.safe_dump(canonical_manifest, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            assert_failure(v11_template, canonical_manifest_path, folder / "canonical-failure", "template/manifest mismatch")

            compatibility_manifest = yaml.safe_load(LESSON_V10_MANIFEST.read_text(encoding="utf-8"))
            compatibility_manifest["template"]["version"] = "1.0.1"
            compatibility_manifest["template"]["file"] = "template.docx"
            compatibility_manifest["template"]["compatibility_entries"] = []
            compatibility_manifest_path = folder / "compatibility-mismatch.yaml"
            compatibility_manifest_path.write_text(
                yaml.safe_dump(compatibility_manifest, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            assert_failure(compatibility, compatibility_manifest_path, folder / "compatibility-failure", "template/manifest mismatch")

            wrong_template, wrong_manifest = write_lesson_manifest_package(
                folder / "wrong-fingerprint",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
                lambda data: data["fingerprint"].update({"sha256": "0" * 64, "value": "0" * 64}),
            )
            assert_failure(wrong_template, wrong_manifest, folder / "wrong-fingerprint-failure", "template fingerprint mismatch")

            missing_bookmark, missing_manifest = write_lesson_manifest_package(
                folder / "missing-bookmark",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
            )

            def remove_bookmark(root) -> None:
                start = bookmark_start(root, "lp_course_name")
                end = bookmark_end(root, start.get(qn("w:id")))
                start.getparent().remove(start)
                end.getparent().remove(end)

            patch_docx_document_xml(missing_bookmark, remove_bookmark)
            update_manifest_fingerprint(missing_manifest, missing_bookmark)
            assert_failure(missing_bookmark, missing_manifest, folder / "missing-bookmark-failure", "bookmark")

            broken_structure, broken_manifest = write_lesson_manifest_package(
                folder / "broken-structure",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
            )

            def remove_main_table_row(root) -> None:
                body = root.find(qn("w:body"))
                table = next(node for node in body if node.tag == qn("w:tbl"))
                rows = [node for node in table if node.tag == qn("w:tr")]
                rows[0].getparent().remove(rows[0])

            patch_docx_document_xml(broken_structure, remove_main_table_row)
            update_manifest_fingerprint(broken_manifest, broken_structure)
            assert_failure(
                broken_structure,
                broken_manifest,
                folder / "broken-structure-failure",
                "main table row count mismatch",
            )

    def test_explicit_template_fingerprint_cannot_be_skipped_by_real_cli(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        v11_template = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"

        def tamper_zip_member(path: Path) -> None:
            temporary = path.with_suffix(path.suffix + ".tmp")
            with zipfile.ZipFile(path, "r") as source_zip, zipfile.ZipFile(
                temporary, "w", zipfile.ZIP_DEFLATED
            ) as target_zip:
                for info in source_zip.infolist():
                    target_zip.writestr(info, source_zip.read(info.filename))
                target_zip.writestr("qa/fingerprint-tamper.txt", b"fingerprint regression")
            temporary.replace(path)

        def run_generator(template: Path, manifest: Path, output: Path) -> subprocess.CompletedProcess[str]:
            return run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--template",
                str(template),
                "--manifest",
                str(manifest),
                "--skip-template-validation",
            )

        def run_output_validator(template: Path, manifest: Path, output: Path) -> subprocess.CompletedProcess[str]:
            return run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
                "--template-path",
                str(template),
                "--manifest",
                str(manifest),
                "--skip-validation",
            )

        with tempfile.TemporaryDirectory(prefix="lesson-package-fingerprint-skips-") as temp_name:
            folder = Path(temp_name)
            tampered_template, tampered_manifest = write_lesson_manifest_package(
                folder / "tampered",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
            )
            manifest_data = yaml.safe_load(tampered_manifest.read_text(encoding="utf-8"))
            expected_hash = manifest_data["fingerprint"]["sha256"]
            tamper_zip_member(tampered_template)
            self.assertNotEqual(file_sha256(tampered_template), expected_hash)

            generate_output = folder / "generate-skip"
            generate_result = run_generator(tampered_template, tampered_manifest, generate_output)
            generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
            self.assertNotEqual(generate_result.returncode, 0)
            self.assertIn("fingerprint mismatch", generate_diagnostic)
            self.assertFalse(list(generate_output.glob("*.docx")))
            generate_qa = generate_output / "qa-report.json"
            if generate_qa.exists():
                self.assertNotIn(json.loads(generate_qa.read_text(encoding="utf-8"))["status"], {"passed", "skipped"})

            output_skip_dir = folder / "output-skip"
            output_result = run_output_validator(tampered_template, tampered_manifest, output_skip_dir)
            output_diagnostic = (output_result.stdout + output_result.stderr).lower()
            self.assertNotEqual(output_result.returncode, 0)
            self.assertIn("fingerprint mismatch", output_diagnostic)
            output_qa = output_skip_dir / "qa-report.json"
            self.assertFalse(output_qa.exists())

            valid_template, valid_manifest = write_lesson_manifest_package(
                folder / "valid",
                LESSON_V11_MANIFEST,
                v11_template,
                "1.1.1",
            )
            valid_manifest_data = yaml.safe_load(valid_manifest.read_text(encoding="utf-8"))
            self.assertEqual(file_sha256(valid_template), valid_manifest_data["fingerprint"]["sha256"])

            valid_generate_output = folder / "valid-generate-skip"
            valid_generate_result = run_generator(valid_template, valid_manifest, valid_generate_output)
            self.assertEqual(
                valid_generate_result.returncode,
                0,
                valid_generate_result.stdout + valid_generate_result.stderr,
            )
            valid_generate_report = json.loads(
                (valid_generate_output / "qa-report.json").read_text(encoding="utf-8")
            )
            self.assertEqual(valid_generate_report["status"], "skipped")
            self.assertIn(
                "Template validation skipped by explicit flag.",
                valid_generate_report["warnings"],
            )
            self.assertTrue(list(valid_generate_output.glob("*.docx")))

            valid_output_skip_dir = folder / "valid-output-skip"
            valid_output_result = run_output_validator(valid_template, valid_manifest, valid_output_skip_dir)
            self.assertEqual(
                valid_output_result.returncode,
                0,
                valid_output_result.stdout + valid_output_result.stderr,
            )
            valid_output_report = json.loads(
                (valid_output_skip_dir / "qa-report.json").read_text(encoding="utf-8")
            )
            self.assertEqual(valid_output_report["status"], "skipped")
            self.assertIn(
                "Output validation skipped by explicit flag.",
                valid_output_report["warnings"],
            )

    def test_legacy_manifest_rejects_all_semantic_metadata_by_real_cli(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        canonical = LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"
        base_manifest = yaml.safe_load(LESSON_V10_MANIFEST.read_text(encoding="utf-8"))

        def mutate_path(data, *path, value=None):
            if path[0] == "anchors" and "anchors" not in data:
                data["anchors"] = {}
            current = data
            for key in path[:-1]:
                current = current[key]
            current[path[-1]] = value

        mutations = (
            ("anchors-required", ("anchors", "required"), ["lp_title"], "anchors.required"),
            ("anchors-containers", ("anchors", "containers"), {}, "anchors.containers"),
            ("anchors-type", ("anchors", "type"), "word_bookmark", "anchors.type"),
            ("fixed-bookmark", ("fields", "title", "bookmark"), "lp_title", "fields.title.bookmark"),
            ("implementation-mode", ("fields", "implementation", "mode"), "anchored_cells", "fields.implementation.mode"),
            ("implementation-stages", ("fields", "implementation", "stages"), [], "fields.implementation.stages"),
            ("implementation-bookmarks", ("fields", "implementation", "bookmarks"), [], "fields.implementation.bookmarks"),
            ("reflection-mode", ("fields", "reflection", "mode"), "anchored_cells", "fields.reflection.mode"),
            ("reflection-bookmarks", ("fields", "reflection", "bookmarks"), [], "fields.reflection.bookmarks"),
            ("evaluation-bookmark", ("fields", "evaluation", "bookmark"), "lp_evaluation", "fields.evaluation.bookmark"),
            ("evaluation-target-conflict", ("fields", "evaluation", "target"), "table_cell", "fields.evaluation.target"),
            ("fixed-container", ("fields", "title", "container"), "document_paragraph", "fields.title.container"),
            ("implementation-table", ("fields", "implementation", "table"), 0, "fields.implementation.table"),
            ("mixed-definition", ("fields", "implementation", "table"), 0, "fields.implementation.table"),
        )
        for label, path, value, expected in mutations:
            with self.subTest(name=label), tempfile.TemporaryDirectory(prefix=f"lesson-legacy-contract-{label}-") as temp_name:
                folder = Path(temp_name)
                manifest = deepcopy(base_manifest)
                mutate_path(manifest, *path, value=value)
                manifest_path = folder / "manifest.yaml"
                manifest_path.write_text(
                    yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
                    encoding="utf-8",
                )
                validate_result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--json",
                )
                validate_diagnostic = (validate_result.stdout + validate_result.stderr).lower()
                self.assertNotEqual(validate_result.returncode, 0)
                self.assertIn(expected.lower(), validate_diagnostic)
                output = folder / "generated"
                generate_result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
                self.assertNotEqual(generate_result.returncode, 0)
                self.assertIn(expected.lower(), generate_diagnostic)
                self.assertFalse(output.exists())

    def test_semantic_composite_manifest_rejects_conflicting_extra_keys_by_real_cli(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        canonical = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
        base_manifest = yaml.safe_load(LESSON_V11_MANIFEST.read_text(encoding="utf-8"))
        mutations = (
            ("implementation-target", lambda data: data["fields"]["implementation"].__setitem__("target", "table_cell"), "fields.implementation contains unsupported key target"),
            ("implementation-rows", lambda data: data["fields"]["implementation"].__setitem__("rows", []), "fields.implementation contains unsupported key rows"),
            ("implementation-bookmark", lambda data: data["fields"]["implementation"].__setitem__("bookmark", "lp_impl"), "fields.implementation contains unsupported key bookmark"),
            ("stage-row", lambda data: data["fields"]["implementation"]["stages"][2].__setitem__("row", 18), "fields.implementation.stages[2] contains unsupported key row"),
            ("stage-target", lambda data: data["fields"]["implementation"]["stages"][2].__setitem__("target", "table_cell"), "fields.implementation.stages[2] contains unsupported key target"),
            ("reflection-target", lambda data: data["fields"]["reflection"].__setitem__("target", "nested_table"), "fields.reflection contains unsupported key target"),
            ("reflection-rows", lambda data: data["fields"]["reflection"].__setitem__("rows", []), "fields.reflection contains unsupported key rows"),
            ("reflection-container", lambda data: data["fields"]["reflection"].__setitem__("container", "cell"), "fields.reflection contains unsupported key container"),
            ("fixed-table", lambda data: data["fields"]["course_name"].__setitem__("table", 0), "fields.course_name contains unsupported key table"),
            ("title-paragraph", lambda data: data["fields"]["title"].__setitem__("paragraph", 0), "fields.title contains unsupported key paragraph"),
            ("evaluation-row", lambda data: data["fields"]["evaluation"].__setitem__("row", 12), "fields.evaluation contains unsupported key row"),
            ("unexpected-key", lambda data: data["fields"].__setitem__("unexpected_key", {}), "fields contains unsupported key unexpected_key"),
        )
        for label, mutate, expected in mutations:
            with self.subTest(name=label), tempfile.TemporaryDirectory(prefix=f"lesson-semantic-extra-{label}-") as temp_name:
                folder = Path(temp_name)
                manifest = deepcopy(base_manifest)
                mutate(manifest)
                manifest_path = folder / "manifest.yaml"
                manifest_path.write_text(
                    yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
                    encoding="utf-8",
                )
                validate_result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--json",
                )
                validate_diagnostic = (validate_result.stdout + validate_result.stderr).lower()
                self.assertNotEqual(validate_result.returncode, 0)
                self.assertIn(expected.lower(), validate_diagnostic)
                output = folder / "generated"
                generate_result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
                self.assertNotEqual(generate_result.returncode, 0)
                self.assertIn(expected.lower(), generate_diagnostic)
                self.assertFalse(output.exists())

    def test_v11_fixed_field_target_and_mode_contract_is_strict_before_generation(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        canonical = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
        base_manifest = yaml.safe_load(LESSON_V11_MANIFEST.read_text(encoding="utf-8"))
        mutations = (
            ("title-target", "fields.title.target", "table_cell", "document_paragraph"),
            ("title-mode", "fields.title.mode", "replace_single_paragraph", "replace_text_preserve_style"),
            ("course-target", "fields.course_name.target", "document_paragraph", "table_cell"),
            ("course-mode", "fields.course_name.mode", "replace_paragraphs", "replace_single_paragraph"),
            ("teaching-content-typo", "fields.teaching_content.mode", "replace_paragrphs", "replace_paragraphs"),
            ("teaching-content-unknown", "fields.teaching_content.mode", "unknown", "replace_paragraphs"),
            ("student-base-target", "fields.student_base.target", "nested_table", "table_cell"),
            ("resources-mode", "fields.resources.mode", "replace_single_paragraph", "replace_paragraphs"),
            ("evaluation-target", "fields.evaluation.target", "table_cell", "nested_table"),
            ("evaluation-mode", "fields.evaluation.mode", "replace_paragraphs", "nested_table"),
        )
        for label, path, actual, expected in mutations:
            with self.subTest(name=label), tempfile.TemporaryDirectory(prefix=f"lesson-field-contract-{label}-") as temp_name:
                folder = Path(temp_name)
                manifest = deepcopy(base_manifest)
                field_name, key = path.split(".", 1)[1].split(".")
                manifest["fields"][field_name][key] = actual
                manifest_path = folder / "manifest.yaml"
                manifest_path.write_text(
                    yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
                    encoding="utf-8",
                )
                validate_result = run_script(
                    LESSON / "scripts" / "validate_template.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--json",
                )
                validate_diagnostic = (validate_result.stdout + validate_result.stderr).lower()
                self.assertNotEqual(validate_result.returncode, 0)
                self.assertIn(path.lower(), validate_diagnostic)
                self.assertIn(expected.lower(), validate_diagnostic)
                output = folder / "generated"
                generate_result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(canonical),
                    "--manifest",
                    str(manifest_path),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                generate_diagnostic = (generate_result.stdout + generate_result.stderr).lower()
                self.assertNotEqual(generate_result.returncode, 0)
                self.assertIn(path.lower(), generate_diagnostic)
                self.assertIn(expected.lower(), generate_diagnostic)
                self.assertFalse(output.exists())

    def test_v11_bookmark_ids_require_ascii_decimal_digits_in_template_and_output(self) -> None:
        sys.modules.pop("bookmark_utils", None)
        sys.path.insert(0, str(LESSON / "scripts"))
        from bookmark_utils import BOOKMARK_ID_PATTERN

        for value in ("0", "1", "69", "9999"):
            self.assertIsNotNone(BOOKMARK_ID_PATTERN.fullmatch(value))
        invalid_values = (
            ("fullwidth", "１２"),
            ("arabic-indic", "١٢"),
            ("mixed", "1٢"),
            ("question", "??"),
            ("mixed-symbol", "1?"),
            ("plus", "+12"),
            ("leading-space", " 12"),
            ("trailing-space", "12 "),
            ("empty", ""),
        )
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        canonical = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
        with tempfile.TemporaryDirectory(prefix="lesson-package-ascii-id-") as temp_name:
            folder = Path(temp_name)
            pristine = folder / "pristine"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(pristine),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)

            def set_hours_id(root, value: str):
                start = bookmark_start(root, "lp_hours")
                end = bookmark_end(root, start.get(qn("w:id")))
                start.set(qn("w:id"), value)
                end.set(qn("w:id"), value)

            for label, invalid_id in invalid_values:
                with self.subTest(name=label):
                    custom = folder / f"custom-{label}.docx"
                    shutil.copy2(canonical, custom)
                    patch_docx_document_xml(custom, lambda root, value=invalid_id: set_hours_id(root, value))
                    template_manifest = write_manifest_for_modified_template(
                        folder,
                        LESSON_V11_MANIFEST,
                        custom,
                        canonical,
                    )
                    template_result = run_script(
                        LESSON / "scripts" / "validate_template.py",
                        "--template",
                        str(custom),
                        "--manifest",
                        str(template_manifest),
                        "--json",
                    )
                    self.assertNotEqual(template_result.returncode, 0)
                    self.assertIn("ASCII", (template_result.stdout + template_result.stderr))

                    output_dir = folder / f"output-{label}"
                    shutil.copytree(pristine, output_dir)
                    patch_docx_document_xml(sorted(output_dir.glob("*.docx"))[0], lambda root, value=invalid_id: set_hours_id(root, value))
                    output_result = run_script(
                        LESSON / "scripts" / "validate_output.py",
                        "--input-json",
                        str(source),
                        "--output-dir",
                        str(output_dir),
                        "--manifest",
                        str(LESSON_V11_MANIFEST),
                    )
                    self.assertNotEqual(output_result.returncode, 0)
                    self.assertIn("ASCII", (output_result.stdout + output_result.stderr))

    def test_template_only_cli_resolution_and_mismatch_guards(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        with tempfile.TemporaryDirectory(prefix="lesson-package-cli-resolution-") as temp_name:
            folder = Path(temp_name)

            default_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(folder / "default-v11"),
            )
            self.assertEqual(default_result.returncode, 0, default_result.stderr or default_result.stdout)
            default_report = json.loads((folder / "default-v11" / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(default_report["template_version"], "1.1.1")
            self.assertEqual(default_report["anchor_mode"], "word_bookmark")

            for label, template in (
                ("canonical-v10", LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"),
                ("compatibility-v10", LESSON / "assets" / "lesson-plan-template.docx"),
                ("canonical-v11", LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"),
                ("canonical-v111", LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.1" / "template.docx"),
            ):
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(template),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(folder / label),
                )
                self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
                report = json.loads((folder / label / "qa-report.json").read_text(encoding="utf-8"))
                expected_version = "1.1.0" if label == "canonical-v11" else ("1.1.1" if label == "canonical-v111" else "1.0.0")
                expected_mode = "word_bookmark" if label in {"canonical-v11", "canonical-v111"} else "legacy_coordinates"
                self.assertEqual(report["template_version"], expected_version)
                self.assertEqual(report["anchor_mode"], expected_mode)

            for label, template, manifest in (
                (
                    "v10-v11-mismatch",
                    LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx",
                    LESSON_V11_MANIFEST,
                ),
                (
                    "v11-v10-mismatch",
                    LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx",
                    LESSON_V10_MANIFEST,
                ),
            ):
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--template",
                    str(template),
                    "--manifest",
                    str(manifest),
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(folder / label),
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn("template/manifest mismatch", result.stderr.lower())

            custom = folder / "custom-without-manifest.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx", custom)
            patch_docx_document_xml(
                custom,
                lambda root: root.xpath(
                    ".//w:t[text()='课程名称']",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"},
                )[0].__setattr__("text", "课程名"),
            )
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--template",
                str(custom),
                "--tasks-json",
                str(source),
                "--output-dir",
                str(folder / "custom-without-manifest-output"),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("custom template requires a matching --manifest", result.stderr.lower())


class GradebookTotalRuleTests(unittest.TestCase):
    def test_total_rule_matches_exactly_with_zero_and_nonzero_skill_weights(self) -> None:
        sys.modules.pop("package_common", None)
        sys.path.insert(0, str(GRADE / "scripts"))
        from package_common import calculate_expected_total, source_total_matches, validate_source_totals

        no_skill_weights = {"regular": 0.6, "theory": 0.4, "skill": 0.0}
        no_skill = {"regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0}
        no_skill_expected = calculate_expected_total(no_skill, no_skill_weights)
        self.assertEqual(no_skill_expected, 87)
        self.assertTrue(source_total_matches(87.0000000001, no_skill_expected))
        self.assertFalse(source_total_matches(87.4, no_skill_expected))
        validate_source_totals([no_skill], no_skill_weights)

        skill_weights = {"regular": 0.5, "theory": 0.3, "skill": 0.2}
        skill = {"regular": 91.0, "theory": 90.0, "skill": 90.0, "total": 91.0}
        skill_expected = calculate_expected_total(skill, skill_weights)
        self.assertEqual(skill_expected, 91)
        validate_source_totals([skill], skill_weights)

        self.assertFalse(source_total_matches(88, no_skill_expected))
        self.assertFalse(source_total_matches(86, no_skill_expected))
        with self.assertRaisesRegex(ValueError, "Source total mismatch"):
            validate_source_totals([{**no_skill, "total": 88}], no_skill_weights)
        with self.assertRaisesRegex(ValueError, "Source total mismatch"):
            validate_source_totals([{**no_skill, "total": 86}], no_skill_weights)
        with self.assertRaisesRegex(ValueError, "Source total mismatch"):
            validate_source_totals([{**no_skill, "total": 87.4}], no_skill_weights)


class GradebookPowerShellContractTests(unittest.TestCase):
    def test_com_path_uses_same_rounding_preflight_and_exact_output_contract(self) -> None:
        script = (GRADE / "scripts" / "generate_gradebook.ps1").read_text(encoding="utf-8-sig")
        self.assertIn("function Excel-Round", script)
        self.assertIn("function Format-Percentage-Label", script)
        self.assertIn("[System.MidpointRounding]::AwayFromZero", script)
        self.assertIn("function Assert-ManifestCompatibility", script)
        self.assertIn("Assert-ManifestCompatibility $ManifestData", script)
        self.assertIn("function Source-Total-Matches", script)
        self.assertIn("function Assert-HalfPointRegularScores", script)
        self.assertIn("function Assert-NormalizedInput", script)
        self.assertIn("validate_input.py", script)
        self.assertIn("Assert-NormalizedInput $normalizedInput", script)
        self.assertIn("function Assert-SourceTotals", script)
        self.assertIn("Assert-HalfPointRegularScores $students", script)
        self.assertIn("Assert-SourceTotals $students $meta", script)
        self.assertIn("'--output-file'", script)
        self.assertIn("excel_named_range", script)
        self.assertIn("function Build-One-NamedRange", script)
        self.assertIn("function Get-ManagedRanges", script)
        self.assertIn("Names.Item", script)
        self.assertIn("Rebuild-ManagedNamesAfterColumnDelete", script)
        self.assertIn("Set-ManagedWorkbookName", script)
        self.assertIn("--identity-only", script)
        self.assertIn("--named-range-runtime-preflight", script)
        self.assertIn("Assert-ManagedRuntimeContract", script)
        self.assertIn("Named value write offset out of bounds", script)
        self.assertIn("Named formula write offset out of bounds", script)
        self.assertNotIn("abs_tol=1.0", script)

    def test_local_com_integration_script_is_repeatable_and_has_skip_boundary(self) -> None:
        script = (GRADE / "tests" / "run_com_integration.ps1").read_text(encoding="utf-8-sig")
        self.assertIn("New-Object -ComObject Excel.Application", script)
        self.assertIn("excel_named_range", script)
        self.assertIn("preserved_named_range_count", script)
        self.assertIn("status = 'skipped'", script)


class WorkflowContractTests(unittest.TestCase):
    def test_template_package_ci_runs_on_main_push(self) -> None:
        workflow = (ROOT / ".github" / "workflows" / "template-package-ci.yml").read_text(encoding="utf-8")
        workflow_data = yaml.safe_load(workflow)
        events = workflow_data.get("on", workflow_data.get(True, {}))
        self.assertIn("main", events["push"]["branches"])
        self.assertIn("master", events["push"]["branches"])
        self.assertEqual(events["schedule"][0]["cron"], "17 3 * * 1")
        self.assertEqual(
            workflow_data["concurrency"]["cancel-in-progress"],
            "${{ github.event_name == 'pull_request' || github.event_name == 'push' }}",
        )

        jobs = workflow_data["jobs"]
        for job_name in (
            "classify-changes",
            "documentation-checks",
            "package-contracts",
            "template-tooling",
            "template-lesson",
            "template-gradebook",
            "template-release",
            "ci-gate",
        ):
            self.assertIn(job_name, jobs)
        self.assertEqual(jobs["classify-changes"]["runs-on"], "ubuntu-latest")
        self.assertEqual(jobs["package-contracts"]["runs-on"], "ubuntu-latest")
        self.assertEqual(jobs["ci-gate"]["runs-on"], "ubuntu-latest")
        for job_name in ("template-tooling", "template-lesson", "template-gradebook", "template-release"):
            self.assertEqual(jobs[job_name]["timeout-minutes"], 30)
            setup_python = next(
                step
                for step in jobs[job_name]["steps"]
                if step.get("uses") == "actions/setup-python@v5"
            )
            self.assertEqual(setup_python["with"]["cache"], "pip")

        tooling_text = str(jobs["template-tooling"]).lower()
        lesson_text = str(jobs["template-lesson"]).lower()
        gradebook_text = str(jobs["template-gradebook"]).lower()
        release_text = str(jobs["template-release"]).lower()
        self.assertIn("libreoffice", tooling_text)
        self.assertNotIn("libreoffice", lesson_text)
        self.assertIn("libreoffice", gradebook_text)
        self.assertIn("libreoffice", release_text)
        self.assertNotIn("pip install --upgrade pip", workflow.lower())

        release_workflow = (ROOT / ".github" / "workflows" / "template-release.yml").read_text(encoding="utf-8")
        self.assertIn("runs-on: macos-14", release_workflow)
        self.assertNotIn("ubuntu-latest", release_workflow)
        self.assertIn("cancel-in-progress: false", release_workflow)
        self.assertNotIn("cancel-in-progress: true", release_workflow)

    def test_template_package_ci_docs_only_fast_path_contract(self) -> None:
        workflow = (ROOT / ".github" / "workflows" / "template-package-ci.yml").read_text(encoding="utf-8")
        workflow_data = yaml.safe_load(workflow)
        jobs = workflow_data["jobs"]
        events = workflow_data.get("on", workflow_data.get(True, {}))

        self.assertFalse(events["pull_request"] and "paths" in events["pull_request"])
        self.assertNotIn("paths", events["push"])
        self.assertIn("workflow_dispatch", events)
        self.assertIn("schedule", events)
        self.assertIn("concurrency", workflow)
        self.assertIn("cancel-in-progress", workflow)

        classifier = (ROOT / ".github" / "scripts" / "classify_ci_changes.py").read_text(encoding="utf-8")
        for output in (
            "docs_only",
            "run_docs",
            "run_lesson",
            "run_gradebook",
            "run_tooling",
            "run_release",
            "run_package_contracts",
            "force_full",
            "changed_files",
            "classification",
            "reason",
        ):
            self.assertIn(output, classifier)
        for allowed_path in (
            "README.md",
            "docs/",
            "教案生成器/简介.md",
            "平时成绩记分册生成器/简介.md",
            "多Agent兼容规范.md",
        ):
            self.assertIn(allowed_path, classifier)
        self.assertIn("workflow_dispatch", classifier)
        self.assertIn("schedule", classifier)
        self.assertIn("unknown or ambiguous", classifier)

        for job_name in ("classify-changes", "documentation-checks", "package-contracts", "ci-gate"):
            self.assertIn(job_name, jobs)
        self.assertEqual(jobs["classify-changes"]["runs-on"], "ubuntu-latest")
        self.assertEqual(jobs["documentation-checks"]["runs-on"], "ubuntu-latest")
        self.assertEqual(jobs["ci-gate"]["runs-on"], "ubuntu-latest")

        documentation = "\n".join(
            step.get("run", "")
            for step in jobs["documentation-checks"]["steps"]
        )
        self.assertIn("git diff --check", documentation)
        self.assertIn("urlsplit", documentation)
        self.assertIn("install.py", documentation)
        self.assertIn("version_from_manifest", documentation)
        self.assertNotIn("libreoffice", documentation.lower())
        self.assertNotIn("pip install", documentation.lower())

        for heavy_job, output in (
            ("template-tooling", "run_tooling"),
            ("template-lesson", "run_lesson"),
            ("template-gradebook", "run_gradebook"),
            ("template-release", "run_release"),
        ):
            self.assertIn("always()", jobs[heavy_job]["if"])
            self.assertIn(output, jobs[heavy_job]["if"])
            self.assertIn("force_full", jobs[heavy_job]["if"])

        self.assertIn("validate_package_contracts.py", "\n".join(step.get("run", "") for step in jobs["package-contracts"]["steps"]))
        for heavy_job in ("template-tooling", "template-gradebook", "template-release"):
            heavy_steps = "\n".join(step.get("run", "") for step in jobs[heavy_job]["steps"])
            self.assertNotIn("tests/validate_template_packages.py", heavy_steps)
        self.assertIn("package-contracts", jobs["template-tooling"]["needs"])
        self.assertIn("package-contracts", jobs["template-release"]["needs"])
        self.assertIn("needs.package-contracts.result == 'success'", jobs["template-tooling"]["if"])
        self.assertIn("needs.package-contracts.result == 'success'", jobs["template-release"]["if"])
        self.assertEqual(jobs["template-lesson"]["needs"], "classify-changes")
        self.assertEqual(jobs["template-gradebook"]["needs"], "classify-changes")

        self.assertEqual(
            set(jobs["ci-gate"]["needs"]),
            {
                "classify-changes",
                "documentation-checks",
                "package-contracts",
                "template-tooling",
                "template-lesson",
                "template-gradebook",
                "template-release",
            },
        )
        gate = "\n".join(step.get("run", "") for step in jobs["ci-gate"]["steps"])
        self.assertIn("skipped", gate)
        self.assertIn("RUN_PACKAGE", gate)
        self.assertIn("CHANGED_FILES", gate)
        self.assertIn("CI Gate: success", gate)


@unittest.skipUnless(soffice_path(), "LibreOffice is required for XLS package tests")
class GradebookTemplatePackageTests(unittest.TestCase):
    def make_source(
        self,
        folder: Path,
        skill: bool = False,
        count: int = 2,
        leading_zero: bool = False,
        total_delta: float = 0.0,
        regular_override: float | None = None,
        regular_pct: float | None = None,
        theory_pct: float | None = None,
        skill_pct: float | None = None,
    ) -> Path:
        folder.mkdir(parents=True, exist_ok=True)
        xlsx = folder / "课程成绩单.xlsx"
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "成绩单"
        regular_pct = (0.5 if skill else 0.6) if regular_pct is None else regular_pct
        theory_pct = (0.3 if skill else 0.4) if theory_pct is None else theory_pct
        skill_pct = (0.2 if skill else 0.0) if skill_pct is None else skill_pct
        sheet["A2"] = f"课程名称:软件测试实训 教师:张老师 上课班级:软件技术2401班 成绩项目比例:技能成绩{skill_pct * 100:g}% 理论成绩{theory_pct * 100:g}% 平时成绩{regular_pct * 100:g}%"
        sheet["A3"] = "开课学期:2025-2026-2"
        headers = ["学号", "姓名", "平时成绩", "理论成绩"] + (["技能成绩"] if skill else []) + ["总成绩"]
        for col, value in enumerate(headers, start=1):
            sheet.cell(4, col).value = value
        rows = []
        for index in range(count):
            regular = regular_override if regular_override is not None and index == 0 else [86.5, 91.0, 100.0, 0.0][index % 4]
            theory = [88.0, 90.0, 100.0, 0.0][index % 4]
            skill_score = [92.0, 90.0, 100.0, 0.0][index % 4]
            total = math.floor(regular * regular_pct + theory * theory_pct + skill_score * skill_pct + 0.5)
            if index == 0:
                total += total_delta
            student_id = "0012345678" if leading_zero and index == 0 else f"240101{index + 1:03d}"
            values = [student_id, f"学生{index + 1}", regular, theory] + ([skill_score] if skill else []) + [total]
            rows.append(values)
        for row, values in enumerate(rows, start=5):
            for col, value in enumerate(values, start=1):
                sheet.cell(row, col).value = value
            sheet.cell(row, 1).number_format = "@"
        workbook.save(xlsx)
        subprocess.run(
            [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(folder), str(xlsx)],
            check=True,
            capture_output=True,
        )
        return folder / "课程成绩单.xls"

    def test_canonical_template_and_compatibility_entry(self) -> None:
        result = run_script(GRADE / "scripts" / "validate_template.py", "--json")
        self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
        report = json.loads(result.stdout)
        self.assertEqual(report["template_version"], "1.1.0")
        self.assertEqual(report["checks"]["named_ranges_xlsx"]["actual_count"], 24)
        self.assertEqual(report["checks"]["named_ranges_xls"]["actual_count"], 24)
        self.assertEqual(report["checks"]["structure"]["anchor_mode"], "excel_named_range")
        self.assertEqual(report["checks"]["structure"]["rows"], 52)
        self.assertEqual(report["checks"]["structure"]["columns"], 17)

        for template, expected_mode in (
            (
                GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls",
                "legacy_coordinates",
            ),
            (GRADE / "assets" / "平时成绩记分册模板.xls", "legacy_coordinates"),
        ):
            legacy = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--json",
            )
            self.assertEqual(legacy.returncode, 0, legacy.stderr or legacy.stdout)
            legacy_report = json.loads(legacy.stdout)
            self.assertEqual(legacy_report["template_version"], "1.0.0")
            self.assertEqual(legacy_report["checks"]["structure"].get("anchor_mode", expected_mode), expected_mode)

    def test_invalid_input_is_rejected(self) -> None:
        sys.modules.pop("package_common", None)
        sys.path.insert(0, str(GRADE / "scripts"))
        from package_common import validate_input

        bad = {"term": "2025", "course": "软件测试", "teacher": "张老师", "class_name": "一班", "weights": {"regular": 1, "theory": 0, "skill": 0}, "students": [{"id": "bad"}]}
        with self.assertRaises(ValueError):
            validate_input(bad)

    def test_python_generator_zero_skill_and_qa(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-test-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            output.mkdir()
            final_output = output / f"{source.parent.name}-平时成绩记分册.xls"
            final_output.write_bytes(b"old-formal-output")
            (output / "qa-report.json").write_text("{\"status\": \"old\"}\n", encoding="utf-8")
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            self.assertTrue((output / "qa-report.json").exists())
            self.assertEqual(len(list(output.glob("*.xls"))), 1)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["template_id"], "course-gradebook")
            self.assertEqual(report["template_version"], "1.1.0")
            self.assertEqual(report["generator_version"], "1.1.0")
            self.assertEqual(report["engine"], "libreoffice-openpyxl")
            self.assertFalse(report["custom_template"])
            self.assertEqual(report["validation_skipped"], [])
            self.assertEqual(report["status"], "passed")
            generated = next(output.glob("*.xls"))
            self.assertEqual(report["output_file"], generated.name)
            self.assertEqual(report["output_dir"], str(output.resolve()))
            self.assertEqual(report["qa_report"], str((output / "qa-report.json").resolve()))
            self.assertTrue(Path(report["output_dir"]).is_dir())
            self.assertTrue(Path(report["qa_report"]).is_file())
            report_text = (output / "qa-report.json").read_text(encoding="utf-8")
            for temporary_marker in (
                "gradebook-run-",
                "gradebook-com-run-",
                "validation-",
                "candidate-",
                "TemporaryDirectory",
                "AppData\\Local\\Temp",
            ):
                self.assertNotIn(temporary_marker, report_text)
            self.assertEqual(report["files_checked"], 1)
            self.assertEqual(report["anchor_mode"], "excel_named_range")
            self.assertEqual(report["named_range_variant"], "without_skill")
            self.assertEqual(report["required_named_range_count"], 21)
            self.assertEqual(report["preserved_named_range_count"], 21)

    def test_python_generator_preserves_fractional_weight_headers(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-fractional-weights-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, skill=True, regular_pct=0.333, theory_pct=0.333, skill_pct=0.334)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            xlsx_dir = folder / "xlsx-output"
            xlsx_dir.mkdir()
            generated = next(output.glob("*.xls"))
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(generated)],
                check=True,
                capture_output=True,
            )
            workbook = load_workbook(next(xlsx_dir.glob("*.xlsx")), data_only=False)
            sheet = workbook["平时成绩"]
            self.assertEqual(sheet["D3"].value, "平时成绩(33.3%)")
            self.assertEqual(sheet["M3"].value, "理论成绩(33.3%)")
            self.assertEqual(sheet["O3"].value, "技能成绩（33.4%）")

    def test_output_validation_rejects_theory_score_mismatch(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-theory-mismatch-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(output.glob("*.xls"))

            xlsx_dir = folder / "tamper-xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(generated)],
                check=True,
                capture_output=True,
            )
            tampered_xlsx = xlsx_dir / f"{generated.stem}.xlsx"
            workbook = load_workbook(tampered_xlsx)
            workbook["平时成绩"]["M5"] = 87.0
            workbook.save(tampered_xlsx)
            tampered_dir = folder / "tampered"
            tampered_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(tampered_dir), str(tampered_xlsx)],
                check=True,
                capture_output=True,
            )
            tampered = output / "tampered.xls"
            shutil.copy2(tampered_dir / f"{generated.stem}.xls", tampered)

            normalized = folder / "normalized.json"
            normalized.write_text(
                json.dumps(
                    {
                        "term": "2025-2026-2",
                        "course": "软件测试实训",
                        "teacher": "张老师",
                        "class_name": "软件技术2401班",
                        "weights": {"regular": 0.6, "theory": 0.4, "skill": 0.0},
                        "students": [
                            {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0},
                            {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 0.0, "total": 91.0},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(normalized),
                "--output-dir",
                str(output),
                "--output-file",
                str(tampered),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("theory score mismatch", result.stderr)

    def test_output_validation_rejects_non_target_sheet_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-non-target-mismatch-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(output.glob("*.xls"))

            xlsx_dir = folder / "tamper-xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(generated)],
                check=True,
                capture_output=True,
            )
            tampered_xlsx = xlsx_dir / f"{generated.stem}.xlsx"
            workbook = load_workbook(tampered_xlsx)
            workbook["Sheet1"]["B1"] = "被篡改的受保护工作表内容"
            workbook.save(tampered_xlsx)
            tampered_dir = folder / "tampered"
            tampered_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(tampered_dir), str(tampered_xlsx)],
                check=True,
                capture_output=True,
            )
            tampered = output / "tampered.xls"
            shutil.copy2(tampered_dir / f"{generated.stem}.xls", tampered)

            normalized = folder / "normalized.json"
            normalized.write_text(
                json.dumps(
                    {
                        "term": "2025-2026-2",
                        "course": "软件测试实训",
                        "teacher": "张老师",
                        "class_name": "软件技术2401班",
                        "weights": {"regular": 0.6, "theory": 0.4, "skill": 0.0},
                        "students": [
                            {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0},
                            {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 0.0, "total": 91.0},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(normalized),
                "--output-dir",
                str(output),
                "--output-file",
                str(tampered),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Protected worksheet changed: Sheet1", result.stderr)

    def test_python_generator_skill_and_leading_zero_id(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-skill-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, skill=True, leading_zero=True)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertTrue(report["checks"]["skill_enabled"])
            self.assertEqual(report["anchor_mode"], "excel_named_range")
            self.assertEqual(report["named_range_variant"], "with_skill")
            self.assertEqual(report["required_named_range_count"], 24)
            self.assertEqual(report["preserved_named_range_count"], 24)
            self.assertEqual(report["checks"]["structure"]["columns"], 17)
            self.assertEqual(report["checks"]["students"][0]["status"], "passed")
            generated = next(output.glob("*.xls"))
            self.assertEqual(report["output_file"], generated.name)
            self.assertEqual(report["files_checked"], 1)
            self.assertNotIn("0012345678", json.dumps(report, ensure_ascii=False))

    def test_python_generator_ignores_unrelated_xls_files(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-unrelated-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            output.mkdir()
            template = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            shutil.copy2(template, output / "unrelated-a.xls")
            shutil.copy2(template, output / "unrelated-b.xls")
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = [path for path in output.glob("*.xls") if path.name.startswith(folder.name)]
            self.assertEqual(len(generated), 1)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["output_file"], generated[0].name)
            self.assertEqual(report["files_checked"], 1)
            self.assertEqual(report["checks"]["file_count"]["actual"], 1)

    def test_rejects_source_total_mismatch_positive_one(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-total-plus-one-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, total_delta=1.0)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Source total mismatch", result.stderr)
            self.assertFalse(output.exists())

    def test_rejects_source_total_mismatch_negative_one(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-total-minus-one-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, total_delta=-1.0)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Source total mismatch", result.stderr)
            self.assertFalse(output.exists())

    def test_rejects_fractional_source_total_before_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-total-fraction-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, total_delta=0.4)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Source total mismatch", result.stderr)
            self.assertIn("received 87.4", result.stderr)
            self.assertFalse(output.exists())

    def test_rejects_fractional_regular_score_before_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-regular-fraction-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, regular_override=89.2)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("students[0].regular must use 0.5-point increments; received 89.2.", result.stderr)
            self.assertFalse(output.exists())

    def test_custom_template_allows_writable_values_but_preserves_formatting(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-values-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            sheet = workbook["平时成绩"]
            for cell, value in {
                "C2": "自定义学期",
                "G2": "自定义课程",
                "L2": "自定义教师",
                "O2": "自定义班级",
                "D3": "平时成绩(99%)",
                "M3": "理论成绩(1%)",
                "O3": "技能成绩（0%）",
            }.items():
                sheet[cell] = value
            workbook.save(xlsx)
            custom_dir = folder / "custom"
            custom_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(custom_dir), str(xlsx)],
                check=True,
                capture_output=True,
            )
            custom = custom_dir / "template.xls"
            custom_manifest = write_manifest_for_modified_template(
                folder / "custom-package",
                GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml",
                custom,
                canonical,
            )
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(custom_manifest),
                "--json",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)

            formatted_xlsx = folder / "formatted-template.xlsx"
            formatted_workbook = load_workbook(xlsx)
            formatted_sheet = formatted_workbook["平时成绩"]
            formatted_font = copy(formatted_sheet["C2"].font)
            formatted_font.sz = (formatted_font.sz or 11) + 1
            formatted_sheet["C2"].font = formatted_font
            formatted_workbook.save(formatted_xlsx)
            formatted_dir = folder / "formatted"
            formatted_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(formatted_dir), str(formatted_xlsx)],
                check=True,
                capture_output=True,
            )
            formatted = formatted_dir / "formatted-template.xls"
            formatted_manifest = write_manifest_for_modified_template(
                folder / "formatted-package",
                GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml",
                formatted,
                canonical,
            )
            formatted_result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(formatted),
                "--manifest",
                str(formatted_manifest),
                "--json",
            )
            self.assertNotEqual(formatted_result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", formatted_result.stdout)

    def test_custom_template_rejects_protected_target_sheet_settings(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-target-protection-guard-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            sheet = workbook["平时成绩"]
            sheet.protection.sheet = True
            sheet.protection.set_password("secret")
            workbook.security.lockStructure = True
            workbook.save(xlsx)
            custom_dir = folder / "custom"
            custom_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(custom_dir), str(xlsx)],
                check=True,
                capture_output=True,
            )
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom_dir / "template.xls"),
                "--manifest",
                str(write_manifest_for_modified_template(
                    folder / "custom-package",
                    GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml",
                    custom_dir / "template.xls",
                    canonical,
                )),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)

    def test_custom_template_rejects_protected_font_family_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-font-family-guard-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            changed_font = copy(workbook["平时成绩"]["A1"].font)
            changed_font.name = "Arial"
            workbook["平时成绩"]["A1"].font = changed_font
            workbook.save(xlsx)
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(xlsx),
                "--manifest",
                str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)

    def _assert_custom_template_rejects_font_fallback(self, font_name: str) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-font-fallback-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            custom = xlsx_dir / "template.xlsx"
            patch_xlsx_cell_font(custom, "C2", font_name)
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)
            self.assertIn("target_cell_formats.C2.font", result.stdout)

    def test_custom_template_rejects_dejavu_fallback_for_simsun(self) -> None:
        self._assert_custom_template_rejects_font_fallback("DejaVu Sans")

    def test_custom_template_rejects_liberation_fallback_for_simsun(self) -> None:
        self._assert_custom_template_rejects_font_fallback("Liberation Serif")

    def test_output_validation_rejects_protected_font_fallback(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-output-font-fallback-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            generated_dir = folder / "generated"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(generated_dir),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(generated_dir.glob("*.xls"))
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(generated)],
                check=True,
                capture_output=True,
            )
            tampered = xlsx_dir / f"{generated.stem}.xlsx"
            patch_xlsx_cell_font(tampered, "C2", "DejaVu Sans")
            patch_xlsx_cell_font(tampered, "O2", "Liberation Serif")
            normalized = folder / "normalized.json"
            normalized.write_text(
                json.dumps(
                    {
                        "term": "2025-2026-2",
                        "course": "软件测试实训",
                        "teacher": "张老师",
                        "class_name": "软件技术2401班",
                        "weights": {"regular": 0.6, "theory": 0.4, "skill": 0.0},
                        "students": [
                            {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0},
                            {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 0.0, "total": 91.0},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(normalized),
                "--output-dir",
                str(xlsx_dir),
                "--output-file",
                str(tampered),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("target sheet formatting mismatch at C2", result.stderr)
            self.assertIn("target sheet formatting mismatch at O2", result.stderr)
            self.assertIn("font", result.stderr)

    def test_output_validation_rejects_xls_roundtrip_font_fallback(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-output-xls-font-fallback-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, skill=True)
            generated_dir = folder / "generated"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(generated_dir),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(generated_dir.glob("*.xls"))
            generated_xlsx = convert_with_soffice(generated, folder / "generated-xlsx", "xlsx")

            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            canonical_xlsx = convert_with_soffice(canonical, folder / "baseline-source", "xlsx")
            controlled_baseline_xlsx = convert_with_soffice(
                convert_with_soffice(canonical_xlsx, folder / "baseline-xls", "xls"),
                folder / "baseline-final",
                "xlsx",
            )
            protected_cells = (
                "A1", "A3", "D3", "D4", "E4", "F4", "G4", "H4", "I4", "J4", "K4",
                "C5", "D5", "E5", "M5", "Q5",
            )
            normalized = folder / "normalized.json"
            normalized.write_text(
                json.dumps(
                    {
                        "term": "2025-2026-2",
                        "course": "软件测试实训",
                        "teacher": "张老师",
                        "class_name": "软件技术2401班",
                        "weights": {"regular": 0.5, "theory": 0.3, "skill": 0.2},
                        "students": [
                            {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 92.0, "total": 88.0},
                            {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 90.0, "total": 91.0},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            for index, font_name in enumerate(("DejaVu Sans", "Liberation Serif")):
                with self.subTest(font=font_name):
                    validation_dir = folder / f"validation-{index}"
                    validation_dir.mkdir()
                    validation_file = validation_dir / "tampered.xls"
                    accepted_result = {}

                    def accepts(candidate: Path, address: str) -> bool:
                        shutil.copy2(candidate, validation_file)
                        candidate_result = run_script(
                            GRADE / "scripts" / "validate_output.py",
                            "--input-json",
                            str(normalized),
                            "--output-dir",
                            str(validation_dir),
                            "--output-file",
                            str(validation_file),
                        )
                        if (
                            candidate_result.returncode != 0
                            and f"target sheet formatting mismatch at {address}" in candidate_result.stderr
                            and "font" in candidate_result.stderr
                        ):
                            accepted_result["result"] = candidate_result
                            return True
                        return False

                    tampered_xls, address = find_roundtrip_font_tamper(
                        generated_xlsx,
                        controlled_baseline_xlsx,
                        font_name,
                        protected_cells,
                        folder / f"tampered-{index}",
                        f"output-{index}",
                        accept=accepts,
                    )
                    result = accepted_result["result"]
                    self.assertNotEqual(result.returncode, 0)
                    self.assertIn(f"target sheet formatting mismatch at {address}", result.stderr)
                    self.assertIn("font", result.stderr)

    def test_custom_template_xls_roundtrip_rejects_protected_and_metadata_font_fallback(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-xls-font-fallback-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            canonical_xlsx = convert_with_soffice(canonical, folder / "baseline-source", "xlsx")
            controlled_baseline_xlsx = convert_with_soffice(
                convert_with_soffice(canonical_xlsx, folder / "baseline-xls", "xls"),
                folder / "baseline-final",
                "xlsx",
            )
            protected_cells = (
                "A1", "A3", "D3", "D4", "E4", "F4", "G4", "H4", "I4", "J4", "K4",
                "C5", "D5", "E5", "M5", "Q5",
            )
            metadata_cells = ("C2", "G2", "L2", "O2")
            cases = (
                ("protected-dejavu", "DejaVu Sans", protected_cells),
                ("protected-liberation", "Liberation Serif", protected_cells),
                ("metadata-dejavu", "DejaVu Sans", metadata_cells),
                ("metadata-liberation", "Liberation Serif", metadata_cells),
            )
            for label, font_name, candidates in cases:
                with self.subTest(case=label):
                    tampered_xls, address = find_roundtrip_font_tamper(
                        canonical_xlsx,
                        controlled_baseline_xlsx,
                        font_name,
                        candidates,
                        folder / label,
                        label,
                    )
                    result = run_script(
                        GRADE / "scripts" / "validate_template.py",
                        "--template",
                        str(tampered_xls),
                        "--manifest",
                        str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                        "--json",
                    )
                    self.assertNotEqual(result.returncode, 0)
                    self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)
                    self.assertIn(address, result.stdout)
                    self.assertIn("font", result.stdout)

    def test_python_generator_skill_xls_roundtrip_passes_font_guard(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-skill-font-baseline-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, skill=True)
            output = folder / "output"
            output.mkdir()
            final_output = output / f"{source.parent.name}-平时成绩记分册.xls"
            final_output.write_bytes(b"old-formal-output")
            (output / "qa-report.json").write_text("{\"status\": \"old\"}\n", encoding="utf-8")
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "passed")
            self.assertTrue(report["checks"]["skill_enabled"])

    def test_canonical_libreoffice_roundtrip_passes_font_guard(self) -> None:
        canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
        result = run_script(
            GRADE / "scripts" / "validate_template.py",
            "--template",
            str(canonical),
            "--json",
        )
        self.assertEqual(result.returncode, 0, result.stderr or result.stdout)

    def test_custom_template_rejects_print_header_footer_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-print-header-guard-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            workbook["平时成绩"].oddHeader.center.text = "自定义打印页眉"
            workbook.save(xlsx)
            custom_dir = folder / "custom"
            custom_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(custom_dir), str(xlsx)],
                check=True,
                capture_output=True,
            )
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom_dir / "template.xls"),
                "--manifest",
                str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)

    def test_custom_template_rejects_non_target_sheet_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-non-target-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            workbook["Sheet1"]["B1"] = "被修改的非目标工作表内容"
            workbook["Sheet3"]["A1"] = "=1+1"
            workbook.save(xlsx)
            custom_dir = folder / "custom"
            custom_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(custom_dir), str(xlsx)],
                check=True,
                capture_output=True,
            )
            custom = custom_dir / "template.xls"
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom),
                "--manifest",
                str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)

    def test_custom_template_rejects_regular_item_header_change(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-custom-header-") as temp_name:
            folder = Path(temp_name)
            canonical = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(canonical)],
                check=True,
                capture_output=True,
            )
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            workbook["平时成绩"]["E4"] = "被修改的常规项目"
            workbook.save(xlsx)
            custom_dir = folder / "custom"
            custom_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(custom_dir), str(xlsx)],
                check=True,
                capture_output=True,
            )
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom_dir / "template.xls"),
                "--manifest",
                str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Custom template changed protected workbook structure or formatting", result.stdout)

    def test_output_validation_rejects_target_sheet_formatting_changes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-format-mismatch-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(output.glob("*.xls"))

            xlsx_dir = folder / "tamper-xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(generated)],
                check=True,
                capture_output=True,
            )
            tampered_xlsx = xlsx_dir / f"{generated.stem}.xlsx"
            workbook = load_workbook(tampered_xlsx)
            tampered_font = copy(workbook["平时成绩"]["C5"].font)
            tampered_font.sz = (tampered_font.sz or 10) + 1
            tampered_font.name = "Arial"
            workbook["平时成绩"]["C5"].font = tampered_font
            workbook["平时成绩"]["E4"] = "被篡改的常规项目"
            workbook["平时成绩"]["F2"] = "被篡改的受保护内容"
            workbook["平时成绩"]["D5"] = 101.5
            workbook["平时成绩"].protection.sheet = True
            workbook["平时成绩"].protection.set_password("secret")
            workbook.security.lockStructure = True
            workbook["平时成绩"].page_margins.left = (workbook["平时成绩"].page_margins.left or 0) + 1
            workbook["平时成绩"].oddHeader.center.text = "被篡改的打印页眉"
            workbook.save(tampered_xlsx)
            tampered_dir = folder / "tampered"
            tampered_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(tampered_dir), str(tampered_xlsx)],
                check=True,
                capture_output=True,
            )
            tampered = output / "tampered.xls"
            shutil.copy2(tampered_dir / f"{generated.stem}.xls", tampered)

            normalized = folder / "normalized.json"
            normalized.write_text(
                json.dumps(
                    {
                        "term": "2025-2026-2",
                        "course": "软件测试实训",
                        "teacher": "张老师",
                        "class_name": "软件技术2401班",
                        "weights": {"regular": 0.6, "theory": 0.4, "skill": 0.0},
                        "students": [
                            {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0},
                            {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 0.0, "total": 91.0},
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            result = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(normalized),
                "--output-dir",
                str(output),
                "--output-file",
                str(tampered),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("target sheet formatting mismatch", result.stderr)
            self.assertIn("target sheet protected value mismatch at E4", result.stderr)
            self.assertIn("target sheet protected value mismatch at F2", result.stderr)
            self.assertIn("target sheet print settings mismatch", result.stderr)
            self.assertIn("target sheet protection settings mismatch", result.stderr)
            self.assertNotIn("101.5", result.stderr)
            self.assertNotIn("101.5", (output / "qa-report.json").read_text(encoding="utf-8"))

    def test_output_validation_rejects_extra_blank_student_rows(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-extra-blank-row-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            legacy_template = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(legacy_template),
                "--output-dir",
                str(output),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            generated = next(output.glob("*.xls"))
            xlsx_dir = folder / "tamper-xlsx"
            xlsx_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(generated)],
                check=True,
                capture_output=True,
            )
            tampered_xlsx = xlsx_dir / f"{generated.stem}.xlsx"
            workbook = load_workbook(tampered_xlsx)
            sheet = workbook["平时成绩"]
            for column in range(1, 16):
                sheet.cell(7, column)._style = copy(sheet.cell(6, column)._style)
            sheet.row_dimensions[7].height = sheet.row_dimensions[6].height
            workbook.save(tampered_xlsx)
            tampered_dir = folder / "tampered"
            tampered_dir.mkdir()
            subprocess.run(
                [soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(tampered_dir), str(tampered_xlsx)],
                check=True,
                capture_output=True,
            )
            tampered = output / "tampered.xls"
            shutil.copy2(tampered_dir / f"{generated.stem}.xls", tampered)
            result = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(ROOT / "tests" / "fixtures" / "gradebook-input.json"),
                "--output-dir",
                str(output),
                "--output-file",
                str(tampered),
                "--template-path",
                str(legacy_template),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Output student row extent mismatch", result.stderr)

    def test_legacy_output_dir_with_multiple_candidates_fails_explicitly(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-multiple-candidates-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            output.mkdir()
            template = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            shutil.copy2(template, output / "candidate-a.xls")
            shutil.copy2(template, output / "candidate-b.xls")
            result = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(ROOT / "tests" / "fixtures" / "gradebook-input.json"),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Expected one generated XLS file, got 2", result.stderr)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["output_file"], "")
            self.assertEqual(report["files_checked"], 0)
            self.assertEqual(report["checks"]["file_count"]["actual"], 2)

    def test_python_compatibility_template_and_skipped_validation_leave_qa_metadata(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-compat-qa-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(GRADE / "assets" / "平时成绩记分册模板.xls"),
                "--output-dir",
                str(output),
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "skipped")
            self.assertTrue(report["custom_template"])
            self.assertEqual(report["engine"], "libreoffice-openpyxl")
            self.assertEqual(report["validation_skipped"], ["output"])
            self.assertIn("Custom template selected", " ".join(report["warnings"]))

    def test_python_generator_expands_beyond_template_rows(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-many-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, count=50)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(len(report["checks"]["students"]), 50)
            self.assertEqual(report["checks"]["named_ranges"]["preserved_named_range_count"], 21)
            self.assertEqual(report["checks"]["named_ranges"]["xlsx"]["locations"]["gb_data_table"]["max_row"], 54)

    def test_named_range_builder_preserves_raw_xls_and_roundtrip_inventory(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-named-builder-") as temp_name:
            folder = Path(temp_name)
            package_a = folder / "package-a"
            package_b = folder / "package-b"
            canonical_hashes = {
                "v1.0": file_sha256(GRADE_V10_TEMPLATE),
                "v1.1": file_sha256(GRADE_V11_TEMPLATE),
            }
            package_a.mkdir()
            (package_a / "marker.txt").write_bytes(b"old-package-marker")
            expected_pdf_pages = None
            for package in (package_a, package_b):
                result = run_script(
                    GRADE / "scripts" / "build_named_range_template.py",
                    "--source",
                    str(GRADE_V10_TEMPLATE),
                    "--output-dir",
                    str(package),
                    "--force",
                )
                self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
                payload = yaml.safe_load(result.stdout)
                self.assertEqual(payload["named_range_count"], 24)
                self.assertGreater(payload["pdf_signature"][0], 0)
                if expected_pdf_pages is None:
                    expected_pdf_pages = payload["pdf_signature"][0]
                else:
                    self.assertEqual(payload["pdf_signature"][0], expected_pdf_pages)
            self.assertFalse((package_a / "marker.txt").exists())

            (package_b / "marker.txt").write_bytes(b"rollback-marker")
            old_package_hashes = package_file_hashes(package_b)
            fault_env = os.environ.copy()
            fault_env["GRADEBOOK_TEST_FAIL_DIRECTORY_SWAP"] = "1"
            failed_swap = run_script(
                GRADE / "scripts" / "build_named_range_template.py",
                "--source",
                str(GRADE_V10_TEMPLATE),
                "--output-dir",
                str(package_b),
                "--force",
                env=fault_env,
            )
            self.assertNotEqual(failed_swap.returncode, 0)
            self.assertIn("directory swap", (failed_swap.stdout + failed_swap.stderr).lower())
            self.assertEqual(package_file_hashes(package_b), old_package_hashes)
            self.assertTrue((package_b / "marker.txt").exists())
            residual = [
                path
                for path in folder.iterdir()
                if path.name.startswith(f".{package_b.name}.")
                and (path.name.endswith(".stage") or path.name.endswith(".backup"))
            ]
            self.assertEqual(residual, [])
            self.assertEqual(file_sha256(GRADE_V10_TEMPLATE), canonical_hashes["v1.0"])
            self.assertEqual(file_sha256(GRADE_V11_TEMPLATE), canonical_hashes["v1.1"])
            sys.modules.pop("package_common", None)
            sys.modules.pop("named_range_utils", None)
            sys.modules.pop("xls_named_range_utils", None)
            sys.path.insert(0, str(GRADE / "scripts"))
            from named_range_utils import compare_named_range_inventories, validate_named_range_inventory
            from xls_named_range_utils import validate_xls_named_range_inventory

            manifest = yaml.safe_load(GRADE_V11_MANIFEST.read_text(encoding="utf-8"))
            inventories = []
            for package in (package_a, package_b):
                raw = validate_xls_named_range_inventory(package / "template.xls", manifest["anchors"], "with_skill")
                roundtrip_xlsx = convert_with_soffice(package / "template.xls", folder / f"{package.name}-xlsx", "xlsx")
                roundtrip = validate_named_range_inventory(
                    load_workbook(roundtrip_xlsx, data_only=False),
                    manifest["anchors"],
                    "with_skill",
                )
                self.assertEqual(raw["errors"], [])
                self.assertEqual(roundtrip["errors"], [])
                self.assertEqual(compare_named_range_inventories(raw, roundtrip), [])
                inventories.append(raw)
            self.assertEqual(inventories[0]["locations"], inventories[1]["locations"])
            self.assertEqual(
                file_sha256(package_a / "template.xls"),
                file_sha256(package_b / "template.xls"),
                "v1.1 builder must reproduce a byte-stable canonical XLS package",
            )

    def test_named_range_builder_preserves_complete_v11_source_and_rejects_partial_or_wrong_sources(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-named-builder-contract-") as temp_name:
            folder = Path(temp_name)
            preserved = folder / "preserved"
            result = run_script(
                GRADE / "scripts" / "build_named_range_template.py",
                "--source",
                str(GRADE_V11_TEMPLATE),
                "--output-dir",
                str(preserved),
                "--force",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            self.assertEqual(file_sha256(preserved / "template.xls"), file_sha256(GRADE_V11_TEMPLATE))

            faults = (
                (
                    "partial",
                    lambda path: patch_xlsx_named_range(path, "gb_term", remove=True),
                    "partial managed-name inventory",
                ),
                (
                    "wrong-destination",
                    lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'平时成绩'!$A$1"),
                    "invalid managed-name inventory",
                ),
                (
                    "regular-seven",
                    lambda path: patch_xlsx_named_range(path, "gb_regular_items", attr_text="'平时成绩'!$D$5:$J$52"),
                    "8 columns",
                ),
                (
                    "duplicate-physical",
                    lambda path: patch_xlsx_named_range(path, "gb_course", attr_text="'平时成绩'!$C$2"),
                    "share one physical destination",
                ),
            )
            for label, mutate, expected_error in faults:
                with self.subTest(builder_fault=label):
                    tampered = tamper_xls_named_range(folder, GRADE_V11_TEMPLATE, mutate, f"builder-{label}")
                    output = folder / f"rejected-{label}"
                    rejected = run_script(
                        GRADE / "scripts" / "build_named_range_template.py",
                        "--source",
                        str(tampered),
                        "--output-dir",
                        str(output),
                        "--force",
                    )
                    self.assertNotEqual(rejected.returncode, 0)
                    self.assertIn(expected_error.lower(), (rejected.stdout + rejected.stderr).lower())
                    self.assertFalse((output / "template.xls").exists())

    def test_named_range_builder_rejects_overlapping_packages_before_soffice(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-builder-path-safety-") as temp_name:
            folder = Path(temp_name)
            gradebook_copy = folder / "course-gradebook-generator"
            shutil.copytree(
                GRADE,
                gradebook_copy,
                ignore=shutil.ignore_patterns("__pycache__", "*.pyc"),
            )
            template_root = gradebook_copy / "assets" / "templates"
            package_root = template_root / "course-gradebook"
            v10_package = package_root / "v1.0.0"
            v11_package = package_root / "v1.1.0"
            v10_template = v10_package / "template.xls"
            v11_template = v11_package / "template.xls"
            builder = gradebook_copy / "scripts" / "build_named_range_template.py"

            custom_parent = folder / "custom-tree"
            custom_package = custom_parent / "input-package"
            shutil.copytree(v10_package, custom_package)
            custom_template = custom_package / "template.xls"
            canonical_hashes = package_file_hashes(template_root)
            custom_hashes = package_file_hashes(custom_parent)
            blocked_soffice_env = os.environ.copy()
            blocked_soffice_env["PATH"] = str(folder / "no-soffice")

            def exchange_artifacts(root: Path) -> list[Path]:
                return sorted(
                    path
                    for path in root.rglob("*")
                    if path.name.endswith((".stage", ".backup"))
                )

            def reject(label: str, source: Path, output: Path, expected: str) -> None:
                result = run_script(
                    builder,
                    "--source",
                    str(source),
                    "--output-dir",
                    str(output),
                    "--force",
                    env=blocked_soffice_env,
                )
                diagnostic = result.stdout + result.stderr
                self.assertNotEqual(result.returncode, 0, label)
                self.assertIn(expected.lower(), diagnostic.lower(), diagnostic)
                self.assertEqual(package_file_hashes(template_root), canonical_hashes, label)
                self.assertEqual(package_file_hashes(custom_parent), custom_hashes, label)
                self.assertEqual(exchange_artifacts(folder), [], label)

            reject(
                "canonical-v10-in-place",
                v10_template,
                v10_package,
                "source template package",
            )
            reject(
                "canonical-common-parent",
                v10_template,
                package_root,
                "source template package",
            )
            reject(
                "canonical-higher-parent",
                v10_template,
                template_root,
                "source template package",
            )
            reject(
                "canonical-v10-with-custom-source",
                custom_template,
                v10_package,
                "canonical v1.0 package",
            )
            reject(
                "canonical-v11-nested",
                custom_template,
                v11_package / "subpackage",
                "canonical v1.1 package",
            )
            reject(
                "custom-source-in-place",
                custom_template,
                custom_package,
                "source template package",
            )
            reject(
                "custom-source-descendant",
                custom_template,
                custom_package / "generated-v1.1",
                "source template package",
            )
            reject(
                "custom-source-ancestor",
                custom_template,
                custom_parent,
                "source template package",
            )

            canonical_v10_hash = file_sha256(v10_template)
            canonical_v11_hash = file_sha256(v11_template)
            canonical_rebuild = run_script(
                builder,
                "--source",
                str(v10_template),
                "--output-dir",
                str(v11_package),
                "--force",
            )
            self.assertEqual(canonical_rebuild.returncode, 0, canonical_rebuild.stderr or canonical_rebuild.stdout)
            canonical_payload = yaml.safe_load(canonical_rebuild.stdout)
            self.assertEqual(canonical_payload["named_range_count"], 24)
            self.assertEqual(file_sha256(v10_template), canonical_v10_hash)
            self.assertEqual(file_sha256(v11_template), canonical_v11_hash)
            rebuilt_manifest = yaml.safe_load((v11_package / "manifest.yaml").read_text(encoding="utf-8"))
            self.assertEqual(rebuilt_manifest["fingerprint"]["sha256"], file_sha256(v11_template))
            self.assertEqual(exchange_artifacts(folder), [])

            sibling_output = folder / "sibling-output-package"
            sibling_build = run_script(
                builder,
                "--source",
                str(custom_template),
                "--output-dir",
                str(sibling_output),
                "--force",
            )
            self.assertEqual(sibling_build.returncode, 0, sibling_build.stderr or sibling_build.stdout)
            sibling_payload = yaml.safe_load(sibling_build.stdout)
            self.assertEqual(sibling_payload["named_range_count"], 24)
            sibling_manifest = yaml.safe_load((sibling_output / "manifest.yaml").read_text(encoding="utf-8"))
            self.assertEqual(sibling_manifest["fingerprint"]["sha256"], file_sha256(sibling_output / "template.xls"))
            self.assertEqual(package_file_hashes(custom_parent), custom_hashes)
            self.assertEqual(exchange_artifacts(folder), [])

    def test_controlled_v11_baseline_handles_roundtrip_normalization_and_protected_tamper(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-controlled-baseline-") as temp_name:
            folder = Path(temp_name)
            soffice = soffice_path()
            self.assertIsNotNone(soffice)
            scripts_path = str(GRADE / "scripts")
            for module_name in (
                "validate_template",
                "named_range_template_baseline",
                "package_common",
                "named_range_contracts",
                "named_range_utils",
                "xls_named_range_utils",
            ):
                sys.modules.pop(module_name, None)
            if scripts_path in sys.path:
                sys.path.remove(scripts_path)
            sys.path.insert(0, scripts_path)
            from named_range_template_baseline import build_controlled_v11_baseline
            from package_common import load_manifest
            from validate_template import _signature_differences, _workbook_signature

            base_manifest = load_manifest(GRADE_V10_MANIFEST)
            canonical_xlsx = convert_with_soffice(
                GRADE_V10_TEMPLATE,
                folder / "direct-v10-xlsx",
                "xlsx",
            )
            committed_xlsx = convert_with_soffice(
                GRADE_V11_TEMPLATE,
                folder / "direct-v11-xlsx",
                "xlsx",
            )
            canonical_signature = _workbook_signature(
                load_workbook(canonical_xlsx, data_only=False),
                base_manifest,
            )
            committed_signature = _workbook_signature(
                load_workbook(committed_xlsx, data_only=False),
                base_manifest,
            )
            for signature in (canonical_signature, committed_signature):
                signature.pop("named_ranges", None)

            simulated_roundtrip_xlsx = folder / "simulated-v10-roundtrip.xlsx"
            simulated_workbook = load_workbook(canonical_xlsx, data_only=False)
            simulated_sheet = simulated_workbook[base_manifest["structure"]["worksheet"]]
            simulated_sheet.column_dimensions["B"].width = round(
                float(committed_signature["column_widths"]["B"] or 0) + 0.2,
                1,
            )
            simulated_workbook.save(simulated_roundtrip_xlsx)
            simulated_signature = _workbook_signature(
                load_workbook(simulated_roundtrip_xlsx, data_only=False),
                base_manifest,
            )
            simulated_signature.pop("named_ranges", None)
            old_model_differences = _signature_differences(
                simulated_signature,
                committed_signature,
            )
            self.assertTrue(
                any(item["path"] == "column_widths.B" for item in old_model_differences),
                old_model_differences,
            )

            controlled = build_controlled_v11_baseline(
                folder / "controlled-v11",
                str(soffice),
            )
            controlled_signature = _workbook_signature(
                controlled.controlled_workbook,
                base_manifest,
            )
            controlled_signature.pop("named_ranges", None)
            self.assertEqual(
                _signature_differences(controlled_signature, committed_signature),
                [],
            )

            tampered_xlsx = folder / "tampered-column-width.xlsx"
            shutil.copy2(committed_xlsx, tampered_xlsx)
            tampered_workbook = load_workbook(tampered_xlsx, data_only=False)
            tampered_sheet = tampered_workbook[base_manifest["structure"]["worksheet"]]
            tampered_sheet.column_dimensions["B"].width = round(
                float(tampered_sheet.column_dimensions["B"].width or 0) + 1.0,
                1,
            )
            tampered_workbook.save(tampered_xlsx)
            tampered_xls = convert_with_soffice(
                tampered_xlsx,
                folder / "tampered-column-width-xls",
                "xls",
            )
            custom_template, custom_manifest = write_gradebook_manifest(
                folder / "custom-tamper",
                GRADE_V11_MANIFEST,
                GRADE_V11_TEMPLATE,
            )
            shutil.copy2(tampered_xls, custom_template)
            manifest_data = yaml.safe_load(custom_manifest.read_text(encoding="utf-8"))
            digest = file_sha256(custom_template)
            manifest_data["fingerprint"]["sha256"] = digest
            manifest_data["fingerprint"]["value"] = digest
            custom_manifest.write_text(
                yaml.safe_dump(manifest_data, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            rejected = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(custom_template),
                "--manifest",
                str(custom_manifest),
                "--json",
            )
            self.assertNotEqual(rejected.returncode, 0, rejected.stdout + rejected.stderr)
            tamper_report = json.loads(rejected.stdout)
            differences = tamper_report["checks"]["protected_signature_differences"]
            self.assertTrue(
                any(str(item.get("path", "")).startswith("column_widths.") for item in differences),
                differences,
            )

    def test_named_range_manifest_contract_is_closed_and_minor_matrix_is_strict(self) -> None:
        sys.modules.pop("package_common", None)
        sys.modules.pop("named_range_contracts", None)
        sys.path.insert(0, str(GRADE / "scripts"))
        from package_common import validate_manifest_contract

        valid = yaml.safe_load(GRADE_V11_MANIFEST.read_text(encoding="utf-8"))
        validate_manifest_contract(valid)

        for generator_version in ("1.1.0", "1.1.1"):
            candidate = deepcopy(valid)
            candidate["generator"]["version"] = generator_version
            validate_manifest_contract(candidate)
        for generator_version in ("1.1.01", "1.01.0", "1.1.０"):
            candidate = deepcopy(valid)
            candidate["generator"]["version"] = generator_version
            with self.assertRaisesRegex(ValueError, "generator.version"):
                validate_manifest_contract(candidate)
        mismatched_generator = deepcopy(valid)
        mismatched_generator["generator"]["version"] = "1.0.9"
        with self.assertRaisesRegex(ValueError, "generator.version"):
            validate_manifest_contract(mismatched_generator)

        regular_count_fault = deepcopy(valid)
        regular_count_fault["validation"]["regular_item_count"] = 7
        with self.assertRaisesRegex(ValueError, "closed generator contract"):
            validate_manifest_contract(regular_count_fault)

        with tempfile.TemporaryDirectory(prefix="grade-package-regular-count-cli-") as temp_name:
            folder = Path(temp_name)
            source_folder = folder / "source"
            source_folder.mkdir()
            source = self.make_source(source_folder, skill=True)
            custom_template, custom_manifest = write_gradebook_manifest(
                folder / "custom",
                GRADE_V11_MANIFEST,
                GRADE_V11_TEMPLATE,
            )
            manifest_data = yaml.safe_load(custom_manifest.read_text(encoding="utf-8"))
            manifest_data["validation"]["regular_item_count"] = 7
            custom_manifest.write_text(
                yaml.safe_dump(manifest_data, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(custom_template),
                "--manifest",
                str(custom_manifest),
                "--output-dir",
                str(folder / "output"),
                "--skip-template-validation",
                "--skip-output-validation",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("closed generator contract", (result.stdout + result.stderr).lower())

        missing_required = deepcopy(valid)
        missing_required["anchors"]["required"] = list(missing_required["anchors"]["required"][:-1])
        with self.assertRaisesRegex(ValueError, "anchors.required"):
            validate_manifest_contract(missing_required)

        legacy_metadata = deepcopy(valid)
        legacy_metadata["structure"]["metadata"] = {"term": "C2"}
        with self.assertRaisesRegex(ValueError, "structure may contain only"):
            validate_manifest_contract(legacy_metadata)

        incomplete_layout = deepcopy(valid)
        del incomplete_layout["layout"]["columns"]["total_score"]
        with self.assertRaisesRegex(ValueError, "layout must map every"):
            validate_manifest_contract(incomplete_layout)

        unsupported_minor = deepcopy(valid)
        unsupported_minor["template"]["version"] = "1.2.0"
        with self.assertRaisesRegex(ValueError, "Unsupported template minor"):
            validate_manifest_contract(unsupported_minor)

        legacy_with_named_metadata = yaml.safe_load(GRADE_V10_MANIFEST.read_text(encoding="utf-8"))
        legacy_with_named_metadata["anchors"] = {"mode": "excel_named_range"}
        with self.assertRaisesRegex(ValueError, "must not declare named-range"):
            validate_manifest_contract(legacy_with_named_metadata)

        semantic_faults = (
            ("swapped-term", lambda manifest: manifest["fields"]["term"].update(name="gb_course")),
            ("swapped-teacher", lambda manifest: manifest["fields"]["teacher"].update(name="gb_term")),
            (
                "swapped-formula-names",
                lambda manifest: manifest["fields"]["formula_columns_with_skill"].update(
                    names=["gb_regular_weighted_col", "gb_theory_weighted_col", "gb_total_score_col", "gb_skill_weighted_col"]
                ),
            ),
            ("wrong-mode", lambda manifest: manifest["fields"]["regular_scores"].update(mode="number")),
            ("extra-field-attribute", lambda manifest: manifest["fields"]["term"].update(extra="reject")),
            ("missing-field-attribute", lambda manifest: manifest["fields"]["term"].pop("max_chars")),
            ("unknown-field-key", lambda manifest: manifest["fields"].update(unknown={"target": "named_range"})),
            (
                "swapped-serial",
                lambda manifest: manifest["layout"]["columns"].update(serial="gb_teacher"),
            ),
            (
                "duplicate-semantic-target",
                lambda manifest: manifest["layout"]["columns"].update(student_id="gb_teacher"),
            ),
            (
                "variant-contract",
                lambda manifest: manifest["anchors"]["variants"]["without_skill"]["forbidden"].clear(),
            ),
        )
        for label, mutate in semantic_faults:
            with self.subTest(contract_fault=label):
                candidate = deepcopy(valid)
                mutate(candidate)
                with self.assertRaisesRegex(ValueError, "(contract|layout|variant|field|names|semantic)"):
                    validate_manifest_contract(candidate)

        version_faults = (
            ("major-from-manifest", "2.0.0", 2),
            ("major-with-declared-support", "2.1.0", 2),
            ("unsupported-minor", "1.2.0", 1),
            ("leading-zero-major", "01.1.0", 1),
            ("leading-zero-minor", "1.01.0", 1),
            ("leading-zero-patch", "1.1.00", 1),
            ("unicode-digits", "１.１.０", 1),
        )
        for label, version, declared_major in version_faults:
            with self.subTest(version_fault=label):
                candidate = deepcopy(valid)
                candidate["template"]["version"] = version
                candidate["generator"]["supported_major"] = declared_major
                with self.assertRaisesRegex(ValueError, "(semantic|Unsupported template)"):
                    validate_manifest_contract(candidate)

        legacy_faults = (
            ("field-name", lambda manifest: manifest["fields"]["term"].update(name="gb_term")),
            ("formula-names", lambda manifest: manifest["fields"]["formula_columns_with_skill"].update(names=["gb_total_score_col"])),
            ("required-named-ranges", lambda manifest: manifest["validation"].update(required_named_ranges=["gb_term"])),
            ("variants", lambda manifest: manifest.update(variants={"with_skill": {}})),
            ("definitions", lambda manifest: manifest.update(definitions={"gb_term": {}})),
            ("layout", lambda manifest: manifest.update(layout={"data_table": "gb_data_table"})),
        )
        legacy = yaml.safe_load(GRADE_V10_MANIFEST.read_text(encoding="utf-8"))
        for label, mutate in legacy_faults:
            with self.subTest(legacy_fault=label):
                candidate = deepcopy(legacy)
                mutate(candidate)
                with self.assertRaisesRegex(ValueError, "(Legacy v1.0|named-range|contract)"):
                    validate_manifest_contract(candidate)

    def test_named_range_variants_and_dynamic_capacity_are_real_outputs(self) -> None:
        cases = ((False, 1, 52, 21), (True, 48, 52, 24), (False, 49, 53, 21), (True, 100, 104, 24))
        for skill, count, expected_last_row, expected_count in cases:
            with self.subTest(skill=skill, count=count), tempfile.TemporaryDirectory(
                prefix=f"grade-package-named-capacity-{count}-"
            ) as temp_name:
                folder = Path(temp_name)
                source = self.make_source(folder, skill=skill, count=count)
                output = folder / "output"
                result = run_script(
                    GRADE / "scripts" / "generate_gradebook.py",
                    "--source",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
                report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
                self.assertEqual(report["status"], "passed")
                self.assertEqual(report["named_range_variant"], "with_skill" if skill else "without_skill")
                self.assertEqual(report["preserved_named_range_count"], expected_count)
                self.assertEqual(
                    report["checks"]["named_ranges"]["xlsx"]["locations"]["gb_data_table"]["max_row"],
                    expected_last_row,
                )
                self.assertEqual(
                    report["checks"]["named_ranges"]["xlsx"]["locations"]["gb_template_row"]["max_row"],
                    5,
                )
                if not skill:
                    self.assertNotIn("gb_skill_score_col", report["checks"]["named_ranges"]["xlsx"]["locations"])
                    self.assertEqual(report["checks"]["structure"]["columns"], 15)

    def test_named_range_template_faults_are_rejected_by_real_validator(self) -> None:
        faults = (
            ("missing", lambda path: patch_xlsx_named_range(path, "gb_term", remove=True), "Missing managed named range"),
            (
                "invalid-name",
                lambda path: patch_xlsx_named_range(path, "gb_term", new_name="gb_非法"),
                ("Invalid managed named range", "Unexpected managed named range"),
            ),
            ("unknown", lambda path: patch_xlsx_named_range(path, "gb_unknown", attr_text="'平时成绩'!$A$1"), "Unexpected managed named range"),
            ("wrong-destination", lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'平时成绩'!$A$1"), "gb_term"),
            ("hidden", lambda path: patch_xlsx_named_range(path, "gb_term", hidden=True), "hidden"),
            ("local-scope", lambda path: patch_xlsx_named_range(path, "gb_term", local_sheet_id=0), "workbook scoped"),
            ("broken", lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="#REF!"), "broken"),
            (
                "external",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'[external.xlsx]Sheet1'!$A$1"),
                "broken",
            ),
            (
                "constant",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="42"),
                "worksheet range",
            ),
            (
                "formula",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="SUM('平时成绩'!$A$1)"),
                "invalid A1",
            ),
            (
                "union",
                lambda path: patch_xlsx_named_range(
                    path,
                    "gb_term",
                    attr_text="'平时成绩'!$A$1,'平时成绩'!$B$1",
                ),
                "broken",
            ),
            (
                "wrong-sheet",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'Sheet1'!$A$1"),
                "one worksheet",
            ),
            (
                "wrong-destination",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'平时成绩'!$Z$1"),
                "gb_term",
            ),
            (
                "merged-cell-not-top-left",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'平时成绩'!$D$2"),
                "top-left",
            ),
            (
                "cell-two-columns",
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'平时成绩'!$C$2:$D$2"),
                "single cell",
            ),
            (
                "regular-seven-columns",
                lambda path: patch_xlsx_named_range(path, "gb_regular_items", attr_text="'平时成绩'!$D$5:$J$52"),
                "8 columns",
            ),
            (
                "duplicate-physical-destination",
                lambda path: patch_xlsx_named_range(path, "gb_course", attr_text="'平时成绩'!$C$2"),
                "share one physical destination",
            ),
            (
                "template-row-mismatch",
                lambda path: patch_xlsx_named_range(path, "gb_template_row", attr_text="'平时成绩'!$A$6:$Q$6"),
                "gb_template_row",
            ),
            (
                "data-row-mismatch",
                lambda path: patch_xlsx_named_range(path, "gb_theory_score_col", attr_text="'平时成绩'!$M$6:$M$52"),
                "same data rows",
            ),
        )
        for label, mutate, expected_error in faults:
            with self.subTest(fault=label), tempfile.TemporaryDirectory(prefix=f"grade-package-name-fault-{label}-") as temp_name:
                folder = Path(temp_name)
                tampered = tamper_xls_named_range(folder, GRADE_V11_TEMPLATE, mutate, label)
                result = run_script(
                    GRADE / "scripts" / "validate_template.py",
                    "--template",
                    str(tampered),
                    "--manifest",
                    str(GRADE_V11_MANIFEST),
                    "--json",
                )
                self.assertNotEqual(result.returncode, 0)
                diagnostic = result.stdout + result.stderr
                expected_errors = (expected_error,) if isinstance(expected_error, str) else expected_error
                self.assertTrue(
                    any(expected.lower() in diagnostic.lower() for expected in expected_errors),
                    diagnostic,
                )

    def test_named_range_runtime_preflight_rejects_matching_fingerprint_faults(self) -> None:
        faults = (
            (
                "regular-seven-columns",
                lambda path: patch_xlsx_named_range(path, "gb_regular_items", attr_text="'平时成绩'!$D$5:$J$52"),
                "8 columns",
            ),
            (
                "template-row-mismatch",
                lambda path: patch_xlsx_named_range(path, "gb_template_row", attr_text="'平时成绩'!$A$6:$Q$6"),
                "gb_template_row",
            ),
            (
                "data-row-mismatch",
                lambda path: patch_xlsx_named_range(path, "gb_theory_score_col", attr_text="'平时成绩'!$M$6:$M$52"),
                "same data rows",
            ),
            (
                "duplicate-physical-destination",
                lambda path: patch_xlsx_named_range(path, "gb_course", attr_text="'平时成绩'!$C$2"),
                "share one physical destination",
            ),
        )
        for label, mutate, expected_error in faults:
            with self.subTest(runtime_fault=label), tempfile.TemporaryDirectory(
                prefix=f"grade-package-runtime-fault-{label}-"
            ) as temp_name:
                folder = Path(temp_name)
                tampered = tamper_xls_named_range(folder, GRADE_V11_TEMPLATE, mutate, f"runtime-{label}")
                custom_template, custom_manifest = write_gradebook_manifest(
                    folder / "custom",
                    GRADE_V11_MANIFEST,
                    GRADE_V11_TEMPLATE,
                )
                shutil.copy2(tampered, custom_template)
                manifest = yaml.safe_load(custom_manifest.read_text(encoding="utf-8"))
                digest = file_sha256(custom_template)
                manifest["fingerprint"]["sha256"] = digest
                manifest["fingerprint"]["value"] = digest
                custom_manifest.write_text(
                    yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
                    encoding="utf-8",
                )
                result = run_script(
                    GRADE / "scripts" / "validate_template.py",
                    "--template",
                    str(custom_template),
                    "--manifest",
                    str(custom_manifest),
                    "--named-range-runtime-preflight",
                    "--json",
                )
                self.assertNotEqual(result.returncode, 0)
                diagnostic = result.stdout + result.stderr
                self.assertIn("runtime preflight", diagnostic.lower())
                self.assertIn(expected_error.lower(), diagnostic.lower())

    def test_com_named_range_runtime_preflight_cannot_be_skipped_by_either_flag(self) -> None:
        if os.name != "nt":
            self.skipTest("Excel COM is only available on Windows")
        probe = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-Command",
                "$excel = New-Object -ComObject Excel.Application; $excel.Quit(); [System.Runtime.Interopservices.Marshal]::ReleaseComObject($excel) | Out-Null",
            ],
            cwd=ROOT,
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            check=False,
        )
        if probe.returncode != 0:
            self.skipTest("Microsoft Excel COM is unavailable on this machine")
        with tempfile.TemporaryDirectory(prefix="grade-package-com-runtime-preflight-") as temp_name:
            folder = Path(temp_name)
            tampered = tamper_xls_named_range(
                folder,
                GRADE_V11_TEMPLATE,
                lambda path: patch_xlsx_named_range(
                    path,
                    "gb_regular_items",
                    attr_text="'平时成绩'!$D$5:$J$52",
                ),
                "com-regular-seven",
            )
            custom_template, custom_manifest = write_gradebook_manifest(
                folder / "custom",
                GRADE_V11_MANIFEST,
                GRADE_V11_TEMPLATE,
            )
            shutil.copy2(tampered, custom_template)
            manifest = yaml.safe_load(custom_manifest.read_text(encoding="utf-8"))
            digest = file_sha256(custom_template)
            manifest["fingerprint"]["sha256"] = digest
            manifest["fingerprint"]["value"] = digest
            custom_manifest.write_text(
                yaml.safe_dump(manifest, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            source = self.make_source(folder, skill=True)
            output = folder / "com-output"
            env = os.environ.copy()
            env["CODEX_PYTHON"] = str(PYTHON)
            result = subprocess.run(
                [
                    "powershell",
                    "-NoProfile",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(GRADE / "scripts" / "generate_gradebook.ps1"),
                    "-SourcePath",
                    str(source),
                    "-OutputDir",
                    str(output),
                    "-TemplatePath",
                    str(custom_template),
                    "-ManifestPath",
                    str(custom_manifest),
                    "-SkipTemplateValidation",
                    "-SkipOutputValidation",
                ],
                cwd=ROOT,
                env=env,
                text=True,
                encoding="utf-8",
                errors="replace",
                capture_output=True,
                check=False,
            )
            self.assertNotEqual(result.returncode, 0)
            diagnostic = result.stdout + result.stderr
            self.assertIn("runtime preflight", diagnostic.lower())
            self.assertIn("8 columns", diagnostic.lower())
            self.assertFalse(output.exists() and list(output.glob("*.xls")))

            offset_template, offset_manifest = write_gradebook_manifest(
                folder / "offset-custom",
                GRADE_V11_MANIFEST,
                GRADE_V11_TEMPLATE,
            )
            offset_data = yaml.safe_load(offset_manifest.read_text(encoding="utf-8"))
            offset_data["validation"]["regular_item_count"] = 9
            offset_manifest.write_text(
                yaml.safe_dump(offset_data, allow_unicode=True, sort_keys=False),
                encoding="utf-8",
            )
            offset_output = folder / "offset-output"
            offset_result = subprocess.run(
                [
                    "powershell",
                    "-NoProfile",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(GRADE / "scripts" / "generate_gradebook.ps1"),
                    "-SourcePath",
                    str(source),
                    "-OutputDir",
                    str(offset_output),
                    "-TemplatePath",
                    str(offset_template),
                    "-ManifestPath",
                    str(offset_manifest),
                    "-SkipTemplateValidation",
                    "-SkipOutputValidation",
                ],
                cwd=ROOT,
                env=env,
                text=True,
                encoding="utf-8",
                errors="replace",
                capture_output=True,
                check=False,
            )
            self.assertNotEqual(offset_result.returncode, 0)
            offset_diagnostic = offset_result.stdout + offset_result.stderr
            self.assertTrue(
                "regular_item_count" in offset_diagnostic.lower()
                or "closed generator contract" in offset_diagnostic.lower()
            )
            self.assertFalse(offset_output.exists() and list(offset_output.glob("*.xls")))

    def test_named_range_output_fault_reports_use_consistent_top_level_fields(self) -> None:
        fields = (
            "missing_named_ranges",
            "duplicate_named_ranges",
            "invalid_named_range_names",
            "unexpected_named_ranges",
            "scope_errors",
            "broken_named_ranges",
            "destination_errors",
            "shape_errors",
            "relationship_errors",
        )
        faults = (
            ("missing", lambda path: patch_xlsx_named_range(path, "gb_term", remove=True), "missing_named_ranges"),
            (
                "broken",
                lambda path: patch_xlsx_named_range(path, "gb_data_table", attr_text="#REF!"),
                ("broken_named_ranges", "missing_named_ranges", "shape_errors", "destination_errors"),
            ),
            ("unknown", lambda path: patch_xlsx_named_range(path, "gb_unknown", attr_text="'平时成绩'!$A$1"), "unexpected_named_ranges"),
            (
                "without-skill-variant-residual",
                lambda path: patch_xlsx_named_range(
                    path,
                    "gb_skill_score_col",
                    attr_text="'平时成绩'!$O$5:$O$52",
                ),
                "unexpected_named_ranges",
            ),
        )
        for label, mutate, field in faults:
            with self.subTest(fault=label), tempfile.TemporaryDirectory(prefix=f"grade-package-output-name-fault-{label}-") as temp_name:
                folder = Path(temp_name)
                source = self.make_source(folder)
                generated = folder / "generated"
                result = run_script(
                    GRADE / "scripts" / "generate_gradebook.py",
                    "--source",
                    str(source),
                    "--output-dir",
                    str(generated),
                    "--skip-output-validation",
                )
                self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
                pristine = next(generated.glob("*.xls"))
                tampered = tamper_xls_named_range(folder, pristine, mutate, label)
                validation_dir = folder / "validation"
                validation_dir.mkdir()
                validation_file = validation_dir / "tampered.xls"
                shutil.copy2(tampered, validation_file)
                normalized = folder / "normalized.json"
                normalized.write_text(json.dumps(gradebook_normalized_input(), ensure_ascii=False), encoding="utf-8")
                validation = run_script(
                    GRADE / "scripts" / "validate_output.py",
                    "--input-json",
                    str(normalized),
                    "--output-dir",
                    str(validation_dir),
                    "--output-file",
                    str(validation_file),
                )
                self.assertNotEqual(validation.returncode, 0)
                self.assertIn("named range", (validation.stdout + validation.stderr).lower())
                report = json.loads((validation_dir / "qa-report.json").read_text(encoding="utf-8"))
                checks = report["checks"]["named_ranges"]
                for name in fields:
                    self.assertEqual(report[name], checks[name], name)
                expected_fields = (field,) if isinstance(field, str) else field
                self.assertTrue(any(report[name] for name in expected_fields))

    def test_named_range_success_report_has_empty_diagnostics_at_both_levels(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-named-success-report-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, skill=True)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            checks = report["checks"]["named_ranges"]
            for field in (
                "missing_named_ranges",
                "duplicate_named_ranges",
                "invalid_named_range_names",
                "unexpected_named_ranges",
                "scope_errors",
                "broken_named_ranges",
                "destination_errors",
                "shape_errors",
                "relationship_errors",
            ):
                self.assertEqual(report[field], [])
                self.assertEqual(checks[field], [])

    def test_fingerprint_is_enforced_before_both_skip_paths(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-fingerprint-skip-") as temp_name:
            folder = Path(temp_name)
            custom_template, custom_manifest = write_gradebook_manifest(
                folder / "custom",
                GRADE_V11_MANIFEST,
                GRADE_V11_TEMPLATE,
            )
            tampered_xlsx = convert_with_soffice(custom_template, folder / "tamper-source", "xlsx")
            patch_xlsx_named_range(tampered_xlsx, "gb_term", attr_text="'平时成绩'!$A$1")
            tampered_xls = convert_with_soffice(tampered_xlsx, folder / "tampered-template", "xls")
            shutil.copy2(tampered_xls, custom_template)
            source = self.make_source(folder, skill=True)
            generate = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(custom_template),
                "--manifest",
                str(custom_manifest),
                "--output-dir",
                str(folder / "generate-output"),
                "--skip-template-validation",
                "--skip-output-validation",
            )
            self.assertNotEqual(generate.returncode, 0)
            self.assertIn("fingerprint mismatch", generate.stderr.lower())

            pristine_output = folder / "pristine-output"
            pristine_result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(pristine_output),
                "--skip-output-validation",
            )
            self.assertEqual(pristine_result.returncode, 0, pristine_result.stderr or pristine_result.stdout)
            normalized = folder / "normalized.json"
            normalized.write_text(json.dumps(gradebook_normalized_input(skill=True), ensure_ascii=False), encoding="utf-8")
            validation_dir = folder / "skip-validation-output"
            validation_dir.mkdir()
            generated = next(pristine_output.glob("*.xls"))
            validation_file = validation_dir / generated.name
            shutil.copy2(generated, validation_file)
            validate = run_script(
                GRADE / "scripts" / "validate_output.py",
                "--input-json",
                str(normalized),
                "--output-dir",
                str(validation_dir),
                "--output-file",
                str(validation_file),
                "--template-path",
                str(custom_template),
                "--manifest",
                str(custom_manifest),
                "--skip-validation",
            )
            self.assertNotEqual(validate.returncode, 0)
            self.assertIn("fingerprint mismatch", validate.stderr.lower())

    def test_matching_custom_v11_package_can_skip_with_explicit_report(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-matching-custom-skip-") as temp_name:
            folder = Path(temp_name)
            custom_template, custom_manifest = write_gradebook_manifest(
                folder / "custom",
                GRADE_V11_MANIFEST,
                GRADE_V11_TEMPLATE,
            )
            source = self.make_source(folder, skill=True)
            output = folder / "output"
            output.mkdir()
            final_output = output / f"{source.parent.name}-平时成绩记分册.xls"
            final_output.write_bytes(b"old-formal-output")
            (output / "qa-report.json").write_text("{\"status\": \"old\"}\n", encoding="utf-8")
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(custom_template),
                "--manifest",
                str(custom_manifest),
                "--output-dir",
                str(output),
                "--skip-template-validation",
                "--skip-output-validation",
            )
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "skipped")
            self.assertEqual(report["validation_skipped"], ["template", "output"])
            generated = next(output.glob("*.xls"))
            self.assertEqual(report["output_dir"], str(output.resolve()))
            self.assertEqual(report["output_file"], generated.name)
            self.assertEqual(report["qa_report"], str((output / "qa-report.json").resolve()))
            self.assertTrue(Path(report["output_dir"]).is_dir())
            self.assertTrue(Path(report["qa_report"]).is_file())
            report_text = (output / "qa-report.json").read_text(encoding="utf-8")
            for temporary_marker in ("gradebook-run-", "gradebook-com-run-", "validation-", "candidate-"):
                self.assertNotIn(temporary_marker, report_text)

    def test_python_output_path_collisions_are_rejected_before_generation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-python-path-safety-") as temp_name:
            folder = Path(temp_name)
            source_folder = folder / "source"
            source_folder.mkdir()
            source = self.make_source(source_folder, skill=True)
            collision_cases = (
                ("source", source, source.parent),
                ("template", GRADE_V11_TEMPLATE, folder / "template-output"),
                ("canonical-v10", GRADE_V10_TEMPLATE, folder / "v10-output"),
                ("compatibility", GRADE / "assets" / "平时成绩记分册模板.xls", folder / "compat-output"),
            )
            protected_hashes = {path: file_sha256(path) for _, path, _ in collision_cases}
            for label, target, output_dir in collision_cases:
                result = run_script(
                    GRADE / "scripts" / "generate_gradebook.py",
                    "--source",
                    str(source),
                    "--output-dir",
                    str(output_dir),
                    "--output-file",
                    str(target),
                    "--skip-template-validation",
                    "--skip-output-validation",
                )
                self.assertNotEqual(result.returncode, 0, label)
                self.assertTrue(
                    any(fragment in (result.stdout + result.stderr).lower() for fragment in ("must not", "inside --output-dir")),
                    result.stdout + result.stderr,
                )
            for path, digest in protected_hashes.items():
                self.assertEqual(file_sha256(path), digest)

    def test_python_custom_package_paths_are_all_protected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-python-custom-paths-") as temp_name:
            folder = Path(temp_name)
            package, actual_template, manifest_path, protected = write_declared_template_gradebook_package(
                folder / "custom"
            )
            source = self.make_source(folder / "source", skill=True)
            package_hashes_before = package_file_hashes(package)
            output_dir = folder / "outputs"

            for label, target in (
                ("actual-template", protected["actual_template"]),
                ("declared-template", protected["declared_template"]),
                ("base-template", protected["base_template"]),
            ):
                result = run_script(
                    GRADE / "scripts" / "generate_gradebook.py",
                    "--source",
                    str(source),
                    "--template",
                    str(actual_template),
                    "--manifest",
                    str(manifest_path),
                    "--output-dir",
                    str(folder),
                    "--output-file",
                    str(target),
                    "--skip-template-validation",
                    "--skip-output-validation",
                )
                self.assertNotEqual(result.returncode, 0, label)
                diagnostic = (result.stdout + result.stderr).lower()
                self.assertIn("must not", diagnostic, label)
                if label == "actual-template":
                    self.assertIn("template file", diagnostic, label)
                else:
                    self.assertIn("declared template package file", diagnostic, label)

            for label, target in protected.items():
                result = run_script(
                    GRADE / "scripts" / "generate_gradebook.py",
                    "--source",
                    str(source),
                    "--template",
                    str(actual_template),
                    "--manifest",
                    str(manifest_path),
                    "--output-dir",
                    str(folder),
                    "--output-file",
                    str(output_dir / f"safe-{label}.xls"),
                    "--qa-report",
                    str(target),
                    "--skip-template-validation",
                    "--skip-output-validation",
                )
                self.assertNotEqual(result.returncode, 0, f"qa-{label}")
                diagnostic = (result.stdout + result.stderr).lower()
                if label == "actual_template":
                    self.assertIn("qa report must not overwrite the template file", diagnostic, label)
                else:
                    self.assertIn("declared template package file", diagnostic, label)

            inside_package = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(actual_template),
                "--manifest",
                str(manifest_path),
                "--output-dir",
                str(package / "generated"),
                "--skip-template-validation",
                "--skip-output-validation",
            )
            self.assertNotEqual(inside_package.returncode, 0)
            self.assertIn(
                "output directory must not be inside the selected template package",
                (inside_package.stdout + inside_package.stderr).lower(),
            )
            self.assertEqual(package_file_hashes(package), package_hashes_before)

    def test_com_output_collisions_and_failure_cleanup_are_real(self) -> None:
        if os.name != "nt":
            self.skipTest("Excel COM is only available on Windows")
        probe = subprocess.run(
            [
                "powershell",
                "-NoProfile",
                "-Command",
                "$excel = New-Object -ComObject Excel.Application; $excel.Quit(); [System.Runtime.Interopservices.Marshal]::ReleaseComObject($excel) | Out-Null",
            ],
            cwd=ROOT,
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            check=False,
        )
        if probe.returncode != 0:
            self.skipTest("Microsoft Excel COM is unavailable on this machine")

        with tempfile.TemporaryDirectory(prefix="grade-package-com-safety-") as temp_name:
            folder = Path(temp_name)
            source_folder = folder / "source"
            source_folder.mkdir()
            source = self.make_source(source_folder, skill=True)
            env = os.environ.copy()
            env["CODEX_PYTHON"] = str(PYTHON)

            def run_com(
                *args: str,
                skip_output_validation: bool = True,
                skip_template_validation: bool = True,
            ) -> subprocess.CompletedProcess[str]:
                command = [
                    "powershell",
                    "-NoProfile",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(GRADE / "scripts" / "generate_gradebook.ps1"),
                    "-SourcePath",
                    str(source),
                ]
                if skip_template_validation:
                    command.append("-SkipTemplateValidation")
                if skip_output_validation:
                    command.append("-SkipOutputValidation")
                command.extend(args)
                return subprocess.run(
                    command,
                    cwd=ROOT,
                    env=env,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    capture_output=True,
                    check=False,
                )

            collision_cases = (
                ("source", source, source.parent),
                ("template", GRADE_V11_TEMPLATE, folder / "template-output"),
                ("canonical-v10", GRADE_V10_TEMPLATE, folder / "v10-output"),
                ("compatibility", GRADE / "assets" / "平时成绩记分册模板.xls", folder / "compat-output"),
            )
            protected_hashes = {path: file_sha256(path) for _, path, _ in collision_cases}
            for label, target, output_dir in collision_cases:
                result = run_com("-OutputDir", str(output_dir), "-OutputFile", str(target))
                self.assertNotEqual(result.returncode, 0, label)
                self.assertIn("must not", (result.stdout + result.stderr).lower(), label)
            for path, digest in protected_hashes.items():
                self.assertEqual(file_sha256(path), digest)

            custom_package, actual_template, custom_manifest, protected = write_declared_template_gradebook_package(
                folder / "custom-package"
            )
            custom_hashes_before = package_file_hashes(custom_package)
            for label, target in protected.items():
                result = run_com(
                    "-TemplatePath",
                    str(actual_template),
                    "-ManifestPath",
                    str(custom_manifest),
                    "-OutputDir",
                    str(folder),
                    "-OutputFile",
                    str(folder / f"com-safe-{label}.xls"),
                    "-QaReportPath",
                    str(target),
                )
                self.assertNotEqual(result.returncode, 0, f"qa-{label}")
                diagnostic = (result.stdout + result.stderr).lower()
                if label == "actual_template":
                    self.assertIn("qa report must not overwrite the template file", diagnostic, label)
                else:
                    self.assertIn("declared template package file", diagnostic, label)
            inside_package = run_com(
                "-TemplatePath",
                str(actual_template),
                "-ManifestPath",
                str(custom_manifest),
                "-OutputDir",
                str(custom_package / "generated"),
            )
            self.assertNotEqual(inside_package.returncode, 0)
            self.assertIn(
                "output directory must not be inside the selected template package",
                (inside_package.stdout + inside_package.stderr).lower(),
            )
            self.assertEqual(package_file_hashes(custom_package), custom_hashes_before)

            passed_output = folder / "com-passed-output"
            passed_output.mkdir()
            passed_final = passed_output / f"{source.parent.name}-平时成绩记分册.xls"
            passed_final.write_bytes(b"old-formal-output")
            passed_qa = passed_output / "qa-report.json"
            passed_qa.write_text("{\"status\": \"old\"}\n", encoding="utf-8")
            passed = run_com(
                "-OutputDir",
                str(passed_output),
                "-OutputFile",
                str(passed_final),
                "-QaReportPath",
                str(passed_qa),
                skip_output_validation=False,
                skip_template_validation=False,
            )
            self.assertEqual(passed.returncode, 0, passed.stderr or passed.stdout)
            passed_report = json.loads(passed_qa.read_text(encoding="utf-8"))
            self.assertEqual(passed_report["status"], "passed")
            self.assertEqual(passed_report["output_dir"], str(passed_output.resolve()))
            self.assertEqual(passed_report["output_file"], passed_final.name)
            self.assertEqual(passed_report["qa_report"], str(passed_qa.resolve()))
            passed_text = passed_qa.read_text(encoding="utf-8")
            for temporary_marker in ("gradebook-com-run-", "validation-", "candidate-"):
                self.assertNotIn(temporary_marker, passed_text)

            skipped_output = folder / "com-skipped-output"
            skipped_output.mkdir()
            skipped_final = skipped_output / f"{source.parent.name}-平时成绩记分册.xls"
            skipped_final.write_bytes(b"old-formal-output")
            skipped_qa = skipped_output / "qa-report.json"
            skipped_qa.write_text("{\"status\": \"old\"}\n", encoding="utf-8")
            skipped = run_com(
                "-OutputDir",
                str(skipped_output),
                "-OutputFile",
                str(skipped_final),
                "-QaReportPath",
                str(skipped_qa),
            )
            self.assertEqual(skipped.returncode, 0, skipped.stderr or skipped.stdout)
            skipped_report = json.loads(skipped_qa.read_text(encoding="utf-8"))
            self.assertEqual(skipped_report["status"], "skipped")
            self.assertEqual(skipped_report["output_dir"], str(skipped_output.resolve()))
            self.assertEqual(skipped_report["output_file"], skipped_final.name)
            self.assertEqual(skipped_report["qa_report"], str(skipped_qa.resolve()))

            failure_source_folder = folder / "failure-source"
            failure_source_folder.mkdir()
            failure_source = self.make_source(failure_source_folder, skill=True, total_delta=1.0)
            output = folder / "failure-output"
            output.mkdir()
            old_output = output / "failure-source-平时成绩记分册.xls"
            unrelated_a = output / "unrelated-a.xls"
            unrelated_b = output / "unrelated-b.xls"
            old_output.write_bytes(b"old-formal-output")
            shutil.copy2(GRADE_V10_TEMPLATE, unrelated_a)
            shutil.copy2(GRADE_V10_TEMPLATE, unrelated_b)
            old_hash = file_sha256(old_output)
            unrelated_hashes = {path: file_sha256(path) for path in (unrelated_a, unrelated_b)}
            result = subprocess.run(
                [
                    "powershell",
                    "-NoProfile",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-File",
                    str(GRADE / "scripts" / "generate_gradebook.ps1"),
                    "-SourcePath",
                    str(failure_source),
                    "-OutputDir",
                    str(output),
                    "-SkipTemplateValidation",
                    "-SkipOutputValidation",
                ],
                cwd=ROOT,
                env=env,
                text=True,
                encoding="utf-8",
                errors="replace",
                capture_output=True,
                check=False,
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertEqual(file_sha256(old_output), old_hash)
            for path, digest in unrelated_hashes.items():
                self.assertEqual(file_sha256(path), digest)
            self.assertEqual(list(output.glob(".*.tmp")), [])

    def test_regular_score_boundaries_are_exact(self) -> None:
        sys.modules.pop("generate_gradebook", None)
        sys.path.insert(0, str(GRADE / "scripts"))
        from generate_gradebook import generate_regular_scores

        for target in (0.0, 86.5, 100.0):
            scores = generate_regular_scores(target, f"boundary-{target}")
            self.assertEqual(len(scores), 8)
            self.assertAlmostEqual(sum(scores) / len(scores), target, places=7)
            self.assertTrue(all(0 <= score <= 100 for score in scores))
            self.assertTrue(all(abs(score * 2 - round(score * 2)) < 1e-7 for score in scores))
        self.assertEqual(len(generate_regular_scores(86.5, "four-items", item_count=4)), 4)

    def test_structure_breaking_xls_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-broken-") as temp_name:
            folder = Path(temp_name)
            source = GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls"
            xlsx_dir = folder / "xlsx"
            xlsx_dir.mkdir()
            subprocess.run([soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(source)], check=True, capture_output=True)
            xlsx = xlsx_dir / "template.xlsx"
            workbook = load_workbook(xlsx)
            workbook["平时成绩"]["A3"] = "坏标签"
            workbook.save(xlsx)
            broken_dir = folder / "broken"
            broken_dir.mkdir()
            subprocess.run([soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(broken_dir), str(xlsx)], check=True, capture_output=True)
            broken = broken_dir / "template.xls"
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(broken),
                "--manifest",
                str(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml"),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("required header", result.stdout)

    def test_incompatible_manifest_major_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-manifest-") as temp_name:
            folder = Path(temp_name)
            template = folder / "template.xls"
            shutil.copy2(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls", template)
            manifest = folder / "manifest.yaml"
            manifest_text = (GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml").read_text(encoding="utf-8")
            manifest.write_text(manifest_text.replace("version: 1.0.0", "version: 2.0.0", 1), encoding="utf-8")
            result = run_script(GRADE / "scripts" / "validate_template.py", "--template", str(template), "--manifest", str(manifest), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Unsupported template major version", result.stdout)

    def test_manifest_loading_failures_are_clear(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-manifest-errors-") as temp_name:
            folder = Path(temp_name)
            template = folder / "template.xls"
            shutil.copy2(GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "template.xls", template)
            missing = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(folder / "missing.yaml"),
                "--json",
            )
            self.assertNotEqual(missing.returncode, 0)
            self.assertIn("No such file", missing.stdout)

            malformed = folder / "malformed.yaml"
            malformed.write_text("template: [", encoding="utf-8")
            result = run_script(GRADE / "scripts" / "validate_template.py", "--manifest", str(malformed), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("errors", result.stdout)

            manifest_text = (GRADE / "assets" / "templates" / "course-gradebook" / "v1.0.0" / "manifest.yaml").read_text(encoding="utf-8")
            missing_version = folder / "missing-version.yaml"
            missing_version.write_text(manifest_text.replace("  version: 1.0.0\n", "", 1), encoding="utf-8")
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(missing_version),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("semantic template version", result.stdout)

            missing_template = folder / "missing-template.yaml"
            missing_template.write_text(manifest_text.replace("file: template.xls", "file: missing.xls", 1), encoding="utf-8")
            result = run_script(GRADE / "scripts" / "validate_template.py", "--manifest", str(missing_template), "--json")
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Template not found", result.stdout)

            missing_canonical = folder / "missing-canonical.yaml"
            missing_canonical.write_text(manifest_text.replace("file: template.xls", "file: missing.xls", 1), encoding="utf-8")
            result = run_script(
                GRADE / "scripts" / "validate_template.py",
                "--template",
                str(template),
                "--manifest",
                str(missing_canonical),
                "--json",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("Canonical template not found", result.stdout)

    def test_formula_error_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-formula-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            result = run_script(GRADE / "scripts" / "generate_gradebook.py", "--source", str(source), "--output-dir", str(output))
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            output_file = next(output.glob("*.xls"))
            xlsx_dir = folder / "broken-xlsx"
            xlsx_dir.mkdir()
            subprocess.run([soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(output_file)], check=True, capture_output=True)
            xlsx = xlsx_dir / f"{output_file.stem}.xlsx"
            workbook = load_workbook(xlsx)
            workbook["平时成绩"]["L5"] = "=#REF!"
            workbook.save(xlsx)
            broken_dir = folder / "broken-xls"
            broken_dir.mkdir()
            subprocess.run([soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(broken_dir), str(xlsx)], check=True, capture_output=True)
            broken = broken_dir / f"{output_file.stem}.xls"
            output_file.unlink()
            broken.replace(output_file)
            input_json = folder / "input.json"
            input_json.write_text(json.dumps({
                "term": "2025-2026-2",
                "course": "软件测试实训",
                "teacher": "张老师",
                "class_name": "软件技术2401班",
                "weights": {"regular": 0.6, "theory": 0.4, "skill": 0.0},
                "students": [
                    {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0},
                    {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 0.0, "total": 91.0},
                ],
            }, ensure_ascii=False), encoding="utf-8")
            result = run_script(GRADE / "scripts" / "validate_output.py", "--input-json", str(input_json), "--output-dir", str(output))
            self.assertNotEqual(result.returncode, 0)
            self.assertIn("formula error", result.stderr)

    def test_student_count_mismatch_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-count-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder)
            output = folder / "output"
            result = run_script(GRADE / "scripts" / "generate_gradebook.py", "--source", str(source), "--output-dir", str(output))
            self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
            output_file = next(output.glob("*.xls"))
            xlsx_dir = folder / "broken-xlsx"
            xlsx_dir.mkdir()
            subprocess.run([soffice_path(), "--headless", "--convert-to", "xlsx", "--outdir", str(xlsx_dir), str(output_file)], check=True, capture_output=True)
            xlsx = xlsx_dir / f"{output_file.stem}.xlsx"
            workbook = load_workbook(xlsx)
            workbook["平时成绩"].delete_rows(6, 1)
            workbook.save(xlsx)
            broken_dir = folder / "broken-xls"
            broken_dir.mkdir()
            subprocess.run([soffice_path(), "--headless", "--convert-to", "xls", "--outdir", str(broken_dir), str(xlsx)], check=True, capture_output=True)
            broken = broken_dir / f"{output_file.stem}.xls"
            output_file.unlink()
            broken.replace(output_file)
            input_json = folder / "input.json"
            input_json.write_text(json.dumps({
                "term": "2025-2026-2",
                "course": "软件测试实训",
                "teacher": "张老师",
                "class_name": "软件技术2401班",
                "weights": {"regular": 0.6, "theory": 0.4, "skill": 0.0},
                "students": [
                    {"id": "240101001", "name": "学生1", "regular": 86.5, "theory": 88.0, "skill": 0.0, "total": 87.0},
                    {"id": "240101002", "name": "学生2", "regular": 91.0, "theory": 90.0, "skill": 0.0, "total": 91.0},
                ],
            }, ensure_ascii=False), encoding="utf-8")
            result = run_script(GRADE / "scripts" / "validate_output.py", "--input-json", str(input_json), "--output-dir", str(output))
            self.assertNotEqual(result.returncode, 0)

    def test_python_double_skip_runs_raw_preflight_before_candidate_creation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-python-raw-skip-") as temp_name:
            folder = Path(temp_name)
            tampered = tamper_xls_named_range(
                folder,
                GRADE_V11_TEMPLATE,
                lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="'平时成绩'!$C$1"),
                "raw-location",
            )
            custom_template, custom_manifest = write_gradebook_manifest(
                folder / "custom",
                GRADE_V11_MANIFEST,
                tampered,
            )
            source = self.make_source(folder, skill=True)
            output = folder / "output"
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--template",
                str(custom_template),
                "--manifest",
                str(custom_manifest),
                "--output-dir",
                str(output),
                "--skip-template-validation",
                "--skip-output-validation",
            )
            self.assertNotEqual(result.returncode, 0)
            diagnostic = result.stdout + result.stderr
            self.assertIn("runtime preflight", diagnostic.lower())
            self.assertIn("gb_term", diagnostic)
            self.assertFalse(output.exists() and list(output.glob("*.xls")))

    def test_skip_validation_requires_real_output_and_raw_named_ranges(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-skip-truth-") as temp_name:
            folder = Path(temp_name)
            input_json = folder / "input.json"
            input_json.write_text(
                json.dumps(gradebook_normalized_input(skill=True), ensure_ascii=False),
                encoding="utf-8",
            )

            def run_skip(output_dir: Path, output_file: Path) -> subprocess.CompletedProcess[str]:
                return run_script(
                    GRADE / "scripts" / "validate_output.py",
                    "--input-json",
                    str(input_json),
                    "--output-dir",
                    str(output_dir),
                    "--output-file",
                    str(output_file),
                    "--manifest",
                    str(GRADE_V11_MANIFEST),
                    "--template-path",
                    str(GRADE_V11_TEMPLATE),
                    "--skip-validation",
                )

            missing_dir = folder / "missing"
            missing = run_skip(missing_dir, missing_dir / "missing.xls")
            self.assertNotEqual(missing.returncode, 0)
            self.assertTrue((missing_dir / "qa-report.json").exists())
            self.assertIn("not found", (missing_dir / "qa-report.json").read_text(encoding="utf-8"))

            empty_dir = folder / "empty"
            empty_dir.mkdir()
            empty_file = empty_dir / "empty.xls"
            empty_file.write_bytes(b"")
            empty = run_skip(empty_dir, empty_file)
            self.assertNotEqual(empty.returncode, 0)
            self.assertIn("empty", (empty_dir / "qa-report.json").read_text(encoding="utf-8").lower())

            directory_dir = folder / "directory"
            directory_dir.mkdir()
            directory_file = directory_dir / "directory.xls"
            directory_file.mkdir()
            directory = run_skip(directory_dir, directory_file)
            self.assertNotEqual(directory.returncode, 0)
            self.assertIn("not a file", (directory_dir / "qa-report.json").read_text(encoding="utf-8"))

            outside_dir = folder / "outside-check"
            outside_dir.mkdir()
            outside = run_skip(outside_dir, folder / "outside.xls")
            self.assertNotEqual(outside.returncode, 0)
            self.assertIn("inside --output-dir", outside.stderr)

            (folder / "pristine-source").mkdir()
            source = self.make_source(folder / "pristine-source", skill=True)
            generated_dir = folder / "pristine"
            generated = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(generated_dir),
                "--skip-output-validation",
            )
            self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)
            pristine = next(generated_dir.glob("*.xls"))
            for label, mutate, expected in (
                ("missing-names", lambda path: patch_xlsx_named_range(path, "gb_term", remove=True), ("missing managed named range",)),
                ("broken-name", lambda path: patch_xlsx_named_range(path, "gb_term", attr_text="#REF!"), ("broken", "absolute worksheet rectangle")),
            ):
                candidate_dir = folder / label
                candidate_dir.mkdir()
                candidate = tamper_xls_named_range(candidate_dir, pristine, mutate, label)
                validation = run_skip(candidate_dir, candidate)
                self.assertNotEqual(validation.returncode, 0)
                report = json.loads((candidate_dir / "qa-report.json").read_text(encoding="utf-8"))
                self.assertEqual(report["status"], "failed")
                diagnostic = (validation.stdout + validation.stderr).lower()
                self.assertTrue(any(fragment in diagnostic for fragment in expected), diagnostic)

    def test_builder_uses_fixed_canonical_v10_baseline_for_layout_faults(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-builder-baseline-") as temp_name:
            folder = Path(temp_name)

            def tamper_xls(source: Path, label: str, mutate) -> Path:
                xlsx = convert_with_soffice(source, folder / f"{label}-xlsx", "xlsx")
                workbook = load_workbook(xlsx)
                mutate(workbook)
                workbook.save(xlsx)
                return convert_with_soffice(xlsx, folder / f"{label}-xls", "xls")

            from openpyxl.styles import Side

            def change_font(workbook) -> None:
                cell = workbook["平时成绩"]["C2"]
                font = copy(cell.font)
                font.name = "Arial"
                cell.font = font

            def change_border(workbook) -> None:
                cell = workbook["平时成绩"]["E4"]
                border = copy(cell.border)
                border.bottom = Side(style="thick")
                cell.border = border

            faults = (
                ("title", lambda workbook: setattr(workbook["平时成绩"]["A1"], "value", "被篡改标题")),
                ("font", change_font),
                ("border", change_border),
                (
                    "width",
                    lambda workbook: setattr(
                        workbook["平时成绩"].column_dimensions["A"],
                        "width",
                        workbook["平时成绩"].column_dimensions["A"].width + 4,
                    ),
                ),
                ("orientation", lambda workbook: setattr(workbook["平时成绩"].page_setup, "orientation", "portrait")),
                ("non-target", lambda workbook: setattr(workbook["Sheet1"]["A1"], "value", "被篡改非目标工作表")),
                ("protection", lambda workbook: setattr(workbook["平时成绩"].protection, "sheet", True)),
            )
            for label, mutate in faults:
                tampered = tamper_xls(GRADE_V10_TEMPLATE, label, mutate)
                output = folder / f"rejected-{label}"
                result = run_script(
                    GRADE / "scripts" / "build_named_range_template.py",
                    "--source",
                    str(tampered),
                    "--output-dir",
                    str(output),
                    "--force",
                )
                self.assertNotEqual(result.returncode, 0, label)
                self.assertIn("canonical v1.0", (result.stdout + result.stderr).lower(), label)
                self.assertFalse((output / "template.xls").exists(), label)

            v11_tampered = tamper_xls(GRADE_V11_TEMPLATE, "v11-title", lambda workbook: setattr(workbook["平时成绩"]["A1"], "value", "被篡改标题"))
            v11_output = folder / "rejected-v11-title"
            v11_result = run_script(
                GRADE / "scripts" / "build_named_range_template.py",
                "--source",
                str(v11_tampered),
                "--output-dir",
                str(v11_output),
                "--force",
            )
            self.assertNotEqual(v11_result.returncode, 0)
            self.assertIn("canonical v1.0", (v11_result.stdout + v11_result.stderr).lower())

    def test_named_range_capacity_faults_are_rejected_and_exact_100_passes(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-capacity-faults-") as temp_name:
            folder = Path(temp_name)

            def input_file(path: Path, skill: bool, count: int) -> Path:
                path.write_text(json.dumps(gradebook_normalized_input(skill=skill, count=count), ensure_ascii=False), encoding="utf-8")
                return path

            def ranges_for(variant: str, last_row: int) -> dict[str, str]:
                if variant == "with_skill":
                    return {
                        "gb_data_table": f"'平时成绩'!$A$5:$Q${last_row}",
                        "gb_serial_col": f"'平时成绩'!$A$5:$A${last_row}",
                        "gb_student_id_col": f"'平时成绩'!$B$5:$B${last_row}",
                        "gb_student_name_col": f"'平时成绩'!$C$5:$C${last_row}",
                        "gb_regular_items": f"'平时成绩'!$D$5:$K${last_row}",
                        "gb_regular_weighted_col": f"'平时成绩'!$L$5:$L${last_row}",
                        "gb_theory_score_col": f"'平时成绩'!$M$5:$M${last_row}",
                        "gb_theory_weighted_col": f"'平时成绩'!$N$5:$N${last_row}",
                        "gb_skill_score_col": f"'平时成绩'!$O$5:$O${last_row}",
                        "gb_skill_weighted_col": f"'平时成绩'!$P$5:$P${last_row}",
                        "gb_total_score_col": f"'平时成绩'!$Q$5:$Q${last_row}",
                    }
                return {
                    "gb_data_table": f"'平时成绩'!$A$5:$O${last_row}",
                    "gb_serial_col": f"'平时成绩'!$A$5:$A${last_row}",
                    "gb_student_id_col": f"'平时成绩'!$B$5:$B${last_row}",
                    "gb_student_name_col": f"'平时成绩'!$C$5:$C${last_row}",
                    "gb_regular_items": f"'平时成绩'!$D$5:$K${last_row}",
                    "gb_regular_weighted_col": f"'平时成绩'!$L$5:$L${last_row}",
                    "gb_theory_score_col": f"'平时成绩'!$M$5:$M${last_row}",
                    "gb_theory_weighted_col": f"'平时成绩'!$N$5:$N${last_row}",
                    "gb_total_score_col": f"'平时成绩'!$O$5:$O${last_row}",
                }

            def validate_fault(label: str, source_output: Path, data: Path, mutate) -> None:
                candidate_dir = folder / label
                candidate_dir.mkdir()
                candidate = tamper_xls_named_range(candidate_dir, source_output, mutate, label)
                result = run_script(
                    GRADE / "scripts" / "validate_output.py",
                    "--input-json",
                    str(data),
                    "--output-dir",
                    str(candidate_dir),
                    "--output-file",
                    str(candidate),
                    "--manifest",
                    str(GRADE_V11_MANIFEST),
                    "--template-path",
                    str(GRADE_V11_TEMPLATE),
                )
                self.assertNotEqual(result.returncode, 0, label)
                self.assertIn("named range", (result.stdout + result.stderr).lower(), label)

            source2 = self.make_source(folder / "source2", skill=True, count=2)
            output2 = folder / "output2"
            generated2 = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source2),
                "--output-dir",
                str(output2),
                "--skip-output-validation",
            )
            self.assertEqual(generated2.returncode, 0, generated2.stderr or generated2.stdout)
            output2_file = next(output2.glob("*.xls"))
            input2 = input_file(folder / "input2.json", True, 2)
            expanded2 = ranges_for("with_skill", 100)
            validate_fault(
                "expand-2-to-100",
                output2_file,
                input2,
                lambda path: [patch_xlsx_named_range(path, name, attr_text=address) for name, address in expanded2.items()],
            )
            validate_fault(
                "single-column-extra-row",
                output2_file,
                input2,
                lambda path: patch_xlsx_named_range(path, "gb_theory_score_col", attr_text="'平时成绩'!$M$5:$M$53"),
            )

            source49 = self.make_source(folder / "source49", skill=False, count=49)
            output49 = folder / "output49"
            generated49 = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source49),
                "--output-dir",
                str(output49),
                "--skip-output-validation",
            )
            self.assertEqual(generated49.returncode, 0, generated49.stderr or generated49.stdout)
            output49_file = next(output49.glob("*.xls"))
            input49 = input_file(folder / "input49.json", False, 49)
            expanded49 = ranges_for("without_skill", 80)
            validate_fault(
                "expand-49-to-80",
                output49_file,
                input49,
                lambda path: [patch_xlsx_named_range(path, name, attr_text=address) for name, address in expanded49.items()],
            )
            validate_fault(
                "table-one-row-short",
                output49_file,
                input49,
                lambda path: patch_xlsx_named_range(path, "gb_data_table", attr_text="'平时成绩'!$A$5:$O$52"),
            )

            for skill in (False, True):
                source100 = self.make_source(folder / ("source100-skill" if skill else "source100-no-skill"), skill=skill, count=100)
                output100 = folder / ("output100-skill" if skill else "output100-no-skill")
                generated100 = run_script(
                    GRADE / "scripts" / "generate_gradebook.py",
                    "--source",
                    str(source100),
                    "--output-dir",
                    str(output100),
                )
                self.assertEqual(generated100.returncode, 0, generated100.stderr or generated100.stdout)
                report = json.loads((output100 / "qa-report.json").read_text(encoding="utf-8"))
                self.assertEqual(report["status"], "passed")
                self.assertEqual(report["checks"]["named_ranges"]["xlsx"]["locations"]["gb_data_table"]["max_row"], 104)

    def test_python_failure_preserves_existing_output_and_unrelated_xls(self) -> None:
        with tempfile.TemporaryDirectory(prefix="grade-package-python-atomic-") as temp_name:
            folder = Path(temp_name)
            source = self.make_source(folder, total_delta=1.0)
            output = folder / "output"
            output.mkdir()
            old_output = output / f"{folder.name}-平时成绩记分册.xls"
            old_output.write_bytes(b"old-formal-output")
            unrelated = output / "unrelated.xls"
            unrelated.write_bytes(b"unrelated")
            old_hash = file_sha256(old_output)
            unrelated_hash = file_sha256(unrelated)
            result = run_script(
                GRADE / "scripts" / "generate_gradebook.py",
                "--source",
                str(source),
                "--output-dir",
                str(output),
                "--skip-template-validation",
                "--skip-output-validation",
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertEqual(file_sha256(old_output), old_hash)
            self.assertEqual(file_sha256(unrelated), unrelated_hash)
            self.assertEqual(list(output.glob("*.tmp")), [])


if __name__ == "__main__":
    unittest.main()
