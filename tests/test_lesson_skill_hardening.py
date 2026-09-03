from __future__ import annotations

from contextlib import contextmanager, redirect_stderr, redirect_stdout
import importlib
import importlib.util
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import Mock, patch


ROOT = Path(__file__).resolve().parents[1]
LESSON = ROOT / "教案生成器" / "lesson-plan-docx-generator"
SCRIPTS = LESSON / "scripts"
_MISSING = object()
_LESSON_MODULE_NAMES = (
    "path_safety",
    "semantic_bookmarks",
    "content_contract",
    "bookmark_utils",
    "content_quality",
    "package_common",
    "validate_template",
    "render_qa",
    "validate_output",
    "generate_lesson_plans",
    "build_template_patch",
    "install",
    "install_adapters",
    "record_visual_inspection",
    "validate_visual_inspection",
)


@contextmanager
def isolated_lesson_imports():
    """Load legacy script modules without publishing process-global imports."""

    original_path = sys.path[:]
    saved_modules = {name: sys.modules.get(name, _MISSING) for name in _LESSON_MODULE_NAMES}
    try:
        for name in _LESSON_MODULE_NAMES:
            sys.modules.pop(name, None)
        sys.path.insert(0, str(SCRIPTS))
        modules = {
            name: importlib.import_module(name)
            for name in _LESSON_MODULE_NAMES
        }
        yield modules
    finally:
        sys.path[:] = original_path
        for name, previous in saved_modules.items():
            if previous is _MISSING:
                sys.modules.pop(name, None)
            else:
                sys.modules[name] = previous


with isolated_lesson_imports() as _lesson_modules:
    build_template_patch = _lesson_modules["build_template_patch"]
    content_quality = _lesson_modules["content_quality"]
    lesson_install = _lesson_modules["install"]
    install_adapters = _lesson_modules["install_adapters"]
    render_qa = _lesson_modules["render_qa"]
    validate_output = _lesson_modules["validate_output"]
    lesson_generator = _lesson_modules["generate_lesson_plans"]
    CHECK_NAMES = _lesson_modules["record_visual_inspection"].CHECK_NAMES
    write_visual_inspection_evidence = _lesson_modules["record_visual_inspection"].write_visual_inspection_evidence
    validate_visual_inspection = _lesson_modules["validate_visual_inspection"].validate_visual_inspection


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


lesson_acceptance = load_module(ROOT / "tests" / "lesson_acceptance.py", "lesson_acceptance_contracts")


class LessonSkillHardeningTests(unittest.TestCase):
    def test_lesson_intake_and_reference_instruction_contract(self) -> None:
        skill = (LESSON / "SKILL.md").read_text(encoding="utf-8")
        prompt = (LESSON / "通用提示词.md").read_text(encoding="utf-8")
        agents = (LESSON / "AGENTS.md").read_text(encoding="utf-8")
        openai = (LESSON / "agents" / "openai.yaml").read_text(encoding="utf-8")
        canonical = "\n".join((skill, prompt, agents, openai))

        for field in ("course_name", "major", "audience", "total_hours"):
            self.assertIn(field, canonical)
        self.assertRegex(canonical, r"default_hours\s*=\s*2")
        self.assertIn("教材", canonical)
        self.assertIn("recommended, not required", canonical)
        self.assertRegex(canonical, r"一次(?:性)?(?:.*)确认")
        self.assertIn("course_reference_pool", canonical)
        self.assertIn("reference_reusable", canonical)
        self.assertIn("references", canonical)
        self.assertIn("resources", canonical)

        for forbidden_second_prompt in ("outline", "输出目录", "是否开始生成 DOCX"):
            self.assertRegex(
                canonical,
                rf"(?:不得再次询问|不再询问)[^\n]*{re.escape(forbidden_second_prompt)}",
            )
        self.assertIn("同一课内部", canonical)
        self.assertIn("禁止为了降低课程重复率编造", canonical)

        self.assertIn("INTAKE_PENDING", canonical)
        self.assertIn("INTAKE_CONFIRMED", canonical)
        for label in (
            "课程名称",
            "专业",
            "授课对象",
            "总课时",
            "理论课时",
            "实践课时",
            "理论与实践组织方式",
            "单课课时",
            "使用教材",
            "辅助参考资料",
            "是否同时生成实践任务工单",
        ):
            self.assertIn(label, canonical)
        for rule in ("待确认", "50/50", "当前理解 / 如不准确请修改"):
            self.assertIn(rule, canonical)
        openai_prompt = openai.split("default_prompt:", 1)[1]
        for internal_key in ("course_name", "major", "audience", "total_hours", "theory_hours", "practice_hours"):
            self.assertNotIn(internal_key, openai_prompt)

    def test_lesson_intake_machine_contract_matches_canonical_facades(self) -> None:
        contract_path = LESSON / "docs" / "intake-contract-v2.1.1.json"
        contract = json.loads(contract_path.read_text(encoding="utf-8"))
        self.assertEqual(contract["version"], "2.1.1")
        self.assertEqual(contract["ui_language"], "zh-CN")
        self.assertEqual(contract["states"], ["INTAKE_PENDING", "INTAKE_CONFIRMED"])
        self.assertEqual(contract["source_semantics"], ["EXPLICIT", "DERIVED", "INFERRED", "PENDING"])
        self.assertEqual(contract["default_hours"], 2)
        fields = {item["internal"]: item for item in contract["confirmation_fields"]}
        for field in ("course_name", "major", "audience", "total_hours"):
            self.assertTrue(fields[field]["required"])
            self.assertEqual(fields[field]["unknown_state"], "PENDING")
        self.assertTrue(fields["total_hours"]["explicit_only"])
        for field in ("major", "audience"):
            self.assertEqual(fields[field]["inference_label"], "当前理解 / 如不准确请修改")
        for field in ("theory_hours", "practice_hours", "delivery_mode", "practice_work_orders"):
            self.assertEqual(fields[field]["unknown_state"], "PENDING")
        self.assertTrue(fields["textbook"]["recommended"])
        self.assertEqual(contract["pure_course_normalization"]["theory_only"]["practice_hours"], 0)
        self.assertEqual(contract["pure_course_normalization"]["practice_only"]["theory_hours"], 0)
        self.assertIn("docx_generation", contract["forbidden_before_confirmation"])
        self.assertIn("docx_generation", contract["post_confirmation_questions_forbidden"])
        self.assertIn("intake_pending", contract["user_visible_errors"])
        self.assertIn("hours_conflict", contract["user_visible_errors"])
        self.assertIn("课程基本信息尚未确认", contract["user_visible_errors"]["intake_pending"])
        self.assertEqual(
            (LESSON / "manifest.yaml").read_text(encoding="utf-8").count("version: 2.1.1"),
            1,
        )

    def test_lesson_installer_doctor_tracks_facade_staleness_and_replacement(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-installer-doctor-") as temp_name:
            folder = Path(temp_name)
            source = folder / "source"
            shutil.copytree(LESSON, source)
            shared_source = ROOT / "schemas" / "shared" / "practice-task-contract.schema.json"
            shared_target = source / "schemas" / "shared" / "practice-task-contract.schema.json"
            shared_target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(shared_source, shared_target)
            skills_root = folder / "skills"

            missing = lesson_install.inspect_installation(source, skills_root)
            self.assertEqual(missing["status"], "missing")
            lesson_install.install(source, skills_root)
            target = skills_root / lesson_install.SKILL_NAME
            current = lesson_install.inspect_installation(source, skills_root)
            self.assertEqual(current["status"], "current", current)
            self.assertTrue((target / lesson_install.INSTALL_MANIFEST).is_file())

            openai_source = source / "agents" / "openai.yaml"
            openai_source.write_text(openai_source.read_text(encoding="utf-8") + "\n# facade fingerprint probe\n", encoding="utf-8")
            stale = lesson_install.inspect_installation(source, skills_root)
            self.assertEqual(stale["status"], "stale", stale)
            lesson_install.install(source, skills_root, replace=True)
            replaced = lesson_install.inspect_installation(source, skills_root)
            self.assertEqual(replaced["status"], "current", replaced)

            (target / "agents" / "openai.yaml").write_text(
                (target / "agents" / "openai.yaml").read_text(encoding="utf-8") + "\n# installed mutation\n",
                encoding="utf-8",
            )
            inconsistent = lesson_install.inspect_installation(source, skills_root)
            self.assertEqual(inconsistent["status"], "inconsistent", inconsistent)

    def _acceptance_fixture(self, folder: Path, lesson_count: int = 3) -> tuple[Path, Path, Path, dict]:
        lessons = []
        for index in range(1, lesson_count + 1):
            lessons.append(
                {
                    "id": f"L{index:02d}",
                    "unit": f"项目{(index - 1) // 2 + 1}",
                    "teaching_content": {"topic": f"主题{index}"},
                    "implementation": {
                        "stages": [
                            {
                                "teacher_actions": [f"教师动作{index}"],
                                "student_actions": [f"学生动作{index}"],
                            }
                        ]
                    },
                    "evaluation": {"remarks": {"practice": f"评价备注{index}"}},
                    "reflection": {"summary": f"反思{index}"},
                }
            )
        data = {
            "course_name": "验收测试课程",
            "major": "软件技术",
            "audience": "高职一年级",
            "content_contract_version": "2.0",
            "total_hours": lesson_count * 2,
            "lessons": lessons,
        }
        links = [
            {
                "from": f"L{index:02d}",
                "to": f"L{index + 1:02d}",
                "status": "passed",
                "same_unit": data["lessons"][index - 1]["unit"] == data["lessons"][index]["unit"],
            }
            for index in range(1, lesson_count)
        ]
        qa = {
            "status": "passed",
            "template_id": "lesson-plan",
            "template_version": "1.1.2",
            "template_path": "lesson-plan/v1.1.2/template.docx",
            "content_contract_version": "2.0",
            "files_checked": lesson_count,
            "errors": [],
            "validation": {"template": True, "output": True},
            "checks": {
                "file_count": {"expected": lesson_count, "actual": lesson_count},
                "total_hours": {"expected": lesson_count * 2, "actual": lesson_count * 2},
                "lessons": [{"file_index": index, "errors": []} for index in range(1, lesson_count + 1)],
                "anchors": {
                    "mode": "word_bookmark",
                    "required": 2,
                    "preserved": 2,
                    "missing": [],
                    "duplicates": [],
                    "invalid_names": [],
                    "unexpected_names": [],
                    "invalid_ids": [],
                    "boundary_errors": [],
                },
            },
            "content_quality": {
                "status": "passed",
                "errors": [],
                "exact_duplicates": [],
                "adjacent_exact_duplicates": [],
                "item_duplicates": [],
                "adjacent_item_duplicates": [],
                "frequency_item_duplicates": [],
                "adjacent_similarity_pairs": [],
                "repeated_sentences": [],
                "whole_lesson_similarity_pairs": [],
                "field_similarity_pairs": [],
                "structural_similarity_pairs": [],
                "implementation_duplicates": [],
                "adjacent_implementation_exact_duplicates": [],
                "implementation_similarity_pairs": [],
                "implementation_structural_similarity_pairs": [],
                "evaluation_remark_duplicates": [],
                "reference_provenance": {
                    "reuse_policy": "reference_reusable",
                    "cross_lesson_reuse": "allowed",
                    "same_lesson_duplicates": [],
                    "invalid_resource_only": [],
                },
                "progression": {"status": "passed", "sequence_links": links},
            },
            "render": {
                "status": "passed",
                "files_checked": lesson_count,
                "page_count": lesson_count * 2,
                "page_counts": {f"L{index:02d}.docx": 2 for index in range(1, lesson_count + 1)},
            },
        }
        input_path = folder / "content-v2.json"
        output = folder / "output"
        output.mkdir()
        qa_path = output / "qa-report.json"
        report_dir = folder / "acceptance-report"
        input_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        qa_path.write_text(json.dumps(qa, ensure_ascii=False), encoding="utf-8")
        for index in range(1, lesson_count + 1):
            (output / f"L{index:02d}.docx").write_bytes(f"docx-{index}".encode("ascii"))
        return input_path, output, qa_path, data

    def test_acceptance_report_schema_metadata_and_read_only_inventory(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-acceptance-schema-") as temp_name:
            folder = Path(temp_name)
            input_path, output, qa_path, _data = self._acceptance_fixture(folder)
            before = {path: path.read_bytes() for path in output.iterdir()}
            report = lesson_acceptance.build_acceptance_report(
                input_path,
                output,
                qa_path,
                source_type="synthetic_fixture",
                report_dir=folder / "report",
                master_commit="a" * 40,
            )
            self.assertEqual(lesson_acceptance.validate_report_schema(report), [])
            from jsonschema import Draft202012Validator

            schema = json.loads((ROOT / "docs" / "lesson-acceptance-report.schema.json").read_text(encoding="utf-8"))
            self.assertEqual(list(Draft202012Validator(schema).iter_errors(report)), [])
            self.assertEqual(report["metadata"]["content_contract_version"], "2.0")
            self.assertEqual(report["metadata"]["template_version"], "v1.1.2")
            self.assertEqual(report["metadata"]["lesson_count"], 3)
            self.assertEqual(report["structural_hard_gates"]["status"], "PASS")
            self.assertEqual(report["final_status"], "PENDING_MANUAL_REVIEW")
            self.assertEqual(report["visual_review"]["status"], "not_executed")
            self.assertTrue(report["metadata"]["output_inventory_fingerprint"])
            self.assertEqual({path: path.read_bytes() for path in output.iterdir()}, before)

            with self.assertRaisesRegex(ValueError, "must not overlap"):
                lesson_acceptance.build_acceptance_report(
                    input_path,
                    output,
                    qa_path,
                    source_type="synthetic_fixture",
                    report_dir=output / "acceptance-report",
                )

    def test_acceptance_reuses_existing_metrics_and_expands_only_review_links(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-acceptance-metrics-") as temp_name:
            folder = Path(temp_name)
            input_path, output, qa_path, _data = self._acceptance_fixture(folder)
            qa = json.loads(qa_path.read_text(encoding="utf-8"))
            qa["content_quality"]["whole_lesson_similarity_pairs"] = [{"from": "L01", "to": "L02", "score": 0.31}]
            qa["content_quality"]["field_similarity_pairs"] = [{"field": "teaching_content", "score": 0.22}]
            qa["content_quality"]["implementation_similarity_pairs"] = [{"field": "teacher_actions", "score": 0.19}]
            qa["content_quality"]["evaluation_remark_duplicates"] = [{"field": "evaluation.remarks.practice", "score": 0.12}]
            qa_path.write_text(json.dumps(qa, ensure_ascii=False), encoding="utf-8")
            report = lesson_acceptance.build_acceptance_report(
                input_path,
                output,
                qa_path,
                source_type="synthetic_fixture",
                report_dir=folder / "report",
            )
            evidence = report["content_quality_evidence"]
            self.assertEqual(
                set(evidence["detector_counts"]),
                {
                    "exact_duplicates",
                    "adjacent_exact_duplicates",
                    "item_duplicates",
                    "adjacent_item_duplicates",
                    "frequency_item_duplicates",
                    "adjacent_similarity_pairs",
                    "repeated_sentences",
                    "whole_lesson_similarity_pairs",
                    "field_similarity_pairs",
                    "structural_similarity_pairs",
                    "implementation_duplicates",
                    "adjacent_implementation_exact_duplicates",
                    "implementation_similarity_pairs",
                    "implementation_structural_similarity_pairs",
                    "evaluation_remark_duplicates",
                },
            )
            self.assertEqual(evidence["similarity"]["whole_lesson"]["max"], 0.31)
            self.assertEqual(evidence["similarity"]["fields"]["teaching_content"]["count"], 1)
            self.assertEqual(evidence["similarity"]["implementation"]["count"], 1)
            self.assertNotIn('"threshold":', json.dumps(evidence, ensure_ascii=False).lower())
            self.assertEqual(report["sequence_review"]["physical_transition_count"], 2)
            self.assertEqual(report["sequence_review"]["summary"]["boundaries"], 1)
            self.assertTrue(all(item["boundary"] or item["status"] != "PASS" for item in report["sequence_review"]["attention_transitions"]))

    def test_acceptance_known_software_modeling_boundaries_use_existing_sequence_links(self) -> None:
        lessons = [{"id": f"L{index:02d}", "unit": f"项目{(index - 1) // 4 + 1}"} for index in range(1, 33)]
        data = {"total_hours": 64, "lessons": lessons}
        links = [
            {"from": f"L{index:02d}", "to": f"L{index + 1:02d}", "status": "passed", "same_unit": (index - 1) // 4 == index // 4}
            for index in range(1, 32)
        ]
        links[3]["requires_agent_review"] = True
        report = lesson_acceptance.sequence_review(
            data,
            {"content_quality": {"progression": {"sequence_links": links}}},
        )
        self.assertEqual(report["physical_transition_count"], 31)
        self.assertTrue(report["known_software_modeling_64h"])
        self.assertEqual(
            [(item["from"], item["to"]) for item in report["known_boundaries"]],
            list(lesson_acceptance.KNOWN_SOFTWARE_MODELING_64H_BOUNDARIES),
        )
        self.assertEqual(report["physical_transitions"][3]["status"], "REVIEW")

    def test_acceptance_manual_layers_and_negative_control_catalog_are_explicit(self) -> None:
        catalog = lesson_acceptance.negative_control_catalog()
        self.assertEqual(
            {item["id"] for item in catalog},
            {
                "nursing_sql_contamination",
                "database_patient_bp",
                "copied_teacher_actions",
                "mechanical_scores",
                "generic_fabricated_reference",
                "short_remark",
                "v21_reference_placeholder",
                "v21_textbook_overlap",
                "v21_same_lesson_reference_id",
                "v21_unresolved_reference_id",
                "v21_resource_only_reference",
                "v21_delivery_hour_mismatch",
                "v21_practice_link_mismatch",
            },
        )
        controls = lesson_acceptance.negative_controls()
        self.assertEqual(controls["status"], "not_executed")
        self.assertEqual(controls["transaction_safety"]["candidate_cleanup"], "not_executed")
        self.assertEqual(controls["transaction_safety"]["old_output_preserved"], "not_executed")
        self.assertEqual(
            lesson_acceptance._final_status(
                {"status": "PASS"},
                {"status": "not_executed"},
                {"status": "not_executed"},
                {"status": "not_executed"},
                {"status": "failed"},
            ),
            "FAILED",
        )
        scope = lesson_acceptance.course_scope_review({"lessons": [{"unit": "P1"}, {"unit": "P2"}]})
        self.assertFalse(scope["hard_gate"])
        self.assertEqual(scope["allowed_classifications"], ["CORE", "EXTENSION", "POSSIBLE_SCOPE_DRIFT"])
        self.assertIn("PENDING_MANUAL_REVIEW", lesson_acceptance.acceptance_markdown({
            "final_status": "PENDING_MANUAL_REVIEW",
            "metadata": {"course": "C", "major": "M", "audience": "A", "lesson_count": 0, "total_hours": 0, "source_type": "mixed", "master_commit": "unknown", "content_contract_version": "2.0", "template_version": "v1.1.2", "input_sha256": "x", "qa_report_sha256": "x", "output_inventory_fingerprint": None, "render_status": "not_executed", "visual_status": "not_executed"},
            "structural_hard_gates": {"status": "FAIL", "gates": []},
            "content_quality_evidence": {"status": "not_available", "detector_counts": {}},
            "sequence_review": {"physical_transition_count": 0, "summary": {}},
            "visual_review": {"status": "not_executed"},
            "teaching_design_review": {"status": "not_executed"},
            "teacher_usability": {"status": "not_executed"},
            "negative_controls": {"status": "not_executed"},
        }))

    def test_hardening_import_loader_restores_process_global_state(self) -> None:
        original_path = sys.path[:]
        original_modules = {name: sys.modules.get(name, _MISSING) for name in _LESSON_MODULE_NAMES}
        with isolated_lesson_imports():
            self.assertIn(str(SCRIPTS), sys.path)
        self.assertEqual(sys.path, original_path)
        for name, previous in original_modules.items():
            if previous is _MISSING:
                self.assertNotIn(name, sys.modules)
            else:
                self.assertIs(sys.modules.get(name), previous)

    def test_intra_lesson_calibration_covers_domain_positives_and_negatives(self) -> None:
        cases = content_quality.intra_lesson_coherence_calibration()
        self.assertGreaterEqual(len(cases), 8)
        self.assertTrue(all(item["expected"] == item["actual"] for item in cases), cases)

    def test_implementation_item_coherence_checks_all_fields_and_bounded_diagnostics(self) -> None:
        def lesson(domain: str, **stage_overrides):
            if domain == "nursing":
                task = "完成患者生命体征测量与护理记录"
                body = ["测量患者血压、体温并判断异常", "记录护理结果并形成生命体征测量记录"]
                deliverable = "生命体征测量记录与护理结果"
            else:
                task = "完成数据库查询与索引维护"
                body = ["编写 SQL 查询语句并分析查询结果", "形成数据库查询报告"]
                deliverable = "数据库查询报告"
            stage = {
                "id": "task_implementation",
                "label": "任务实施",
                "modality": "小组实训",
                "content": [body[0], body[1]],
                "teacher_actions": ["教师巡视"],
                "student_actions": ["学生讨论"],
                "objective": body[0],
            }
            stage.update(stage_overrides)
            return {
                "lesson_id": "L01",
                "task": task,
                "teaching_content": body,
                "key_point": {"content": [body[0]]},
                "difficult_point": {"content": [body[1]]},
                "progression": {"deliverable": deliverable},
                "implementation": [stage],
            }

        cases = (
            ("nursing mixed content", lesson("nursing", content=["完成血压测量", "记录护理结果", "使用 SQL 创建数据库索引"]), "content", 2),
            ("nursing mixed teacher action", lesson("nursing", teacher_actions=["教师巡视", "演示数据库索引优化"]), "teacher_actions", 1),
            ("database mixed student action", lesson("database", student_actions=["完成 SQL 查询", "测量患者血压"]), "student_actions", 1),
            ("nursing domain label", lesson("nursing", label="数据库索引优化"), "label", 0),
            ("nursing domain modality", lesson("nursing", modality="SQL数据库查询实训"), "modality", 0),
        )
        for name, payload, field, item_index in cases:
            with self.subTest(case=name):
                report = content_quality._intra_lesson_coherence([payload], ["L01"])
                failures = [
                    item
                    for item in report["failures"]
                    if item.get("failed_gate") == "implementation_item_coherence"
                ]
                failure = next(item for item in failures if item["field"] == field and item["item_index"] == item_index)
                self.assertEqual(failure["status"], "failed")
                self.assertEqual(failure["reason"], "substantive item is disconnected from lesson main semantic component")
                self.assertEqual(set(failure["diagnostic"]), {"preview", "text_sha256"})
                value = payload["implementation"][0][field]
                value = value if isinstance(value, str) else value[item_index]
                if len(str(value)) > content_quality.DIAGNOSTIC_PREVIEW_MAX_CHARS:
                    self.assertNotIn(str(value), failure["diagnostic"]["preview"])

        for name, payload in (
            ("nursing positive", lesson("nursing", content=["完成血压测量", "记录护理结果"])),
            ("database positive", lesson("database", content=["编写 SQL 查询", "分析查询结果"])),
        ):
            with self.subTest(case=name):
                report = content_quality._intra_lesson_coherence([payload], ["L01"])
                self.assertEqual(report["status"], "passed", report)
                items = report["lessons"][0]["implementation"]["stages"][0]["items"]
                self.assertTrue(any(item["generic"] for item in items))
                self.assertTrue(all(item["status"] == "passed" for item in items))

    def test_intra_lesson_gate_b_requires_the_main_semantic_component(self) -> None:
        fixture = json.loads(
            (ROOT / "tests" / "fixtures" / "lesson-plan-content-v2-it.json").read_text(
                encoding="utf-8"
            )
        )
        lesson = next(item for item in fixture["lessons"] if item["lesson_id"] == "L03")
        report = content_quality._intra_lesson_coherence([lesson], ["L03"])
        self.assertEqual(report["status"], "passed", report)
        lesson_report = report["lessons"][0]
        self.assertEqual(
            set(lesson_report["main_component"]),
            {"task_connected", "deliverable_connected", "members"},
        )
        self.assertTrue(lesson_report["main_component"]["task_connected"])
        self.assertTrue(lesson_report["main_component"]["deliverable_connected"])
        self.assertNotIn("value", lesson_report["core_nodes"][0])
        self.assertIn("diagnostic", lesson_report["core_nodes"][0])
        self.assertIn("preview", lesson_report["core_nodes"][0]["diagnostic"])
        self.assertIn("text_sha256", lesson_report["core_nodes"][0]["diagnostic"])
        gate_b = report["lessons"][0]["gate_b"]
        self.assertTrue(gate_b["shared_cluster"]["cross_body_bridge"])
        self.assertTrue(gate_b["shared_cluster"]["same_component"])

    def test_installer_dry_run_and_replace_are_transactional(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-installer-hardening-") as temp_name:
            root = Path(temp_name) / "skills"
            lesson_install.install(LESSON, root, dry_run=True)
            self.assertFalse(root.exists())
            lesson_install.install(LESSON, root)
            target = root / lesson_install.SKILL_NAME
            (target / "user-marker.txt").write_text("old", encoding="utf-8")
            lesson_install.install(LESSON, root, replace=True)
            backups = list(root.glob("lesson-plan-docx-generator_backup_*"))
            self.assertEqual(len(backups), 1)
            self.assertEqual((backups[0] / "user-marker.txt").read_text(encoding="utf-8"), "old")

    def test_lesson_installer_checks_source_and_staged_runtime_inventory_before_commit(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-installer-inventory-") as temp_name:
            folder = Path(temp_name)
            source = folder / "source"
            shutil.copytree(LESSON, source)
            (source / "scripts" / "content_quality.py").unlink()
            target_root = folder / "missing-source-target"
            with self.assertRaisesRegex(FileNotFoundError, "content_quality.py"):
                lesson_install.install(source, target_root)
            self.assertFalse(target_root.exists())

            target_root = folder / "missing-stage-target"
            real_verify_stage = lesson_install._verify_stage

            def verify_stage_with_missing_runtime(source_path, stage_path):
                (Path(stage_path) / "scripts" / "content_quality.py").unlink()
                return real_verify_stage(source_path, stage_path)

            with patch.object(lesson_install, "_verify_stage", side_effect=verify_stage_with_missing_runtime):
                with self.assertRaisesRegex(RuntimeError, "staged installation is missing required file"):
                    lesson_install.install(LESSON, target_root)
            self.assertFalse((target_root / lesson_install.SKILL_NAME).exists())
            self.assertEqual(list(target_root.glob(f".{lesson_install.SKILL_NAME}.stage-*")), [])

    def test_installer_commit_failure_restores_previous_target(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-installer-rollback-") as temp_name:
            root = Path(temp_name) / "skills"
            lesson_install.install(LESSON, root)
            target = root / lesson_install.SKILL_NAME
            marker = target / "rollback-marker.txt"
            marker.write_text("keep", encoding="utf-8")
            real_replace = lesson_install.os.replace
            calls = 0

            def fail_commit(source, destination):
                nonlocal calls
                calls += 1
                if calls == 2:
                    raise OSError("injected commit failure")
                return real_replace(source, destination)

            with patch.object(lesson_install.os, "replace", side_effect=fail_commit):
                with self.assertRaisesRegex(RuntimeError, "previous installation restored"):
                    lesson_install.install(LESSON, root, replace=True)
            self.assertEqual(marker.read_text(encoding="utf-8"), "keep")

    def test_installer_cleanup_warning_preserves_success_and_primary_failure(self) -> None:
        warning = "cleanup failed for staging directory: locked; residual path: stage"
        with tempfile.TemporaryDirectory(prefix="lesson-installer-cleanup-success-") as temp_name:
            root = Path(temp_name) / "skills"
            diagnostics = io.StringIO()
            with patch.object(lesson_install, "_remove_stage", return_value=warning):
                with redirect_stderr(diagnostics):
                    target = lesson_install.install(LESSON, root)
            self.assertTrue(target.is_dir())
            self.assertIn("WARNING:", diagnostics.getvalue())
            self.assertIn("residual path:", diagnostics.getvalue())

        with tempfile.TemporaryDirectory(prefix="lesson-installer-cleanup-failure-") as temp_name:
            root = Path(temp_name) / "skills"
            diagnostics = io.StringIO()
            with patch.object(lesson_install, "_verify_stage", side_effect=RuntimeError("injected verify failure")):
                with patch.object(lesson_install, "_remove_stage", return_value=warning):
                    with redirect_stderr(diagnostics):
                        with self.assertRaisesRegex(RuntimeError, "injected verify failure"):
                            lesson_install.install(LESSON, root)
            self.assertIn("WARNING:", diagnostics.getvalue())
            self.assertIn("residual path:", diagnostics.getvalue())

    def test_adapter_transaction_failure_injections_a_through_e_are_independent(self) -> None:
        real_replace = install_adapters.os.replace

        with tempfile.TemporaryDirectory(prefix="lesson-adapter-transaction-") as temp_name:
            folder = Path(temp_name)

            with self.subTest(injection="A-stage validation"):
                root = folder / "a"
                root.mkdir()
                target = root / "existing.txt"
                target.write_bytes(b"old")
                with patch.object(
                    install_adapters,
                    "_assert_no_symlink_ancestor",
                    side_effect=OSError("injected staging validation failure"),
                ):
                    with self.assertRaisesRegex(RuntimeError, "injected staging validation failure"):
                        install_adapters._apply_files({target: b"new"}, root)
                self.assertEqual(target.read_bytes(), b"old")
                self.assertEqual(list(root.glob(".lesson-adapters.stage-*")), [])

            with self.subTest(injection="B-backup"):
                root = folder / "b"
                root.mkdir()
                target = root / "existing.txt"
                target.write_bytes(b"old")
                calls = 0

                def fail_backup(source, destination):
                    nonlocal calls
                    calls += 1
                    if calls == 1:
                        raise OSError("injected backup failure")
                    return real_replace(source, destination)

                with patch.object(install_adapters.os, "replace", side_effect=fail_backup):
                    with self.assertRaisesRegex(RuntimeError, "injected backup failure"):
                        install_adapters._apply_files({target: b"new"}, root)
                self.assertEqual(target.read_bytes(), b"old")
                self.assertEqual(list(root.glob("*.backup_*")), [])

            with self.subTest(injection="C-file commit"):
                root = folder / "c"
                root.mkdir()
                target = root / "existing.txt"
                target.write_bytes(b"old")
                calls = 0

                def fail_file_commit(source, destination):
                    nonlocal calls
                    calls += 1
                    if calls == 2:
                        raise OSError("injected file commit failure")
                    return real_replace(source, destination)

                with patch.object(install_adapters.os, "replace", side_effect=fail_file_commit):
                    with self.assertRaisesRegex(RuntimeError, "previous files restored"):
                        install_adapters._apply_files({target: b"new"}, root)
                self.assertEqual(target.read_bytes(), b"old")
                self.assertEqual(list(root.glob("*.backup_*")), [])

            with self.subTest(injection="D-engine swap"):
                root = folder / "d"
                engine = root / install_adapters.ENGINE_NAME
                engine.mkdir(parents=True)
                old_runtime = engine / "scripts" / "old.py"
                old_runtime.parent.mkdir()
                old_runtime.write_bytes(b"old")
                target = engine / "scripts" / "new.py"
                calls = 0

                def fail_engine_swap(source, destination):
                    nonlocal calls
                    calls += 1
                    if calls == 2:
                        raise OSError("injected engine swap failure")
                    return real_replace(source, destination)

                with patch.object(install_adapters.os, "replace", side_effect=fail_engine_swap):
                    with self.assertRaisesRegex(RuntimeError, "previous files restored"):
                        install_adapters._apply_files({target: b"new"}, root, replace_engine=engine)
                self.assertEqual(old_runtime.read_bytes(), b"old")
                self.assertFalse((engine / "scripts" / "new.py").exists())
                self.assertEqual(list(root.glob(f"{install_adapters.ENGINE_NAME}.backup_*")), [])

            with self.subTest(injection="E-rollback accumulation"):
                root = folder / "e"
                root.mkdir()
                first = root / "a.txt"
                second = root / "b.txt"
                first.write_bytes(b"old-a")
                second.write_bytes(b"old-b")
                calls = 0

                def fail_commit_and_restore(source, destination):
                    nonlocal calls
                    calls += 1
                    if calls == 4:
                        raise OSError("injected commit failure")
                    if calls in (5, 6):
                        raise OSError(f"injected rollback restore failure {calls}")
                    return real_replace(source, destination)

                real_remove = install_adapters._remove_path

                def fail_first_removal(path):
                    if path == first:
                        raise OSError("injected rollback remove failure")
                    return real_remove(path)

                with patch.object(install_adapters.os, "replace", side_effect=fail_commit_and_restore):
                    with patch.object(install_adapters, "_remove_path", side_effect=fail_first_removal):
                        with self.assertRaisesRegex(RuntimeError, "injected commit failure") as raised:
                            install_adapters._apply_files({first: b"new-a", second: b"new-b"}, root)
                message = str(raised.exception)
                self.assertIn("rollback failures:", message)
                self.assertIn("injected rollback remove failure", message)
                self.assertIn("injected rollback restore failure 5", message)
                self.assertIn("injected rollback restore failure 6", message)
                self.assertIn(str(first), message)
                self.assertIn(str(second), message)
                self.assertEqual(first.read_bytes(), b"new-a")
                self.assertFalse(second.exists())
                self.assertEqual(len(list(root.glob("*.backup_*"))), 2)
                self.assertEqual(list(root.glob(".lesson-adapters.stage-*")), [])

    def test_adapter_stage_cleanup_warning_preserves_success_and_primary_error(self) -> None:
        warning = "cleanup failed for adapter staging directory: locked; residual path: stage"
        real_replace = install_adapters.os.replace
        with tempfile.TemporaryDirectory(prefix="lesson-adapter-cleanup-") as temp_name:
            root = Path(temp_name)
            target = root / "new.txt"
            diagnostics = io.StringIO()
            with patch.object(install_adapters, "_cleanup_path", return_value=warning):
                with redirect_stderr(diagnostics):
                    install_adapters._apply_files({target: b"new"}, root)
            self.assertEqual(target.read_bytes(), b"new")
            self.assertIn("WARNING:", diagnostics.getvalue())
            self.assertIn("residual path:", diagnostics.getvalue())

            old = root / "old.txt"
            old.write_bytes(b"old")
            calls = 0

            def fail_commit(source, destination):
                nonlocal calls
                calls += 1
                if calls == 2:
                    raise OSError("injected primary commit failure")
                return real_replace(source, destination)

            diagnostics = io.StringIO()
            with patch.object(install_adapters.os, "replace", side_effect=fail_commit):
                with patch.object(install_adapters, "_cleanup_path", return_value=warning):
                    with redirect_stderr(diagnostics):
                        with self.assertRaisesRegex(RuntimeError, "injected primary commit failure"):
                            install_adapters._apply_files({old: b"new"}, root)
            self.assertEqual(old.read_bytes(), b"old")
            self.assertIn("WARNING:", diagnostics.getvalue())
            self.assertIn("residual path:", diagnostics.getvalue())

    def test_lesson_adapter_engine_modes_preserve_full_runtime_and_fail_closed(self) -> None:
        def engine_snapshot(engine: Path) -> dict[str, bytes]:
            return {
                str(path.relative_to(engine)): path.read_bytes()
                for path in engine.rglob("*")
                if path.is_file() and not path.is_symlink()
            }

        def tree_snapshot(root: Path) -> dict[str, bytes]:
            return {
                str(path.relative_to(root)): path.read_bytes()
                for path in root.rglob("*")
                if path.is_file() and not path.is_symlink()
            }

        with tempfile.TemporaryDirectory(prefix="lesson-adapter-engine-modes-") as temp_name:
            folder = Path(temp_name)
            minimal = folder / "minimal"
            install_adapters.install(LESSON, minimal, adapters=["all"])
            minimal_engine = minimal / install_adapters.ENGINE_NAME
            self.assertEqual(install_adapters._detect_existing_engine_mode(minimal_engine), "minimal")
            minimal_before = engine_snapshot(minimal_engine)
            install_adapters.install(LESSON, minimal, adapters=["all"])
            self.assertEqual(engine_snapshot(minimal_engine), minimal_before)

            full = folder / "full"
            output = io.StringIO()
            with redirect_stdout(output):
                install_adapters.install(LESSON, full, adapters=["all"], copy_engine=True)
            full_engine = full / install_adapters.ENGINE_NAME
            self.assertEqual(install_adapters._detect_existing_engine_mode(full_engine, LESSON), "full-current")
            full_before = engine_snapshot(full_engine)
            agents = full / "AGENTS.md"
            agents.write_text("project-owned notes\n", encoding="utf-8")
            output = io.StringIO()
            with redirect_stdout(output):
                install_adapters.install(LESSON, full, adapters=["agents"])
            self.assertIn("preserved-full-engine=", output.getvalue())
            self.assertEqual(engine_snapshot(full_engine), full_before)
            agents_text = agents.read_text(encoding="utf-8")
            self.assertIn("project-owned notes", agents_text)
            self.assertIn(install_adapters.MARKER_START, agents_text)
            self.assertIn(".lesson-plan-docx-generator/scripts/generate_lesson_plans.py", agents_text)
            self.assertNotIn("the full Lesson runtime (install with --copy-engine)", agents_text)

            legacy = folder / "legacy-full"
            install_adapters.install(LESSON, legacy, adapters=["all"], copy_engine=True)
            legacy_engine = legacy / install_adapters.ENGINE_NAME
            (legacy_engine / install_adapters.ENGINE_STATE_FILE).unlink()
            self.assertEqual(install_adapters._detect_existing_engine_mode(legacy_engine, LESSON), "full-stale")
            legacy_before = tree_snapshot(legacy)
            with self.assertRaisesRegex(ValueError, "A full Lesson engine is installed but is older/different"):
                install_adapters.install(LESSON, legacy, adapters=["agents"])
            self.assertEqual(tree_snapshot(legacy), legacy_before)
            install_adapters.install(LESSON, legacy, adapters=["all"], copy_engine=True, replace=True)
            self.assertEqual(install_adapters._detect_existing_engine_mode(legacy_engine, LESSON), "full-current")

            stale = full_engine / "scripts" / "obsolete-helper.py"
            stale.write_text("stale", encoding="utf-8")
            install_adapters.install(LESSON, full, adapters=["all"], copy_engine=True, replace=True)
            self.assertFalse(stale.exists())
            self.assertEqual(install_adapters._detect_existing_engine_mode(full_engine, LESSON), "full-current")
            install_adapters.install(LESSON, full, adapters=["all"], copy_engine=True, replace=True)
            self.assertEqual(install_adapters._detect_existing_engine_mode(full_engine, LESSON), "full-current")

            modified_source = folder / "modified-source"
            shutil.copytree(LESSON, modified_source)
            runtime_source = modified_source / "scripts" / "content_quality.py"
            runtime_source.write_text(runtime_source.read_text(encoding="utf-8") + "\n# source fingerprint probe\n", encoding="utf-8")
            self.assertEqual(install_adapters._detect_existing_engine_mode(full_engine, modified_source), "full-stale")
            before = tree_snapshot(full)
            with self.assertRaisesRegex(ValueError, "A full Lesson engine is installed but is older/different"):
                install_adapters.install(modified_source, full, adapters=["agents"])
            self.assertEqual(tree_snapshot(full), before)
            install_adapters.install(modified_source, full, adapters=["all"], copy_engine=True, replace=True)
            self.assertEqual(install_adapters._detect_existing_engine_mode(full_engine, modified_source), "full-current")

            installed_runtime = full_engine / "scripts" / "content_quality.py"
            installed_runtime.write_text(installed_runtime.read_text(encoding="utf-8") + "\n# installed mutation\n", encoding="utf-8")
            self.assertEqual(install_adapters._detect_existing_engine_mode(full_engine, modified_source), "inconsistent")

            inconsistent = folder / "inconsistent"
            inconsistent_engine = inconsistent / install_adapters.ENGINE_NAME
            inconsistent_engine.mkdir(parents=True)
            (inconsistent_engine / "SKILL.md").write_text("partial", encoding="utf-8")
            before = tree_snapshot(inconsistent)
            with self.assertRaisesRegex(ValueError, "inconsistent"):
                install_adapters.install(LESSON, inconsistent, adapters=["agents"])
            self.assertEqual(tree_snapshot(inconsistent), before)
            self.assertEqual(install_adapters._detect_existing_engine_mode(inconsistent_engine, LESSON), "inconsistent")

            incomplete_full = folder / "incomplete-full"
            install_adapters.install(LESSON, incomplete_full, adapters=["all"], copy_engine=True)
            incomplete_engine = incomplete_full / install_adapters.ENGINE_NAME
            (incomplete_engine / "scripts" / "content_quality.py").unlink()
            before = tree_snapshot(incomplete_full)
            self.assertEqual(install_adapters._detect_existing_engine_mode(incomplete_engine, LESSON), "inconsistent")
            with self.assertRaisesRegex(ValueError, "inconsistent"):
                install_adapters.install(LESSON, incomplete_full, adapters=["agents"])
            self.assertEqual(tree_snapshot(incomplete_full), before)

    def test_adapter_namespace_markers_and_aider_fail_closed(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-adapter-hardening-") as temp_name:
            target = Path(temp_name) / "project"
            install_adapters.install(LESSON, target, adapters=["all"])
            agents = (target / "AGENTS.md").read_text(encoding="utf-8")
            self.assertEqual(agents.count(install_adapters.MARKER_START), 1)
            self.assertIn(".lesson-plan-docx-generator/SKILL.md", agents)
            self.assertNotIn(".lesson-plan-docx-generator/scripts/generate_lesson_plans.py", agents)
            self.assertIn("--copy-engine", agents)
            self.assertTrue((target / ".lesson-plan-docx-generator" / "SKILL.md").is_file())
            self.assertTrue((target / ".lesson-plan-docx-generator" / "CONVENTIONS.md").is_file())
            for minimal_name in ("SKILL.md", "通用提示词.md", "AGENTS.md", "CONVENTIONS.md"):
                minimal_text = (target / ".lesson-plan-docx-generator" / minimal_name).read_text(encoding="utf-8")
                self.assertNotIn("scripts/generate_lesson_plans.py", minimal_text)
                self.assertNotIn("docs/content-contract-v2.md", minimal_text)
                self.assertNotIn("examples/tasks.example.json", minimal_text)
                self.assertNotRegex(minimal_text, r"(?:scripts|docs|examples|schemas|assets)/")
            full_target = Path(temp_name) / "full-project"
            install_adapters.install(LESSON, full_target, adapters=["all"], copy_engine=True)
            full_agents = (full_target / "AGENTS.md").read_text(encoding="utf-8")
            self.assertIn(".lesson-plan-docx-generator/scripts/generate_lesson_plans.py", full_agents)
            stale = full_target / ".lesson-plan-docx-generator" / "obsolete-helper.py"
            stale.write_text("stale", encoding="utf-8")
            install_adapters.install(LESSON, full_target, adapters=["all"], copy_engine=True, replace=True)
            self.assertFalse(stale.exists())
            self.assertTrue(any(path.name.startswith(".lesson-plan-docx-generator.backup_") for path in full_target.iterdir()))
            install_adapters.install(LESSON, target, adapters=["agents"])
            self.assertEqual((target / "AGENTS.md").read_text(encoding="utf-8").count(install_adapters.MARKER_START), 1)
            aider = target / ".aider.conf.yml"
            before = aider.read_bytes()
            aider.write_text("read: {unsafe: true}\n", encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "complex"):
                install_adapters.install(LESSON, target, adapters=["aider"])
            self.assertEqual(aider.read_text(encoding="utf-8"), "read: {unsafe: true}\n")
            self.assertNotEqual(before, aider.read_bytes())

    def test_lesson_marker_merge_is_fail_closed_and_idempotent(self) -> None:
        start = install_adapters.MARKER_START
        end = install_adapters.MARKER_END
        empty = install_adapters.merge_marker_section("", "payload")
        self.assertEqual(empty.count(start), 1)
        self.assertEqual(empty.count(end), 1)

        existing = f"unrelated\n{start}\nold\n{end}\nsuffix\n"
        replaced = install_adapters.merge_marker_section(existing, "new")
        self.assertEqual(replaced.count(start), 1)
        self.assertEqual(replaced.count(end), 1)
        self.assertIn("unrelated", replaced)
        self.assertIn("suffix", replaced)
        self.assertEqual(
            install_adapters.merge_marker_section(replaced, "new"),
            replaced,
        )

        malformed_cases = {
            "missing end": f"{start}\nold\n",
            "duplicate blocks": f"{start}\na\n{end}\n{start}\nb\n{end}",
            "nested start": f"{start}\nouter\n{start}\ninner\n{end}",
            "duplicate end": f"{start}\nold\n{end}\n{end}",
            "end before start": f"{end}\n{start}\nold",
        }
        for name, value in malformed_cases.items():
            with self.subTest(name=name):
                with self.assertRaisesRegex(ValueError, "refusing to mutate"):
                    install_adapters.merge_marker_section(value, "replacement")

    def test_adapter_runtime_reference_closure_for_minimal_and_full_modes(self) -> None:
        runtime_path = re.compile(
            r"(?:(?:<skill>|lesson-plan-docx-generator|course-gradebook-generator|"
            r"\.lesson-plan-docx-generator|\.course-gradebook-generator)[/\\])?"
            r"(?:scripts|docs|examples|schemas|assets)/[^\s`\"'<>（）(),，。；：！？]+"
        )

        def references(root: Path) -> list[str]:
            values: list[str] = []
            for name in ("SKILL.md", "通用提示词.md", "AGENTS.md", "CONVENTIONS.md"):
                values.extend(
                    item.rstrip(".,;:!?")
                    for item in runtime_path.findall((root / name).read_text(encoding="utf-8"))
                )
            return values

        with tempfile.TemporaryDirectory(prefix="lesson-adapter-closure-") as temp_name:
            folder = Path(temp_name)
            minimal = folder / "minimal"
            install_adapters.install(LESSON, minimal, adapters=["all"])
            self.assertEqual(references(minimal / install_adapters.ENGINE_NAME), [])

            full = folder / "full"
            install_adapters.install(LESSON, full, adapters=["all"], copy_engine=True)
            for reference in references(full / install_adapters.ENGINE_NAME):
                self.assertTrue((full / reference).is_file(), reference)
            self.assertTrue((full / install_adapters.ENGINE_NAME / "scripts" / "generate_lesson_plans.py").is_file())
            self.assertTrue((full / install_adapters.ENGINE_NAME / "schemas" / "lesson-plan-input.schema.json").is_file())

            gradebook = load_module(
                ROOT / "平时成绩记分册生成器" / "course-gradebook-generator" / "scripts" / "install_adapters.py",
                "gradebook_runtime_closure_hardening",
            )
            gradebook_root = ROOT / "平时成绩记分册生成器" / "course-gradebook-generator"
            grade_minimal = folder / "grade-minimal"
            for name in gradebook.ADAPTER_PATHS:
                for relative in gradebook.ADAPTER_PATHS[name]:
                    gradebook.copy_file(gradebook_root, grade_minimal, relative, False, False, copy_engine=False)
            gradebook.copy_engine(gradebook_root, grade_minimal, False, False, full=False)
            self.assertEqual(references(grade_minimal / gradebook.ENGINE_NAME), [])

            grade_full = folder / "grade-full"
            for name in gradebook.ADAPTER_PATHS:
                for relative in gradebook.ADAPTER_PATHS[name]:
                    gradebook.copy_file(gradebook_root, grade_full, relative, False, False, copy_engine=True)
            gradebook.copy_engine(gradebook_root, grade_full, False, False, full=True)
            for reference in references(grade_full / gradebook.ENGINE_NAME):
                self.assertTrue((grade_full / reference).is_file(), reference)
            self.assertTrue((grade_full / gradebook.ENGINE_NAME / "scripts" / "generate_gradebook.py").is_file())
            self.assertTrue((grade_full / gradebook.ENGINE_NAME / "schemas" / "gradebook-input.schema.json").is_file())

    def test_gradebook_target_containment_allows_parent_and_sibling(self) -> None:
        gradebook = load_module(
            ROOT / "平时成绩记分册生成器" / "course-gradebook-generator" / "scripts" / "install_adapters.py",
            "gradebook_containment_hardening",
        )
        with tempfile.TemporaryDirectory(prefix="gradebook-containment-") as temp_name:
            folder = Path(temp_name)
            source = folder / "project" / "course-gradebook-generator"
            source.mkdir(parents=True)
            with self.assertRaises(ValueError):
                gradebook.copy_engine(source, source, False, True, full=True)
            with self.assertRaises(ValueError):
                gradebook.copy_engine(source, source / "tmp" / "project", False, True, full=True)

            parent = source.parent
            gradebook.copy_engine(source, parent, False, True, full=True)
            gradebook.copy_engine(source, folder / "sibling", False, True, full=True)

            alias = folder / "alias"
            try:
                alias.symlink_to(source, target_is_directory=True)
            except (OSError, NotImplementedError):
                pass
            else:
                with self.assertRaises(ValueError):
                    gradebook.copy_engine(source, alias / "nested", False, True, full=True)

    def test_successful_install_reports_cleanup_residue_without_failing(self) -> None:
        source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
        with tempfile.TemporaryDirectory(prefix="lesson-cleanup-success-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            argv = [
                "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--skip-template-validation",
                "--skip-output-validation",
            ]
            diagnostics = io.StringIO()
            with patch.object(
                lesson_generator,
                "_cleanup_path",
                return_value="cleanup failed for candidate directory: candidate is locked; residual path: candidate",
            ):
                with patch.dict(os.environ, {"LESSON_ALLOW_UNSAFE_VALIDATION_SKIP": "1"}):
                    with patch.object(sys, "argv", argv):
                        with redirect_stderr(diagnostics):
                            lesson_generator.main()
            self.assertTrue(output.is_dir())
            self.assertIn("WARNING:", diagnostics.getvalue())
            self.assertIn("residual path:", diagnostics.getvalue())

    def test_gradebook_shared_marker_and_aider_coexistence(self) -> None:
        gradebook = load_module(
            ROOT / "平时成绩记分册生成器" / "course-gradebook-generator" / "scripts" / "install_adapters.py",
            "gradebook_install_adapters_hardening",
        )
        with tempfile.TemporaryDirectory(prefix="lesson-gradebook-coexist-") as temp_name:
            target = Path(temp_name) / "project"
            install_adapters.install(LESSON, target, adapters=["agents", "aider"])
            gradebook_root = ROOT / "平时成绩记分册生成器" / "course-gradebook-generator"
            gradebook.copy_engine(gradebook_root, target, False, False)
            gradebook.copy_file(gradebook_root, target, "AGENTS.md", False, False)
            first_agents = (target / "AGENTS.md").read_bytes()
            backups_before = set(target.glob("AGENTS.md.backup_*"))
            gradebook.copy_file(gradebook_root, target, "AGENTS.md", False, False)
            self.assertEqual(first_agents, (target / "AGENTS.md").read_bytes())
            self.assertEqual(backups_before, set(target.glob("AGENTS.md.backup_*")))
            agents = (target / "AGENTS.md").read_text(encoding="utf-8")
            self.assertTrue((target / ".course-gradebook-generator" / "SKILL.md").is_file())
            self.assertTrue((target / ".course-gradebook-generator" / "通用提示词.md").is_file())
            self.assertIn(".course-gradebook-generator/SKILL.md", agents)
            gradebook.copy_file(gradebook_root, target, ".aider.conf.yml", False, False)
            aider = (target / ".aider.conf.yml").read_text(encoding="utf-8")
            self.assertIn("lesson-plan-docx-generator:start", agents)
            self.assertIn("course-gradebook-generator:start", agents)
            self.assertIn(".lesson-plan-docx-generator/SKILL.md", aider)
            self.assertIn(".course-gradebook-generator/SKILL.md", aider)

            default_target = Path(temp_name) / "default-project"
            with patch.object(sys, "argv", ["install_adapters.py", "--target-dir", str(default_target)]):
                gradebook.main()
            default_agents = default_target / "AGENTS.md"
            self.assertTrue((default_target / ".course-gradebook-generator" / "SKILL.md").is_file())
            self.assertIn(".course-gradebook-generator/SKILL.md", default_agents.read_text(encoding="utf-8"))
            default_backups = set(default_target.glob("*.backup_*"))
            with patch.object(sys, "argv", ["install_adapters.py", "--target-dir", str(default_target)]):
                gradebook.main()
            self.assertEqual(default_backups, set(default_target.glob("*.backup_*")))

    def test_visual_evidence_validator_rejects_stale_output(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-visual-hardening-") as temp_name:
            output = Path(temp_name) / "output"
            output.mkdir()
            lesson = output / "教案01_示例.docx"
            lesson.write_bytes(b"docx-bytes")
            qa_report = output / "qa-report.json"
            qa_report.write_text(
                json.dumps(
                    {
                        "status": "passed",
                        "output_dir": str(output.resolve()),
                        "validation": {"output": True},
                        "render": {"page_counts": {lesson.name: 2}},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            checks = {name: "passed" for name in CHECK_NAMES}
            evidence_path = output / "visual-inspection.json"
            for invalid_pages in (
                {},
                {lesson.name: []},
                {lesson.name: [0]},
                {lesson.name: [-1]},
                {lesson.name: [1.5]},
                {lesson.name: [True]},
                [1],
            ):
                with self.subTest(invalid_pages=invalid_pages):
                    with self.assertRaises(ValueError):
                        write_visual_inspection_evidence(
                            output_dir=output,
                            qa_report=qa_report,
                            destination=evidence_path,
                            status="passed",
                            inspected_pages=invalid_pages,
                            checks=checks,
                            notes="invalid direct API input",
                        )
            self.assertFalse(evidence_path.exists())
            write_visual_inspection_evidence(
                output_dir=output,
                qa_report=qa_report,
                destination=evidence_path,
                status="passed",
                inspected_pages={lesson.name: [1, 2]},
                checks=checks,
                notes="Agent inspected representative pages.",
            )
            validate_visual_inspection(output, qa_report, evidence_path)
            valid_evidence = evidence_path.read_bytes()
            empty_evidence = json.loads(valid_evidence.decode("utf-8"))
            empty_evidence["inspected_files"] = []
            empty_evidence["inspected_pages"] = {}
            evidence_path.write_text(json.dumps(empty_evidence, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "at least one inspected"):
                validate_visual_inspection(output, qa_report, evidence_path)
            evidence_path.write_bytes(valid_evidence)
            lesson.write_bytes(b"changed")
            with self.assertRaisesRegex(ValueError, "stale"):
                validate_visual_inspection(output, qa_report, evidence_path)

    def test_xml_template_patch_escapes_visible_text_and_reopens(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-template-patch-hardening-") as temp_name:
            folder = Path(temp_name)
            source = folder / "source.docx"
            target = folder / "patched.docx"
            shutil.copy2(LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.2" / "template.docx", source)
            build_template_patch.build_patch(source, target, [("职业价值观", "A&B <测试>")])
            self.assertTrue(target.is_file())
            from docx import Document

            document = Document(target)
            visible = [paragraph.text for paragraph in document.paragraphs]
            def cell_text(cell):
                return [cell.text, *[value for table in cell.tables for row in table.rows for nested in row.cells for value in cell_text(nested)]]

            visible.extend(value for table in document.tables for row in table.rows for cell in row.cells for value in cell_text(cell))
            self.assertIn("A&B <测试>", "\n".join(visible))

    def test_xml_template_patch_does_not_cross_paragraph_boundaries(self) -> None:
        xml = (
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:body><w:p><w:r><w:t>跨</w:t></w:r></w:p>'
            '<w:p><w:r><w:t>段</w:t></w:r></w:p></w:body></w:document>'
        ).encode("utf-8")
        with self.assertRaisesRegex(ValueError, "found 0"):
            build_template_patch._replace_document_text(xml, [("跨段", "不应替换")])

    def test_render_timeout_is_reported_as_failed_qa(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-render-timeout-") as temp_name:
            output = Path(temp_name)
            (output / "教案01_示例.docx").write_bytes(b"docx")
            timeout = subprocess.TimeoutExpired(["soffice"], 1)
            with patch.object(render_qa, "find_renderer", return_value="soffice"):
                with patch.object(render_qa.subprocess, "run", side_effect=timeout):
                    report = render_qa.render_docx_directory(output, timeout=1)
            self.assertEqual(report["status"], "failed")
            self.assertIn("timed out after 1s", report["errors"][0])

    def test_output_and_render_qa_reject_docx_symlinks_without_opening_or_rendering(self) -> None:
        data = json.loads(
            (ROOT / "tests" / "fixtures" / "lesson-plan-input.json").read_text(encoding="utf-8")
        )
        with tempfile.TemporaryDirectory(prefix="lesson-symlink-qa-") as temp_name:
            output = Path(temp_name) / "output"
            output.mkdir()
            expected_name = validate_output.lesson_filename(
                1, data["lessons"][0]["unit"], data["lessons"][0]["task"]
            )
            fake_docx = Mock()
            fake_docx.name = expected_name
            fake_docx.is_symlink.return_value = True
            with patch.object(validate_output.Path, "glob", return_value=[fake_docx]):
                with patch.object(validate_output, "Document", side_effect=AssertionError("DOCX target opened")) as document:
                    with self.assertRaisesRegex(RuntimeError, "target was not opened"):
                        validate_output.validate_output_dir(output, data)
            document.assert_not_called()

            fake_render_docx = Mock()
            fake_render_docx.name = "lesson.docx"
            fake_render_docx.is_symlink.return_value = True
            with patch.object(render_qa.Path, "glob", return_value=[fake_render_docx]):
                with patch.object(render_qa, "find_renderer", return_value="soffice"):
                    with patch.object(render_qa.subprocess, "run") as run:
                        report = render_qa.render_docx_directory(output)
            self.assertEqual(report["status"], "failed")
            self.assertIn("not rendered or opened", report["errors"][0])
            run.assert_not_called()


if __name__ == "__main__":
    unittest.main()
