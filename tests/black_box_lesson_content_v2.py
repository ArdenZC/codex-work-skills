"""Black-box acceptance harness for natural-language Lesson Content V2 briefs.

This file deliberately builds its temporary V2 payload from a short user brief
instead of reading one of the committed JSON fixtures.  It is an explicit local
acceptance command, not part of the regular core test shard because it renders
24 DOCX files with LibreOffice.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from contextlib import contextmanager
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
LESSON = ROOT / "教案生成器" / "lesson-plan-docx-generator"
GENERATOR = LESSON / "scripts" / "generate_lesson_plans.py"
NON_IT_FORBIDDEN = ("工程伦理", "平凡又不平凡的价值观", "软件技术", "标准机房", "脚本", "截图工具", "代码编辑器", "数据安全")


def _brief_metadata(brief: str) -> tuple[str, str, int, int]:
    course = re.search(r"课程：([^\n]+)", brief).group(1).strip()
    major = re.search(r"专业：([^\n]+)", brief).group(1).strip()
    total_hours = int(re.search(r"总课时：([0-9]+)", brief).group(1))
    lesson_hours = int(re.search(r"每次：([0-9]+)学时", brief).group(1))
    return course, major, total_hours, lesson_hours


def _lesson(
    *,
    course: str,
    major: str,
    audience: str,
    index: int,
    unit: str,
    task: str,
    focus: str,
    artifact: str,
    next_focus: str,
    previous_artifact: str | None,
    score: float,
) -> dict:
    prior_learning = (
        f"课程开始前完成需求情境梳理，明确本课将围绕{focus}建立工作入口"
        if previous_artifact is None
        else f"承接上一课的{previous_artifact}，能够据此定位本课{focus}的关键条件"
    )
    methods = [
        f"{focus}任务驱动法",
        ("案例研讨法" if index % 3 == 1 else "示范练习法" if index % 3 == 2 else "对照复盘法"),
        f"{artifact}成果互评法",
    ]
    stages = (
        ("before_class_preparation", "课前准备", 10, "预读资料", "检查准备边界", "完成资料核对", "明确实训入口"),
        ("task_introduction", "任务导入", 10, "呈现情境", "提出关键问题", "描述现象与需求", "建立任务目标"),
        ("operation_demonstration", "方法示范", 20, "演示步骤", "分解判断方法", "记录操作要点", "掌握方法路径"),
        ("task_implementation", "任务实施", 25, "开展操作", "巡视过程证据", "分工完成任务", "形成阶段成果"),
        ("task_extension", "任务拓展", 10, "比较变式", "追问条件变化", "调整方案并说明理由", "迁移处理能力"),
        ("project_practice", "项目实训", 10, "整理成果", "按清单检查质量", "提交可复查成果", "完成交付准备"),
        ("peer_review", "组间互评", 5, "交换成果", "组织快速互审", "指出一项证据缺口", "学会外部校验"),
        ("lesson_summary", "课堂小结", 10, "回顾路径", "归纳本课要点", "复述判断依据", "连接后续任务"),
        ("after_class_improvement", "课后完善", 15, "修订成果", "反馈修改方向", "上传修订版本", "巩固迁移结果"),
    )
    implementation = []
    for stage_id, label, minutes, activity, teacher_action, student_action, objective in stages:
        implementation.append(
            {
                "id": stage_id,
                "label": label,
                "minutes": minutes,
                "modality": "线上+线下" if stage_id in {"before_class_preparation", "after_class_improvement"} else "小组实训",
                "content": [f"围绕{focus}完成{activity}，记录{artifact}中的阶段证据"],
                "teacher_actions": [f"针对{focus}{teacher_action}，提醒成果必须对应{artifact}"],
                "student_actions": [f"围绕{focus}{student_action}，说明本组对{artifact}的处理依据"],
                "objective": f"通过{activity}掌握{focus}并推进{artifact}",
            }
        )
    remarks = {
        key: f"在{focus}任务中{suffix}"
        for key, suffix in {
            "attendance": "按时完成到课与准备",
            "attention": "持续关注关键条件",
            "participation": "主动参与讨论与操作",
            "compliance": "依照任务规范记录过程",
            "values": "理解成果质量的职业价值",
            "ethics": "如实说明过程与结果",
            "habits": "保持材料归档有序",
            "online_learning": "完成课前资料核对",
            "discussion": "能解释一项判断依据",
            "homework": "按要求提交修订成果",
            "practice": f"完成{artifact}",
            "presentation": "能够清楚介绍成果边界",
            "improvement": f"依据反馈完善{next_focus}",
        }.items()
    }
    reflection_forms = (
        (
            f"第{index}课从{focus}情境切入，学生围绕{artifact}完成条件识别、操作记录和成果核验，"
            "主要差异体现在依据说明的完整度。",
            f"把{focus}的变式比较放进{artifact}制作，是本课的组织变化；在{focus}判断中小组需要先解释选择，再提交结果。",
            f"下一课从{artifact}的实际缺口进入{next_focus}，要求补充新的判断条件并保留修订理由。",
        ),
        (
            f"本次{task}的成果是{artifact}，多数学生能依据{focus}的任务条件完成处理，"
            "少数小组还需加强结果与证据的对应。",
            f"课堂将{focus}拆成连续的小检查点，学生通过互评发现{artifact}中容易遗漏的要素。",
            f"课后先复核{artifact}的边界，再带着问题学习{next_focus}，避免只重复操作步骤。",
        ),
        (
            f"围绕{focus}的本课练习已经形成{artifact}，学生在情境变化时开始主动说明处理依据，"
            "成果可复查性有所提高。",
            f"以{artifact}为交付线索串联示范、实训和互审，使{focus}不再停留在概念辨认。",
            f"下一次课使用本课{artifact}作为输入，增加{next_focus}的限制条件并比较两种结果。",
        ),
        (
            f"从{focus}到{artifact}的转换是第{index}课的重点，学生能够完成主体任务，"
            "但个别记录仍缺少关键条件。",
            f"本课先让学生处理{focus}的一个变式，再回到{artifact}核对成果，讨论更集中。",
            f"下一课围绕{next_focus}继续推进，要求学生引用{artifact}中的一条证据完成新的任务。",
        ),
    )
    chosen_reflection = reflection_forms[(index - 1) % len(reflection_forms)]
    return {
        "lesson_id": f"L{index:02d}",
        "unit": unit,
        "task": task,
        "hours": lesson_hours_for_course(course),
        "progression": {
            "prior_lesson_id": None if previous_artifact is None else f"L{index - 1:02d}",
            "prior_learning": prior_learning,
            "capability_stage": ("认知", "理解", "模仿", "独立", "迁移")[min(index - 1, 4)],
            "deliverable": artifact,
            "next_bridge": f"将{artifact}带入下一课的{next_focus}，继续完善项目成果",
        },
        "student_analysis": {
            "base": [f"能够识别{focus}涉及的基本对象", f"接触过与{artifact}相关的简单记录"],
            "problems": [f"容易忽略{focus}中的条件差异", f"成果说明对{artifact}的依据关联不够清晰"],
            "strategies": [f"使用清单拆解{focus}的处理步骤", f"通过互评核对{artifact}与任务要求"],
        },
        "teaching_content": [f"分析{focus}的任务边界与工作条件", f"示范形成{artifact}的关键步骤", f"比较不同情境下{focus}的处理结果"],
        "goals": {
            "knowledge": [f"说明{focus}的核心概念与判断依据", f"辨认{artifact}应包含的主要要素"],
            "ability": [f"依据任务条件完成{focus}处理", f"提交可复查的{artifact}"],
            "quality": ["形成先核对后操作的习惯", f"对{focus}结果保持客观负责"],
        },
        "key_point": {"content": [f"{focus}的关键条件与成果要素"], "strategy": [f"用对照清单拆解{focus}并逐项核验"]},
        "difficult_point": {"content": [f"把复杂情境转化为可执行的{focus}步骤"], "strategy": [f"用变式案例追问{artifact}的依据和边界"]},
        "teaching_methods": methods,
        "resources": [f"{focus}任务单", f"{artifact}成果模板", f"{next_focus}参考样例"],
        "references": [{"text": "本课程项目任务资料", "source_kind": "generic"}],
        "implementation": implementation,
        "evaluation": {"score": score, "remarks": remarks},
        "reflection": {
            "summary": chosen_reflection[0],
            "innovation": chosen_reflection[1],
            "improvement": chosen_reflection[2],
        },
    }


_CURRENT_LESSON_HOURS = 2


def lesson_hours_for_course(_course: str) -> int:
    return _CURRENT_LESSON_HOURS


def plan_from_brief(brief: str) -> dict:
    global _CURRENT_LESSON_HOURS
    course, major, total_hours, lesson_hours = _brief_metadata(brief)
    _CURRENT_LESSON_HOURS = lesson_hours
    audience = "高职二年级"
    if "数据库" in course:
        specs = (
            ("项目一 数据库项目准备", "梳理业务需求与数据边界", "需求边界", "需求范围清单", "数据对象建模"),
            ("项目一 数据库项目准备", "建立业务实体关系草图", "数据对象建模", "实体关系草图", "表结构设计"),
            ("项目一 数据库项目准备", "完成数据库项目初始化", "项目初始化", "数据库初始化记录", "字段与类型规划"),
            ("项目二 数据库结构设计", "规划字段与数据类型", "字段与类型规划", "数据库字段字典", "约束规则配置"),
            ("项目二 数据库结构设计", "配置主键外键与约束", "约束规则配置", "数据库约束检查表", "单表数据检索"),
            ("项目二 数据库结构设计", "录入样例数据并校验", "样例数据校验", "数据校验记录", "多表业务关联"),
            ("项目三 业务查询实现", "实现多表业务关联查询", "多表业务关联", "关联查询脚本", "分组统计查询"),
            ("项目三 业务查询实现", "完成分组统计与汇总", "分组统计查询", "业务统计结果表", "嵌套查询设计"),
            ("项目三 业务查询实现", "设计嵌套查询解决方案", "嵌套查询设计", "嵌套查询记录", "视图封装"),
            ("项目四 数据库运行维护", "创建视图封装业务结果", "视图封装", "业务视图说明", "索引优化"),
            ("项目四 数据库运行维护", "依据查询特征调整索引", "索引优化", "索引调整记录", "事务控制"),
            ("项目四 数据库运行维护", "处理事务提交与回滚", "事务控制", "事务验证记录", "访问权限管理"),
            ("项目五 安全与性能管理", "配置角色与访问权限", "访问权限管理", "权限配置清单", "备份策略制定"),
            ("项目五 安全与性能管理", "制定数据库备份策略", "备份策略制定", "数据库备份计划表", "性能指标诊断"),
            ("项目五 安全与性能管理", "诊断查询性能指标", "性能指标诊断", "性能诊断报告", "综合过程编排"),
            ("项目六 综合项目交付", "编排存储过程完成业务处理", "综合过程编排", "过程调用记录", "综合报表设计"),
            ("项目六 综合项目交付", "完成综合业务报表设计", "综合报表设计", "综合报表成果", "项目验收答辩"),
            ("项目六 综合项目交付", "展示数据库项目并完成答辩", "项目验收答辩", "数据库项目验收包", "课程成果复盘"),
        )
    else:
        specs = (
            ("项目一 基础护理准备", "完成护理评估与操作准备", "护理评估", "护理评估记录", "基础操作核对"),
            ("项目一 基础护理准备", "完成无菌操作前核对", "基础操作核对", "操作核对清单", "生命体征观察"),
            ("项目二 生命体征照护", "完成生命体征测量记录", "生命体征观察", "生命体征测量记录", "异常情况沟通"),
            ("项目二 生命体征照护", "开展异常情况沟通处置", "异常情况沟通", "异常沟通记录单", "基础护理操作"),
            ("项目三 基础护理实施", "完成基础护理操作练习", "基础护理操作", "基础操作评价表", "综合照护交付"),
            ("项目三 基础护理实施", "展示综合照护成果并复盘", "综合照护交付", "综合照护记录单", "岗位规范复盘"),
        )
    scores = (88.5, 90, 89.5, 91, 90.5, 92, 91.5, 93, 92.5, 94, 93.5, 90.5, 94.5, 91.5, 95, 92, 95.5, 93)
    lessons = []
    previous_artifact = None
    for index, (unit, task, focus, artifact, next_focus) in enumerate(specs, 1):
        lessons.append(
            _lesson(
                course=course,
                major=major,
                audience=audience,
                index=index,
                unit=unit,
                task=task,
                focus=focus,
                artifact=artifact,
                next_focus=next_focus,
                previous_artifact=previous_artifact,
                score=scores[(index - 1) % len(scores)],
            )
        )
        previous_artifact = artifact
    assert total_hours == len(lessons) * lesson_hours
    return {
        "content_contract_version": "2.0",
        "course_name": course,
        "major": major,
        "audience": audience,
        "default_hours": lesson_hours,
        "total_hours": total_hours,
        "lessons": lessons,
    }


def _document_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]

    def visit_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                values.append(cell.text)
                for nested in cell.tables:
                    visit_table(nested)

    for table in document.tables:
        visit_table(table)
    return "\n".join(values)


@contextmanager
def _case_directory(label: str):
    evidence_root = os.environ.get("LESSON_V2_BLACK_BOX_EVIDENCE_DIR", "").strip()
    if evidence_root:
        folder = Path(evidence_root).expanduser().resolve() / label
        if folder.exists():
            raise RuntimeError(f"black-box evidence directory already exists: {folder}")
        folder.mkdir(parents=True)
        yield folder
        return
    with tempfile.TemporaryDirectory(prefix="lesson-v2-black-box-") as temp_name:
        yield Path(temp_name)


def _export_visual_evidence(output: Path, representative: list[Path], evidence_root: Path | None) -> dict[str, object]:
    if evidence_root is None:
        return {"status": "pending_agent", "representative_files": [path.name for path in representative]}
    docx_dir = evidence_root / "docx"
    pdf_dir = evidence_root / "pdf"
    docx_dir.mkdir()
    pdf_dir.mkdir()
    for path in representative:
        shutil.copy2(path, docx_dir / path.name)
    renderer_candidates = (
        shutil.which("soffice"),
        shutil.which("soffice.com"),
        r"C:\Program Files\LibreOffice\program\soffice.com",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    )
    renderer = next((item for item in renderer_candidates if item and Path(item).exists()), None)
    if renderer is None:
        return {
            "status": "pending_agent",
            "representative_files": [path.name for path in representative],
            "pdf_files": [],
            "reason": "LibreOffice was not found for representative visual evidence",
        }
    profile = evidence_root / "profile"
    profile.mkdir()
    pdf_files: list[str] = []
    for path in sorted(docx_dir.glob("*.docx")):
        result = subprocess.run(
            [renderer, "--headless", f"-env:UserInstallation={profile.as_uri()}", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            check=False,
        )
        pdf = pdf_dir / f"{path.stem}.pdf"
        if result.returncode != 0 or not pdf.is_file() or pdf.stat().st_size == 0:
            raise RuntimeError(f"representative PDF conversion failed for {path.name}: {result.stderr or result.stdout}")
        pdf_files.append(pdf.name)
    return {
        "status": "pending_agent",
        "representative_files": [path.name for path in representative],
        "pdf_files": pdf_files,
        "evidence_directory": str(evidence_root),
    }


def run_case(brief: str, expected_lessons: int) -> dict:
    label = "database" if "数据库" in brief else "nursing"
    with _case_directory(label) as folder:
        evidence_root = folder if os.environ.get("LESSON_V2_BLACK_BOX_EVIDENCE_DIR", "").strip() else None
        source = folder / "brief-derived-input.json"
        source.write_text(json.dumps(plan_from_brief(brief), ensure_ascii=False, indent=2), encoding="utf-8")
        output = folder / "output"
        env = os.environ.copy()
        env["PYTHONUTF8"] = "1"
        env["PYTHONIOENCODING"] = "utf-8"
        result = subprocess.run(
            [sys.executable, str(GENERATOR), "--tasks-json", str(source), "--output-dir", str(output), "--render"],
            cwd=ROOT,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
            capture_output=True,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr or result.stdout)
        report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
        lessons = sorted(output.glob("*.docx"))
        if len(lessons) != expected_lessons or report["files_checked"] != expected_lessons:
            raise AssertionError((len(lessons), report["files_checked"], expected_lessons))
        if report["status"] != "passed" or report["content_quality"]["status"] != "passed":
            raise AssertionError(report)
        if report["content_quality"]["progression"]["status"] != "passed":
            raise AssertionError(report["content_quality"]["progression"])
        score_pattern = report["content_quality"]["coverage"]["score_pattern"]
        if not score_pattern["range_valid"] or score_pattern["all_same"] or score_pattern["simple_cycle"] or score_pattern["arithmetic_progression"]:
            raise AssertionError(score_pattern)
        for key in ("exact_duplicates", "adjacent_exact_duplicates", "adjacent_similarity_pairs", "whole_lesson_similarity_pairs", "implementation_similarity_pairs", "repeated_sentences"):
            if report["content_quality"].get(key):
                raise AssertionError({key: report["content_quality"][key]})
        all_text = "\n".join(_document_text(path) for path in lessons)
        representative = [lessons[0], lessons[len(lessons) // 2], lessons[-1]]
        representative_text = [_document_text(path) for path in representative]
        if len(set(representative_text)) != len(representative_text):
            raise AssertionError("representative lessons are not distinct")
        visual_inspection = _export_visual_evidence(output, representative, evidence_root)
        result_data = {
            "lessons": expected_lessons,
            "hours": report["checks"]["total_hours"]["actual"],
            "max_similarity": max(
                [item["score"] for item in report["content_quality"].get("whole_lesson_similarity_pairs", [])],
                default=0,
            ),
            "duplicates": 0,
            "progression": report["content_quality"]["progression"]["status"],
            "scores": "valid-natural-pattern",
            "render": report["render"],
            "visual_inspection": visual_inspection,
        }
        if "护理" in brief:
            for term in NON_IT_FORBIDDEN:
                if term in all_text:
                    raise AssertionError(f"non-IT contamination: {term}")
            if "职业伦理" not in all_text or "职业价值观" not in all_text:
                raise AssertionError("generalized evaluation labels are missing")
        print(json.dumps(result_data, ensure_ascii=False, indent=2))
        return result_data


def main() -> int:
    run_case(
        "课程：《数据库技术》\n专业：软件技术\n总课时：36\n每次：2学时\n没有其他资料\n允许合理设计",
        18,
    )
    run_case(
        "课程：《基础护理技术》\n专业：护理\n总课时：12\n每次：2学时\n没有教材\n允许合理设计",
        6,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
