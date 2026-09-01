"""Synthetic Lesson Content 2.1 acceptance for three representative courses.

This is an executable acceptance harness, not a deterministic Agent-behavior
test.  It uses the existing projectized V2 lesson authoring fixture, converts
the resulting course plan to the 2.1 delivery/material/practice contract, and
then invokes the real DOCX generator and output validator.
"""

from __future__ import annotations

import copy
import json
import os
from pathlib import Path
import subprocess
import sys
import tempfile

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from tests.synthetic_lesson_content_v2_acceptance import synthetic_plan_from_brief


LESSON = ROOT / "教案生成器" / "lesson-plan-docx-generator"
GENERATOR = LESSON / "scripts" / "generate_lesson_plans.py"


COURSES = (
    ("database", "数据库技术", "软件技术", "高职软件技术二年级", "数据库课程标准相关章节"),
    ("nursing", "基础护理技术", "护理", "高职护理二年级", "基础护理技术课程标准相关章节"),
    ("accounting", "会计凭证与账簿实训", "大数据与会计", "高职会计二年级", "会计核算课程标准相关章节"),
)


def _brief(course: str, major: str) -> str:
    # The inherited synthetic authoring fixture defines the complete course
    # before this harness selects the four-lesson 8-hour acceptance slice.
    source_hours = 36 if "数据库" in course else 12
    return f"课程：《{course}》\n专业：{major}\n总课时：{source_hours}\n每次：2学时\n没有教材\n允许合理设计"


def _practice_task(course: str, lessons: list[dict[str, object]]) -> dict[str, object]:
    return {
        "contract_version": "1.0",
        "course_name": course,
        "practice_hours": 4,
        "granularity": "per_task",
        "tasks": [
            {
                "task_id": "PT-01",
                "project_id": "P-01",
                "title": f"完成{course}综合实训成果",
                "lesson_ids": [lesson["lesson_id"] for lesson in lessons[2:]],
                "practice_hours": 4,
                "scenario": f"在{course}岗位情境中完成一项可复查的综合练习。",
                "objectives": ["按步骤完成岗位任务", "留下能够复核的过程证据"],
                "required_inputs": ["课程案例", "前序课堂成果"],
                "tools_or_materials": ["实训环境", "任务记录表"],
                "steps": ["核对情境和输入", "实施任务并记录证据", "整理成果并复盘"],
                "deliverables": ["综合实训成果", "过程记录"],
                "acceptance_criteria": ["成果与任务要求对应", "关键步骤有过程证据"],
                "safety_or_compliance": ["遵守课程所在专业的操作规范"],
            }
        ],
    }


def _to_v21(old: dict[str, object], *, course: str, major: str, audience: str, reference_title: str) -> dict[str, object]:
    data = copy.deepcopy(old)
    lessons = data["lessons"][:4]
    assert isinstance(lessons, list)
    data.update(
        {
            "content_contract_version": "2.1",
            "course_name": course,
            "major": major,
            "audience": audience,
            "default_hours": 2,
            "total_hours": 8,
            "delivery_plan": {"mode": "split_lessons", "total_hours": 8, "theory_hours": 4, "practice_hours": 4},
            "course_materials": {"textbook": None},
            "reference_pool": [
                {
                    "reference_id": "REF-CORE",
                    "reference_type": "standard",
                    "title": reference_title,
                    "source_kind": "generic",
                }
            ],
            "artifact_plan": {"lesson_plans": True, "practice_work_orders": False},
        }
    )
    for index, lesson in enumerate(lessons):
        lesson["hours"] = 2
        lesson["lesson_type"] = "practice" if index >= 2 else "theory"
        lesson["theory_hours"] = 0 if index >= 2 else 2
        lesson["practice_hours"] = 2 if index >= 2 else 0
        lesson["reference_ids"] = ["REF-CORE"]
        lesson["practice_task_ids"] = ["PT-01"] if index >= 2 else []
        lesson.pop("references", None)
    data["lessons"] = lessons
    data["practice_task_contract"] = _practice_task(course, lessons)
    data["outline"] = [
        {
            "lesson_id": lesson["lesson_id"],
            "unit": lesson["unit"],
            "task": lesson["task"],
            "lesson_type": lesson["lesson_type"],
            "hours": lesson["hours"],
            "theory_hours": lesson["theory_hours"],
            "practice_hours": lesson["practice_hours"],
            "prior_learning": lesson["progression"]["prior_learning"],
            "capability_stage": lesson["progression"]["capability_stage"],
            "deliverable": lesson["progression"]["deliverable"],
            "next_bridge": lesson["progression"]["next_bridge"],
            "practice_task_ids": lesson["practice_task_ids"],
        }
        for lesson in lessons
    ]
    return data


def _document_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]

    def visit_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                values.append(cell.text)
                for nested in cell.tables:
                    visit_table(nested)

    for table in document.tables:
        visit_table(table)
    return "\n".join(values)


def _run_case(
    root: Path,
    label: str,
    course: str,
    major: str,
    audience: str,
    reference_title: str,
    *,
    render: bool,
) -> dict[str, object]:
    old = synthetic_plan_from_brief(_brief(course, major))
    payload = _to_v21(old, course=course, major=major, audience=audience, reference_title=reference_title)
    with tempfile.TemporaryDirectory(prefix=f"lesson-v21-{label}-", dir=str(root)) as temp_name:
        folder = Path(temp_name)
        source = folder / "tasks.json"
        source.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        output = folder / "output"
        env = os.environ.copy()
        env["PYTHONUTF8"] = "1"
        env["PYTHONIOENCODING"] = "utf-8"
        env["TEMP"] = str(folder)
        env["TMP"] = str(folder)
        command = [sys.executable, str(GENERATOR), "--tasks-json", str(source), "--output-dir", str(output)]
        if render:
            command.append("--render")
        result = subprocess.run(
            command,
            cwd=ROOT,
            env=env,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError(f"{label}: {result.stderr or result.stdout}")
        report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
        if report["status"] != "passed" or report["content_quality"]["status"] != "passed":
            raise AssertionError({"label": label, "qa": report})
        provenance = report["content_quality"]["reference_provenance"]
        if provenance["cross_lesson_reuse"] != "allowed" or provenance["reuse_frequency"]["REF-CORE"] != 4:
            raise AssertionError({"label": label, "reference_provenance": provenance})
        documents = sorted(output.glob("*.docx"))
        if len(documents) != 4:
            raise AssertionError({"label": label, "docx_count": len(documents)})
        representative = [documents[0], documents[1], documents[-1]]
        representative_text = [_document_text(path) for path in representative]
        if len(set(representative_text)) != 3:
            raise AssertionError(f"{label}: representative DOCX files are not distinct")
        handoff = json.loads((output / "practice-task-contract.json").read_text(encoding="utf-8"))
        if handoff["practice_hours"] != 4 or handoff["tasks"][0]["lesson_ids"] != ["L03", "L04"]:
            raise AssertionError({"label": label, "handoff": handoff})
        return {
            "course": course,
            "lessons": len(documents),
            "hours": report["checks"]["total_hours"]["actual"],
            "reference_reuse": provenance["reuse_frequency"],
            "practice_handoff": report["content_quality"]["practice_handoff"]["status"],
            "render_requested": render,
            "render": report.get("render", {}),
            "representative_docx": [path.name for path in representative],
        }


def main() -> int:
    import argparse

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--work-dir", type=Path, required=True, help="task-owned temporary root on a non-system volume")
    parser.add_argument("--render", action="store_true", help="also run the generator's optional LibreOffice render step")
    args = parser.parse_args()
    args.work_dir.mkdir(parents=True, exist_ok=True)
    results = [_run_case(args.work_dir, *case, render=args.render) for case in COURSES]
    print(json.dumps({"status": "passed", "cases": results}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
