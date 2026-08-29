"""Synthetic acceptance harness for Lesson Content V2 course briefs.

This file deliberately builds a deterministic temporary V2 payload from a short
brief instead of reading one of the committed JSON fixtures. It is synthetic
acceptance evidence, not a true Agent-authored E2E, and is not part of the
regular core test shard because it renders 24 DOCX files with LibreOffice.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from contextlib import contextmanager
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
LESSON = ROOT / "教案生成器" / "lesson-plan-docx-generator"
GENERATOR = LESSON / "scripts" / "generate_lesson_plans.py"
NON_IT_FORBIDDEN = ("工程伦理", "平凡又不平凡的价值观", "软件技术", "标准机房", "脚本", "截图工具", "代码编辑器", "数据安全")
SYNTHETIC_LESSON_MARKERS = (
    "先核对输入边界再开始操作",
    "把条件分支写进处理记录",
    "用反例检验关键判断",
    "让每个步骤都留下可复查证据",
    "按风险高低安排处理顺序",
    "把异常现象与原始条件对应",
    "用小样本先验证处理假设",
    "在结果提交前进行交叉核对",
    "区分事实记录与个人推测",
    "为关键结论保留依据来源",
    "遇到变式时先说明调整理由",
    "把操作规范转成检查清单",
    "用同伴复核发现遗漏条件",
    "对边界结果进行单独标记",
    "把修订前后的差异说清楚",
    "在交付前确认成果可以复现",
    "用对照结果支持质量判断",
    "把后续改进落实到具体动作",
)


def _brief_metadata(brief: str) -> tuple[str, str, int, int]:
    course = re.search(r"课程：([^\n]+)", brief).group(1).strip()
    major = re.search(r"专业：([^\n]+)", brief).group(1).strip()
    total_hours = int(re.search(r"总课时：([0-9]+)", brief).group(1))
    lesson_hours = int(re.search(r"每次：([0-9]+)学时", brief).group(1))
    return course, major, total_hours, lesson_hours


def _lesson(
    *,
    course: str,
    major: str,
    audience: str,
    index: int,
    unit: str,
    task: str,
    focus: str,
    artifact: str,
    next_focus: str,
    previous_artifact: str | None,
    score: float,
) -> dict:
    lesson_marker = SYNTHETIC_LESSON_MARKERS[index - 1]
    prior_learning = (
        f"课程开始前完成需求情境梳理，明确本课将围绕{focus}建立工作入口"
        if previous_artifact is None
        else f"承接上一课的{previous_artifact}，能够据此定位本课{focus}的关键条件"
    )
    methods = [
        f"{focus}任务驱动法",
        ("案例研讨法" if index % 3 == 1 else "示范练习法" if index % 3 == 2 else "对照复盘法"),
        f"{artifact}成果互评法",
    ]
    stages = (
        ("before_class_preparation", "课前准备", 10, "预读资料", "检查准备边界", "完成资料核对", "明确实训入口"),
        ("task_introduction", "任务导入", 10, "呈现情境", "提出关键问题", "描述现象与需求", "建立任务目标"),
        ("operation_demonstration", "方法示范", 20, "演示步骤", "分解判断方法", "记录操作要点", "掌握方法路径"),
        ("task_implementation", "任务实施", 25, "开展操作", "巡视过程证据", "分工完成任务", "形成阶段成果"),
        ("task_extension", "任务拓展", 10, "比较变式", "追问条件变化", "调整方案并说明理由", "迁移处理能力"),
        ("project_practice", "项目实训", 10, "整理成果", "按清单检查质量", "提交可复查成果", "完成交付准备"),
        ("peer_review", "组间互评", 5, "交换成果", "组织快速互审", "指出一项证据缺口", "学会外部校验"),
        ("lesson_summary", "课堂小结", 10, "回顾路径", "归纳本课要点", "复述判断依据", "连接后续任务"),
        ("after_class_improvement", "课后完善", 15, "修订成果", "反馈修改方向", "上传修订版本", "巩固迁移结果"),
    )
    implementation = []
    for stage_id, label, minutes, activity, teacher_action, student_action, objective in stages:
        implementation.append(
            {
                "id": stage_id,
                "label": label,
                "minutes": minutes,
                "modality": "线上+线下" if stage_id in {"before_class_preparation", "after_class_improvement"} else "小组实训",
                "content": [f"{lesson_marker}；围绕{focus}完成{activity}，记录{artifact}中的阶段证据"],
                "teacher_actions": [f"{lesson_marker}；针对{focus}{teacher_action}，提醒成果必须对应{artifact}"],
                "student_actions": [f"{lesson_marker}；围绕{focus}{student_action}，说明本组对{artifact}的处理依据"],
                "objective": f"{lesson_marker}；通过{activity}掌握{focus}并推进{artifact}",
            }
        )
    remarks = {
        key: f"在{focus}任务中{suffix}，本课落实{lesson_marker}"
        for key, suffix in {
            "attendance": "按时完成到课与准备",
            "attention": "持续关注关键条件",
            "participation": "主动参与讨论与操作",
            "compliance": "依照任务规范记录过程",
            "values": "理解成果质量的职业价值",
            "ethics": "如实说明过程与结果",
            "habits": "保持材料归档有序",
            "online_learning": "完成课前资料核对",
            "discussion": "能解释一项判断依据",
            "homework": "按要求提交修订成果",
            "practice": f"完成{artifact}",
            "presentation": "能够清楚介绍成果边界",
            "improvement": f"依据反馈完善{next_focus}",
        }.items()
    }
    reflection_forms = (
        (
            f"第{index}课从{focus}情境切入，学生围绕{artifact}完成条件识别、操作记录和成果核验，"
            "主要差异体现在依据说明的完整度。",
            f"把{focus}的变式比较放进{artifact}制作，是本课的组织变化；在{focus}判断中小组需要先解释选择，再提交结果。",
            f"下一课从{artifact}的实际缺口进入{next_focus}，要求补充新的判断条件并保留修订理由。",
        ),
        (
            f"本次{task}的成果是{artifact}，多数学生能依据{focus}的任务条件完成处理，"
            "少数小组还需加强结果与证据的对应。",
            f"课堂将{focus}拆成连续的小检查点，学生通过互评发现{artifact}中容易遗漏的要素。",
            f"课后先复核{artifact}的边界，再带着问题学习{next_focus}，避免只重复操作步骤。",
        ),
        (
            f"围绕{focus}的本课练习已经形成{artifact}，学生在情境变化时开始主动说明处理依据，"
            "成果可复查性有所提高。",
            f"以{artifact}为交付线索串联示范、实训和互审，使{focus}不再停留在概念辨认。",
            f"下一次课使用本课{artifact}作为输入，增加{next_focus}的限制条件并比较两种结果。",
        ),
        (
            f"从{focus}到{artifact}的转换是第{index}课的重点，学生能够完成主体任务，"
            "但个别记录仍缺少关键条件。",
            f"本课先让学生处理{focus}的一个变式，再回到{artifact}核对成果，讨论更集中。",
            f"下一课围绕{next_focus}继续推进，要求学生引用{artifact}中的一条证据完成新的任务。",
        ),
    )
    chosen_reflection = reflection_forms[(index - 1) % len(reflection_forms)]
    chosen_reflection = tuple(f"{text}本课特别落实{lesson_marker}。" for text in chosen_reflection)
    quality_principles = (
        "先核对输入边界再开始操作",
        "把条件分支写进处理记录",
        "用反例检验关键判断",
        "让每个步骤都留下可复查证据",
        "按风险高低安排处理顺序",
        "把异常现象与原始条件对应",
        "用小样本先验证处理假设",
        "在结果提交前进行交叉核对",
        "区分事实记录与个人推测",
        "为关键结论保留依据来源",
        "遇到变式时先说明调整理由",
        "把操作规范转成检查清单",
        "用同伴复核发现遗漏条件",
        "对边界结果进行单独标记",
        "把修订前后的差异说清楚",
        "在交付前确认成果可以复现",
        "用对照结果支持质量判断",
        "把后续改进落实到具体动作",
    )
    difficulty_strategies = (
        "先用对照案例区分条件，再逐项核验成果依据",
        "把输入、处理和结果拆开，分别检查遗漏",
        "让小组交换成果，按反例清单追查边界",
        "先复现一条异常路径，再回看处理记录",
        "用风险排序决定先验证哪一个关键条件",
        "把现象与证据逐一配对，拒绝只写结论",
        "先设计最小验证样例，再扩展到完整任务",
        "设置提交前的交叉检查点，逐项确认结果",
        "将事实、推断和待确认信息分栏记录",
        "要求每项结论附带来源和复核方式",
        "改变一个条件后比较前后结果并说明原因",
        "把规范要求转换成可勾选的检查清单",
        "安排同伴按不同角色复核同一份成果",
        "单独抽取边界样例，验证特殊结果处理",
        "标记修订前后的差异，说明每次调整目的",
        "用另一组环境复现成果，确认步骤可迁移",
        "将对照结果和判断标准并排呈现",
        "把改进意见转换成下一次可执行的动作",
    )
    ability_patterns = (
        "从输入边界出发，按清单完成{focus}处理",
        "先拆分业务对象，再依据条件完成{focus}",
        "对照正常与异常样例，完成{focus}判断",
        "沿着操作链记录{focus}的每个处理节点",
        "按风险优先顺序处理{focus}并留下结果",
        "把现象和证据配对，说明{focus}处理结论",
        "用最小样例验证{focus}的关键假设",
        "在交付前逐项复核{focus}结果",
        "区分事实和推断后说明{focus}处理",
        "为{focus}结论补充来源和依据",
        "改变一个条件，比较{focus}前后结果",
        "将规范要求转成{focus}检查步骤",
        "交换成果后按角色复核{focus}",
        "单独抽取边界样例验证{focus}",
        "标注修改点并解释{focus}调整",
        "换一组条件复现{focus}成果",
        "用对照结果支持{focus}判断",
        "把改进意见落到下一次{focus}动作",
    )
    delivery_patterns = (
        "按成果清单整理并提交可复查的{artifact}",
        "把对象关系画清后提交{artifact}",
        "用异常样例核对并归档{artifact}",
        "沿操作节点整理一份可追溯的{artifact}",
        "按风险顺序检查后交付{artifact}",
        "将现象、步骤和证据合并为{artifact}",
        "以最小验证样例为依据形成{artifact}",
        "完成提交前复核并发布{artifact}",
        "把事实判断分栏记录到{artifact}",
        "为每项结论补全来源后形成{artifact}",
        "比较条件变化后修订{artifact}",
        "依据规范清单完善{artifact}",
        "完成角色互审后确认{artifact}",
        "单列边界样例并补充到{artifact}",
        "标注修订差异后重新提交{artifact}",
        "在另一组条件下复现并校正{artifact}",
        "用对照结果说明并提交{artifact}",
        "把改进动作落实到新版{artifact}",
    )
    key_content_patterns = (
        "输入边界、处理条件与成果要素的对应关系",
        "业务对象拆分和关系确认的关键依据",
        "正常样例与异常样例的差异识别方法",
        "操作节点记录与过程证据的完整要求",
        "风险优先级和处理顺序的判断标准",
        "现象、步骤、证据和结论的配对规则",
        "最小验证样例与假设检验的关系",
        "提交前复核和成果质量的检查要点",
        "事实记录、推断说明和待确认项的区分",
        "结论来源、依据链和成果可信度的关系",
        "单一条件变化对结果的影响方式",
        "规范条目转化为可执行检查步骤的方法",
        "角色互审中不同观察角度的互补作用",
        "边界样例对特殊结果的验证作用",
        "修订差异与调整理由的清晰表达",
        "跨条件复现对成果可迁移性的检验",
        "对照结果支撑质量判断的证据关系",
        "改进动作与后续成果衔接的落实要求",
    )
    key_strategy_patterns = (
        "先画出输入边界，再用清单核对每个成果要素",
        "先建立对象关系，再检查条件是否闭合",
        "用一组正例和一组反例对照确认判断依据",
        "沿操作节点记录证据，最后回看过程是否完整",
        "先按风险排序，再决定验证和修订顺序",
        "把现象逐条绑定到步骤和证据，避免只写结论",
        "用最小样例试跑假设，再扩展到完整任务",
        "设置提交前检查点，逐项确认成果可复查",
        "用两栏记录事实与推断，再补充待确认信息",
        "要求每项结论标注来源，并安排同伴复核",
        "只改变一个条件，比较前后结果并解释差异",
        "把规范条目改写成可勾选的操作检查单",
        "交换成果并分配不同角色，覆盖不同观察角度",
        "单独抽取边界样例，检查特殊结果处理",
        "并排展示修改前后内容，说明每处调整目的",
        "换一组条件复现成果，确认步骤仍然有效",
        "将对照结果和判断标准并排呈现再下结论",
        "把反馈意见写成下一课可以执行的具体动作",
    )
    difficulty_content_patterns = (
        "把输入边界的复杂条件转成可以逐项检查的处理路径",
        "把多个业务对象的关系转成可复核的判断顺序",
        "把正反样例的差异转成可执行的判定规则",
        "把连续操作中的隐含节点转成可追溯记录",
        "把多项风险约束转成有先后次序的处理步骤",
        "把零散现象与证据转成完整的结果说明",
        "把抽象假设转成可以反复运行的最小验证任务",
        "把交付前的遗漏转成明确的复核清单",
        "把事实和推断混杂的表述转成两类记录",
        "把缺少来源的结论转成有证据链的成果",
        "把条件变化带来的差异转成可比较的结果",
        "把规范要求转成不会遗漏步骤的操作路径",
        "把单一视角的检查转成多角色互审流程",
        "把普通样例覆盖不到的情况转成边界验证",
        "把模糊的修改过程转成前后差异说明",
        "把一次性成果转成可以跨条件复现的流程",
        "把孤立的对照结果转成有标准支撑的判断",
        "把笼统的改进意见转成下一次可执行动作",
    )
    knowledge_patterns = (
        "说清输入边界包含哪些条件以及它们的作用",
        "说明业务对象之间如何建立关系并保持一致",
        "辨别正例和反例分别对应什么判断结果",
        "解释过程记录为什么需要覆盖每个节点",
        "概括风险优先级如何影响处理顺序",
        "说明现象、步骤和证据怎样共同支持结论",
        "理解最小样例如何用来检验一个假设",
        "掌握交付前复核应覆盖的质量检查点",
        "区分事实、推断以及仍需确认的信息",
        "说明来源和依据如何支撑成果可信度",
        "分析单一条件变化会怎样影响结果",
        "理解规范条目如何变成可执行步骤",
        "说明不同角色的观察角度如何互补",
        "识别边界样例为什么需要单独验证",
        "理解修改差异和调整理由的对应关系",
        "说明跨条件复现如何检验流程迁移性",
        "解释对照结果怎样支撑质量判断",
        "理解改进动作如何衔接后续成果",
    )
    knowledge_artifact_patterns = (
        "列出输入边界和成果要素的对应项",
        "画出对象关系并标注关键关联",
        "整理正反样例的判定依据表",
        "列出过程节点和必备证据项",
        "制作风险排序与处理步骤对照表",
        "整理现象、步骤、证据和结论关系",
        "写出最小样例及其预期判断结果",
        "列明提交前需要复核的成果要素",
        "把事实、推断和待确认项分开记录",
        "为各项结论补充对应的来源说明",
        "整理条件变化前后的结果对照",
        "把规范要求整理成执行检查单",
        "形成不同角色的互审分工表",
        "单独汇总边界样例及处理结果",
        "列出修改前后差异和调整原因",
        "形成跨条件复现所需的步骤说明",
        "整理对照结果和判断标准的关系",
        "形成下一步改进动作清单",
    )
    quality_principle = quality_principles[index - 1]
    difficulty_strategy = difficulty_strategies[index - 1]
    ability_pattern = ability_patterns[index - 1].format(focus=focus)
    delivery_pattern = delivery_patterns[index - 1].format(artifact=artifact)
    key_content_pattern = key_content_patterns[index - 1]
    key_strategy_pattern = key_strategy_patterns[index - 1]
    difficulty_content_pattern = difficulty_content_patterns[index - 1]
    knowledge_pattern = knowledge_patterns[index - 1]
    knowledge_artifact_pattern = knowledge_artifact_patterns[index - 1]
    return {
        "lesson_id": f"L{index:02d}",
        "unit": unit,
        "task": task,
        "hours": lesson_hours_for_course(course),
        "progression": {
            "prior_lesson_id": None if previous_artifact is None else f"L{index - 1:02d}",
            "prior_learning": prior_learning,
            "capability_stage": ("认知", "理解", "模仿", "独立", "迁移")[min(index - 1, 4)],
            "deliverable": artifact,
            "next_bridge": f"将{artifact}带入下一课的{next_focus}，继续完善项目成果",
        },
        "student_analysis": {
            "base": [f"能够{knowledge_pattern}，并识别{focus}涉及的基本对象", f"接触过{knowledge_artifact_pattern}对应的简单记录"],
            "problems": [f"容易在{key_content_pattern}中遗漏条件", f"成果说明常缺少{key_strategy_pattern}形成的依据关联"],
            "strategies": [f"采用{difficulty_strategy}来拆解本课处理步骤", f"按照{key_strategy_pattern}互评并核对任务要求"],
        },
        "teaching_content": [
            f"{lesson_marker}；分析{focus}的任务边界与工作条件",
            f"{lesson_marker}；示范形成{artifact}的关键步骤",
            f"{lesson_marker}；比较不同情境下{focus}的处理结果",
        ],
        "goals": {
            "knowledge": [f"{knowledge_pattern}，并联系{focus}的核心依据", f"{knowledge_artifact_pattern}，对应{artifact}"],
            "ability": [ability_pattern, delivery_pattern],
            "quality": [f"{quality_principle}，并在任务中留下可追溯依据", f"对{focus}结果保持{quality_principle}"],
        },
        "key_point": {"content": [f"{key_content_pattern}，对应{focus}的成果要求"], "strategy": [f"{key_strategy_pattern}，再核对{artifact}"]},
        "difficult_point": {"content": [f"{difficulty_content_pattern}，最后落到{focus}步骤"], "strategy": [f"{difficulty_strategy}，再回看{artifact}的依据和边界"]},
        "teaching_methods": methods,
        "resources": [f"{focus}任务单", f"{artifact}成果模板", f"{next_focus}参考样例"],
        "references": [{"text": "本课程项目任务资料", "source_kind": "generic"}],
        "implementation": implementation,
        "evaluation": {"score": score, "remarks": remarks},
        "reflection": {
            "summary": chosen_reflection[0],
            "innovation": chosen_reflection[1],
            "improvement": chosen_reflection[2],
        },
    }


_CURRENT_LESSON_HOURS = 2


def lesson_hours_for_course(_course: str) -> int:
    return _CURRENT_LESSON_HOURS


def synthetic_plan_from_brief(brief: str) -> dict:
    global _CURRENT_LESSON_HOURS
    course, major, total_hours, lesson_hours = _brief_metadata(brief)
    _CURRENT_LESSON_HOURS = lesson_hours
    audience = "高职二年级"
    if "数据库" in course:
        specs = (
            ("项目一 数据库项目准备", "梳理业务需求与数据边界", "需求边界", "需求范围清单", "数据对象建模"),
            ("项目一 数据库项目准备", "建立业务实体关系草图", "数据对象建模", "实体关系草图", "表结构设计"),
            ("项目一 数据库项目准备", "完成数据库项目初始化", "项目初始化", "数据库初始化记录", "字段与类型规划"),
            ("项目二 数据库结构设计", "规划字段与数据类型", "字段与类型规划", "数据库字段字典", "约束规则配置"),
            ("项目二 数据库结构设计", "配置主键外键与约束", "约束规则配置", "数据库约束检查表", "单表数据检索"),
            ("项目二 数据库结构设计", "录入样例数据并校验", "样例数据校验", "数据校验记录", "多表业务关联"),
            ("项目三 业务查询实现", "实现多表业务关联查询", "多表业务关联", "关联查询脚本", "分组统计查询"),
            ("项目三 业务查询实现", "完成分组统计与汇总", "分组统计查询", "业务统计结果表", "嵌套查询设计"),
            ("项目三 业务查询实现", "设计嵌套查询解决方案", "嵌套查询设计", "嵌套查询记录", "视图封装"),
            ("项目四 数据库运行维护", "创建视图封装业务结果", "视图封装", "业务视图说明", "索引优化"),
            ("项目四 数据库运行维护", "依据查询特征调整索引", "索引优化", "索引调整记录", "事务控制"),
            ("项目四 数据库运行维护", "处理事务提交与回滚", "事务控制", "事务验证记录", "访问权限管理"),
            ("项目五 安全与性能管理", "配置角色与访问权限", "访问权限管理", "权限配置清单", "备份策略制定"),
            ("项目五 安全与性能管理", "制定数据库备份策略", "备份策略制定", "数据库备份计划表", "性能指标诊断"),
            ("项目五 安全与性能管理", "诊断查询性能指标", "性能指标诊断", "性能诊断报告", "综合过程编排"),
            ("项目六 综合项目交付", "编排存储过程完成业务处理", "综合过程编排", "过程调用记录", "综合报表设计"),
            ("项目六 综合项目交付", "完成综合业务报表设计", "综合报表设计", "综合报表成果", "项目验收答辩"),
            ("项目六 综合项目交付", "展示数据库项目并完成答辩", "项目验收答辩", "数据库项目验收包", "课程成果复盘"),
        )
    else:
        specs = (
            ("项目一 基础护理准备", "完成护理评估与操作准备", "护理评估", "护理评估记录", "基础操作核对"),
            ("项目一 基础护理准备", "完成无菌操作前核对", "基础操作核对", "操作核对清单", "生命体征观察"),
            ("项目二 生命体征照护", "完成生命体征测量记录", "生命体征观察", "生命体征测量记录", "异常情况沟通"),
            ("项目二 生命体征照护", "开展异常情况沟通处置", "异常情况沟通", "异常沟通记录单", "基础护理操作"),
            ("项目三 基础护理实施", "完成基础护理操作练习", "基础护理操作", "基础操作评价表", "综合照护交付"),
            ("项目三 基础护理实施", "展示综合照护成果并复盘", "综合照护交付", "综合照护记录单", "岗位规范复盘"),
        )
    scores = (88.5, 90, 89.5, 91, 90.5, 92, 91.5, 93, 92.5, 94, 93.5, 90.5, 94.5, 91.5, 95, 92, 95.5, 93)
    lessons = []
    previous_artifact = None
    for index, (unit, task, focus, artifact, next_focus) in enumerate(specs, 1):
        lessons.append(
            _lesson(
                course=course,
                major=major,
                audience=audience,
                index=index,
                unit=unit,
                task=task,
                focus=focus,
                artifact=artifact,
                next_focus=next_focus,
                previous_artifact=previous_artifact,
                score=scores[(index - 1) % len(scores)],
            )
        )
        previous_artifact = artifact
    assert total_hours == len(lessons) * lesson_hours
    return {
        "content_contract_version": "2.0",
        "course_name": course,
        "major": major,
        "audience": audience,
        "default_hours": lesson_hours,
        "total_hours": total_hours,
        "lessons": lessons,
    }


def _document_text(path: Path) -> str:
    document = Document(path)
    values = [paragraph.text for paragraph in document.paragraphs]

    def visit_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                values.append(cell.text)
                for nested in cell.tables:
                    visit_table(nested)

    for table in document.tables:
        visit_table(table)
    return "\n".join(values)


@contextmanager
def _case_directory(label: str):
    evidence_root = os.environ.get("LESSON_V2_SYNTHETIC_EVIDENCE_DIR", "").strip()
    if evidence_root:
        folder = Path(evidence_root).expanduser().resolve() / label
        if folder.exists():
            raise RuntimeError(f"synthetic evidence directory already exists: {folder}")
        folder.mkdir(parents=True)
        yield folder
        return
    with tempfile.TemporaryDirectory(prefix="lesson-v2-synthetic-") as temp_name:
        yield Path(temp_name)


def _export_visual_evidence(output: Path, representative: list[Path], evidence_root: Path | None) -> dict[str, object]:
    if evidence_root is None:
        return {"status": "not_exported", "scope": "synthetic_acceptance", "representative_files": [path.name for path in representative]}
    docx_dir = evidence_root / "docx"
    pdf_dir = evidence_root / "pdf"
    docx_dir.mkdir()
    pdf_dir.mkdir()
    for path in representative:
        shutil.copy2(path, docx_dir / path.name)
    renderer_candidates = (
        shutil.which("soffice"),
        shutil.which("soffice.com"),
        r"C:\Program Files\LibreOffice\program\soffice.com",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    )
    renderer = next((item for item in renderer_candidates if item and Path(item).exists()), None)
    if renderer is None:
        return {
            "status": "not_executed",
            "scope": "synthetic_acceptance",
            "representative_files": [path.name for path in representative],
            "pdf_files": [],
            "reason": "LibreOffice was not found for representative visual evidence",
        }
    profile = evidence_root / "profile"
    profile.mkdir()
    pdf_files: list[str] = []
    for path in sorted(docx_dir.glob("*.docx")):
        result = subprocess.run(
            [renderer, "--headless", f"-env:UserInstallation={profile.as_uri()}", "--convert-to", "pdf", "--outdir", str(pdf_dir), str(path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            check=False,
        )
        pdf = pdf_dir / f"{path.stem}.pdf"
        if result.returncode != 0 or not pdf.is_file() or pdf.stat().st_size == 0:
            raise RuntimeError(f"representative PDF conversion failed for {path.name}: {result.stderr or result.stdout}")
        pdf_files.append(pdf.name)
    return {
        "status": "rendered_for_external_inspection",
        "scope": "synthetic_acceptance",
        "representative_files": [path.name for path in representative],
        "pdf_files": pdf_files,
        "evidence_directory": str(evidence_root),
    }


def run_case(brief: str, expected_lessons: int) -> dict:
    label = "database" if "数据库" in brief else "nursing"
    with _case_directory(label) as folder:
        evidence_root = folder if os.environ.get("LESSON_V2_SYNTHETIC_EVIDENCE_DIR", "").strip() else None
        source = folder / "brief-derived-input.json"
        source.write_text(json.dumps(synthetic_plan_from_brief(brief), ensure_ascii=False, indent=2), encoding="utf-8")
        output = folder / "output"
        env = os.environ.copy()
        env["PYTHONUTF8"] = "1"
        env["PYTHONIOENCODING"] = "utf-8"
        result = subprocess.run(
            [sys.executable, str(GENERATOR), "--tasks-json", str(source), "--output-dir", str(output), "--render"],
            cwd=ROOT,
            text=True,
            encoding="utf-8",
            errors="replace",
            env=env,
            capture_output=True,
            check=False,
        )
        if result.returncode != 0:
            raise RuntimeError(result.stderr or result.stdout)
        report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
        lessons = sorted(output.glob("*.docx"))
        if len(lessons) != expected_lessons or report["files_checked"] != expected_lessons:
            raise AssertionError((len(lessons), report["files_checked"], expected_lessons))
        if report["status"] != "passed" or report["content_quality"]["status"] != "passed":
            raise AssertionError(report)
        if report["content_quality"]["progression"]["status"] != "passed":
            raise AssertionError(report["content_quality"]["progression"])
        score_pattern = report["content_quality"]["coverage"]["score_pattern"]
        if not score_pattern["range_valid"] or score_pattern["all_same"] or score_pattern["simple_cycle"] or score_pattern["arithmetic_progression"]:
            raise AssertionError(score_pattern)
        for key in ("exact_duplicates", "adjacent_exact_duplicates", "adjacent_similarity_pairs", "whole_lesson_similarity_pairs", "implementation_similarity_pairs", "repeated_sentences"):
            if report["content_quality"].get(key):
                raise AssertionError({key: report["content_quality"][key]})
        all_text = "\n".join(_document_text(path) for path in lessons)
        representative = [lessons[0], lessons[len(lessons) // 2], lessons[-1]]
        representative_text = [_document_text(path) for path in representative]
        if len(set(representative_text)) != len(representative_text):
            raise AssertionError("representative lessons are not distinct")
        visual_inspection = _export_visual_evidence(output, representative, evidence_root)
        result_data = {
            "lessons": expected_lessons,
            "hours": report["checks"]["total_hours"]["actual"],
            "max_similarity": max(
                [item["score"] for item in report["content_quality"].get("whole_lesson_similarity_pairs", [])],
                default=0,
            ),
            "duplicates": 0,
            "progression": report["content_quality"]["progression"]["status"],
            "scores": "valid-natural-pattern",
            "render": report["render"],
            "visual_inspection": visual_inspection,
        }
        if "护理" in brief:
            for term in NON_IT_FORBIDDEN:
                if term in all_text:
                    raise AssertionError(f"non-IT contamination: {term}")
            if "职业伦理" not in all_text or "职业价值观" not in all_text:
                raise AssertionError("generalized evaluation labels are missing")
        print(json.dumps(result_data, ensure_ascii=False, indent=2))
        return result_data


def main() -> int:
    run_case(
        "课程：《数据库技术》\n专业：软件技术\n总课时：36\n每次：2学时\n没有其他资料\n允许合理设计",
        18,
    )
    run_case(
        "课程：《基础护理技术》\n专业：护理\n总课时：12\n每次：2学时\n没有教材\n允许合理设计",
        6,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
