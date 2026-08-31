from __future__ import annotations

import copy
from contextlib import contextmanager, redirect_stderr
import hashlib
import io
import importlib.util
import json
import os
import shutil
import sys
import tempfile
import types
import unittest
import unicodedata
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.table import _Cell


ROOT = Path(__file__).resolve().parents[1]
LESSON = ROOT / "教案生成器" / "lesson-plan-docx-generator"
SCRIPTS = LESSON / "scripts"
V10_TEMPLATE = LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "template.docx"
V10_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.0.0" / "manifest.yaml"
V110_TEMPLATE = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "template.docx"
V110_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.0" / "manifest.yaml"
V112_TEMPLATE = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.2" / "template.docx"
V112_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.2" / "manifest.yaml"
V111_TEMPLATE = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.1" / "template.docx"
V111_MANIFEST = LESSON / "assets" / "templates" / "lesson-plan" / "v1.1.1" / "manifest.yaml"

_LESSON_NAMESPACE = "lesson_v2"
_MISSING = object()
_LESSON_DEPENDENCIES = {
    "semantic_bookmarks": (),
    "content_contract": (),
    "bookmark_utils": ("semantic_bookmarks",),
    "content_quality": ("content_contract",),
    "path_safety": (),
    "render_qa": (),
    "record_visual_inspection": ("path_safety",),
    "package_common": ("semantic_bookmarks", "content_contract"),
    "validate_template": ("bookmark_utils", "package_common"),
    "validate_output": (
        "bookmark_utils",
        "content_contract",
        "content_quality",
        "package_common",
        "path_safety",
        "render_qa",
    ),
    "generate_lesson_plans": (
        "bookmark_utils",
        "content_contract",
        "content_quality",
        "package_common",
        "path_safety",
        "validate_output",
        "validate_template",
    ),
}

_lesson_package = types.ModuleType(_LESSON_NAMESPACE)
_lesson_package.__path__ = [str(SCRIPTS)]
sys.modules[_LESSON_NAMESPACE] = _lesson_package


def _load_lesson_module(name: str):
    """Load Lesson modules without publishing their generic import names."""

    fullname = f"{_LESSON_NAMESPACE}.{name}"
    existing = sys.modules.get(fullname)
    if existing is not None:
        return existing
    for dependency in _LESSON_DEPENDENCIES.get(name, ()):
        _load_lesson_module(dependency)

    spec = importlib.util.spec_from_file_location(fullname, SCRIPTS / f"{name}.py")
    if spec is None or spec.loader is None:
        raise ImportError(f"Unable to load Lesson module: {name}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[fullname] = module
    aliases = {name, *_LESSON_DEPENDENCIES.get(name, ())}
    saved = {alias: sys.modules.get(alias, _MISSING) for alias in aliases}
    try:
        for alias in aliases:
            sys.modules[alias] = sys.modules[f"{_LESSON_NAMESPACE}.{alias}"]
        spec.loader.exec_module(module)
    except Exception:
        sys.modules.pop(fullname, None)
        raise
    finally:
        for alias, previous in saved.items():
            if previous is _MISSING:
                sys.modules.pop(alias, None)
            else:
                sys.modules[alias] = previous
    return module


@contextmanager
def isolated_generic_imports(scripts: Path, module_names: tuple[str, ...]):
    """Exercise a legacy package import while restoring process-global state."""

    original_path = sys.path[:]
    saved_modules = {name: sys.modules.get(name, _MISSING) for name in module_names}
    try:
        for name in module_names:
            sys.modules.pop(name, None)
        sys.path.insert(0, str(scripts))
        yield
    finally:
        sys.path[:] = original_path
        for name, previous in saved_modules.items():
            if previous is _MISSING:
                sys.modules.pop(name, None)
            else:
                sys.modules[name] = previous


lesson_bookmark_utils = _load_lesson_module("bookmark_utils")
lesson_content_contract = _load_lesson_module("content_contract")
lesson_content_quality = _load_lesson_module("content_quality")
lesson_package_common = _load_lesson_module("package_common")
lesson_path_safety = _load_lesson_module("path_safety")
lesson_visual_inspection = _load_lesson_module("record_visual_inspection")
lesson_validate_template = _load_lesson_module("validate_template")
lesson_output = _load_lesson_module("validate_output")
lesson_generator = _load_lesson_module("generate_lesson_plans")

bookmark_boundary_locations = lesson_bookmark_utils.bookmark_boundary_locations
bookmark_parent_cell = lesson_bookmark_utils.bookmark_parent_cell
bookmark_parent_paragraph = lesson_bookmark_utils.bookmark_parent_paragraph
find_bookmark = lesson_bookmark_utils.find_bookmark
format_evaluation_values = lesson_content_contract.format_evaluation_values
format_implementation = lesson_content_contract.format_implementation
format_reflection = lesson_content_contract.format_reflection
lesson_content_field_values = lesson_content_contract.lesson_content_field_values
lesson_header_values = lesson_content_contract.lesson_header_values
ContentQualityError = lesson_content_quality.ContentQualityError
assess_content_quality = lesson_content_quality.assess_content_quality
validate_content_quality = lesson_content_quality.validate_content_quality
detect_non_it_contamination = lesson_content_quality.detect_non_it_contamination
atomic_commit_candidate = lesson_generator.atomic_commit_candidate
atomic_commit_candidate_with_external_qa = lesson_generator.atomic_commit_candidate_with_external_qa
field_bookmark = lesson_package_common.field_bookmark
implementation_bookmarks = lesson_package_common.implementation_bookmarks
load_manifest = lesson_package_common.load_manifest
reflection_bookmarks = lesson_package_common.reflection_bookmarks
required_bookmarks = lesson_package_common.required_bookmarks
score_breakdown = lesson_package_common.score_breakdown
assert_output_path_safe = lesson_path_safety.assert_output_path_safe
paths_equal = lesson_path_safety.paths_equal
paths_overlap = lesson_path_safety.paths_overlap
filesystem_case_sensitive = lesson_path_safety.filesystem_case_sensitive
write_visual_inspection_evidence = lesson_visual_inspection.write_visual_inspection_evidence
manifest_field_text = lesson_output.manifest_field_text


def run_script(script: Path, *args: str) -> "subprocess.CompletedProcess[str]":
    import subprocess

    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    if "--skip-template-validation" in args or "--skip-output-validation" in args:
        # Existing compatibility tests exercise the legacy switches explicitly;
        # the dedicated unsafe-skip test overrides this to assert fail-closed.
        env.setdefault("LESSON_ALLOW_UNSAFE_VALIDATION_SKIP", "1")
    return subprocess.run(
        [sys.executable, str(script), *args],
        cwd=ROOT,
        env=env,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        check=False,
    )


def load_fixture(name: str) -> dict:
    return json.loads((ROOT / "tests" / "fixtures" / name).read_text(encoding="utf-8"))


def reference_reuse_fixture(references: list[dict[str, str]]) -> dict:
    """Build the six-lesson reference reuse regression fixture from the real V2 baseline."""

    data = copy.deepcopy(load_fixture("lesson-plan-content-v2-it.json"))
    data["lessons"] = data["lessons"][:6]
    data["total_hours"] = 12
    for lesson in data["lessons"]:
        lesson["references"] = copy.deepcopy(references)
    return data


def write_payload(folder: Path, payload: dict, name: str = "tasks.json") -> Path:
    source = folder / name
    source.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return source


def document_text(document: Document) -> str:
    values = [paragraph.text for paragraph in document.paragraphs]

    def visit_table(table) -> None:
        for row in table.rows:
            for cell in row.cells:
                values.append(cell.text)
                for nested in cell.tables:
                    visit_table(nested)

    for table in document.tables:
        visit_table(table)
    for section in document.sections:
        values.append(section.header._element.text or "")
        values.append(section.footer._element.text or "")
    return "\n".join(values)


def bookmark_text(document: Document, name: str) -> str:
    record = find_bookmark(document, name)
    if record is None:
        raise AssertionError(f"missing bookmark {name}")
    cell_element = bookmark_parent_cell(document, record)
    if cell_element is not None:
        return _Cell(cell_element, document).text.strip()
    paragraph_element = bookmark_parent_paragraph(document, record)
    if paragraph_element is not None:
        return paragraph_element.text.strip()
    raise AssertionError(f"bookmark {name} has no text container")


class LessonContentV2Mixin:
    def test_lesson_and_gradebook_imports_are_isolated_in_one_interpreter(self) -> None:
        grade_scripts = ROOT / "平时成绩记分册生成器" / "course-gradebook-generator" / "scripts"
        grade_modules = ("package_common", "named_range_contracts", "named_range_utils", "xls_named_range_utils")
        lesson_modules = (
            "package_common",
            "semantic_bookmarks",
            "content_contract",
            "bookmark_utils",
            "content_quality",
            "path_safety",
            "render_qa",
            "validate_template",
            "validate_output",
            "generate_lesson_plans",
        )
        original_path = sys.path[:]

        with isolated_generic_imports(grade_scripts, grade_modules):
            import package_common as gradebook_package_common

            self.assertTrue(hasattr(gradebook_package_common, "calculate_expected_total"))
            self.assertIsNot(gradebook_package_common, lesson_package_common)
            self.assertEqual(sys.path[1:], original_path)

        self.assertEqual(sys.path, original_path)
        self.assertIs(sys.modules[f"{_LESSON_NAMESPACE}.package_common"], lesson_package_common)

        with isolated_generic_imports(SCRIPTS, lesson_modules):
            import package_common as generic_lesson_package_common

            self.assertTrue(hasattr(generic_lesson_package_common, "validate_content_v2_input"))
            self.assertIsNot(generic_lesson_package_common, lesson_package_common)
            self.assertEqual(sys.path[1:], original_path)

        self.assertEqual(sys.path, original_path)
        self.assertIs(sys.modules[f"{_LESSON_NAMESPACE}.package_common"], lesson_package_common)

    def test_sparse_input_fails_closed_through_real_generator(self) -> None:
        sparse = {
            "course_name": "软件测试实训",
            "lessons": [{"unit": "项目一 测试准备", "task": "完成环境检查", "hours": 2}],
        }
        with tempfile.TemporaryDirectory(prefix="lesson-v2-sparse-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, sparse)
            output = folder / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            self.assertIn(
                "Legacy sparse lesson content is no longer accepted for production generation.",
                result.stderr,
            )
            self.assertIn("Regenerate tasks JSON using the Lesson Content V2 Skill workflow.", result.stderr)
            self.assertFalse(output.exists())

    def test_schema_and_classroom_time_contracts_fail_through_real_generator(self) -> None:
        cases = (
            ("missing-progression", lambda payload: payload["lessons"][0].pop("progression"), "Input schema validation failed"),
            ("missing-stage", lambda payload: payload["lessons"][0]["implementation"].pop(), "Input schema validation failed"),
            (
                "wrong-minutes",
                lambda payload: payload["lessons"][0]["implementation"][1].__setitem__("minutes", 11),
                "in-class implementation minutes must equal hours*45",
            ),
            ("empty-teaching-content", lambda payload: payload["lessons"][0].__setitem__("teaching_content", []), "Input schema validation failed"),
            (
                "future-prior",
                lambda payload: payload["lessons"][1]["progression"].__setitem__("prior_lesson_id", "L03"),
                "prior_lesson_id must reference an earlier lesson",
            ),
            (
                "self-prior",
                lambda payload: payload["lessons"][1]["progression"].__setitem__("prior_lesson_id", "L02"),
                "prior_lesson_id must reference an earlier lesson",
            ),
            (
                "unknown-prior",
                lambda payload: payload["lessons"][1]["progression"].__setitem__("prior_lesson_id", "missing"),
                "prior_lesson_id must reference an earlier lesson",
            ),
        )
        for label, mutate, expected in cases:
            with self.subTest(case=label), tempfile.TemporaryDirectory(prefix=f"lesson-v2-contract-{label}-") as temp_name:
                folder = Path(temp_name)
                payload = load_fixture("lesson-plan-input.json")
                mutate(payload)
                source = write_payload(folder, payload)
                output = folder / "output"
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                )
                self.assertNotEqual(result.returncode, 0)
                self.assertIn(expected, result.stderr)
                self.assertFalse(output.exists())

    def test_out_of_class_minutes_use_dynamic_sanity_bounds(self) -> None:
        base = load_fixture("lesson-plan-input.json")

        def one_lesson(hours: int, classroom: tuple[int, ...]) -> dict:
            payload = copy.deepcopy(base)
            payload["lessons"] = [payload["lessons"][0]]
            payload["lessons"][0]["hours"] = hours
            payload["total_hours"] = hours
            for stage, minutes in zip(payload["lessons"][0]["implementation"][1:-1], classroom):
                stage["minutes"] = minutes
            return payload

        for hours, classroom in (
            (1, (5, 10, 10, 5, 5, 5, 5)),
            (2, (10, 20, 25, 10, 10, 5, 10)),
            (4, (20, 40, 50, 20, 20, 10, 20)),
        ):
            with self.subTest(hours=hours):
                lesson_generator.validate_content_v2_input(one_lesson(hours, classroom))

        before = one_lesson(2, (10, 20, 25, 10, 10, 5, 10))
        before["lessons"][0]["implementation"][0]["minutes"] = 999
        with self.assertRaisesRegex(
            ValueError,
            r"lesson_id=L01 stage=before_class_preparation actual=999 limit=90",
        ):
            lesson_generator.validate_content_v2_input(before)

        after = one_lesson(2, (10, 20, 25, 10, 10, 5, 10))
        after["lessons"][0]["implementation"][-1]["minutes"] = 999
        with self.assertRaisesRegex(
            ValueError,
            r"lesson_id=L01 stage=after_class_improvement actual=999 limit=90",
        ):
            lesson_generator.validate_content_v2_input(after)

        total = one_lesson(1, (5, 10, 10, 5, 5, 5, 5))
        total["lessons"][0]["implementation"][0]["minutes"] = 50
        total["lessons"][0]["implementation"][-1]["minutes"] = 50
        with self.assertRaisesRegex(
            ValueError,
            r"lesson_id=L01 stage=out_of_class_total actual=100 limit=90",
        ):
            lesson_generator.validate_content_v2_input(total)

        zero_middle = one_lesson(1, (5, 10, 10, 5, 5, 5, 5))
        zero_middle["lessons"][0]["implementation"][1]["minutes"] = 0
        zero_middle["lessons"][0]["implementation"][2]["minutes"] = 15
        with self.assertRaisesRegex(ValueError, "must be positive for in-class stages"):
            lesson_generator.validate_content_v2_input(zero_middle)

    def test_lesson_hours_schema_accepts_whole_numbers_and_rejects_fractions(self) -> None:
        base = load_fixture("lesson-plan-input.json")
        for value in (1, 2, 4, 12, 1.0, 2.0, "1", "2", "2.0", "12.0"):
            with self.subTest(valid=value):
                payload = copy.deepcopy(base)
                payload["default_hours"] = value
                lesson_generator.validate_content_v2_input(payload)
        for value in (0.5, 1.5, 2.25, "0.5", "1.5", "2.25", 0, "0", "0.0", -1, "-1", "abc", "", " 2", "2 ", "NaN", "Infinity"):
            with self.subTest(invalid=value):
                payload = copy.deepcopy(base)
                payload["default_hours"] = value
                with self.assertRaises(ValueError):
                    lesson_generator.validate_content_v2_input(payload)

    def test_evaluation_score_boundaries_use_half_point_contract(self) -> None:
        cases = ((84.5, False), (85, True), (85.5, True), (96, True), (96.5, False))
        with tempfile.TemporaryDirectory(prefix="lesson-v2-score-boundaries-") as temp_name:
            folder = Path(temp_name)
            for score, should_pass in cases:
                with self.subTest(score=score):
                    payload = load_fixture("lesson-plan-input.json")
                    payload["lessons"][0]["evaluation"]["score"] = score
                    source = write_payload(folder, payload, f"score-{str(score).replace('.', '_')}.json")
                    output = folder / f"output-{str(score).replace('.', '_')}"
                    result = run_script(
                        LESSON / "scripts" / "generate_lesson_plans.py",
                        "--tasks-json",
                        str(source),
                        "--output-dir",
                        str(output),
                    )
                    if should_pass:
                        self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
                        self.assertTrue((output / "qa-report.json").exists())
                    else:
                        self.assertNotEqual(result.returncode, 0)
                        self.assertIn("between 85 and 96 in 0.5-point increments", result.stderr)
                        self.assertFalse(output.exists())

    def test_content_quality_detects_real_failure_categories(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")
        cases = (
            (
                "exact student analysis",
                lambda data: data["lessons"][1]["student_analysis"].__setitem__(
                    "base", copy.deepcopy(data["lessons"][0]["student_analysis"]["base"])
                ),
                lambda report: any(item["field"] == "student_analysis.base" for item in report["exact_duplicates"]),
            ),
            (
                "similar reflection",
                lambda data: data["lessons"][1]["reflection"].__setitem__(
                    "summary",
                    data["lessons"][0]["reflection"]["summary"][:-5] + "并补充接口证据分析。",
                ),
                lambda report: any(item["field"] == "reflection.summary" for item in report["field_similarity_pairs"]),
            ),
            (
                "implementation duplicate",
                lambda data: [
                    lesson["implementation"][1]["teacher_actions"].__setitem__(
                        slice(None), ["教师根据本课任务检查需求依据并追问关键证据"]
                    )
                    for lesson in data["lessons"][:3]
                ],
                lambda report: any("task_introduction.teacher_actions" in item["field"] for item in report["implementation_duplicates"]),
            ),
            (
                "repeated sentence",
                lambda data: [
                    lesson["teaching_content"].__setitem__(
                        0, "学生按照任务清单核对关键证据并记录判定依据"
                    )
                    for lesson in data["lessons"][:3]
                ],
                lambda report: bool(report["repeated_sentences"]),
            ),
            (
                "old boilerplate",
                lambda data: data["lessons"][0]["reflection"].__setitem__(
                    "summary", data["lessons"][0]["reflection"]["summary"] + "任务驱动、教师示范、分组实训、过程评价"
                ),
                lambda report: bool(report["boilerplate_hits"]),
            ),
            (
                "identical progression",
                lambda data: [lesson["progression"].__setitem__("capability_stage", "独立") for lesson in data["lessons"]],
                lambda report: not report["progression"]["valid_variety"],
            ),
            (
                "simple score pattern",
                lambda data: [lesson["evaluation"].__setitem__("score", 90) for lesson in data["lessons"]],
                lambda report: report["coverage"]["score_pattern"]["all_same"],
            ),
        )
        for label, mutate, assertion in cases:
            with self.subTest(case=label):
                payload = copy.deepcopy(base)
                mutate(payload)
                report = assess_content_quality(payload)
                self.assertEqual(report["status"], "failed")
                self.assertTrue(assertion(report), report)

    def test_content_quality_calibration_rejects_copy_rewrites_but_allows_real_difference(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")

        def recursive_replace(value, replacements):
            if isinstance(value, str):
                for old, new in replacements:
                    value = value.replace(old, new)
                return value
            if isinstance(value, list):
                return [recursive_replace(item, replacements) for item in value]
            if isinstance(value, dict):
                return {key: recursive_replace(item, replacements) for key, item in value.items()}
            return value

        def copied_pair() -> dict:
            data = copy.deepcopy(base)
            data["lessons"] = data["lessons"][:2]
            data["total_hours"] = 4
            data["lessons"][1] = copy.deepcopy(data["lessons"][0])
            data["lessons"][1]["lesson_id"] = "L02"
            data["lessons"][1]["unit"] = "项目二 数据边界分析"
            data["lessons"][1]["task"] = "设计字段边界用例"
            data["lessons"][1]["progression"]["prior_lesson_id"] = "L01"
            data["lessons"][1]["evaluation"]["score"] = 90
            return data

        cases = (
            ("copy+task rename", lambda data: None, "adjacent_similarity_pairs"),
            (
                "10 percent wording",
                lambda data: data["lessons"].__setitem__(
                    1,
                    recursive_replace(data["lessons"][1], [("测试", "检验")]),
                ),
                "field_similarity_pairs",
            ),
            (
                "20 percent wording",
                lambda data: data["lessons"].__setitem__(
                    1,
                    recursive_replace(
                        data["lessons"][1],
                        [("测试", "检验"), ("记录", "登记"), ("检查", "核验"), ("证据", "依据"), ("任务", "项目任务")],
                    ),
                ),
                "adjacent_similarity_pairs",
            ),
            (
                "synonym and word order",
                lambda data: data["lessons"].__setitem__(
                    1,
                    recursive_replace(
                        data["lessons"][1],
                        [
                            ("完成环境检查并留下可追溯证据", "留下可追溯证据并完成环境核验"),
                            ("依据需求边界安排测试任务", "按需求边界编排检验任务"),
                            ("核验工具版本、账号权限与测试数据", "复核工具版本、账户权限以及检验数据"),
                        ],
                    ),
                ),
                "adjacent_similarity_pairs",
            ),
            (
                "same flow with renamed nouns",
                lambda data: data["lessons"].__setitem__(
                    1,
                    recursive_replace(
                        data["lessons"][1],
                        [("环境", "接口"), ("配置", "请求"), ("风险", "异常"), ("测试", "验证")],
                    ),
                ),
                "adjacent_similarity_pairs",
            ),
        )
        for label, mutate, expected_category in cases:
            with self.subTest(case=label):
                payload = copied_pair()
                mutate(payload)
                report = assess_content_quality(payload)
                self.assertEqual(report["status"], "failed", report)
                self.assertTrue(report[expected_category], report)
                for record in report[expected_category]:
                    self.assertEqual(len(record["lessons"]), 2)
                    self.assertIn("score", record)
                    self.assertIn("top_fragments", record)

        distinct = copy.deepcopy(base)
        distinct["lessons"] = distinct["lessons"][:2]
        distinct["total_hours"] = 4
        distinct_report = assess_content_quality(distinct)
        self.assertEqual(distinct_report["status"], "passed", distinct_report)
        thresholds = distinct_report["coverage"]["similarity_thresholds"]
        self.assertLess(thresholds["adjacent_whole_lesson"], thresholds["whole_lesson"])
        self.assertLess(thresholds["adjacent_field"], thresholds["field"])

    def test_content_quality_checks_expanded_fields_and_adjacent_implementation(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")

        def pair() -> dict:
            data = copy.deepcopy(base)
            data["lessons"] = data["lessons"][:2]
            data["total_hours"] = 4
            return data

        field_paths = (
            ("teaching_content", ("teaching_content",)),
            ("goals.knowledge", ("goals", "knowledge")),
            ("goals.ability", ("goals", "ability")),
            ("goals.quality", ("goals", "quality")),
            ("teaching_methods", ("teaching_methods",)),
            ("progression.prior_learning", ("progression", "prior_learning")),
            ("progression.deliverable", ("progression", "deliverable")),
            ("progression.next_bridge", ("progression", "next_bridge")),
        )
        for field_name, path in field_paths:
            with self.subTest(field=field_name):
                data = pair()
                left = data["lessons"][0]
                right = data["lessons"][1]
                source = left
                target = right
                for part in path[:-1]:
                    source = source[part]
                    target = target[part]
                target[path[-1]] = copy.deepcopy(source[path[-1]])
                report = assess_content_quality(data)
                self.assertTrue(
                    any(item["field"] == field_name for item in report["exact_duplicates"]),
                    (field_name, report),
                )

        data = pair()
        data["lessons"][1]["implementation"] = copy.deepcopy(data["lessons"][0]["implementation"])
        report = assess_content_quality(data)
        self.assertEqual(report["status"], "failed", report)
        self.assertTrue(report["adjacent_implementation_exact_duplicates"], report)
        self.assertTrue(
            any(
                item["stage"] == "task_introduction"
                and item["lessons"] == ["L01", "L02"]
                for item in report["implementation_similarity_pairs"]
            ),
            report,
        )

    def test_content_quality_adversarial_skeletons_and_allowed_reuse(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")

        def add_topic_entities(data: dict) -> None:
            for index, lesson in enumerate(data["lessons"], 1):
                topic = f"主题{index}"
                artifact = f"成果{index}"
                lesson["unit"] = f"项目{index} {topic}"
                lesson["task"] = f"完成{topic}处理"
                lesson["progression"]["deliverable"] = artifact

        implementation_skeleton = copy.deepcopy(base)
        add_topic_entities(implementation_skeleton)
        for index, lesson in enumerate(implementation_skeleton["lessons"], 1):
            topic, artifact = f"主题{index}", f"成果{index}"
            stage = lesson["implementation"][3]
            stage["content"] = [f"围绕{topic}分析任务边界并记录{artifact}的关键证据"]
            stage["teacher_actions"] = [f"教师引导学生围绕{topic}逐项核对操作条件并记录判断依据"]
            stage["student_actions"] = [f"学生根据{topic}完成条件核对并提交{artifact}"]
            stage["objective"] = f"通过{topic}完成任务并形成{artifact}"
        implementation_report = assess_content_quality(implementation_skeleton)
        self.assertEqual(implementation_report["status"], "failed", implementation_report)
        self.assertTrue(implementation_report["implementation_structural_similarity_pairs"], implementation_report)

        skeleton_cases = {
            "goals": lambda lesson, topic: lesson["goals"].update(
                {
                    "knowledge": [f"掌握{topic}基本方法", f"分析{topic}关键条件"],
                    "ability": [f"完成{topic}操作并说明依据", f"提交{topic}成果"],
                    "quality": [f"形成规范{topic}意识", f"保持{topic}记录习惯"],
                }
            ),
            "student_analysis": lambda lesson, topic: lesson["student_analysis"].update(
                {
                    "base": [f"了解{topic}基本对象", f"接触{topic}简单记录"],
                    "problems": [f"容易忽略{topic}条件", f"说明{topic}依据不清"],
                    "strategies": [f"用清单拆解{topic}步骤", f"通过互评核对{topic}结果"],
                }
            ),
            "reflection": lambda lesson, topic: lesson.__setitem__(
                "reflection",
                {
                    "summary": f"围绕{topic}完成本课任务并形成成果",
                    "innovation": f"将{topic}拆成连续检查点组织课堂",
                    "improvement": f"下一课继续完善{topic}的判断依据",
                },
            ),
        }
        for field_name, mutate in skeleton_cases.items():
            data = copy.deepcopy(base)
            add_topic_entities(data)
            for index, lesson in enumerate(data["lessons"], 1):
                mutate(lesson, f"主题{index}")
            report = assess_content_quality(data)
            self.assertEqual(report["status"], "failed", (field_name, report))
            self.assertTrue(
                any(item["field"].startswith(field_name + ".") for item in report["structural_similarity_pairs"]),
                (field_name, report),
            )

        distinct = copy.deepcopy(base)
        distinct["lessons"] = distinct["lessons"][:2]
        distinct["total_hours"] = 4
        self.assertEqual(assess_content_quality(distinct)["status"], "passed")

        allowed_reuse = copy.deepcopy(base)
        for lesson in allowed_reuse["lessons"]:
            lesson["teaching_methods"][0] = "任务驱动法"
            lesson["resources"][0] = "数据库客户端工具"
        allowed_report = assess_content_quality(allowed_reuse)
        self.assertEqual(allowed_report["status"], "passed", allowed_report)

        short_item = copy.deepcopy(base)
        for lesson in short_item["lessons"]:
            lesson["goals"]["quality"][0] = "形成质量意识"
        short_report = assess_content_quality(short_item)
        self.assertEqual(short_report["status"], "failed", short_report)
        self.assertTrue(any(item["field"] == "goals.quality" for item in short_report["frequency_item_duplicates"]))

        adjacent_item = copy.deepcopy(base)
        adjacent_item["lessons"][1]["teaching_content"][0] = adjacent_item["lessons"][0]["teaching_content"][0]
        adjacent_report = assess_content_quality(adjacent_item)
        self.assertEqual(adjacent_report["status"], "failed", adjacent_report)
        self.assertTrue(any(item["field"] == "teaching_content" for item in adjacent_report["adjacent_item_duplicates"]))

    def test_progression_coherence_and_score_calibration_are_deterministic(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")
        connected = copy.deepcopy(base)
        connected["lessons"] = connected["lessons"][:2]
        connected["total_hours"] = 4
        connected_report = assess_content_quality(connected)
        self.assertEqual(connected_report["progression"]["status"], "passed", connected_report)
        self.assertTrue(all(link["status"] == "passed" for link in connected_report["progression"]["links"]))

        boundary = copy.deepcopy(connected)
        boundary["lessons"][1]["unit"] = "项目二 测试执行与交付"
        boundary_report = assess_content_quality(boundary)
        self.assertEqual(boundary_report["progression"]["status"], "passed", boundary_report)
        self.assertFalse(boundary_report["progression"]["links"][0]["same_unit"])
        self.assertGreater(boundary_report["progression"]["links"][0]["threshold"], 0)

        disconnected = copy.deepcopy(base)
        disconnected["lessons"][1]["progression"].update(
            {
                "prior_learning": "完全转入园艺种植的病虫害防治",
                "deliverable": "灌溉阈值记录",
                "next_bridge": "进入植物病害观察",
            }
        )
        disconnected["lessons"][1]["task"] = "完成温室灌溉控制"
        disconnected["lessons"][1]["teaching_content"] = [
            "识别土壤湿度传感器",
            "配置灌溉阈值",
            "形成园艺控制记录",
        ]
        disconnected_report = assess_content_quality(disconnected)
        self.assertEqual(disconnected_report["progression"]["links"][0]["status"], "failed", disconnected_report)
        self.assertGreater(disconnected_report["progression"]["links"][0]["threshold"], 0)

        declared = copy.deepcopy(base)
        declared["lessons"][3]["progression"]["prior_lesson_id"] = "L02"
        declared["lessons"][3]["progression"]["prior_learning"] = "承接L02功能用例集，依据其中的边界条件分析缺陷影响"
        declared_report = assess_content_quality(declared)
        self.assertTrue(declared_report["progression"]["declared_prior_links"], declared_report)
        declared_l04 = next(item for item in declared_report["progression"]["declared_prior_links"] if item["to"] == "L04")
        self.assertEqual(declared_l04["from"], "L02")
        self.assertEqual(declared_l04["to"], "L04")
        self.assertEqual(declared_report["progression"]["sequence_links"][2]["from"], "L03")
        self.assertEqual(declared_report["progression"]["sequence_links"][2]["to"], "L04")

        calibration = lesson_content_quality.progression_calibration()
        self.assertEqual(len(calibration), 11)
        self.assertTrue(all(item["expected"] == item["calibrated_status"] for item in calibration), calibration)
        self.assertTrue(
            all(
                item["effective_status"] == ("passed" if item["passed"] else "failed")
                and item["passed"] == (item["expected"] == "通过")
                and item["raw_score"] == item["evidence_score"]
                for item in calibration
            ),
            calibration,
        )
        margin = lesson_content_quality.progression_calibration_margin()
        self.assertEqual(margin["positive_count"], 5)
        self.assertEqual(margin["hard_negative_count"], 6)
        self.assertEqual(
            margin["negative_maximum"],
            max(item["raw_score"] for item in calibration if item["expected"] == "失败"),
        )
        self.assertEqual(
            margin["margin"],
            round(margin["positive_minimum"] - margin["negative_maximum"], 4),
        )
        for item in calibration:
            anchor = item["signals"]["substantive_anchor"]
            self.assertIn("matched_fragments", anchor)
            self.assertIn("longest_substantive_match", anchor)
            self.assertIn("acronym_matches", anchor)
            self.assertIn("generic_only_matches", anchor)
            if item["expected"] == "失败":
                self.assertEqual(anchor["status"], "failed", item)

        gate_calibration = lesson_content_quality.progression_gate_calibration()
        self.assertGreaterEqual(len(gate_calibration), 12)
        self.assertTrue(
            all(item["expected"] == item["actual"] for item in gate_calibration),
            gate_calibration,
        )
        for case in {
            "generic_operation_composed",
            "generic_design_composed",
            "generic_analysis_composed",
            "generic_check_record_composed",
            "generic_implementation_composed",
            "generic_flow_design_composed",
        }:
            with self.subTest(case=case):
                item = next(record for record in gate_calibration if record["case"] == case)
                forward = item["gates"]["forward_transition"]
                self.assertEqual(forward["current_body_coherence"]["status"], "passed", item)
                self.assertEqual(forward["substantive_anchor"]["status"], "failed", item)
                self.assertEqual(item["actual"]["forward_transition"], "failed", item)
                self.assertEqual(item["actual"]["overall"], "failed", item)

        arithmetic = copy.deepcopy(base)
        for lesson, score in zip(arithmetic["lessons"], (88, 88.5, 89, 89.5, 90, 90.5)):
            lesson["evaluation"]["score"] = score
        arithmetic_report = assess_content_quality(arithmetic)
        self.assertTrue(arithmetic_report["coverage"]["score_pattern"]["arithmetic_progression"])
        self.assertEqual(arithmetic_report["status"], "failed")

        natural = copy.deepcopy(base)
        for lesson, score in zip(natural["lessons"], (89, 90.5, 89.5, 91, 90, 92)):
            lesson["evaluation"]["score"] = score
        natural_report = assess_content_quality(natural)
        self.assertFalse(natural_report["coverage"]["score_pattern"]["arithmetic_progression"])
        self.assertEqual(natural_report["status"], "passed", natural_report)

        cycle = copy.deepcopy(base)
        cycle["lessons"] = []
        for index, source in enumerate(base["lessons"][:4] * 2, 1):
            lesson = copy.deepcopy(source)
            lesson["lesson_id"] = f"L{index:02d}"
            lesson["progression"]["prior_lesson_id"] = None if index == 1 else f"L{index - 1:02d}"
            lesson["evaluation"]["score"] = (88, 89, 90, 91)[(index - 1) % 4]
            cycle["lessons"].append(lesson)
        cycle["total_hours"] = 16
        cycle_report = assess_content_quality(cycle)
        self.assertTrue(cycle_report["coverage"]["score_pattern"]["simple_cycle"], cycle_report)
        self.assertEqual(cycle_report["coverage"]["score_pattern"]["cycle_period"], 4)

        default_override = copy.deepcopy(base)
        default_override["default_hours"] = 1
        lesson_generator.validate_content_v2_input(default_override)

        invalid_stage = copy.deepcopy(base)
        invalid_stage["lessons"][0]["progression"]["capability_stage"] = "熟练"
        with self.assertRaisesRegex(ValueError, "capability_stage"):
            lesson_generator.validate_content_v2_input(invalid_stage)

    def test_progression_gates_do_not_allow_false_inheritance_or_false_bridge(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")

        def pair() -> dict:
            data = copy.deepcopy(base)
            data["lessons"] = data["lessons"][:2]
            data["total_hours"] = 4
            return data

        false_bridge = pair()
        previous, current = false_bridge["lessons"]
        previous["progression"].update(
            deliverable="数据库需求分析报告",
            next_bridge="进入实体关系建模",
        )
        current["progression"].update(
            prior_learning="已完成数据库需求分析报告",
            deliverable="生命体征测量记录",
        )
        current["task"] = "测量患者生命体征"
        current["teaching_content"] = ["体温、脉搏、呼吸", "记录测量结果", "判断异常情况"]
        false_bridge_report = assess_content_quality(false_bridge)
        false_bridge_link = false_bridge_report["progression"]["declared_prior_links"][0]
        self.assertEqual(false_bridge_link["artifact_inheritance"]["status"], "passed")
        self.assertEqual(false_bridge_link["forward_transition"]["status"], "failed")
        self.assertEqual(false_bridge_link["status"], "failed")
        self.assertEqual(false_bridge_report["progression"]["status"], "failed")

        false_inheritance = pair()
        previous, current = false_inheritance["lessons"]
        previous["progression"].update(
            deliverable="数据库需求分析报告",
            next_bridge="根据表结构完成SQL查询",
        )
        current["progression"].update(
            prior_learning="承接护理实训基础",
            deliverable="SQL查询脚本和结果记录",
        )
        current["task"] = "完成SQL查询设计"
        current["teaching_content"] = ["依据SQL查询设计的表结构编写查询语句", "验证查询结果", "整理查询证据"]
        for stage in current["implementation"]:
            stage["content"] = [f"围绕SQL查询设计，{stage['content'][0]}"]
            stage["teacher_actions"] = ["教师检查SQL查询设计与执行条件"]
            stage["student_actions"] = ["完成SQL查询并记录查询结果"]
            stage["objective"] = "形成SQL查询执行依据"
        false_inheritance_report = assess_content_quality(false_inheritance)
        false_inheritance_link = false_inheritance_report["progression"]["declared_prior_links"][0]
        self.assertEqual(false_inheritance_link["artifact_inheritance"]["status"], "failed")
        self.assertEqual(false_inheritance_link["forward_transition"]["status"], "passed")
        self.assertEqual(false_inheritance_link["status"], "failed")

        both_pass = pair()
        previous, current = both_pass["lessons"]
        previous["progression"].update(
            deliverable="E-R图",
            next_bridge="将实体关系转换为关系表",
        )
        current["progression"].update(
            prior_learning="能解释上一课E-R图中的实体及联系",
            deliverable="关系模式设计表",
        )
        current["task"] = "完成关系模式设计"
        current["teaching_content"] = ["主键、外键与关系表映射", "依据E-R图确定字段", "检查关系模式约束"]
        for stage in current["implementation"]:
            stage["content"] = [f"围绕关系模式设计，{stage['content'][0]}"]
            stage["teacher_actions"] = ["教师检查关系模式设计与执行条件"]
            stage["student_actions"] = ["完成关系模式设计并记录查询结果"]
            stage["objective"] = "形成关系模式设计依据"
        both_pass_report = assess_content_quality(both_pass)
        both_pass_link = both_pass_report["progression"]["declared_prior_links"][0]
        self.assertEqual(both_pass_link["artifact_inheritance"]["status"], "passed")
        self.assertEqual(both_pass_link["forward_transition"]["status"], "passed")
        self.assertEqual(both_pass_link["status"], "passed")
        self.assertEqual(both_pass_report["progression"]["status"], "passed")

    def test_nonadjacent_physical_sequence_review_requires_agent_acceptance(self) -> None:
        data = load_fixture("lesson-plan-content-v2-it.json")
        data["lessons"] = data["lessons"][:4]
        data["total_hours"] = 8
        data["lessons"][1]["progression"]["next_bridge"] = "依据功能用例集执行功能测试并分析缺陷影响，后续形成缺陷报告"
        data["lessons"][2]["progression"]["next_bridge"] = "进入园艺病虫害观察"
        data["lessons"][3]["progression"]["prior_lesson_id"] = "L02"
        data["lessons"][3]["progression"]["prior_learning"] = "承接L02功能用例集分析缺陷影响"
        report = assess_content_quality(data)
        self.assertEqual(report["progression"]["status"], "passed", report)
        self.assertTrue(report["progression"]["requires_agent_review"], report)
        review = report["progression"]["agent_review_items"][0]
        self.assertEqual(review["from"], "L03")
        self.assertEqual(review["to"], "L04")
        self.assertEqual(review["declared_prior"], "L02")
        self.assertIn("reason", review)
        self.assertIn("score", review)

    def test_resource_reuse_is_domain_neutral_but_narrative_reuse_still_fails(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")
        resource_cases = (
            ["数据库客户端工具", "任务数据"],
            ["血压计", "护理床"],
            ["原始凭证", "计算器"],
        )
        for resources in resource_cases:
            with self.subTest(resources=resources):
                data = copy.deepcopy(base)
                data["lessons"] = data["lessons"][:3]
                data["total_hours"] = 6
                for lesson in data["lessons"]:
                    lesson["resources"] = list(resources)
                report = assess_content_quality(data)
                self.assertEqual(report["status"], "passed", report)
                self.assertTrue(report["resource_reuse"])

        narrative = copy.deepcopy(base)
        narrative["lessons"] = narrative["lessons"][:3]
        narrative["total_hours"] = 6
        for lesson in narrative["lessons"]:
            lesson["teaching_content"][0] = "学生依据任务清单核对关键证据"
        narrative_report = assess_content_quality(narrative)
        self.assertEqual(narrative_report["status"], "failed", narrative_report)
        self.assertTrue(
            any(item["field"] == "teaching_content" for item in narrative_report["frequency_item_duplicates"]),
            narrative_report,
        )

    def test_unified_reuse_policy_allows_terms_but_keeps_narrative_strict(self) -> None:
        data = load_fixture("lesson-plan-content-v2-it.json")
        repeated_method = "基于案例分析与任务执行的分组研讨教学法"
        for lesson in data["lessons"]:
            lesson["teaching_methods"][0] = repeated_method
            lesson["resources"][0] = "数据库客户端工具"
            lesson["references"] = [{"text": "课程配套教学资源", "source_kind": "generic"}]
            lesson["evaluation"]["remarks"]["attendance"] = "出勤完整且准时"
        report = assess_content_quality(data)
        self.assertEqual(report["status"], "passed", report)
        policies = report["reuse_policy"]["classes"]
        self.assertEqual(policies["narrative_strict"], "narrative_strict")
        allowed = report["reuse_policy"]["allowed_reuse"]
        self.assertTrue(any(item["field"] == "teaching_methods" for item in allowed), allowed)
        self.assertTrue(any(item["field"] == "resources" for item in allowed), allowed)
        self.assertTrue(any(item["field"] == "references" for item in allowed), allowed)
        self.assertTrue(any(item["field"] == "evaluation.remarks.attendance" for item in allowed), allowed)
        self.assertFalse(any(item["field"] == "teaching_methods" for item in report["structural_similarity_pairs"]))

        short_narrative = copy.deepcopy(data)
        for lesson in short_narrative["lessons"]:
            lesson["goals"]["quality"][0] = "规范操作"
        short_report = assess_content_quality(short_narrative)
        self.assertEqual(short_report["status"], "failed", short_report)
        self.assertTrue(
            any(item["field"] == "goals.quality" for item in short_report["frequency_item_duplicates"]),
            short_report,
        )

    def test_structural_entities_from_teaching_content_do_not_hide_a_copied_skeleton(self) -> None:
        data = load_fixture("lesson-plan-content-v2-it.json")
        data["lessons"] = data["lessons"][:2]
        data["total_hours"] = 4
        left, right = data["lessons"]
        left["teaching_content"][0] = "围绕护理脉搏对象分析关键条件并记录护理脉搏结果"
        right["teaching_content"][0] = "围绕会计凭证对象分析关键条件并记录会计凭证结果"
        report = assess_content_quality(data)
        self.assertEqual(report["status"], "failed", report)
        pair = next(
            item for item in report["structural_similarity_pairs"]
            if item["field"] == "teaching_content"
        )
        self.assertGreaterEqual(pair["masked_score"], pair["raw_score"])

    def test_score_cycle_detection_rejects_partial_tails_and_keeps_natural_variation(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")

        def scored_data(scores: tuple[float, ...]) -> dict:
            data = copy.deepcopy(base)
            while len(data["lessons"]) < len(scores):
                extra = copy.deepcopy(base["lessons"][-1])
                extra["lesson_id"] = f"L{len(data['lessons']) + 1:02d}"
                data["lessons"].append(extra)
            data["lessons"] = data["lessons"][: len(scores)]
            data["total_hours"] = len(scores) * 2
            for index, (lesson, score) in enumerate(zip(data["lessons"], scores), 1):
                lesson["evaluation"]["score"] = score
                lesson["progression"]["prior_lesson_id"] = None if index == 1 else f"L{index - 1:02d}"
            return data

        cycle_cases = (
            ((88, 89, 90, 88, 89, 90), 3),
            ((88, 89, 90, 88, 89, 90, 88), 3),
            ((90, 91, 90, 91, 90), 2),
            ((88, 89, 90, 91, 88, 89, 90), 4),
            ((90, 90, 90, 90), 1),
        )
        for scores, period in cycle_cases:
            with self.subTest(scores=scores):
                report = assess_content_quality(scored_data(scores))
                self.assertTrue(report["coverage"]["score_pattern"]["simple_cycle"], report)
                self.assertEqual(report["coverage"]["score_pattern"]["cycle_period"], period)
                self.assertGreater(report["coverage"]["score_pattern"]["cycle_confidence"], 0)
                self.assertGreaterEqual(report["coverage"]["score_pattern"]["full_cycles"], 1)
                self.assertIn("tail_length", report["coverage"]["score_pattern"])
                self.assertIn("tail_fraction", report["coverage"]["score_pattern"])
                self.assertEqual(report["status"], "failed")

        natural_tail_cases = (
            (89, 90, 93, 88, 92, 89, 90),
            (89, 90.5, 89.5, 91, 90, 92, 89, 90.5),
        )
        for scores in natural_tail_cases:
            with self.subTest(natural_tail=scores):
                score_errors: list[str] = []
                pattern = lesson_content_quality._score_pattern(list(scores), score_errors)
                self.assertFalse(pattern["simple_cycle"], pattern)
                self.assertEqual(pattern["cycle_confidence"], 0.0)
                self.assertFalse(any("repeating cycle" in error for error in score_errors))

        natural = copy.deepcopy(base)
        natural_scores = (89, 90.5, 89.5, 91, 90, 92)
        for lesson, score in zip(natural["lessons"], natural_scores):
            lesson["evaluation"]["score"] = score
        natural_report = assess_content_quality(natural)
        self.assertEqual(natural_report["status"], "passed", natural_report)
        self.assertFalse(natural_report["coverage"]["score_pattern"]["simple_cycle"])

        varied = (90, 89.5, 91, 90.5, 92, 91.5, 93)
        score_errors: list[str] = []
        pattern = lesson_content_quality._score_pattern(list(varied), score_errors)
        self.assertFalse(pattern["simple_cycle"])
        self.assertFalse(any("repeating cycle" in error for error in score_errors))
        self.assertFalse(any("outside 85-96" in error for error in score_errors))

    def test_content_quality_failure_is_reported_by_real_output_validator(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-output-quality-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr)
            broken = copy.deepcopy(payload)
            broken["lessons"][1]["student_analysis"]["base"] = copy.deepcopy(
                broken["lessons"][0]["student_analysis"]["base"]
            )
            broken_source = write_payload(folder, broken, "broken-tasks.json")
            result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(broken_source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(result.returncode, 0)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["content_quality"]["status"], "failed")
            self.assertTrue(report["content_quality"]["exact_duplicates"])
            self.assertTrue(any("student_analysis.base" in error for error in report["content_quality"]["errors"]))

    def test_evaluation_remark_failures_are_written_to_real_qa_report(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-evaluation-quality-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)

            duplicate = copy.deepcopy(payload)
            duplicate["lessons"][1]["evaluation"]["remarks"]["practice"] = duplicate["lessons"][0]["evaluation"]["remarks"]["practice"]
            duplicate_source = write_payload(folder, duplicate, "duplicate-remarks.json")
            duplicate_result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(duplicate_source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(duplicate_result.returncode, 0)
            duplicate_report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertTrue(duplicate_report["content_quality"]["evaluation_remark_duplicates"])
            self.assertTrue(
                any(item["field"] == "evaluation.remarks.practice" for item in duplicate_report["content_quality"]["evaluation_remark_duplicates"]),
                duplicate_report,
            )

            dense = copy.deepcopy(payload)
            dense["lessons"][0]["evaluation"]["remarks"]["practice"] = "评价备注" * 12 + "超限"
            dense_source = write_payload(folder, dense, "dense-remarks.json")
            dense_result = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(dense_source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(dense_result.returncode, 0)
            dense_report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertTrue(dense_report["content_quality"]["evaluation_remark_density"])
            density = dense_report["content_quality"]["evaluation_remark_density"][0]
            self.assertEqual(density["field"], "evaluation.remarks.practice")
            self.assertEqual(density["actual_chars"], 50)
            self.assertEqual(density["limit"], 48)

            whitespace = copy.deepcopy(payload)
            whitespace["lessons"][0]["evaluation"]["remarks"]["practice"] = " \t"
            with self.assertRaisesRegex(ValueError, "meaningful"):
                lesson_generator.validate_content_v2_input(whitespace)

    def test_distinct_cross_domain_fixtures_generate_without_contamination(self) -> None:
        fixtures = (
            "lesson-plan-content-v2-it.json",
            "lesson-plan-content-v2-non-it.json",
            "lesson-plan-content-v2-accounting.json",
        )
        with tempfile.TemporaryDirectory(prefix="lesson-v2-cross-domain-") as temp_name:
            folder = Path(temp_name)
            for fixture_name in fixtures:
                output = folder / Path(fixture_name).stem
                result = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(ROOT / "tests" / "fixtures" / fixture_name),
                    "--output-dir",
                    str(output),
                )
                self.assertEqual(result.returncode, 0, result.stderr or result.stdout)
                report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
                self.assertEqual(report["status"], "passed")
                self.assertEqual(report["content_quality"]["status"], "passed")
                self.assertEqual(report["content_quality"]["coverage"]["non_it_contamination"], [])
                self.assertEqual(len(list(output.glob("*.docx"))), 6)
                if "non-it" in fixture_name:
                    text = "\n".join(document_text(Document(path)) for path in output.glob("*.docx"))
                    for term in ("软件技术", "标准机房", "脚本", "截图工具", "代码编辑器", "数据安全"):
                        self.assertNotIn(term, text)
                    for term in ("工程伦理", "平凡又不平凡的价值观"):
                        self.assertNotIn(term, text)
                    for term in ("职业伦理", "职业价值观"):
                        self.assertIn(term, text)

    def test_writer_uses_v2_values_and_keeps_internal_qa_out_of_docx(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        manifest = load_manifest(V111_MANIFEST)

        with tempfile.TemporaryDirectory(prefix="lesson-v2-writer-fidelity-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            for path, lesson in zip(sorted(output.glob("*.docx")), payload["lessons"]):
                document = Document(path)
                table = document.tables[0]
                for name, expected in lesson_header_values(payload, lesson).items():
                    self.assertEqual(bookmark_text(document, field_bookmark(manifest, name)), expected)
                for name, expected in lesson_content_field_values(lesson).items():
                    self.assertEqual(bookmark_text(document, field_bookmark(manifest, name)), expected)
                for group, expected_values in zip(implementation_bookmarks(manifest), format_implementation(lesson["implementation"])):
                    for name, expected in zip(group, expected_values.values()):
                        self.assertEqual(bookmark_text(document, name), expected)
                for name, expected in zip(reflection_bookmarks(manifest), format_reflection(lesson["reflection"])):
                    self.assertEqual(bookmark_text(document, name), expected)
                nested = table.cell(12, 1).tables[0]
                for row_index, expected_values in enumerate(
                    format_evaluation_values(lesson, score_breakdown(lesson["evaluation"]["score"])),
                    start=1,
                ):
                    for cell_index, expected in expected_values.items():
                        self.assertEqual(nested.cell(row_index, cell_index).text.strip(), expected)
                text = document_text(document)
                for internal in ("Content V2", "similarity", "confidence", "source provenance", "资料不足", "根据现有资料推断"):
                    self.assertNotIn(internal, text)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["content_contract_version"], "2.0")
            self.assertEqual(report["content_quality"]["status"], "passed")
            self.assertEqual(report["content_quality"]["coverage"]["non_it_contamination"], [])

    def test_references_require_provenance_without_writing_metadata_to_docx(self) -> None:
        base = load_fixture("lesson-plan-input.json")
        cases = (
            ("generic", "课程配套教学资源", "generic", None, True),
            ("generic-book-title", "《数据库原理与应用》", "generic", None, False),
            ("provided", "《软件测试基础》作者甲，某出版社，2024年第2版", "provided", "用户提供教材目录", True),
            ("generic-isbn", "《软件测试基础》ISBN 978-7-0000-0000-0", "generic", None, False),
            ("generic-standard", "GB/T 0000-2024 软件测试规范", "generic", None, False),
            ("generic-author", "作者甲 某出版社 软件测试教材", "generic", None, False),
            ("verified-missing-evidence", "软件测试公开课程资料", "verified_public", None, False),
            ("verified-arbitrary-evidence", "软件测试公开课程资料", "verified_public", "有人说这是公开资料", False),
            ("verified-url", "软件测试公开课程资料", "verified_public", "https://example.edu/testing", True),
            ("verified-official-locator", "软件测试公开课程资料", "verified_public", "教育部官网第2章节", True),
            ("provided-missing-evidence", "《数据库原理与应用》", "provided", None, False),
            ("provided-file", "《数据库原理与应用》", "provided", "用户上传：数据库教材目录.pdf", True),
        )
        for label, text, source_kind, evidence, should_pass in cases:
            with self.subTest(reference=label):
                payload = copy.deepcopy(base)
                reference = {"text": text, "source_kind": source_kind}
                if evidence is not None:
                    reference["evidence"] = evidence
                payload["lessons"][0]["references"] = [reference]
                try:
                    lesson_generator.validate_content_v2_input(payload)
                except ValueError:
                    self.assertFalse(should_pass)
                else:
                    self.assertTrue(should_pass)

        with tempfile.TemporaryDirectory(prefix="lesson-v2-reference-docx-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, base)
            output = folder / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            text = "\n".join(document_text(Document(path)) for path in output.glob("*.docx"))
            self.assertNotIn("source_kind", text)
            self.assertNotIn("generic", text)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(
                report["content_quality"]["reference_provenance"]["validation_scope"],
                "contract_and_locator_only",
            )

    def test_reference_reuse_fixture_allows_same_source_across_six_lessons(self) -> None:
        cases = (
            (
                "provided",
                [{"text": "核心教材 A", "source_kind": "provided", "evidence": "用户上传教材：核心教材A.pdf"}],
            ),
            (
                "verified_public",
                [{"text": "MySQL 8.0 Reference Manual", "source_kind": "verified_public", "evidence": "https://dev.mysql.com/doc/refman/8.0/en/"}],
            ),
            (
                "generic",
                [{"text": "数据库课程标准相关章节", "source_kind": "generic"}],
            ),
            (
                "provided-and-generic",
                [
                    {"text": "核心教材 A", "source_kind": "provided", "evidence": "用户上传教材：核心教材A.pdf"},
                    {"text": "课程标准 B 相关章节", "source_kind": "generic"},
                ],
            ),
        )
        for label, references in cases:
            with self.subTest(reference_kind=label):
                report = assess_content_quality(reference_reuse_fixture(references))
                self.assertEqual(report["status"], "passed", report)
                provenance = report["reference_provenance"]
                self.assertEqual(provenance["reuse_policy"], "reference_reusable")
                self.assertEqual(provenance["cross_lesson_reuse"], "allowed")
                self.assertEqual(provenance["same_lesson_duplicates"], [])
                self.assertEqual(provenance["invalid_resource_only"], [])
                self.assertFalse(any("references" in error for error in report["errors"]), report)

    def test_reference_reuse_exemption_does_not_leak_into_narrative_or_implementation(self) -> None:
        base = reference_reuse_fixture(
            [{"text": "核心教材 A", "source_kind": "provided", "evidence": "用户上传教材：核心教材A.pdf"}]
        )
        cases = (
            ("teaching_content", lambda lesson: lesson["teaching_content"].__setitem__(0, "相同教学内容重复回归样本")),
            (
                "teacher_actions",
                lambda lesson: lesson["implementation"][3]["teacher_actions"].__setitem__(0, "相同教师动作重复回归样本"),
            ),
            (
                "student_actions",
                lambda lesson: lesson["implementation"][3]["student_actions"].__setitem__(0, "相同学生动作重复回归样本"),
            ),
            ("reflection", lambda lesson: lesson["reflection"].__setitem__("summary", "相同反思摘要重复回归样本")),
        )
        for field, mutate in cases:
            with self.subTest(field=field):
                data = copy.deepcopy(base)
                for lesson in data["lessons"]:
                    mutate(lesson)
                report = assess_content_quality(data)
                self.assertEqual(report["status"], "failed", report)
                duplicate_fields = {
                    item["field"]
                    for item in (
                        report["exact_duplicates"]
                        + report["item_duplicates"]
                        + report["adjacent_item_duplicates"]
                        + report["frequency_item_duplicates"]
                        + report["implementation_duplicates"]
                    )
                }
                self.assertTrue(any(field in item for item in duplicate_fields), duplicate_fields)
                self.assertTrue(any(field in error for error in report["errors"]), report["errors"])
                self.assertFalse(any("references" in error for error in report["errors"]), report)

    def test_reference_source_and_resource_boundary_is_conservative(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")
        resource_only = ("投影仪", "PPT", "MySQL Workbench", "血压计", "数据库服务器", "护理模型", "计算机机房")
        for text in resource_only:
            with self.subTest(resource=text):
                data = copy.deepcopy(base)
                data["lessons"][0]["references"] = [{"text": text, "source_kind": "generic"}]
                report = assess_content_quality(data)
                self.assertEqual(report["status"], "failed", report)
                self.assertEqual(report["reference_provenance"]["invalid_resource_only"], [{"lesson": "L01", "reference": 1}])
                self.assertTrue(any("resource-only" in error for error in report["errors"]), report["errors"])

        document_reference = copy.deepcopy(base)
        document_reference["lessons"][0]["references"] = [
            {
                "text": "MySQL 8.0 Reference Manual",
                "source_kind": "verified_public",
                "evidence": "https://dev.mysql.com/doc/refman/8.0/en/",
            }
        ]
        document_report = assess_content_quality(document_reference)
        self.assertEqual(document_report["status"], "passed", document_report)
        self.assertEqual(document_report["reference_provenance"]["invalid_resource_only"], [])

        provided = copy.deepcopy(base)
        provided["lessons"][0]["references"] = [
            {"text": "用户上传教材", "source_kind": "provided", "evidence": "用户上传：教材.pdf"}
        ]
        self.assertEqual(assess_content_quality(provided)["status"], "passed")

        duplicate = copy.deepcopy(base)
        duplicate["lessons"][0]["references"] = [
            {"text": "数据库课程标准相关章节", "source_kind": "generic"},
            {"text": "数据库课程标准相关章节", "source_kind": "generic"},
        ]
        duplicate_report = assess_content_quality(duplicate)
        self.assertEqual(duplicate_report["status"], "failed", duplicate_report)
        self.assertEqual(
            duplicate_report["reference_provenance"]["same_lesson_duplicates"],
            [{"lesson": "L01", "reference": 2, "duplicate_of": 1}],
        )
        self.assertTrue(any("within the same lesson" in error for error in duplicate_report["errors"]))

        invented_generic = copy.deepcopy(base)
        invented_generic["lessons"][0]["references"] = [
            {"text": "《数据库原理》作者甲，某出版社，2024年第2版 ISBN 978-7-0000-0000-0", "source_kind": "generic"}
        ]
        invented_report = assess_content_quality(invented_generic)
        self.assertEqual(invented_report["status"], "failed", invented_report)
        self.assertEqual(invented_report["reference_provenance"]["invalid_generic"], [{"lesson": "L01", "reference": 1}])

    def test_evaluation_remark_contract_limit_applies_to_all_template_versions(self) -> None:
        base = load_fixture("lesson-plan-content-v2-it.json")
        base["lessons"] = base["lessons"][:2]
        base["total_hours"] = 4
        manifests = (
            ("v1.0.0", V10_MANIFEST),
            ("v1.1.0", V110_MANIFEST),
            ("v1.1.1", V111_MANIFEST),
            ("v1.1.2", V112_MANIFEST),
        )
        for version, manifest_path in manifests:
            with self.subTest(version=version):
                accepted = copy.deepcopy(base)
                accepted["lessons"][0]["evaluation"]["remarks"]["practice"] = "甲" * 48
                accepted_report = assess_content_quality(accepted, load_manifest(manifest_path))
                self.assertEqual(accepted_report["status"], "passed", accepted_report)
                self.assertEqual(accepted_report["coverage"]["evaluation_remark_contract_limit"], 48)

                rejected = copy.deepcopy(base)
                rejected["lessons"][0]["evaluation"]["remarks"]["practice"] = "甲" * 49
                rejected_report = assess_content_quality(rejected, load_manifest(manifest_path))
                self.assertEqual(rejected_report["status"], "failed", rejected_report)
                self.assertTrue(rejected_report["evaluation_remark_density"])
                self.assertEqual(rejected_report["evaluation_remark_density"][0]["limit"], 48)

    def test_non_it_contamination_uses_current_lesson_scope_and_real_output_validation(self) -> None:
        payload = load_fixture("lesson-plan-content-v2-non-it.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-lesson-scope-contamination-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr or generated.stdout)
            target = sorted(output.glob("*.docx"))[0]
            document = Document(target)
            document.tables[0].cell(0, 0).paragraphs[0].add_run("脚本")
            document.save(target)

            validation = run_script(
                LESSON / "scripts" / "validate_output.py",
                "--input-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertNotEqual(validation.returncode, 0)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertIn("脚本", report["content_quality"]["coverage"]["non_it_contamination"])
            self.assertEqual(
                report["content_quality"]["coverage"]["non_it_contamination_scope"],
                "template_or_generator_injected_defaults_only",
            )
            self.assertTrue(any("脚本" in error for error in report["errors"]))

        course_metadata = {"course_name": "基础护理技术", "major": "护理", "audience": "高职护理"}
        allowed_lesson = {"task": "使用护理信息系统脚本查询", "resources": ["护理信息系统"]}
        other_lesson = {"task": "完成生命体征测量", "resources": ["血压计"]}
        self.assertEqual(detect_non_it_contamination(course_metadata, allowed_lesson, "脚本生成结果"), [])
        self.assertEqual(detect_non_it_contamination(course_metadata, other_lesson, "脚本生成结果"), ["脚本"])

    def test_cleanup_failure_preserves_primary_error_and_reports_residual_candidate(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-cleanup-diagnostics-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            argv = [
                "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            ]
            diagnostics = io.StringIO()
            with patch.object(lesson_generator, "validate_content_quality", side_effect=RuntimeError("primary operation failure")):
                with patch.object(lesson_generator, "_remove_path", side_effect=OSError("candidate is locked")):
                    with patch.object(sys, "argv", argv):
                        with redirect_stderr(diagnostics):
                            with self.assertRaisesRegex(RuntimeError, "primary operation failure") as raised:
                                lesson_generator.main()
            self.assertTrue(any("cleanup failed" in note for note in getattr(raised.exception, "__notes__", [])))
            self.assertIn("candidate is locked", diagnostics.getvalue())
            self.assertIn("residual path", diagnostics.getvalue())
            self.assertFalse(output.exists())

    def test_active_content_quality_has_no_legacy_implementation(self) -> None:
        source = (SCRIPTS / "content_quality.py").read_text(encoding="utf-8")
        self.assertNotIn("_assess_content_quality_legacy", source)
        self.assertNotIn("_progression_coherence", source)
        self.assertNotIn("0.035 if same_unit else 0.015", source)

    def test_render_documentation_distinguishes_smoke_from_visual_qa(self) -> None:
        skill_text = (LESSON / "SKILL.md").read_text(encoding="utf-8")
        contract_text = (LESSON / "docs" / "content-contract-v2.md").read_text(encoding="utf-8")
        self.assertIn("render smoke", skill_text.lower())
        self.assertIn("not pagination", contract_text.lower())
        self.assertNotIn("--render 做真实分页检查", skill_text)

    def test_darwin_case_probe_uses_independent_same_device_temp_root(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-darwin-probe-") as temp_name:
            folder = Path(temp_name)
            protected = folder / "protected"
            safe_temp = folder / "safe-temp"
            protected.mkdir()
            safe_temp.mkdir()
            created_under: list[Path] = []
            real_mkdtemp = tempfile.mkdtemp

            def tracked_mkdtemp(*args, **kwargs):
                created_under.append(Path(kwargs["dir"]).resolve())
                return real_mkdtemp(*args, **kwargs)

            lesson_path_safety._DARWIN_CASE_SENSITIVITY.clear()
            with (
                patch.object(lesson_path_safety.sys, "platform", "darwin"),
                patch.object(lesson_path_safety, "IS_WINDOWS", False),
                patch.object(lesson_path_safety.tempfile, "gettempdir", return_value=str(safe_temp)),
                patch.object(lesson_path_safety.tempfile, "mkdtemp", side_effect=tracked_mkdtemp),
            ):
                self.assertFalse(filesystem_case_sensitive(protected))
            self.assertEqual(created_under, [safe_temp.resolve()])
            self.assertEqual(list(protected.iterdir()), [])
            self.assertFalse(paths_overlap(safe_temp, protected))

    def test_darwin_case_probe_failure_is_conservative_and_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-darwin-failed-probe-") as temp_name:
            folder = Path(temp_name)
            protected = folder / "protected"
            safe_temp = folder / "safe-temp"
            protected.mkdir()
            safe_temp.mkdir()
            lesson_path_safety._DARWIN_CASE_SENSITIVITY.clear()
            with (
                patch.object(lesson_path_safety.sys, "platform", "darwin"),
                patch.object(lesson_path_safety, "IS_WINDOWS", False),
                patch.object(lesson_path_safety.tempfile, "gettempdir", return_value=str(safe_temp)),
                patch.object(lesson_path_safety.tempfile, "mkdtemp", side_effect=OSError("read only")),
            ):
                self.assertFalse(filesystem_case_sensitive(protected))
            self.assertEqual(list(protected.iterdir()), [])

    def test_darwin_probe_refuses_protected_ancestor_or_descendant(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-darwin-overlap-probe-") as temp_name:
            protected = Path(temp_name) / "protected"
            descendant = protected / "temp"
            descendant.mkdir(parents=True)
            lesson_path_safety._DARWIN_CASE_SENSITIVITY.clear()
            with (
                patch.object(lesson_path_safety.sys, "platform", "darwin"),
                patch.object(lesson_path_safety, "IS_WINDOWS", False),
                patch.object(lesson_path_safety.tempfile, "gettempdir", return_value=str(descendant)),
                patch.object(lesson_path_safety.tempfile, "mkdtemp") as mocked_probe,
            ):
                self.assertFalse(filesystem_case_sensitive(protected))
            mocked_probe.assert_not_called()
            self.assertEqual(list(descendant.iterdir()), [])

    @unittest.skipUnless(sys.platform == "darwin", "Darwin path aliases are only exercised on macOS")
    def test_macos_unicode_and_case_aliases_are_real_path_contracts(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-路径-") as temp_name:
            folder = Path(temp_name)
            nfc_name = unicodedata.normalize("NFC", "临时目录")
            nfd_name = unicodedata.normalize("NFD", "临时目录")
            actual = folder / nfc_name / "教案生成器" / "Template"
            actual.mkdir(parents=True)
            alias = folder / nfd_name / "教案生成器" / "Template"
            self.assertTrue(paths_equal(actual, alias))
            self.assertTrue(paths_overlap(actual / "输出", alias))
            case_alias = folder / nfc_name / "教案生成器" / "template"
            if not filesystem_case_sensitive(actual):
                self.assertTrue(paths_overlap(actual, case_alias))
            else:
                self.assertFalse(paths_overlap(actual, case_alias))

    @unittest.skipUnless(sys.platform.startswith("linux"), "Linux case semantics are checked on Linux CI")
    def test_linux_path_safety_remains_case_sensitive(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-linux-case-") as temp_name:
            folder = Path(temp_name)
            upper = folder / "CaseSensitive"
            lower = folder / "casesensitive"
            self.assertFalse(paths_equal(upper, lower))
            self.assertFalse(paths_overlap(upper, lower))

    def test_external_qa_report_is_committed_with_output_and_preflight_is_fail_closed(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-external-qa-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            external = folder / "reports" / "qa-report.json"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--qa-report",
                str(external),
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertTrue(external.is_file())
            internal_report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            external_report = json.loads(external.read_text(encoding="utf-8"))
            self.assertEqual(internal_report["status"], "passed")
            self.assertEqual(external_report["status"], "passed")
            self.assertEqual(external_report["qa_report"], str(external.resolve()))
            self.assertNotIn("cleanup failed", result.stderr.lower())

            for label in ("output-descendant", "source"):
                with self.subTest(path=label):
                    bad_output = folder / f"bad-{label}"
                    qa_path = bad_output / "nested" / "qa.json" if label == "output-descendant" else source
                    failed = run_script(
                        LESSON / "scripts" / "generate_lesson_plans.py",
                        "--tasks-json",
                        str(source),
                        "--output-dir",
                        str(bad_output),
                        "--qa-report",
                        str(qa_path),
                    )
                    self.assertNotEqual(failed.returncode, 0)
                    self.assertIn("external qa report", failed.stderr.lower())
                    self.assertFalse(bad_output.exists())

    def test_external_qa_transaction_rolls_back_output_and_report_on_each_exchange_failure(self) -> None:
        for label, fail_calls in (
            ("output-commit", {2}),
            ("external-commit", {4}),
            ("external-backup", {3}),
        ):
            with self.subTest(failure=label), tempfile.TemporaryDirectory(prefix=f"lesson-v2-atomic-{label}-") as temp_name:
                folder = Path(temp_name)
                output = folder / "output"
                output.mkdir()
                (output / "old.txt").write_bytes(b"old-output")
                external = folder / "qa-report.json"
                external.write_bytes(b"old-qa")
                candidate = folder / "candidate"
                candidate.mkdir()
                (candidate / "new.txt").write_bytes(b"new-output")
                external_candidate = folder / "qa-candidate.tmp"
                external_candidate.write_bytes(b"new-qa")
                real_replace = os.replace
                calls = {"count": 0}

                def fail_selected(source, target):
                    calls["count"] += 1
                    if calls["count"] in fail_calls:
                        raise OSError(f"injected {label} failure")
                    return real_replace(source, target)

                with patch.object(lesson_generator.os, "replace", side_effect=fail_selected):
                    with self.assertRaises(OSError):
                        atomic_commit_candidate_with_external_qa(
                            candidate,
                            output,
                            external_candidate,
                            external,
                            backup_existing=True,
                        )
                self.assertEqual((output / "old.txt").read_bytes(), b"old-output")
                self.assertEqual(external.read_bytes(), b"old-qa")
                self.assertFalse((output / "new.txt").exists())
                self.assertEqual(list(folder.glob("_output_backup_*")), [])
                self.assertEqual(list(folder.glob("_qa-report.json_backup_*")), [])
                shutil.rmtree(candidate, ignore_errors=True)
                external_candidate.unlink(missing_ok=True)

    def test_external_qa_transaction_reports_rollback_failure_and_main_cleans_temp_candidate(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-external-rollback-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            external = folder / "reports" / "qa.json"
            first = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--qa-report",
                str(external),
            )
            self.assertEqual(first.returncode, 0, first.stderr)
            before_output = {path.relative_to(output).as_posix(): path.read_bytes() for path in output.rglob("*") if path.is_file()}
            before_qa = external.read_bytes()
            with patch.object(lesson_generator.tempfile, "mkstemp", side_effect=OSError("injected external temp write failure")):
                argv = [
                    "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                    "--qa-report",
                    str(external),
                    "--backup-existing",
                ]
                with patch.object(sys, "argv", argv):
                    with self.assertRaisesRegex(OSError, "external temp write failure"):
                        lesson_generator.main()
            after_output = {path.relative_to(output).as_posix(): path.read_bytes() for path in output.rglob("*") if path.is_file()}
            self.assertEqual(before_output, after_output)
            self.assertEqual(before_qa, external.read_bytes())
            self.assertEqual(list(folder.rglob("*.candidate-*")), [])
            self.assertEqual(list(folder.rglob("*_backup_*")), [])

            candidate = folder / "candidate"
            candidate.mkdir()
            (candidate / "new.txt").write_bytes(b"new")
            external_candidate = folder / "qa-candidate.tmp"
            external_candidate.write_bytes(b"new-qa")
            real_replace = os.replace
            calls = {"count": 0}

            def fail_commit_and_rollback(source_path, target_path):
                calls["count"] += 1
                if calls["count"] in {4, 5}:
                    raise OSError("injected rollback failure")
                return real_replace(source_path, target_path)

            with patch.object(lesson_generator.os, "replace", side_effect=fail_commit_and_rollback):
                with self.assertRaisesRegex(RuntimeError, "rollback failed"):
                    atomic_commit_candidate_with_external_qa(
                        candidate,
                        output,
                        external_candidate,
                        external,
                        backup_existing=True,
                    )
            self.assertEqual(
                {path.relative_to(output).as_posix(): path.read_bytes() for path in output.rglob("*") if path.is_file()},
                before_output,
            )
            self.assertFalse(external.exists())
            shutil.rmtree(candidate, ignore_errors=True)
            external_candidate.unlink(missing_ok=True)
            for path in folder.glob("_*_backup_*"):
                if path.is_dir():
                    shutil.rmtree(path, ignore_errors=True)
                else:
                    path.unlink(missing_ok=True)

    def test_external_qa_backup_cleanup_failure_keeps_successful_publication(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-external-cleanup-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            output.mkdir()
            (output / "old.txt").write_bytes(b"old-output")
            external = folder / "qa-report.json"
            external.write_bytes(b"old-qa")
            candidate = folder / "candidate"
            candidate.mkdir()
            (candidate / "new.txt").write_bytes(b"new-output")
            external_candidate = folder / "qa-candidate.tmp"
            external_candidate.write_bytes(b"new-qa")
            real_remove = lesson_generator._remove_path
            diagnostics = io.StringIO()

            def fail_old_qa_cleanup(path: Path) -> None:
                if path.name.startswith("_qa-report.json_backup_"):
                    raise OSError("injected old QA cleanup failure")
                real_remove(path)

            with patch.object(lesson_generator, "_remove_path", side_effect=fail_old_qa_cleanup):
                with redirect_stderr(diagnostics):
                    backup = atomic_commit_candidate_with_external_qa(
                        candidate,
                        output,
                        external_candidate,
                        external,
                        backup_existing=True,
                    )
            self.assertIsNotNone(backup)
            self.assertEqual((output / "new.txt").read_bytes(), b"new-output")
            self.assertEqual(external.read_bytes(), b"new-qa")
            residual = list(folder.glob("_qa-report.json_backup_*"))
            self.assertEqual(len(residual), 1)
            self.assertIn("WARNING:", diagnostics.getvalue())
            self.assertIn("residual backup path:", diagnostics.getvalue())

    def test_unsafe_validation_skips_require_explicit_test_environment(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-unsafe-skip-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "rejected-output"
            with patch.dict(os.environ, {"LESSON_ALLOW_UNSAFE_VALIDATION_SKIP": "0"}):
                rejected = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(output),
                    "--skip-template-validation",
                    "--skip-output-validation",
                )
            self.assertNotEqual(rejected.returncode, 0)
            self.assertIn("Unsafe validation bypass is disabled.", rejected.stderr)
            self.assertFalse(output.exists())

            allowed_output = folder / "allowed-output"
            with patch.dict(os.environ, {"LESSON_ALLOW_UNSAFE_VALIDATION_SKIP": "1"}):
                allowed = run_script(
                    LESSON / "scripts" / "generate_lesson_plans.py",
                    "--tasks-json",
                    str(source),
                    "--output-dir",
                    str(allowed_output),
                    "--skip-template-validation",
                    "--skip-output-validation",
                )
            self.assertEqual(allowed.returncode, 0, allowed.stderr)
            report = json.loads((allowed_output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["status"], "skipped")
            self.assertEqual(report["validation_skipped"], ["template", "output"])

    def test_content_failure_preserves_existing_output_and_cleans_candidate(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-transaction-quality-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            first = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(first.returncode, 0, first.stderr)
            before = {path.relative_to(output).as_posix(): path.read_bytes() for path in output.rglob("*") if path.is_file()}
            broken = copy.deepcopy(payload)
            broken["lessons"][1]["student_analysis"]["base"] = copy.deepcopy(
                broken["lessons"][0]["student_analysis"]["base"]
            )
            broken_source = write_payload(folder, broken, "broken.json")
            failed = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(broken_source),
                "--output-dir",
                str(output),
                "--backup-existing",
            )
            self.assertNotEqual(failed.returncode, 0)
            after = {path.relative_to(output).as_posix(): path.read_bytes() for path in output.rglob("*") if path.is_file()}
            self.assertEqual(before, after)
            self.assertEqual(list(folder.glob(".output.candidate-*")), [])
            self.assertEqual(list(folder.glob("_output_backup_*")), [])

    def test_atomic_commit_restores_existing_directory_on_exchange_failure(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-transaction-commit-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            output.mkdir()
            (output / "old.txt").write_bytes(b"old")
            candidate = folder / "candidate"
            candidate.mkdir()
            (candidate / "new.txt").write_bytes(b"new")
            real_replace = os.replace
            calls = {"count": 0}

            def fail_second_replace(source: str, target: str) -> None:
                calls["count"] += 1
                if calls["count"] == 2:
                    raise OSError("injected candidate exchange failure")
                real_replace(source, target)

            with patch.object(lesson_generator.os, "replace", side_effect=fail_second_replace):
                with self.assertRaises(OSError):
                    atomic_commit_candidate(candidate, output, backup_existing=True)
            self.assertEqual((output / "old.txt").read_bytes(), b"old")
            self.assertFalse((output / "new.txt").exists())
            shutil.rmtree(candidate)
            for path in folder.glob("_output_backup_*"):
                shutil.rmtree(path, ignore_errors=True)

    def test_atomic_commit_preserves_existing_directory_when_backup_move_fails(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-transaction-backup-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            output.mkdir()
            (output / "old.txt").write_bytes(b"old")
            candidate = folder / "candidate"
            candidate.mkdir()
            (candidate / "new.txt").write_bytes(b"new")

            def fail_replace(_source: str, _target: str) -> None:
                raise OSError("injected backup move failure")

            with patch.object(lesson_generator.os, "replace", side_effect=fail_replace):
                with self.assertRaises(OSError):
                    atomic_commit_candidate(candidate, output, backup_existing=True)
            self.assertEqual((output / "old.txt").read_bytes(), b"old")
            self.assertTrue((candidate / "new.txt").exists())
            shutil.rmtree(candidate)

    def test_generation_midway_failure_never_commits_candidate(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-transaction-generation-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            output.mkdir()
            (output / "old.txt").write_bytes(b"old")
            before = (output / "old.txt").read_bytes()
            original_build = lesson_generator.build_lesson
            calls = {"count": 0}

            def fail_on_second(*args, **kwargs):
                calls["count"] += 1
                if calls["count"] == 2:
                    raise RuntimeError("injected generation failure")
                return original_build(*args, **kwargs)

            argv = [
                "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--backup-existing",
            ]
            with patch.object(lesson_generator, "build_lesson", side_effect=fail_on_second):
                with patch.object(sys, "argv", argv):
                    with self.assertRaises(RuntimeError):
                        lesson_generator.main()
            self.assertEqual((output / "old.txt").read_bytes(), before)
            self.assertEqual(list(folder.glob(".output.candidate-*")), [])
            self.assertEqual(list(folder.glob("_output_backup_*")), [])

    def test_output_qa_failure_never_commits_candidate(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-transaction-output-qa-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            output.mkdir()
            (output / "old.txt").write_bytes(b"old")
            argv = [
                "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--backup-existing",
            ]
            with patch.object(lesson_generator, "validate_outputs", side_effect=RuntimeError("injected output QA failure")):
                with patch.object(sys, "argv", argv):
                    with self.assertRaisesRegex(RuntimeError, "output QA failure"):
                        lesson_generator.main()
            self.assertEqual((output / "old.txt").read_bytes(), b"old")
            self.assertEqual(list(folder.glob(".output.candidate-*")), [])
            self.assertEqual(list(folder.glob("_output_backup_*")), [])

    def test_render_failure_is_reported_without_being_hidden(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        manifest = load_manifest(V112_MANIFEST)
        with tempfile.TemporaryDirectory(prefix="lesson-v2-render-failure-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            output = folder / "output"
            generated = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
            )
            self.assertEqual(generated.returncode, 0, generated.stderr)
            failed_render = {
                "status": "failed",
                "reason": "injected renderer failure",
                "renderer": "injected",
                "files_checked": 2,
                "errors": ["lesson document did not render"],
            }
            with patch.object(lesson_output, "render_docx_directory", return_value=failed_render):
                with self.assertRaisesRegex(RuntimeError, "render QA"):
                    lesson_output.validate_output_dir(
                        output,
                        payload,
                        manifest,
                        output / "qa-report.json",
                        LESSON / "schemas" / "lesson-plan-input.schema.json",
                        template_path=V112_TEMPLATE,
                        render=True,
                    )
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertEqual(report["render"]["status"], "failed")
            self.assertEqual(report["render"]["errors"], failed_render["errors"])

    def test_visual_inspection_evidence_is_explicit_durable_and_fingerprinted(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-visual-evidence-") as temp_name:
            output = Path(temp_name) / "output"
            output.mkdir()
            lesson = output / "教案01_示例.docx"
            lesson.write_bytes(b"explicit visual fixture")
            qa_report = output / "qa-report.json"
            qa_report.write_text('{"status":"passed"}\n', encoding="utf-8")
            destination = output / "visual-inspection.json"
            checks = {name: "passed" for name in lesson_visual_inspection.CHECK_NAMES}
            evidence = write_visual_inspection_evidence(
                output_dir=output,
                qa_report=qa_report,
                destination=destination,
                status="passed",
                inspected_pages={lesson.name: [1, 2]},
                checks=checks,
                notes="Agent inspected representative rendered pages.",
            )
            persisted = json.loads(destination.read_text(encoding="utf-8"))
            self.assertEqual(persisted, evidence)
            self.assertEqual(persisted["status"], "passed")
            self.assertEqual(persisted["inspected_files"], [lesson.name])
            self.assertEqual(persisted["inspected_pages"][lesson.name], [1, 2])
            self.assertEqual(len(persisted["output_fingerprint"]), 64)
            self.assertTrue(persisted["timestamp"].endswith("+00:00"))
            with self.assertRaisesRegex(ValueError, "identify at least one failed check"):
                write_visual_inspection_evidence(
                    output_dir=output,
                    qa_report=qa_report,
                    destination=destination,
                    status="failed",
                    inspected_pages={lesson.name: [1]},
                    checks=checks,
                    notes="",
                )

    def test_visual_inspection_destination_cannot_overwrite_protected_inputs(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-visual-destination-guard-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            output.mkdir()
            lesson = output / "教案01_示例.docx"
            lesson.write_bytes(b"explicit visual fixture")
            qa_report = output / "qa-report.json"
            qa_report.write_text('{"status":"passed"}\n', encoding="utf-8")
            checks = {name: "passed" for name in lesson_visual_inspection.CHECK_NAMES}

            for label, destination in (
                ("qa report", qa_report),
                ("generated docx", lesson),
                ("skill path", LESSON / "SKILL.md"),
            ):
                with self.subTest(destination=label):
                    before_report = qa_report.read_bytes()
                    before_lesson = lesson.read_bytes()
                    with self.assertRaisesRegex(ValueError, "protected"):
                        write_visual_inspection_evidence(
                            output_dir=output,
                            qa_report=qa_report,
                            destination=destination,
                            status="passed",
                            inspected_pages={lesson.name: [1]},
                            checks=checks,
                            notes="",
                        )
                    self.assertEqual(qa_report.read_bytes(), before_report)
                    self.assertEqual(lesson.read_bytes(), before_lesson)

    def test_path_safety_rejects_protected_real_cli_paths_and_allows_sibling(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        with tempfile.TemporaryDirectory(prefix="lesson-v2-paths-") as temp_name:
            folder = Path(temp_name)
            source = write_payload(folder, payload)
            protected_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(LESSON),
            )
            self.assertNotEqual(protected_result.returncode, 0)
            self.assertIn("protected path", protected_result.stderr.lower())

            collision_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(source),
            )
            self.assertNotEqual(collision_result.returncode, 0)
            self.assertIn("protected path", collision_result.stderr.lower())

            sibling = folder / "safe-output"
            safe_result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(sibling),
            )
            self.assertEqual(safe_result.returncode, 0, safe_result.stderr)
            self.assertTrue((sibling / "qa-report.json").exists())

    def test_path_safety_detects_symlink_overlap_when_supported(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-symlink-") as temp_name:
            folder = Path(temp_name)
            alias = folder / "skill-alias"
            try:
                alias.symlink_to(LESSON, target_is_directory=True)
            except (OSError, NotImplementedError):
                self.skipTest("symlinks are unavailable in this Windows environment")
            self.assertTrue(paths_overlap(alias / "output", LESSON))
            with self.assertRaises(ValueError):
                assert_output_path_safe(alias / "output", [LESSON])

    def test_internal_qa_path_accepts_platform_aliases(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-path-alias-") as temp_name:
            folder = Path(temp_name)
            output = folder / "output"
            output.mkdir()
            internal = output / "qa-report.json"
            self.assertTrue(paths_equal(internal, output.resolve() / "qa-report.json"))
            if os.name == "nt":
                import ctypes

                get_short_path = ctypes.windll.kernel32.GetShortPathNameW
                get_short_path.argtypes = [ctypes.c_wchar_p, ctypes.c_wchar_p, ctypes.c_uint32]
                get_short_path.restype = ctypes.c_uint32
                buffer = ctypes.create_unicode_buffer(32768)
                length = get_short_path(str(folder), buffer, len(buffer))
                if length:
                    short_internal = Path(buffer.value) / "output" / "qa-report.json"
                    self.assertTrue(paths_equal(short_internal, internal))

    def test_render_report_is_real_or_explicitly_not_executed(self) -> None:
        with tempfile.TemporaryDirectory(prefix="lesson-v2-render-") as temp_name:
            folder = Path(temp_name)
            source = ROOT / "tests" / "fixtures" / "lesson-plan-input.json"
            output = folder / "output"
            result = run_script(
                LESSON / "scripts" / "generate_lesson_plans.py",
                "--tasks-json",
                str(source),
                "--output-dir",
                str(output),
                "--render",
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            report = json.loads((output / "qa-report.json").read_text(encoding="utf-8"))
            self.assertIn(report["render"]["status"], {"passed", "not_executed"})
            self.assertEqual(report["render"]["scope"], "smoke")
            self.assertEqual(report["render"]["page_count_method"], "pdf_page_object_regex")
            self.assertEqual(report["visual_inspection"]["status"], "not_executed")
            self.assertEqual(report["visual_inspection"]["evidence_file"], "visual-inspection.json")
            if report["render"]["status"] == "passed":
                self.assertEqual(report["render"]["files_checked"], 2)
                self.assertGreater(report["render"]["page_count"], 0)
                self.assertEqual(report["render"]["reason"], "LibreOffice render smoke passed")
            self.assertEqual(list(output.glob("*.pdf")), [])

    def test_template_binaries_keep_master_hashes(self) -> None:
        def digest(path: Path) -> str:
            value = hashlib.sha256()
            value.update(path.read_bytes())
            return value.hexdigest().upper()

        self.assertEqual(digest(V10_TEMPLATE), "11783108468204DD67C9F8EAA1543B67279361ECC842A8B37F8541BDD01D16D5")
        self.assertEqual(digest(V110_TEMPLATE), "5139D6C0D48E6BA23991D4415BD1D7ED594913010B19DACC37415E946A65DE8B")
        self.assertEqual(digest(V111_TEMPLATE), "569B076DE30CD64172EE86F2123C8AE5EA67828F46B51C4280FE32DAF6DE1AD0")
        self.assertEqual(digest(V112_TEMPLATE), "6FFAFA579D3AACBC535BA624D0A6A766644868B0AF1FCE6AC37B20BA8F3D8FC1")

    def test_v112_is_a_visible_text_patch_with_identical_semantic_anchors(self) -> None:
        old_document = Document(V111_TEMPLATE)
        new_document = Document(V112_TEMPLATE)
        old_text = document_text(old_document)
        new_text = document_text(new_document)
        self.assertIn("平凡又不平凡的价值观", old_text)
        self.assertIn("工程伦理", old_text)
        self.assertNotIn("平凡又不平凡的价值观", new_text)
        self.assertNotIn("工程伦理", new_text)
        self.assertIn("职业价值观", new_text)
        self.assertIn("职业伦理", new_text)

        old_manifest = load_manifest(V111_MANIFEST)
        new_manifest = load_manifest(V112_MANIFEST)
        old_names = required_bookmarks(old_manifest)
        self.assertEqual(old_names, required_bookmarks(new_manifest))
        for name in old_names:
            old_record = find_bookmark(old_document, name)
            new_record = find_bookmark(new_document, name)
            self.assertIsNotNone(old_record)
            self.assertIsNotNone(new_record)
            self.assertEqual(old_record.bookmark_id, new_record.bookmark_id)
            self.assertEqual(
                bookmark_boundary_locations(old_document, name),
                bookmark_boundary_locations(new_document, name),
            )

        result = run_script(
            LESSON / "scripts" / "validate_template.py",
            "--template",
            str(V112_TEMPLATE),
            "--manifest",
            str(V112_MANIFEST),
            "--json",
        )
        self.assertEqual(result.returncode, 0, result.stdout + result.stderr)
        report = json.loads(result.stdout)
        self.assertEqual(report["template_version"], "1.1.2")
        self.assertEqual(report["checks"]["visible_text_replacements"][0]["from"], "平凡又不平凡的价值观")
        self.assertEqual(report["errors"], [])

    def test_density_failure_reports_actual_chars_and_limit_without_truncation(self) -> None:
        payload = load_fixture("lesson-plan-input.json")
        payload["lessons"] = [payload["lessons"][0]]
        payload["total_hours"] = 2
        payload["lessons"][0]["teaching_content"] = [
            f"教学内容{i}：围绕测试范围、证据链和交付要求展开" * 8
            for i in range(8)
        ]
        manifest = load_manifest(V111_MANIFEST)
        report = assess_content_quality(payload, manifest)
        self.assertEqual(report["status"], "failed")
        teaching_errors = [item for item in report["density_errors"] if item["field"] == "teaching_content"]
        self.assertTrue(teaching_errors)
        self.assertGreater(teaching_errors[0]["actual_chars"], teaching_errors[0]["limit"])
        self.assertTrue(any("actual_chars=" in message and "limit=" in message for message in report["errors"]))
        with self.assertRaises(ContentQualityError):
            validate_content_quality(payload, manifest)
